from __future__ import annotations

import base64
import html
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import streamlit as st
from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


W_NS = (
    "http://schemas.openxmlformats.org/"
    "wordprocessingml/2006/main"
)

M_NS = (
    "http://schemas.openxmlformats.org/"
    "officeDocument/2006/math"
)

A_NS = (
    "http://schemas.openxmlformats.org/"
    "drawingml/2006/main"
)

R_NS = (
    "http://schemas.openxmlformats.org/"
    "officeDocument/2006/relationships"
)


def iter_block_items(parent):
    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body
    else:
        parent_element = parent._tc

    for child in parent_element.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)

        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def css_length_twips(value):
    if value is None:
        return None

    try:
        points = value.pt
    except Exception:
        return None

    return f"{points:.1f}pt"


def paragraph_alignment(paragraph):
    alignment = paragraph.alignment

    if alignment is None:
        return "left"

    value = str(alignment).upper()

    if "CENTER" in value:
        return "center"

    if "RIGHT" in value:
        return "right"

    if "JUSTIFY" in value:
        return "justify"

    return "left"


def run_to_html(run):
    text = html.escape(
        run.text or ""
    ).replace(
        "\n",
        "<br>"
    )

    if not text:
        return ""

    if run.bold:
        text = f"<strong>{text}</strong>"

    if run.italic:
        text = f"<em>{text}</em>"

    if run.underline:
        text = f"<u>{text}</u>"

    if run.font.superscript:
        text = f"<sup>{text}</sup>"

    if run.font.subscript:
        text = f"<sub>{text}</sub>"

    styles = []

    if run.font.name:
        styles.append(
            f"font-family:{html.escape(run.font.name)}"
        )

    if run.font.size:
        styles.append(
            f"font-size:{run.font.size.pt:.1f}pt"
        )

    if styles:
        text = (
            '<span style="'
            + ";".join(styles)
            + '">'
            + text
            + "</span>"
        )

    return text


def paragraph_has_math(paragraph):
    return bool(
        paragraph._p.xpath(
            ".//m:oMath"
        )
    )


def math_text_from_paragraph(paragraph):
    nodes = paragraph._p.xpath(
        ".//m:oMath//m:t"
    )

    return "".join(
        node.text or ""
        for node in nodes
    )


def paragraph_to_html(paragraph):
    fmt = paragraph.paragraph_format

    styles = [
        f"text-align:{paragraph_alignment(paragraph)}"
    ]

    if fmt.left_indent:
        value = css_length_twips(
            fmt.left_indent
        )
        if value:
            styles.append(
                f"margin-left:{value}"
            )

    if fmt.right_indent:
        value = css_length_twips(
            fmt.right_indent
        )
        if value:
            styles.append(
                f"margin-right:{value}"
            )

    if fmt.space_before:
        value = css_length_twips(
            fmt.space_before
        )
        if value:
            styles.append(
                f"margin-top:{value}"
            )

    if fmt.space_after:
        value = css_length_twips(
            fmt.space_after
        )
        if value:
            styles.append(
                f"margin-bottom:{value}"
            )

    content = "".join(
        run_to_html(run)
        for run in paragraph.runs
    )

    if paragraph_has_math(paragraph):
        math_text = html.escape(
            math_text_from_paragraph(
                paragraph
            )
        )

        if math_text:
            content += (
                '<span class="math-block">'
                + math_text
                + "</span>"
            )

    if not content.strip():
        content = "&nbsp;"

    style_name = (
        paragraph.style.name
        if paragraph.style
        else ""
    )

    css_class = "doc-paragraph"

    if (
        style_name
        and "heading" in style_name.lower()
    ):
        css_class += " doc-heading"

    return (
        f'<p class="{css_class}" '
        f'style="{";".join(styles)}">'
        f"{content}"
        "</p>"
    )


def cell_to_html(cell):
    content = "".join(
        paragraph_to_html(p)
        for p in cell.paragraphs
    )

    return (
        '<td class="doc-cell">'
        + content
        + "</td>"
    )


def table_to_html(table):
    rows = []

    for row in table.rows:
        cells = "".join(
            cell_to_html(cell)
            for cell in row.cells
        )

        rows.append(
            "<tr>"
            + cells
            + "</tr>"
        )

    return (
        '<table class="doc-table">'
        + "".join(rows)
        + "</table>"
    )


def extract_images(docx_bytes):
    images = []

    with ZipFile(
        BytesIO(docx_bytes)
    ) as archive:

        for name in archive.namelist():
            if not name.startswith(
                "word/media/"
            ):
                continue

            data = archive.read(name)

            suffix = Path(name).suffix.lower()

            mime = {
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".gif": "image/gif",
                ".bmp": "image/bmp",
            }.get(
                suffix,
                "application/octet-stream",
            )

            encoded = base64.b64encode(
                data
            ).decode("ascii")

            images.append(
                (
                    Path(name).name,
                    mime,
                    encoded,
                )
            )

    return images


def count_omml(docx_bytes):
    with ZipFile(
        BytesIO(docx_bytes)
    ) as archive:
        xml = archive.read(
            "word/document.xml"
        )

    root = etree.fromstring(xml)

    return len(
        root.xpath(
            ".//m:oMath",
            namespaces={
                "m": M_NS
            },
        )
    )


def build_document_html(
    docx_bytes,
):
    document = Document(
        BytesIO(docx_bytes)
    )

    blocks = []

    for block in iter_block_items(
        document
    ):
        if isinstance(
            block,
            Paragraph,
        ):
            blocks.append(
                paragraph_to_html(
                    block
                )
            )

        elif isinstance(
            block,
            Table,
        ):
            blocks.append(
                table_to_html(
                    block
                )
            )

    images = extract_images(
        docx_bytes
    )

    if images:
        image_html = [
            '<div class="image-section">',
            '<div class="image-title">'
            "Hình ảnh trong tài liệu"
            "</div>",
        ]

        for (
            name,
            mime,
            encoded,
        ) in images:
            image_html.append(
                '<div class="doc-image">'
                f'<img src="data:{mime};base64,{encoded}" '
                'style="max-width:100%;height:auto;">'
                f'<div class="image-caption">{html.escape(name)}</div>'
                "</div>"
            )

        image_html.append(
            "</div>"
        )

        blocks.extend(
            image_html
        )

    body = "\n".join(
        blocks
    )

    return f"""
<!doctype html>
<html>
<head>
<meta charset="utf-8">
<style>

body {{
    margin: 0;
    padding: 28px;
    background: #e9edf2;
    font-family: "Times New Roman", serif;
}}

.document-shell {{
    max-width: 900px;
    margin: 0 auto;
}}

.word-page {{
    background: white;
    min-height: 1120px;
    padding: 70px 75px;
    box-sizing: border-box;
    box-shadow:
        0 2px 8px rgba(0,0,0,.12);
    border: 1px solid #d5d9df;
}}

.doc-paragraph {{
    font-size: 14pt;
    line-height: 1.35;
    margin-top: 0;
    margin-bottom: 6pt;
    white-space: normal;
}}

.doc-heading {{
    font-weight: 700;
}}

.doc-table {{
    width: 100%;
    border-collapse: collapse;
    margin: 12px 0;
}}

.doc-cell {{
    border: 1px solid #444;
    padding: 5px 8px;
    vertical-align: top;
}}

.doc-cell p {{
    margin: 0 0 4px 0;
}}

.math-block {{
    display: inline-block;
    margin: 0 3px;
    padding: 1px 4px;
    border-radius: 3px;
    background: #f6f7f8;
    font-family: "Cambria Math", serif;
}}

.image-section {{
    margin-top: 28px;
    border-top: 1px solid #ddd;
    padding-top: 15px;
}}

.image-title {{
    font-family: Arial, sans-serif;
    font-size: 13px;
    font-weight: bold;
    margin-bottom: 12px;
}}

.doc-image {{
    margin: 12px 0 22px 0;
    text-align: center;
}}

.image-caption {{
    font-family: Arial, sans-serif;
    font-size: 11px;
    color: #666;
    margin-top: 4px;
}}

</style>
</head>

<body>
<div class="document-shell">
<div class="word-page">
{body}
</div>
</div>
</body>
</html>
"""


def main():
    st.set_page_config(
        page_title="Visual Lesson Plan Viewer",
        layout="wide",
    )

    st.title(
        "Visual Lesson Plan Viewer"
    )

    st.caption(
        "Bản thử nghiệm hiển thị trực quan giáo án Word. "
        "Chưa chỉnh sửa tài liệu."
    )

    uploaded = st.file_uploader(
        "Chọn giáo án Word",
        type=("docx",),
        accept_multiple_files=False,
    )

    if uploaded is None:
        st.info(
            "Hãy tải lên một file .docx."
        )
        st.stop()

    content = uploaded.getvalue()

    document = Document(
        BytesIO(content)
    )

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        "Đoạn văn",
        len(document.paragraphs),
    )

    col2.metric(
        "Bảng",
        len(document.tables),
    )

    col3.metric(
        "Hình",
        len(document.inline_shapes),
    )

    col4.metric(
        "Công thức OMML",
        count_omml(content),
    )

    st.divider()

    st.subheader(
        "Xem toàn bộ giáo án"
    )

    viewer_html = build_document_html(
        content
    )

    st.components.v1.html(
        viewer_html,
        height=1200,
        scrolling=True,
    )

    st.caption(
        "Viewer hiện chưa thay đổi file gốc."
    )


if __name__ == "__main__":
    main()
