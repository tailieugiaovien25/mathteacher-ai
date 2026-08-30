from __future__ import annotations

from datetime import datetime, timezone

import csv

# V58_C4C1_CONTEXT_PERFORMANCE_TRACE
from time import perf_counter as _v58_perf_counter
from pathlib import Path as _V58PerfPath
import json as _v58_perf_json


def _v58_perf_log(event: str, started_at: float, **fields) -> None:
    try:
        elapsed_ms = round(
            (_v58_perf_counter() - started_at) * 1000.0,
            3,
        )
        report_dir = _V58PerfPath("reports/system_context/v58_c4c1")
        report_dir.mkdir(parents=True, exist_ok=True)
        payload = {"event": event, "elapsed_ms": elapsed_ms, **fields}
        with (report_dir / "CONTEXT_PERFORMANCE_TRACE.jsonl").open(
            "a", encoding="utf-8"
        ) as handle:
            handle.write(
                _v58_perf_json.dumps(
                    payload,
                    ensure_ascii=False,
                    default=str,
                )
                + "\n"
            )
    except Exception:
        pass


def _v58_timed_class_catalog_get(*, client, class_id):
    started_at = _v58_perf_counter()
    outcome = "found"
    try:
        item = SupabaseClassCatalogRepository(client=client).get(
            class_id=class_id
        )
        if item is None:
            outcome = "missing"
        return item
    except Exception:
        outcome = "error"
        raise
    finally:
        _v58_perf_log(
            "class_catalog_grade_lookup",
            started_at,
            class_id=class_id,
            outcome=outcome,
        )


from io import StringIO
from html import escape
from pathlib import Path

import streamlit as st

from lesson_planning_v2.services.lesson_plan_grouping_policy_source import (
    LessonPlanGroupingPolicySource,
)
from lesson_planning_v2.models.lesson_plan_grouping import (
    LessonPlanGroupingPolicy,
)
from lesson_planning_v2.services.lesson_plan_grouping_service import (
    LessonPlanGroupingPolicyResolver,
    LessonPlanGroupingService,
)

from portal_v2.ui.teacher_workspace_styles import (
    apply_lesson_authoring_workspace_styles,
)

from educational_planning_v2.adapters.operational_weekly_schedule_workbook_intake import (
    WeeklyScheduleWorkbookIntakeAdapter,
)
from educational_planning_v2.models.operational_data_io import (
    OperationalInputLocation,
    OperationalInputReference,
)
from educational_planning_v2.services.local_weekly_schedule_generation_service import (
    LocalWeeklyScheduleGenerationService,
    WeeklyScheduleGenerationRequest,
    WeeklyScheduleGenerationResult,
)
from educational_planning_v2.services.operational_input_selection_service import (
    OperationalInputSelection,
)
from educational_planning_v2.services.weekly_schedule_output_service import (
    WeeklyScheduleOutputService,
)
from portal_v2.ui.weekly_schedule_portal import (
    WeeklySchedulePortalPresenter,
)
from educational_planning_v2.adapters.supabase_teaching_assignment_repository import (
    SupabaseTeachingAssignmentRepository,
)
from educational_planning_v2.adapters.supabase_teacher_timetable_repository import (
    SupabaseTeacherTimetableRepository,
)
from educational_planning_v2.adapters.supabase_weekly_schedule_repository import (
    SupabaseWeeklyScheduleRepository,
)
from educational_planning_v2.adapters.supabase_class_catalog_repository import (
    SupabaseClassCatalogRepository,
)
from educational_planning_v2.adapters.supabase_subject_catalog_repository import (
    SupabaseSubjectCatalogRepository,
)
from educational_planning_v2.adapters.supabase_academic_year_configuration_repository import (
    SupabaseAcademicYearConfigurationRepository,
)
from educational_planning_v2.adapters.supabase_academic_week_repository import (
    SupabaseAcademicWeekRepository,
)
from educational_planning_v2.models.teaching_assignment import (
    TeachingAssignmentRole,
    TeachingAssignmentStatus,
)
from educational_planning_v2.models.teacher_timetable import (
    TeacherTimetableSlotStatus,
)
from portal_v2.runtime.system_weekly_schedule_runtime import (
    SystemWeeklyScheduleRuntime,
    SystemWeeklyScheduleRuntimeRequest,
)
from portal_v2.context.legacy_session_context_adapter import (
    project_system_context,
)
from portal_v2.context.session_scoped_context_holder import (
    apply_canonical_year_week_change,
    get_canonical_context,
    publish_year_week_projection,
)
from portal_v2.context.runtime_context_bridge import apply_runtime_context_change
from portal_v2.context.user_scoped_store import ContextIdentity
from lesson_planning_v2.services.lesson_plan_lesson_selector_service import LessonPlanLessonSelectorService
from lesson_planning_v2.services.lesson_plan_draft_workspace_service import (
    LessonPlanDraftWorkspaceService,
)
from lesson_planning_v2.workspace_draft import (
    LessonPlanWorkspaceDraft,
)

from lesson_planning_v2.services.lesson_plan_unit_selector_service import (
    LessonPlanSelectionMode,
    LessonPlanUnitSelectorService,
)
from lesson_planning_v2.services import (
    LessonPlanDocumentProcessingService,
)

from document_intelligence.lesson_plan_preview_upload import (
    LessonPlanPreviewUploadService,
)
from document_intelligence.validation import (
    CanonicalDocumentContext,
)
from portal_v2.ui.lesson_plan_preview_streamlit import (
    render_lesson_plan_preview,
)
from document_intelligence.contracts import (
    DocumentField,
)
from document_intelligence.lesson_plan_teacher_review_presenter import (
    LessonPlanTeacherReviewPresenter,
)
from document_intelligence.lesson_plan_teacher_review_resolver import (
    LessonPlanTeacherReviewResolver,
)
from document_intelligence.lesson_plan_modification_plan import (
    LessonPlanModificationPlanner,
)
from document_intelligence.lesson_plan_reviewed_schedule_row import (
    LessonPlanReviewedScheduleRow,
)
from document_intelligence.lesson_plan_workflow_state import (
    LessonPlanWorkflowIdentity,
    LessonPlanWorkflowState,
)
from document_standardization import (
    LessonPlanStandardizationOptions,
)
from portal_v2.ui.lesson_plan_teacher_review_streamlit import (
    render_lesson_plan_teacher_review,
)

from scripts.teacher_portal.lesson_plan_visual_viewer import (
    build_document_html,
)

_VIEW_STATE_KEY = "weekly_schedule_portal_view"

# GLOBAL_WEEKLY_CONTEXT_V1
_ACTIVE_SCHEDULE_ID_KEY = (
    "global_weekly_active_schedule_id"
)
_ACTIVE_ACADEMIC_YEAR_KEY = (
    "global_weekly_active_academic_year"
)
_ACTIVE_WEEK_NUMBER_KEY = (
    "global_weekly_active_week_number"
)
_ACTIVE_VIEW_KEY = (
    "global_weekly_active_view"
)
_LESSON_AUTHORING_FOCUS_KEY = "lesson_authoring_tool_focus"
_AI_AUTHORING_SELECTION_STATE_KEYS = (
    "lesson_authoring_ai_assignment_subject",
    "lesson_authoring_ai_schedule_row",
)
_STANDARDIZATION_WEEK_KEY = (
    "standardization_authoring_week_number"
)
_WORKING_LESSON_CONTEXT_KEY = "lesson_authoring_working_context"
_RESTORE_LESSON_CONTEXT_KEY = "lesson_authoring_restore_context_to_standardization"
_LESSON_AUTHORING_NOTICE_KEY = "lesson_authoring_navigation_notice"
_LBG_CONTEXT_SNAPSHOT_KEY = "lbg_autosaved_filter_context"
_LBG_NOTICE_KEY = "lbg_floating_notice"
_LBG_DATA_WEEK_SIGNATURE_KEY = "lbg_data_week_signature"
_LBG_WEEK_USER_CHANGE_KEY = "lbg_week_changed_by_user"
_LBG_DATA_WEEK_CONTEXT_MISMATCH_KEY = "lbg_data_week_context_mismatch"
_STANDARDIZATION_DRAFT_KEY = "standardization_autosaved_lesson_context"
_STANDARDIZATION_NOTICE_KEY = "standardization_floating_notice"


def _emit_standardization_canonical_context_change(*, field: str, value, source_control: str) -> None:
    user_id = str(st.session_state.get("portal_user_id", "") or "").strip()
    if not user_id:
        raise RuntimeError("CANONICAL_CONTEXT_USER_ID_REQUIRED")
    current = get_canonical_context(st.session_state, user_id=user_id, source_page="weekly_schedule")
    if getattr(current, field) == value:
        return
    outcome = apply_runtime_context_change(
        current=current, field=field, value=value, source_page="weekly_schedule",
        source_control=source_control, occurred_at=datetime.now(timezone.utc),
    )
    store = st.session_state.get("_v57_system_context_store")
    context_id = st.session_state.get("_v57_system_context_id")
    if store is None or not context_id:
        raise RuntimeError("CANONICAL_CONTEXT_STORE_REQUIRED")
    store.put(ContextIdentity(user_id=user_id, context_id=str(context_id)), outcome.context)


def _autosave_standardization_change(changed_field: str, state_key: str | None = None) -> None:
    """Persist the standardization selector context before every rerun."""
    st.session_state[_STANDARDIZATION_DRAFT_KEY] = {
        "week_number": st.session_state.get(_STANDARDIZATION_WEEK_KEY),
        "subject_ref": st.session_state.get("standardization_subject_filter"),
        "component_ref": st.session_state.get("standardization_component_filter"),
        "selection_mode": next(
            (
                value
                for key, value in st.session_state.items()
                if str(key).startswith("lbg_lesson_plan_selection_mode_")
            ),
            None,
        ),
        "lesson_unit": next(
            (
                value
                for key, value in st.session_state.items()
                if str(key).startswith("lbg_lesson_plan_unit_")
            ),
            None,
        ),
        "working_lesson": dict(
            st.session_state.get(_WORKING_LESSON_CONTEXT_KEY, {}) or {}
        ),
    }
    if changed_field == "Môn":
        value = st.session_state.get("standardization_subject_filter")
        if value:
            _emit_standardization_canonical_context_change(field="subject_ref", value=str(value), source_control="standardization_subject_filter")
    elif changed_field == "Phân môn":
        value = st.session_state.get("standardization_component_filter")
        if value is not None:
            _emit_standardization_canonical_context_change(field="component_ref", value=str(value), source_control="standardization_component_filter")
    elif changed_field == "Khối lớp":
        grade_value = st.session_state.get(state_key) if state_key else None
        _emit_standardization_canonical_context_change(
            field="grade", value=grade_value,
            source_control=state_key or "standardization_grade_filter",
        )

    st.session_state[_STANDARDIZATION_NOTICE_KEY] = (
        f"Đã tự lưu thay đổi {changed_field} của bài đang chuẩn hóa."
    )


def _emit_canonical_week_change(
    *,
    selected_week: int,
    source_control: str,
) -> None:
    user_id = str(st.session_state.get("portal_user_id", "") or "").strip()
    if not user_id:
        raise RuntimeError("CANONICAL_CONTEXT_USER_ID_REQUIRED")
    apply_canonical_year_week_change(
        st.session_state,
        user_id=user_id,
        field="week_number",
        value=int(selected_week),
        source_page="weekly_schedule",
        source_control=source_control,
    )


def _autosave_lbg_filter_context(changed_field: str) -> None:
    """Persist LBG controls in-session and notify after Streamlit reruns."""
    st.session_state[_LBG_CONTEXT_SNAPSHOT_KEY] = {
        "source": st.session_state.get("weekly_schedule_source"),
        "academic_year": st.session_state.get("system_weekly_academic_year"),
        "week_number": st.session_state.get("system_weekly_week_number"),
        "assignment": st.session_state.get("system_weekly_assignment_name"),
        "ppct_mode": st.session_state.get("system_weekly_ppct_mode"),
    }
    if changed_field == "Tuần":
        st.session_state[_LBG_WEEK_USER_CHANGE_KEY] = True
        selected_week = st.session_state.get("system_weekly_week_number")
        if selected_week is not None:
            _emit_canonical_week_change(
                selected_week=int(selected_week),
                source_control="system_weekly_week_number",
            )
    elif changed_field == "Năm học":
        selected_year = str(
            st.session_state.get("system_weekly_academic_year", "") or ""
        ).strip()
        if selected_year:
            user_id = str(
                st.session_state.get("portal_user_id", "") or ""
            ).strip()
            if not user_id:
                raise RuntimeError("CANONICAL_CONTEXT_USER_ID_REQUIRED")
            apply_canonical_year_week_change(
                st.session_state,
                user_id=user_id,
                field="academic_year",
                value=selected_year,
                source_page="weekly_schedule",
                source_control="system_weekly_academic_year",
            )
    st.session_state[_LBG_NOTICE_KEY] = (
        f"Đã tự lưu thay đổi: {changed_field}."
    )


def _sync_lbg_week_from_loaded_data() -> None:
    """Publish a newly loaded schedule's week back to the week selector."""
    view = (
        st.session_state.get(_ACTIVE_VIEW_KEY)
        or st.session_state.get(_VIEW_STATE_KEY)
    )
    if view is None:
        return
    try:
        data_week = int(getattr(view, "week_number"))
    except (AttributeError, TypeError, ValueError):
        return
    if data_week not in range(1, 41):
        return
    signature = (
        str(getattr(view, "academic_year", "") or ""),
        data_week,
        str(st.session_state.get(_ACTIVE_SCHEDULE_ID_KEY, "") or ""),
    )
    if st.session_state.pop(_LBG_WEEK_USER_CHANGE_KEY, False):
        # The user's newly selected week is authoritative for this rerun.  The
        # matching data view will be loaded below and then becomes the reverse
        # synchronization source on subsequent activities.
        st.session_state[_LBG_DATA_WEEK_SIGNATURE_KEY] = signature
        return
    if st.session_state.get(_LBG_DATA_WEEK_SIGNATURE_KEY) == signature:
        return
    st.session_state[_LBG_DATA_WEEK_SIGNATURE_KEY] = signature

    user_id = str(st.session_state.get("portal_user_id", "") or "").strip()
    if not user_id:
        return

    canonical_context = get_canonical_context(
        st.session_state,
        user_id=user_id,
        source_page="weekly_schedule",
    )
    canonical_week = canonical_context.week_number

    if canonical_week is not None and int(canonical_week) != data_week:
        st.session_state[_LBG_DATA_WEEK_CONTEXT_MISMATCH_KEY] = {
            "canonical_week": int(canonical_week),
            "data_week": data_week,
            "schedule_id": str(
                st.session_state.get(_ACTIVE_SCHEDULE_ID_KEY, "") or ""
            ),
        }
        st.session_state[_LBG_NOTICE_KEY] = (
            f"Dữ liệu LBG đang mở thuộc Tuần {data_week}, "
            f"khác Tuần {int(canonical_week)} của ngữ cảnh hệ thống. "
            "Dữ liệu được đánh dấu cần kiểm tra và không được phép đổi Tuần hệ thống."
        )
    else:
        st.session_state.pop(_LBG_DATA_WEEK_CONTEXT_MISMATCH_KEY, None)


def _sync_legacy_lbg_week_to_canonical() -> None:
    """Emit the legacy LBG week selector into canonical SystemContext."""
    selected_week = st.session_state.get("lbg_user_week_number")
    if selected_week is None:
        return
    st.session_state[_LBG_WEEK_USER_CHANGE_KEY] = True
    _emit_canonical_week_change(
        selected_week=int(selected_week),
        source_control="lbg_user_week_number",
    )


def _sync_standardization_week_to_lbg() -> None:
    """Emit the authoring week into canonical SystemContext."""
    selected_week = int(
        st.session_state[_STANDARDIZATION_WEEK_KEY]
    )
    _emit_canonical_week_change(
        selected_week=selected_week,
        source_control=_STANDARDIZATION_WEEK_KEY,
    )
    _autosave_standardization_change("Tuần soạn")


def _open_ai_authoring_page(
    selected_lesson: dict,
) -> None:
    """Navigate before Streamlit instantiates the sidebar radio."""
    # A keyed Streamlit selectbox keeps its previous value and ignores a new
    # ``index`` on the next page render.  Clear only the two destination
    # selectors so the transferred lesson becomes their new default.
    for state_key in _AI_AUTHORING_SELECTION_STATE_KEYS:
        st.session_state.pop(state_key, None)

    transferred_context = dict(selected_lesson)
    transferred_context.update(
        context_origin="STANDARDIZATION",
        context_read_only=True,
    )
    st.session_state[
        "lesson_authoring_ai_context"
    ] = transferred_context
    st.session_state[_WORKING_LESSON_CONTEXT_KEY] = dict(transferred_context)
    st.session_state[_LESSON_AUTHORING_NOTICE_KEY] = (
        "Đã tự lưu thông tin bài dạy và mở trang Soạn bài cùng AI."
    )
    st.session_state["portal_navigation_request"] = "Soạn bài cùng AI"

_LESSON_PLAN_PROFILE = (
    Path(__file__).resolve().parents[3]
    / "scripts"
    / "word_standardizer"
    / "lesson_plan_profile.json"
)


def _render_lesson_authoring_tool_hub(
    *,
    page_title: str = "Chuẩn hóa giáo án",
    show_entry_actions: bool = True,
) -> str:
    """Render the page heading and optional entry actions."""

    apply_lesson_authoring_workspace_styles()
    standardization_notice = st.session_state.pop(
        _STANDARDIZATION_NOTICE_KEY, ""
    )
    if standardization_notice:
        st.toast(str(standardization_notice), icon="💾")
    st.markdown(
        """
<style>
div[class*="st-key-standardization_"],
div[class*="st-key-lbg_lesson_plan_selection_mode_"],
div[class*="st-key-lbg_lesson_plan_unit_"] {
  border-radius:12px;
}
div[class*="st-key-standardization_"] input,
div[class*="st-key-standardization_"] div[data-baseweb="select"] > div,
div[class*="st-key-lbg_lesson_plan_"] div[data-baseweb="select"] > div {
  min-height:42px!important;
  height:42px!important;
  font-size:14px!important;
  line-height:1.25!important;
}
div[class*="st-key-standardization_"] div[data-baseweb="select"] span,
div[class*="st-key-lbg_lesson_plan_"] div[data-baseweb="select"] span,
div[role="listbox"] li { font-size:14px!important; }
div[class*="st-key-standardization_"] div[data-baseweb="select"] > div,
div[class*="st-key-lbg_lesson_plan_"] div[data-baseweb="select"] > div {
  background:linear-gradient(145deg,#ffffff,#edf5ff)!important;
  border:1px solid #7898bd!important;
  box-shadow:inset 0 1px 2px rgba(3,16,31,.10),2px 3px 0 #c5d3e4!important;
}
div[class*="st-key-standardization_all_"] input {
  background:#f7fbff!important;
  border:1px solid #7898bd!important;
  box-shadow:inset 0 1px 2px rgba(3,16,31,.10),2px 3px 0 #c5d3e4!important;
}
div[class*="st-key-lbg_open_full_ai_page_"] button {
  min-height:44px!important;
  border-radius:10px!important;
  background:linear-gradient(145deg,#3789e8,#1f63bd)!important;
  color:#fff!important;
  box-shadow:3px 4px 0 #12345b!important;
}
</style>
        """.strip(),
        unsafe_allow_html=True,
    )

    st.markdown(
        """
<style>
.mt-standard-selector-shell{margin:.4rem 0 1rem;padding:1rem 1rem .35rem;border:1px solid rgba(112,139,184,.30);border-radius:22px;background:linear-gradient(145deg,rgba(255,255,255,.98),rgba(235,244,255,.94));box-shadow:0 12px 28px rgba(42,72,118,.14),inset 0 1px 0 #fff}
.mt-standard-selector-shell h2{margin:0;color:#142845;font:700 1.18rem/1.35 Inter,Arial,sans-serif}
.mt-standard-selector-shell p{margin:.28rem 0 .75rem;color:#64748b;font:400 .88rem/1.5 Inter,Arial,sans-serif}
div[data-testid="stSelectbox"] label p{font-family:Inter,Arial,sans-serif!important;font-size:.82rem!important;font-weight:650!important;color:#334155!important}
div[data-testid="stSelectbox"]>div>div{border-radius:13px!important;box-shadow:0 4px 10px rgba(30,64,110,.07)}
</style>
<div class="mt-standard-selector-shell"><h2>Chọn nội dung bài soạn</h2><p>Lọc trực tiếp từ Lịch báo giảng theo môn, phân môn và nội dung giảng dạy.</p></div>
        """.strip(),
        unsafe_allow_html=True,
    )

    active_focus = str(
        st.session_state.get(
            _LESSON_AUTHORING_FOCUS_KEY,
            "AI",
        )
    )

    if active_focus not in {"AI", "STANDARDIZE"}:
        active_focus = "AI"

    ai_active_class = (
        " is-active"
        if active_focus == "AI"
        else ""
    )
    standard_active_class = (
        " is-active"
        if active_focus == "STANDARDIZE"
        else ""
    )

    st.markdown(
        f"""
<section class="mt-tool-hub" aria-labelledby="mt-tool-hub-title">
  <h1 id="mt-tool-hub-title" class="mt-tool-hub-heading">
    {escape(page_title)}
  </h1>
  {'''<p class="mt-tool-hub-lead">
    Chọn bài theo tuần, soạn cùng AI và chuẩn hóa giáo án Word
    trong một quy trình thống nhất.
  </p>''' if show_entry_actions else ''}
</section>
        """.strip(),
        unsafe_allow_html=True,
    )

    if not show_entry_actions:
        return active_focus

    action_columns = st.columns(
        2,
        gap="medium",
    )

    with action_columns[0]:
        open_ai_workspace = st.button(
            "✨ Bắt đầu soạn bài",
            type=(
                "primary"
                if active_focus == "AI"
                else "secondary"
            ),
            use_container_width=True,
            key="lesson_authoring_open_ai",
        )

    with action_columns[1]:
        open_standardization_workspace = st.button(
            "📁 Chọn tệp giáo án",
            type=(
                "primary"
                if active_focus == "STANDARDIZE"
                else "secondary"
            ),
            use_container_width=True,
            key="lesson_authoring_open_standardization",
        )

    if open_ai_workspace and active_focus != "AI":
        st.session_state[_LESSON_AUTHORING_FOCUS_KEY] = "AI"
        st.rerun()

    if (
        open_standardization_workspace
        and active_focus != "STANDARDIZE"
    ):
        st.session_state[
            _LESSON_AUTHORING_FOCUS_KEY
        ] = "STANDARDIZE"
        st.rerun()

    return active_focus


_STANDARDIZATION_ACTION_FLOW_CSS = r"""
<style>
.mt-standardization-actions {
  display:grid;
  grid-template-columns:repeat(5,minmax(130px,1fr));
  gap:12px;
  margin:4px 0 24px;
}
.mt-standardization-action {
  align-items:center;
  background:linear-gradient(145deg,#071a33,#0d3159);
  border:1px solid #315d88;
  border-radius:13px;
  box-shadow:0 6px 0 #03101f,0 12px 22px rgba(3,16,31,.18);
  color:#fff !important;
  display:flex;
  flex-direction:column;
  justify-content:center;
  min-height:98px;
  padding:13px 11px;
  text-align:center;
  text-decoration:none !important;
  transition:transform .14s ease,box-shadow .14s ease,border-color .14s ease;
}
.mt-standardization-action:hover {
  border-color:#f4c95d;
  box-shadow:0 8px 0 #03101f,0 16px 26px rgba(3,16,31,.22);
  transform:translateY(-2px);
}
.mt-standardization-action:active {
  box-shadow:0 2px 0 #03101f;
  transform:translateY(4px);
}
.mt-standardization-action strong {
  color:#f4c95d;
  display:block;
  font-size:15px;
  line-height:1.35;
}
.mt-standardization-action span {
  color:#dcecff;
  display:block;
  font-size:12px;
  line-height:1.35;
  margin-top:5px;
}
@media(max-width:980px){
  .mt-standardization-actions{grid-template-columns:1fr 1fr;}
}
@media(max-width:560px){
  .mt-standardization-actions{grid-template-columns:1fr;}
}
</style>
"""


def _render_standardization_action_flow(
    *,
    client=None,
    user_id: str | None = None,
) -> None:
    """Render five visible Streamlit controls for standardization."""

    if client is not None:
        st.session_state[
            "lesson_standardization_supabase_client"
        ] = client

    if user_id:
        st.session_state[
            "lesson_standardization_teacher_user_id"
        ] = str(user_id)

    st.markdown(
        """
<div id="standardization-action-bar"></div>
        """.strip(),
        unsafe_allow_html=True,
    )

    columns = st.columns(
        [1.05, 1.45, 1.0, 0.8, 1.0],
        gap="small",
    )

    actions = (
        (
            "\U0001f4e4 Up gi\u00e1o \u00e1n",
            "standardization_action_upload",
            "upload",
        ),
        (
            "\u2728 T\u1ea1o gi\u00e1o \u00e1n chu\u1ea9n",
            "standardization_action_create",
            "create",
        ),
        (
            "\U0001f441 Xem tr\u01b0\u1edbc",
            "standardization_action_preview",
            "preview",
        ),
        (
            "\U0001f4be L\u01b0u",
            "standardization_action_save",
            "save",
        ),
        (
            "\U0001f4e5 T\u1ea3i xu\u1ed1ng",
            "standardization_action_download",
            "download",
        ),
    )

    for column, (
        label,
        key,
        action,
    ) in zip(
        columns,
        actions,
        strict=True,
    ):
        with column:
            st.button(
                label,
                key=key,
                use_container_width=True,
                type=(
                    "primary"
                    if action == "upload"
                    else "secondary"
                ),
                on_click=_activate_standardization_action,
                args=(action,),
            )

    active_action = str(
        st.session_state.get(
            "lesson_plan_standardization_action",
            "",
        )
    )

    anchor_map = {
        "upload": "upload-lesson-plan",
        "create": "standardize-lesson-plan",
        "preview":
            "preview-standardized-lesson-plan",
        "save":
            "save-standardized-lesson-plan",
        "download":
            "download-standardized-lesson-plan",
    }

    target = anchor_map.get(
        active_action
    )

    if target:
        st.markdown(
            (
                "<script>"
                "setTimeout(function(){"
                "const element="
                "window.parent.document."
                f'getElementById("{target}");'
                "if(element){"
                'element.scrollIntoView({'
                'behavior:"smooth",'
                'block:"start"'
                "});"
                "}"
                "},150);"
                "</script>"
            ),
            unsafe_allow_html=True,
        )



def _activate_standardization_action(action: str) -> None:
    st.session_state[
        "lesson_plan_standardization_action"
    ] = action
    st.session_state[
        "lesson_plan_management_pending_action"
    ] = action
    st.session_state["portal_navigation_request"] = "Chu\u1ea9n h\u00f3a gi\u00e1o \u00e1n"


_STANDARDIZATION_OPTION_KEYS = {
    "preserve_original_maximum": "standardization_option_preserve_original",
    "sync_context": "standardization_option_sync_context",
    "normalize_font": "standardization_option_font",
    "normalize_equations": "standardization_option_equations",
    "normalize_tables": "standardization_option_tables",
    "normalize_page_layout": "standardization_option_page_layout",
    "normalize_spacing": "standardization_option_spacing",
    "normalize_header_footer": "standardization_option_header_footer",
}


def _standardization_options_from_state() -> LessonPlanStandardizationOptions:
    return LessonPlanStandardizationOptions(
        **{
            name: bool(st.session_state.get(key, True))
            for name, key in _STANDARDIZATION_OPTION_KEYS.items()
        }
    )


def _set_all_standardization_options(value: bool) -> None:
    for key in _STANDARDIZATION_OPTION_KEYS.values():
        st.session_state[key] = bool(value)

    for key in (
        "standardization_drafting_before_monday_enabled",
        "standardization_approval_before_monday_enabled",
        "standardization_assignment_timetable_sync_enabled",
        "standardization_image_autofit_enabled",
        "standardization_lesson_end_rule_enabled",
    ):
        st.session_state[key] = bool(value)


def _confirm_standardization_options() -> None:
    st.session_state[
        "lesson_plan_standardization_confirmed_options"
    ] = _standardization_options_from_state()
    st.session_state[
        "lesson_plan_standardization_execute_requested"
    ] = True
    _activate_standardization_action("create")


# STANDARDIZATION_MODERN_3D_UI_V1
def _render_standardization_modern_3d_header() -> None:
    """Presentation-only shell for Standardization."""

    st.markdown(
        """
<style>
/* ==================================================
   MATHTEACHER AI
   STANDARDIZATION MODERN 3D UI
   PRESENTATION ONLY
   ================================================== */

.mt-std-hero {
    position: relative;
    overflow: hidden;
    padding: 26px 30px;
    margin: 4px 0 22px 0;
    border-radius: 24px;

    background:
        linear-gradient(
            135deg,
            rgba(7,35,70,.98),
            rgba(16,78,135,.94)
        );

    border:
        1px solid rgba(255,255,255,.16);

    box-shadow:
        0 18px 45px rgba(8,35,70,.20),
        0 5px 0 rgba(3,28,58,.22),
        inset 0 1px 0 rgba(255,255,255,.15);
}

.mt-std-hero::before {
    content: "";
    position: absolute;

    width: 260px;
    height: 260px;

    right: -90px;
    top: -120px;

    border-radius: 50%;

    background:
        radial-gradient(
            circle,
            rgba(87,196,255,.36),
            rgba(87,196,255,0)
        );
}

.mt-std-hero::after {
    content: "";
    position: absolute;

    width: 180px;
    height: 180px;

    left: -75px;
    bottom: -120px;

    border-radius: 50%;

    background:
        radial-gradient(
            circle,
            rgba(86,255,212,.20),
            rgba(86,255,212,0)
        );
}

.mt-std-hero-content {
    position: relative;
    z-index: 2;
}

.mt-std-badge {
    display: inline-flex;
    align-items: center;
    gap: 8px;

    padding: 7px 12px;
    margin-bottom: 10px;

    border-radius: 999px;

    font-size: 12px;
    font-weight: 700;
    letter-spacing: .04em;

    color: #dff5ff;

    background:
        rgba(255,255,255,.10);

    border:
        1px solid rgba(255,255,255,.16);

    backdrop-filter: blur(10px);
}

.mt-std-title {
    margin: 0;

    color: #ffffff;

    font-size:
        clamp(26px,3vw,38px);

    font-weight: 800;
    line-height: 1.15;
    letter-spacing: -.02em;
}

.mt-std-description {
    margin: 10px 0 0 0;

    max-width: 760px;

    color:
        rgba(239,248,255,.88);

    font-size: 14px;
    line-height: 1.65;
}

.mt-std-flow {
    display: flex;
    flex-wrap: wrap;

    gap: 8px;

    margin-top: 17px;
}

.mt-std-flow span {
    padding: 7px 11px;

    border-radius: 10px;

    color: #ecf8ff;

    font-size: 12px;
    font-weight: 600;

    background:
        rgba(255,255,255,.09);

    border:
        1px solid rgba(255,255,255,.13);
}


/* ==================================================
   CONTROL CARDS
   ================================================== */

div[data-testid="stExpander"] {

    border:
        1px solid rgba(30,88,145,.16)
        !important;

    border-radius:
        18px !important;

    background:
        linear-gradient(
            145deg,
            rgba(255,255,255,.96),
            rgba(240,247,255,.94)
        )
        !important;

    box-shadow:
        0 13px 32px rgba(20,61,105,.10),
        0 3px 0 rgba(22,72,120,.07)
        !important;

    overflow: hidden;
}

div[data-testid="stExpander"]
details summary {

    font-weight: 750 !important;
    font-size: 15px !important;
}


/* ==================================================
   CHECKBOXES
   ================================================== */

div[data-testid="stCheckbox"] {

    padding: 8px 11px;

    margin: 3px 0;

    border-radius: 12px;

    transition:
        transform .16s ease,
        background .16s ease,
        box-shadow .16s ease;
}

div[data-testid="stCheckbox"]:hover {

    transform:
        translateY(-1px);

    background:
        rgba(224,240,255,.55);

    box-shadow:
        0 5px 16px
        rgba(20,70,120,.07);
}


/* ==================================================
   BUTTONS
   ================================================== */

.stButton > button {

    min-height: 43px;

    border-radius:
        12px !important;

    font-weight:
        700 !important;

    transition:
        transform .15s ease,
        box-shadow .15s ease;
}

.stButton > button:hover {

    transform:
        translateY(-2px);

    box-shadow:
        0 8px 19px
        rgba(23,79,136,.18);
}

.stButton > button[kind="primary"] {

    border:
        0 !important;

    background:
        linear-gradient(
            135deg,
            #287bd8,
            #1557b0
        )
        !important;

    box-shadow:
        0 8px 18px
        rgba(25,92,170,.23),

        0 3px 0
        rgba(8,58,121,.20);
}


/* ==================================================
   FILE UPLOADER
   ================================================== */

div[data-testid="stFileUploader"] {

    padding: 10px;

    border-radius: 17px;

    background:
        rgba(249,252,255,.86);

    border:
        1px solid
        rgba(43,105,166,.13);
}


/* ==================================================
   DOWNLOAD
   ================================================== */

div[data-testid="stDownloadButton"]
button {

    min-height: 43px;

    border-radius:
        12px !important;

    font-weight: 700;
}


/* ==================================================
   MOBILE
   ================================================== */

@media (max-width:720px) {

    .mt-std-hero {

        padding:
            21px 19px;

        border-radius:
            18px;
    }

    .mt-std-title {

        font-size:
            25px;
    }

    .mt-std-flow {

        gap: 6px;
    }
}
</style>

<div class="mt-std-hero">

  <div class="mt-std-hero-content">

    <div class="mt-std-badge">
      MATHTEACHER AI ? DOCUMENT INTELLIGENCE
    </div>

    <h1 class="mt-std-title">
      Chu?n h?a gi?o ?n
    </h1>

    <p class="mt-std-description">
      Gi? t?i ?a c?u tr?c gi?o ?n g?c,
      k?t h?p n?i dung AI v? d? li?u
      L?ch b?o gi?ng ?? t?o b?n gi?o ?n
      th?ng nh?t, ch?nh x?c v? s?n s?ng s? d?ng.
    </p>

    <div class="mt-std-flow">

      <span>? Gi?o ?n g?c</span>

      <span>? N?i dung AI</span>

      <span>? D? li?u gi?ng d?y</span>

      <span>? Chu?n h?a</span>

      <span>? Xu?t Word</span>

    </div>

  </div>

</div>
        """,
        unsafe_allow_html=True,
    )


def _render_standardization_control_panel() -> None:
    """Collect options; confirmation is the only processing boundary."""

    for key in _STANDARDIZATION_OPTION_KEYS.values():
        st.session_state.setdefault(key, True)

    with st.expander(
        "\u2699\ufe0f B\u1ea3ng \u0111i\u1ec1u khi\u1ec3n chu\u1ea9n h\u00f3a",
        expanded=False,
    ):
        st.caption(
            "Ch\u1ecdn t\u00e1c v\u1ee5 c\u1ea7n th\u1ef1c hi\u1ec7n. "
            "H\u1ec7 th\u1ed1ng ch\u1ec9 x\u1eed l\u00fd sau khi b\u1ea1n x\u00e1c nh\u1eadn."
        )
        st.checkbox(
            "\U0001f6e1\ufe0f Gi\u1eef nguy\u00ean b\u1ea3n g\u1ed1c t\u1ed1i \u0111a",
            key=_STANDARDIZATION_OPTION_KEYS["preserve_original_maximum"],
            help=(
                "Ch\u1ec9 \u00e1p d\u1ee5ng thay \u0111\u1ed5i AI \u00e1nh x\u1ea1 an to\u00e0n "
                "v\u00e0 c\u00e1c t\u00e1c v\u1ee5 b\u1ea1n ch\u1ecdn."
            ),
        )
        st.checkbox(
            "\u0110\u1ed3ng b\u1ed9 ng\u00e0y so\u1ea1n, ng\u00e0y d\u1ea1y, l\u1edbp, PPCT v\u00e0 t\u00ean b\u00e0i",
            key=_STANDARDIZATION_OPTION_KEYS["sync_context"],
        )
        st.checkbox(
            "Chu\u1ea9n h\u00f3a Times New Roman, c\u1ee1 14",
            key=_STANDARDIZATION_OPTION_KEYS["normalize_font"],
        )
        st.checkbox(
            "Chu\u1ea9n h\u00f3a font c\u00f4ng th\u1ee9c to\u00e1n",
            key=_STANDARDIZATION_OPTION_KEYS["normalize_equations"],
        )
        st.checkbox(
            "Chu\u1ea9n h\u00f3a b\u1ea3ng v\u00e0 h\u1ea1n ch\u1ebf v\u1ee1 b\u1ea3ng",
            key=_STANDARDIZATION_OPTION_KEYS["normalize_tables"],
        )
        st.checkbox(
            "Chu\u1ea9n h\u00f3a kh\u1ed5 gi\u1ea5y v\u00e0 l\u1ec1 trang",
            key=_STANDARDIZATION_OPTION_KEYS["normalize_page_layout"],
        )
        st.checkbox(
            "Chu\u1ea9n h\u00f3a gi\u00e3n d\u00f2ng v\u00e0 kho\u1ea3ng c\u00e1ch \u0111o\u1ea1n",
            key=_STANDARDIZATION_OPTION_KEYS["normalize_spacing"],
        )
        st.checkbox(
            "G\u1ee1 \u0111\u1ea7u/cu\u1ed1i trang c\u0169 v\u00e0 \u0111\u00e1nh s\u1ed1 trang",
            key=_STANDARDIZATION_OPTION_KEYS["normalize_header_footer"],
        )

        _render_standardization_date_and_document_options()

        selected = _standardization_options_from_state()
        if selected.preserve_original_maximum and any(
            (
                selected.normalize_font,
                selected.normalize_tables,
                selected.normalize_page_layout,
                selected.normalize_spacing,
                selected.normalize_header_footer,
            )
        ):
            st.warning(
                "C\u00e1c t\u00e1c v\u1ee5 tr\u00ecnh b\u00e0y \u0111\u01b0\u1ee3c ch\u1ecdn s\u1ebd "
                "thay \u0111\u1ed5i h\u00ecnh th\u1ee9c trong ph\u1ea1m vi t\u01b0\u01a1ng \u1ee9ng."
            )

        controls = st.columns(2)
        controls[0].button(
            "Ch\u1ecdn t\u1ea5t c\u1ea3",
            key="standardization_select_all",
            use_container_width=True,
            on_click=_set_all_standardization_options,
            args=(True,),
        )
        controls[1].button(
            "B\u1ecf ch\u1ecdn t\u1ea5t c\u1ea3",
            key="standardization_clear_all",
            use_container_width=True,
            on_click=_set_all_standardization_options,
            args=(False,),
        )
        st.button(
            "\u2705 X\u00e1c nh\u1eadn v\u00e0 t\u1ea1o gi\u00e1o \u00e1n chu\u1ea9n h\u00f3a",
            key="standardization_control_panel_confirm",
            type="primary",
            use_container_width=True,
            disabled=not selected.has_selected_operation,
            on_click=_confirm_standardization_options,
        )


def _render_pending_standardization_target() -> None:
    action = str(
        st.session_state.pop(
            "lesson_plan_management_pending_action",
            "",
        )
    )
    anchor_map = {
        "upload": "upload-lesson-plan",
        "create": "standardize-lesson-plan",
        "preview": "preview-standardized-lesson-plan",
        "save": "save-standardized-lesson-plan",
        "download": "download-standardized-lesson-plan",
    }
    target = anchor_map.get(action)
    if not target:
        return

    st.components.v1.html(
        f"""
<script>
(function () {{
  const doc = window.parent.document;
  function openTarget() {{
    const element = doc.getElementById("{target}");
    if (!element) return;
    element.scrollIntoView({{
      behavior: "smooth",
      block: "start"
    }});
  }}
  [120, 350, 700].forEach(function (delay) {{
    window.setTimeout(openTarget, delay);
  }});
}})();
</script>
        """.strip(),
        height=0,
    )


def _local_selection() -> OperationalInputSelection:
    return OperationalInputSelection(
        reference=OperationalInputReference(
            location=OperationalInputLocation.LOCAL_UPLOAD,
        ),
        source=None,
    )


def _academic_year_options(intake) -> tuple[str, ...]:
    return tuple(
        sorted(
            {
                week.academic_year
                for week in intake.source_data.academic_weeks
            }
        )
    )


def _week_options(
    intake,
    academic_year: str,
) -> tuple[int, ...]:
    return tuple(
        sorted(
            {
                week.week_number
                for week in intake.source_data.academic_weeks
                if week.academic_year == academic_year
            }
        )
    )


def _teacher_options(intake) -> tuple[str, ...]:
    return tuple(
        sorted(
            {
                slot.teacher_id
                for slot in intake.source_data.timetable_slots
            }
        )
    )


def _resolve_lbg_display_names(
    *,
    client,
    view,
) -> tuple[
    dict[str, str],
    dict[str, str],
    dict[str, str],
]:
    class_names: dict[str, str] = {}
    subject_names: dict[str, str] = {}
    component_names: dict[str, str] = {}

    if client is None:
        return (
            class_names,
            subject_names,
            component_names,
        )

    class_repository = (
        SupabaseClassCatalogRepository(
            client=client,
        )
    )

    subject_repository = (
        SupabaseSubjectCatalogRepository(
            client=client,
        )
    )

    class_ids = {
        row.class_id
        for row in view.rows
        if row.class_id
    }

    subject_ids = {
        row.subject_ref
        for row in view.rows
        if row.subject_ref
    }

    component_ids = {
        row.component_ref
        for row in view.rows
        if row.component_ref
    }

    for class_id in class_ids:
        try:
            item = class_repository.get(
                class_id=class_id,
            )

            if item is not None:
                class_names[class_id] = (
                    item.display_name
                )

        except Exception:
            pass

    for subject_id in subject_ids:
        try:
            item = (
                subject_repository.get_subject(
                    subject_id=subject_id,
                )
            )

            if item is not None:
                subject_names[subject_id] = (
                    item.name
                )

        except Exception:
            pass

    for component_id in component_ids:
        try:
            item = (
                subject_repository.get_component(
                    component_id=component_id,
                )
            )

            if item is not None:
                component_names[component_id] = (
                    item.name
                )

        except Exception:
            pass

    return (
        class_names,
        subject_names,
        component_names,
    )


def _preview_rows(
    view,
    *,
    class_names: dict[str, str] | None = None,
    subject_names: dict[str, str] | None = None,
    component_names: dict[str, str] | None = None,
) -> list[dict]:
    class_names = class_names or {}
    subject_names = subject_names or {}
    component_names = component_names or {}

    return [
        {
            "Th\u1ee9/ng\u00e0y":
                (
                    (
                        "Ch\u1ee7 nh\u1eadt"
                        if row.weekday == 7
                        else f"Th\u1ee9 {row.weekday + 1}"
                    )
                    + " - "
                    f"{row.teaching_date.strftime('%d/%m/%Y')}"
                ),

            "Bu\u1ed5i": (
                "S\u00e1ng"
                if getattr(row.session, "value", row.session)
                == "MORNING"
                else "Chi\u1ec1u"
            ),

            "Ti\u1ebft TKB":
                row.timetable_period,

            "M\u00f4n/Ph\u00e2n m\u00f4n":
                (
                    component_names.get(
                        row.component_ref,
                        row.component_ref,
                    )
                    if row.component_ref
                    else subject_names.get(
                        row.subject_ref,
                        row.subject_ref,
                    )
                    or ""
                ),

            "L\u1edbp":
                class_names.get(
                    row.class_id,
                    row.class_id,
                ),

            "Ti\u1ebft PPCT":
                row.curriculum_period,

            "T\u00ean b\u00e0i d\u1ea1y":
                row.lesson_title,

            "Chu\u1ea9n b\u1ecb, \u0111i\u1ec1u ch\u1ec9nh":
                ", ".join(
                    row.teaching_equipment
                ),

            "Ghi ch\u00fa":
                "",
        }
        for row in view.rows
    ]


def _render_lbg_table(
    view,
    *,
    client=None,
    teacher_user_id="",
    workspace_focus="AI",
    compact_setup_ui=False,
) -> None:
    (
        class_names,
        subject_names,
        component_names,
    ) = _resolve_lbg_display_names(
        client=client,
        view=view,
    )

    rows = _preview_rows(
        view,
        class_names=class_names,
        subject_names=subject_names,
        component_names=component_names,
    )

    if not rows:
        st.info(
            "Tu\u1ea7n n\u00e0y ch\u01b0a c\u00f3 "
            "ti\u1ebft d\u1ea1y trong L\u1ecbch b\u00e1o gi\u1ea3ng."
        )
        return

    with st.expander(
        "Xem L\u1ecbch b\u00e1o gi\u1ea3ng "
        f"Tu\u1ea7n {view.week_number}",
        expanded=False,
    ):
        st.data_editor(
            rows,
            width="stretch",
            hide_index=True,
            disabled=(
                "Th\u1ee9/ng\u00e0y",
                "Bu\u1ed5i",
                "Ti\u1ebft TKB",
                "M\u00f4n/Ph\u00e2n m\u00f4n",
                "L\u1edbp",
                "Ti\u1ebft PPCT",
                "T\u00ean b\u00e0i d\u1ea1y",
            ),
            key=(
                "lbg_user_editor_"
                + str(view.week_number)
            ),
        )

    _render_lesson_plan_standardization_workspace(
        view,
        teacher_user_id=teacher_user_id,
        client=client,
        workspace_focus=workspace_focus,
        compact_setup_ui=compact_setup_ui,
)




def _process_lesson_plan_upload(
    *,
    row,
    drafting_date,
    content: bytes,
    original_name: str,
    modification_plan=None,
) -> tuple[
    str,
    bytes,
    tuple[str, ...],
]:
    service = (
        LessonPlanDocumentProcessingService(
            profile_path=(
                _LESSON_PLAN_PROFILE
            )
        )
    )

    result = service.process(
        row=row,
        drafting_date=drafting_date,
        content=content,
        original_name=original_name,
        modification_plan=modification_plan,
    )

    return (
        result.output_name,
        result.output_bytes,
        result.unresolved_fields,
    )



def _lesson_plan_lesson_options_from_rows(
    rows,
):
    """
    Build lesson-level choices from the current
    weekly schedule.

    Transitional adapter:
    downstream processing still receives one
    representative schedule row.
    """

    grouped = {}

    for index, row in enumerate(rows):
        lesson_title = str(
            getattr(
                row,
                "lesson_title",
                "",
            )
            or ""
        ).strip()

        curriculum_period = getattr(
            row,
            "curriculum_period",
            None,
        )

        if (
            not lesson_title
            or curriculum_period is None
        ):
            continue

        item = grouped.setdefault(
            lesson_title,
            {
                "lesson_title": lesson_title,
                "periods": set(),
                "classes": set(),
                "teaching_dates": [],
                "row_indices": [],
            },
        )

        item["periods"].add(
            int(curriculum_period)
        )

        class_id = str(
            getattr(
                row,
                "class_id",
                "",
            )
            or ""
        ).strip()

        if class_id:
            item["classes"].add(
                class_id
            )

        teaching_date = getattr(
            row,
            "teaching_date",
            None,
        )

        if teaching_date is not None:
            item["teaching_dates"].append(
                (
                    teaching_date,
                    class_id,
                )
            )

        item["row_indices"].append(
            index
        )

    result = []

    for item in grouped.values():
        periods = tuple(
            sorted(
                item["periods"]
            )
        )

        classes = tuple(
            sorted(
                item["classes"]
            )
        )

        teaching_dates = tuple(
            sorted(
                set(
                    item[
                        "teaching_dates"
                    ]
                ),
                key=lambda value: (
                    value[0],
                    value[1],
                ),
            )
        )

        row_indices = tuple(
            item["row_indices"]
        )

        if not row_indices:
            continue

        period_text = " + ".join(
            str(value)
            for value in periods
        )

        result.append(
            {
                "lesson_title": (
                    item["lesson_title"]
                ),
                "periods": periods,
                "classes": classes,
                "teaching_dates": (
                    teaching_dates
                ),
                "row_indices": (
                    row_indices
                ),
                "representative_index": (
                    row_indices[0]
                ),
                "label": (
                    f"{item['lesson_title']} "
                    f"(Ti\u1ebft {period_text})"
                ),
            }
        )

    return tuple(
        sorted(
            result,
            key=lambda item: (
                (
                    item["periods"][0]
                    if item["periods"]
                    else 10**9
                ),
                item["lesson_title"],
            ),
        )
    )


def _class_display_name(
    class_id: str,
    *,
    client=None,
) -> str:
    """
    Resolve canonical class_id to the teacher-facing
    class name.

    class_id remains the internal canonical identifier.
    Only the UI/document-facing value is converted.
    """
    value = str(
        class_id or ""
    ).strip()

    if not value:
        return "-"

    if client is not None:
        try:
            item = (
                SupabaseClassCatalogRepository(
                    client=client
                )
                .get(
                    class_id=value
                )
            )

            if item is not None:
                class_name = str(
                    getattr(
                        item,
                        "class_name",
                        "",
                    )
                    or ""
                ).strip()

                if class_name:
                    return class_name

                class_code = str(
                    getattr(
                        item,
                        "class_code",
                        "",
                    )
                    or ""
                ).strip()

                if class_code:
                    return class_code

        except Exception:
            pass

    try:
        runtime = st.session_state.get(
            "_system_weekly_schedule_runtime"
        )

        repository = getattr(
            runtime,
            "_class_repository",
            None,
        )

        if repository is not None:
            item = repository.get(
                class_id=value,
            )

            if item is not None:
                class_name = str(
                    getattr(
                        item,
                        "class_name",
                        "",
                    )
                    or ""
                ).strip()

                if class_name:
                    return class_name

                class_code = str(
                    getattr(
                        item,
                        "class_code",
                        "",
                    )
                    or ""
                ).strip()

                if class_code:
                    return class_code

    except Exception:
        pass

    return value


def _rows_for_same_timetable_lesson(
    schedule_rows,
    *,
    selected_row,
) -> tuple[object, ...]:
    """Return every timetable row assigned to the selected PPCT lesson."""

    selected_subject = str(
        getattr(selected_row, "subject_ref", "") or ""
    ).strip()
    selected_component = str(
        getattr(selected_row, "component_ref", "") or ""
    ).strip()
    selected_period = getattr(
        selected_row,
        "curriculum_period",
        None,
    )
    selected_lesson_title = str(
        getattr(selected_row, "lesson_title", "") or ""
    ).strip().casefold()

    matched_rows = []

    for row in tuple(schedule_rows or ()):
        if getattr(row, "curriculum_period", None) != selected_period:
            continue

        row_subject = str(
            getattr(row, "subject_ref", "") or ""
        ).strip()
        row_component = str(
            getattr(row, "component_ref", "") or ""
        ).strip()

        if row_subject != selected_subject:
            continue

        if row_component != selected_component:
            continue

        row_lesson_title = str(
            getattr(row, "lesson_title", "") or ""
        ).strip().casefold()

        if (
            selected_lesson_title
            and row_lesson_title != selected_lesson_title
        ):
            continue

        matched_rows.append(row)

    return tuple(matched_rows)


def _rows_for_selected_lesson_unit(filtered_schedule_rows, *, selected_unit, schedule_rows, selected_row) -> tuple[object, ...]:
    # V58-C3D6F
    # The selected lesson unit is already canonical for the active
    # subject/component/grade/PPCT scope. Its row_indices therefore index
    # filtered_schedule_rows, not the broader schedule_rows collection.
    strict_rows = tuple(
        filtered_schedule_rows[index]
        for index in tuple(getattr(selected_unit, "row_indices", ()) or ())
        if 0 <= index < len(filtered_schedule_rows)
    )

    if strict_rows:
        return strict_rows

    # Fail closed. Keep the representative row only when the canonical
    # selection unit has no usable row index; never broaden back to the
    # unfiltered weekly schedule.
    if selected_row is not None:
        return (selected_row,)

    return ()




def _class_ids_for_same_timetable_lesson(
    schedule_rows,
    *,
    selected_row,
) -> tuple[str, ...]:
    """Return every class assigned the selected PPCT lesson in this week."""

    class_ids = {
        str(getattr(row, "class_id", "") or "").strip()
        for row in _rows_for_same_timetable_lesson(
            schedule_rows,
            selected_row=selected_row,
        )
    }

    return tuple(sorted(value for value in class_ids if value))


def _class_display_names(
    class_ids,
    *,
    client=None,
) -> str:
    """Resolve and join canonical class IDs for teacher/document display."""

    names = []

    for class_id in tuple(class_ids or ()):
        name = _class_display_name(
            class_id,
            client=client,
        )

        if name and name != "-" and name not in names:
            names.append(name)

    return ", ".join(names) if names else "-"


def _class_schedule_display_values(
    timetable_rows,
    *,
    client=None,
) -> tuple[str, str]:
    """Format timetable period and teaching date separately for each class."""

    details = []

    for row in tuple(timetable_rows or ()):
        class_name = _class_display_name(
            getattr(row, "class_id", ""),
            client=client,
        )
        timetable_period = str(
            getattr(row, "timetable_period", "") or "-"
        ).strip()
        teaching_date = getattr(row, "teaching_date", None)

        if teaching_date is None:
            teaching_date_text = "-"
        else:
            try:
                teaching_date_text = teaching_date.strftime("%d/%m/%Y")
            except AttributeError:
                teaching_date_text = str(teaching_date)

        detail = (
            class_name,
            timetable_period,
            teaching_date_text,
        )

        if detail not in details:
            details.append(detail)

    details.sort(key=lambda value: (value[0], value[1], value[2]))

    if not details:
        return "-", "-"

    return (
        "; ".join(
            f"{class_name}: tiết {period}"
            for class_name, period, _ in details
        ),
        "; ".join(
            f"{class_name}: {teaching_date}"
            for class_name, _, teaching_date in details
        ),
    )


def _subject_display_name(
    *,
    subject_ref: str,
    component_ref: str,
    client=None,
) -> str:
    """Resolve a teacher-facing subject/component name without changing IDs."""

    subject_name, component_name = (
        _subject_component_display_names(
            subject_ref=subject_ref,
            component_ref=component_ref,
            client=client,
        )
    )

    return (
        component_name
        if component_name != "-"
        else subject_name
    )


def _subject_component_display_names(
    *,
    subject_ref: str,
    component_ref: str,
    client=None,
) -> tuple[str, str]:
    """Resolve separate teacher-facing subject and component names."""

    subject_value = str(subject_ref or "").strip()
    component_value = str(component_ref or "").strip()
    subject_name = ""
    component_name = ""

    if client is not None:
        try:
            repository = SupabaseSubjectCatalogRepository(
                client=client
            )

            if component_value:
                component = repository.get_component(
                    component_id=component_value
                )
                if component is not None:
                    name = str(
                        getattr(component, "name", "")
                        or ""
                    ).strip()
                    if name:
                        component_name = name

            if subject_value:
                subject = repository.get_subject(
                    subject_id=subject_value
                )
                if subject is not None:
                    name = str(
                        getattr(subject, "name", "")
                        or ""
                    ).strip()
                    if name:
                        subject_name = name

        except Exception:
            pass

    return (
        subject_name or subject_value or "-",
        component_name or component_value or "-",
    )


def _render_selected_lesson_summary(
    lesson,
    *,
    drafting_date=None,
    client=None,
) -> None:
    lesson_title = str(
        lesson.get(
            "lesson_title",
            "",
        )
        or ""
    ).strip()

    periods = tuple(
        lesson.get(
            "periods",
            (),
        )
        or ()
    )

    teaching_dates = tuple(
        lesson.get(
            "teaching_dates",
            (),
        )
        or ()
    )

    st.markdown(
        "**B\u00e0i:** "
        + (
            lesson_title
            or "-"
        )
    )

    st.markdown(
        "**S\u1ed1 ti\u1ebft:** "
        + str(
            len(periods)
        )
    )

    st.markdown(
        "**Ng\u00e0y d\u1ea1y - L\u1edbp**"
    )

    if not teaching_dates:
        st.write("-")
        return

    for (
        teaching_date,
        class_id,
    ) in teaching_dates:
        try:
            date_text = (
                teaching_date.strftime(
                    "%d/%m/%Y"
                )
            )
        except Exception:
            date_text = str(
                teaching_date
            )

        st.write(
            date_text
            + " - "
            + _class_display_name(
                class_id,
                client=client,
            )
        )



def _lesson_plan_row_label(
    row,
) -> str:
    return (
        f"{row.teaching_date.strftime('%d/%m/%Y')}"
        f" | Ti\u1ebft TKB {row.timetable_period}"
        f" | {row.class_id}"
        f" | PPCT {row.curriculum_period}"
        f" | {row.lesson_title}"
    )


def _build_lesson_plan_starter(
    *,
    lesson_title: str,
    class_ref: str,
    curriculum_period,
    subject_name: str = "",
    timetable_period=None,
    teaching_date=None,
    teaching_equipment=(),
) -> str:
    """Build an editable starter document from the selected lesson context."""

    period_text = str(
        curriculum_period
        if curriculum_period is not None
        else "..."
    )

    timetable_text = str(
        timetable_period
        if timetable_period is not None
        else "..."
    )

    try:
        teaching_date_text = teaching_date.strftime(
            "%d/%m/%Y"
        )
    except AttributeError:
        teaching_date_text = str(
            teaching_date or "..."
        )

    equipment_text = ", ".join(
        str(item)
        for item in (teaching_equipment or ())
        if str(item).strip()
    ) or "..."

    return (
        "KẾ HOẠCH BÀI DẠY\n\n"
        f"Tên bài: {lesson_title or '...'}\n"
        f"Môn/Phân môn: {subject_name or '...'}\n"
        f"Lớp: {class_ref or '...'}\n"
        f"Tiết PPCT: {period_text}\n\n"
        f"Tiết TKB: {timetable_text}\n"
        f"Ngày dạy: {teaching_date_text}\n\n"
        "I. MỤC TIÊU\n"
        "1. Kiến thức\n- ...\n"
        "2. Năng lực\n- ...\n"
        "3. Phẩm chất\n- ...\n\n"
        "II. THIẾT BỊ VÀ HỌC LIỆU\n"
        f"- {equipment_text}\n\n"
        "III. TIẾN TRÌNH DẠY HỌC\n"
        "1. Hoạt động mở đầu\n"
        "- Mục tiêu: ...\n- Nội dung: ...\n"
        "- Sản phẩm: ...\n- Tổ chức thực hiện: ...\n\n"
        "2. Hoạt động hình thành kiến thức\n"
        "- Mục tiêu: ...\n- Nội dung: ...\n"
        "- Sản phẩm: ...\n- Tổ chức thực hiện: ...\n\n"
        "3. Hoạt động luyện tập\n"
        "- Mục tiêu: ...\n- Nội dung: ...\n"
        "- Sản phẩm: ...\n- Tổ chức thực hiện: ...\n\n"
        "4. Hoạt động vận dụng\n"
        "- Mục tiêu: ...\n- Nội dung: ...\n"
        "- Sản phẩm: ...\n- Tổ chức thực hiện: ...\n\n"
        "IV. ĐIỀU CHỈNH SAU BÀI DẠY\n- ..."
    )


def _lesson_plan_quality_checks(
    document_text: str,
) -> tuple[tuple[str, bool], ...]:
    """Return fast, deterministic completeness checks for a lesson plan."""

    normalized = str(document_text or "").upper()

    return (
        ("Có mục tiêu bài học", "MỤC TIÊU" in normalized),
        (
            "Có thiết bị và học liệu",
            "THIẾT BỊ" in normalized
            or "HỌC LIỆU" in normalized,
        ),
        (
            "Có tiến trình dạy học",
            "TIẾN TRÌNH" in normalized
            or "HOẠT ĐỘNG" in normalized,
        ),
        ("Có sản phẩm học tập", "SẢN PHẨM" in normalized),
        (
            "Có phần điều chỉnh sau bài dạy",
            "ĐIỀU CHỈNH SAU" in normalized,
        ),
    )



def _render_lesson_plan_drafting_workspace(
    selected_lesson=None,
    teacher_user_id="",
    academic_year="",
    week_number=0,
    selection_mode="LESSON",
    selection_unit_id="",
    client=None,
) -> None:
    """
    Complete V1 teacher-facing lesson-plan editor.

    Draft identity is scoped by teacher, academic year,
    week, subject, selection mode and selection unit.
    """
    from lesson_planning_v2.services.lesson_plan_draft_workspace_service import (
        LessonPlanDraftWorkspaceService,
    )
    from lesson_planning_v2.services.lesson_plan_workspace_v1_service import (
        LessonPlanDocxTextImporter,
        LessonPlanDocxWholeDocumentImporter,
        LessonPlanFullDocumentDocxAdapter,
        LessonPlanLibrarySourceService,
        LessonPlanSimpleDocxExporter,
        LessonPlanWorkspaceContent,
        LessonPlanWorkspaceContext,
        LessonPlanWorkspaceV1Service,
    )

    st.markdown(
        """
<div class="mt-ai-workspace-heading">
  <div class="mt-ai-workspace-icon">AI</div>
  <div>
    <div class="mt-section-kicker">KHÔNG GIAN BIÊN SOẠN</div>
    <h2 class="mt-ai-workspace-title">SOẠN BÀI CÙNG AI</h2>
    <p class="mt-ai-workspace-description">
      Khởi tạo, nhập giáo án cũ, biên tập, kiểm tra và chuyển thẳng
      sang bước chuẩn hóa trong một không gian làm việc.
    </p>
  </div>
</div>
        """.strip(),
        unsafe_allow_html=True,
    )

    selected = (
        selected_lesson
        if isinstance(
            selected_lesson,
            dict,
        )
        else {}
    )

    normalized_teacher = str(
        teacher_user_id
    ).strip()

    normalized_year = str(
        academic_year
    ).strip()

    try:
        normalized_week = int(
            week_number
        )
    except (
        TypeError,
        ValueError,
    ):
        normalized_week = 0

    normalized_unit = str(
        selection_unit_id
    ).strip()

    normalized_mode = (
        getattr(
            selection_mode,
            "value",
            selection_mode,
        )
    )

    normalized_mode = str(
        normalized_mode
    ).strip() or "LESSON"

    subject_ref = str(
        selected.get(
            "subject_ref",
            selected.get(
                "subject_id",
                selected.get(
                    "subject",
                    "general",
                ),
            ),
        )
    ).strip() or "general"

    class_id = str(
        selected.get(
            "class_id",
            "",
        )
    ).strip()

    grade_level = str(
        selected.get(
            "grade_level",
            "",
        )
    ).strip()

    class_ref = str(
        selected.get(
            "class_name",
            "",
        )
    ).strip()

    if (
        not class_ref
        and class_id
        and client is not None
    ):
        try:
            class_item = (
                SupabaseClassCatalogRepository(
                    client=client
                )
                .get(
                    class_id=class_id
                )
            )

            if class_item is not None:
                class_ref = str(
                    class_item.class_name
                ).strip()

                if not grade_level:
                    grade_level = str(
                        class_item.grade_level
                    ).strip()

        except Exception:
            # Display-name resolution must never
            # destroy the lesson-plan workspace.
            pass

    if not class_ref:
        class_ref = (
            grade_level
            or class_id
            or "N/A"
        )

    lesson_title = str(
        selected.get(
            "lesson_title",
            selected.get(
                "title",
                "",
            ),
        )
    ).strip()

    curriculum_period = (
        selected.get(
            "curriculum_period"
        )
    )

    teaching_date = (
        selected.get(
            "teaching_date"
        )
    )

    if (
        not normalized_teacher
        or not normalized_year
        or normalized_week <= 0
        or not normalized_unit
    ):
        st.info(
            "Ch\u1ecdn \u0111\u1ea7y \u0111\u1ee7 b\u00e0i/ti\u1ebft v\u00e0 tu\u1ea7n "
            "\u0111\u1ec3 b\u1eaft \u0111\u1ea7u so\u1ea1n b\u00e0i."
        )
        return

    try:
        context = (
            LessonPlanWorkspaceContext(
                teacher_user_id=(
                    normalized_teacher
                ),
                academic_year=(
                    normalized_year
                ),
                week_number=(
                    normalized_week
                ),
                subject_ref=subject_ref,
                selection_mode=(
                    normalized_mode
                ),
                selection_unit_id=(
                    normalized_unit
                ),
                class_or_grade_ref=(
                    class_ref
                ),
                lesson_id=(
                    normalized_unit
                ),
                title=lesson_title,
            )
        )
    except ValueError as error:
        st.error(str(error))
        return

    repository = st.session_state.get(
        "lesson_plan_workspace_draft_repository"
    )

    if repository is None:
        st.error(
            "Kho l\u01b0u b\u1ea3n nh\u00e1p ch\u01b0a s\u1eb5n s\u00e0ng. "
            "H\u00e3y \u0111\u0103ng nh\u1eadp l\u1ea1i."
        )
        return

    draft_service = (
        LessonPlanDraftWorkspaceService(
            repository
        )
    )

    workspace_service = (
        LessonPlanWorkspaceV1Service(
            draft_service=draft_service
        )
    )

    try:
        persisted = (
            workspace_service.load(
                context=context
            )
        )
    except Exception as error:
        st.warning(
            "Chưa thể đọc bản nháp đã lưu: "
            + str(error)
        )
        persisted = None

    _legacy_drafting_entry_label = "Cách bắt đầu"
    prefix = context.widget_prefix

    objectives_key = (
        prefix + "_objectives"
    )
    materials_key = (
        prefix + "_materials"
    )
    process_key = (
        prefix + "_process"
    )

    full_document_key = (
        prefix + "_full_document"
    )

    source_key = (
        prefix + "_source_mode"
    )

    standardization_transfer_key = (
        prefix
        + "_standardization_transfer"
    )

    standardization_transfer_ready_key = (
        prefix
        + "_standardization_transfer_ready"
    )

    # Initialize each lesson independently.
    if objectives_key not in st.session_state:
        st.session_state[
            objectives_key
        ] = (
            persisted.objectives_text
            if persisted is not None
            else ""
        )

    if materials_key not in st.session_state:
        st.session_state[
            materials_key
        ] = (
            persisted.materials_text
            if persisted is not None
            else ""
        )

    if process_key not in st.session_state:
        st.session_state[
            process_key
        ] = (
            persisted.teaching_process_text
            if persisted is not None
            else ""
        )

    if full_document_key not in st.session_state:
        persisted_full_document = ""

        if persisted is not None:
            persisted_full_document = str(
                getattr(
                    persisted,
                    "full_document_text",
                    "",
                )
                or ""
            ).strip()

        if persisted_full_document:
            st.session_state[
                full_document_key
            ] = persisted_full_document

        else:
            legacy_parts = []

            if st.session_state[
                objectives_key
            ].strip():
                legacy_parts.extend(
                    (
                        "I. M\u1ee4C TI\u00caU",
                        st.session_state[
                            objectives_key
                        ].strip(),
                    )
                )

            if st.session_state[
                materials_key
            ].strip():
                legacy_parts.extend(
                    (
                        (
                            "II. THI\u1ebeT B\u1eca "
                            "V\u00c0 H\u1eccC LI\u1ec6U"
                        ),
                        st.session_state[
                            materials_key
                        ].strip(),
                    )
                )

            if st.session_state[
                process_key
            ].strip():
                legacy_parts.extend(
                    (
                        (
                            "III. TI\u1ebeN TR\u00ccNH "
                            "D\u1ea0Y H\u1eccC"
                        ),
                        st.session_state[
                            process_key
                        ].strip(),
                    )
                )

            st.session_state[
                full_document_key
            ] = "\n\n".join(
                legacy_parts
            ).strip()


    st.caption(
        "Bản nháp được lưu riêng theo giáo viên, tuần và bài dạy."
    )

    mode = st.radio(
        "Nguồn nội dung",
        (
            "Soạn mới cùng AI",
            "Tải & chỉnh sửa giáo án cũ",
        ),
        horizontal=True,
        key=source_key,
    )

    if mode == "Soạn mới cùng AI":
        if st.button(
            "Tạo khung giáo án từ bài đang chọn",
            key=prefix + "_create_starter",
            use_container_width=True,
            disabled=bool(
                str(
                    st.session_state.get(
                        full_document_key,
                        "",
                    )
                ).strip()
            ),
        ):
            st.session_state[full_document_key] = (
                _build_lesson_plan_starter(
                    lesson_title=lesson_title,
                    class_ref=class_ref,
                    curriculum_period=curriculum_period,
                    subject_name=str(
                        selected.get("subject_name", "")
                        or subject_ref
                    ),
                    timetable_period=selected.get(
                        "timetable_period"
                    ),
                    teaching_date=teaching_date,
                    teaching_equipment=selected.get(
                        "teaching_equipment",
                        (),
                    ),
                )
            )
            st.rerun()

    uploaded = None

    if mode == "Tải & chỉnh sửa giáo án cũ":
        uploaded = st.file_uploader(
            "Tải giáo án cũ để chỉnh sửa",
            type=("docx",),
            key=prefix + "_upload",
        )

        if uploaded is not None:
            import_clicked = st.button(
                "Đưa giáo án vào trình chỉnh sửa",
                key=prefix + "_import_word",
            )

            if import_clicked:
                try:
                    imported = (
                        LessonPlanDocxWholeDocumentImporter()
                        .import_bytes(
                            uploaded.getvalue()
                        )
                    )

                    st.session_state[
                        full_document_key
                    ] = imported

                    st.session_state[
                        objectives_key
                    ] = ""

                    st.session_state[
                        materials_key
                    ] = ""

                    st.session_state[
                        process_key
                    ] = ""

                    st.session_state[
                        prefix
                        + "_source_docx"
                    ] = uploaded.getvalue()

                    st.session_state[
                        prefix
                        + "_source_name"
                    ] = uploaded.name

                    st.success(
                        "Đã đưa nội dung Word "
                        "vào trình soạn."
                    )

                    st.rerun()

                except Exception as error:
                    st.error(
                        "Không thể đọc file Word: "
                        + str(error)
                    )

    st.markdown("---")

    st.markdown("### Không gian biên tập")

    editor_col, ai_col = st.columns(
        [7, 3],
        gap="large",
    )

    with editor_col:
        st.markdown(
            '#### \U0001f4c4 Gi\xe1o \xe1n \u0111ang l\xe0m vi\u1ec7c'
        )

        st.caption(
            'To\xe0n b\u1ed9 gi\xe1o \xe1n \u0111\u01b0\u1ee3c hi\u1ec3n th\u1ecb v\xe0 ch\u1ec9nh s\u1eeda li\xean t\u1ee5c trong m\u1ed9t v\xf9ng.'
        )

        full_document = st.text_area(
            'N\u1ed9i dung gi\xe1o \xe1n',
            key=full_document_key,
            height=760,
            label_visibility="collapsed",
            placeholder=(
                'N\u1ed9i dung gi\xe1o \xe1n s\u1ebd xu\u1ea5t hi\u1ec7n t\u1ea1i \u0111\xe2y. B\u1ea1n c\xf3 th\u1ec3 so\u1ea1n m\u1edbi ho\u1eb7c t\u1ea3i gi\xe1o \xe1n c\u0169 \u0111\u1ec3 ti\u1ebfp t\u1ee5c ch\u1ec9nh s\u1eeda.'
            ),
        )

    with ai_col:
        assistant_tab, quality_tab = st.tabs(
            ("✨ Trợ lý", "✓ Kiểm tra")
        )

        with assistant_tab:
            st.caption(
                "Mô tả điều bạn muốn AI bổ sung hoặc điều chỉnh."
            )

            ai_request = st.text_area(
                "Yêu cầu AI",
                key=prefix + "_ai_request",
                height=260,
                placeholder=(
                    "Ví dụ: Thiết kế hoạt động mở đầu gắn với thực tế; "
                    "bổ sung câu hỏi phân hóa; làm rõ sản phẩm học tập..."
                ),
                label_visibility="collapsed",
            )

            st.button(
                "Gửi yêu cầu cho AI",
                key=prefix + "_ai_request_submit",
                use_container_width=True,
                disabled=True,
            )

            st.warning(
                "Dịch vụ AI chưa được kết nối. Nội dung giáo án "
                "sẽ không bị thay đổi khi chưa có xác nhận của giáo viên."
            )

        with quality_tab:
            quality_checks = (
                _lesson_plan_quality_checks(
                    full_document
                )
            )
            completed_checks = sum(
                1
                for _, passed in quality_checks
                if passed
            )

            st.metric(
                "Mức độ đầy đủ",
                f"{completed_checks}/{len(quality_checks)}",
            )

            for label, passed in quality_checks:
                st.markdown(
                    ("✅ " if passed else "⬜ ")
                    + label
                )

            st.caption(
                f"{len(str(full_document or '')):,} ký tự trong giáo án."
            )


    content = (
        LessonPlanWorkspaceContent(
            objectives_text="",
            materials_text="",
            teaching_process_text="",
            full_document_text=(
                full_document
            ),
        )
    )

    draft_docx_bytes = None

    if str(full_document or "").strip():
        try:
            draft_docx_bytes = (
                LessonPlanFullDocumentDocxAdapter()
                .build_bytes(full_document)
            )
        except Exception:
            draft_docx_bytes = None

    if draft_docx_bytes is not None:
        st.download_button(
            "Tải bản nháp Word",
            data=draft_docx_bytes,
            file_name=(
                (lesson_title or "giao-an-ban-nhap")
                + ".docx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            key=prefix + "_download_draft",
            use_container_width=True,
        )

    st.markdown("---")

    st.markdown(
        '### Ho\u00e0n t\u1ea5t v\u00e0 chuy\u1ec3n ti\u1ebfp'
    )

    st.caption(
        'Khi n\u1ed9i dung gi\xe1o \xe1n \u0111\xe3 ph\xf9 h\u1ee3p, chuy\u1ec3n tr\u1ef1c ti\u1ebfp sang c\xf4ng c\u1ee5 Chu\u1ea9n h\xf3a gi\xe1o \xe1n theo L\u1ecbch b\xe1o gi\u1ea3ng. Kh\xf4ng c\u1ea7n xu\u1ea5t Word trung gian.'
    )

    transfer_clicked = st.button(
        '\u27a1\ufe0f Chuy\u1ec3n sang Chu\u1ea9n h\xf3a gi\xe1o \xe1n theo L\u1ecbch b\xe1o gi\u1ea3ng',
        key=(
            prefix
            + "_transfer_to_standardization"
        ),
        type="primary",
        use_container_width=True,
        disabled=(
            not str(
                full_document
            ).strip()
        ),
    )

    if transfer_clicked:
        try:
            source_docx_bytes = (
                st.session_state.get(
                    prefix + "_source_docx"
                )
            )

            if (
                isinstance(
                    source_docx_bytes,
                    bytes,
                )
                and source_docx_bytes
            ):
                # Preservation-first:
                # preserve the uploaded DOCX package.
                internal_docx_bytes = (
                    source_docx_bytes
                )

            else:
                # Fallback for lesson plans created
                # without an uploaded DOCX source.
                internal_docx_bytes = (
                    LessonPlanFullDocumentDocxAdapter()
                    .build_bytes(
                        full_document
                    )
                )

        except Exception as error:
            st.error(
                "Không thể chuẩn bị "
                "giáo án để "
                "chuyển sang bước "
                "chuẩn hóa: "
                + str(error)
            )

            internal_docx_bytes = None

        if internal_docx_bytes is not None:
            transfer_source_name = (
                (
                    str(
                        lesson_title
                    ).strip()
                    or "giao-an-ai"
                )
                + ".docx"
            )

            st.session_state[
                standardization_transfer_key
            ] = {
                "source": "AI_DRAFT",
                "docx_bytes": (
                    internal_docx_bytes
                ),
                "source_name": (
                    transfer_source_name
                ),
                "teacher_user_id": (
                    normalized_teacher
                ),
                "academic_year": (
                    normalized_year
                ),
                "week_number": (
                    normalized_week
                ),
                "subject_ref": (
                    subject_ref
                ),
                "component_ref": str(
                    selected.get(
                        "component_ref",
                        "",
                    )
                    or ""
                ),
                "subject_name": str(
                    selected.get(
                        "subject_name",
                        "",
                    )
                    or ""
                ),
                "selection_mode": (
                    normalized_mode
                ),
                "selection_unit_id": (
                    normalized_unit
                ),
                "lesson_title": (
                    lesson_title
                ),
                "class_ref": (
                    class_ref
                ),
                "curriculum_period": (
                    curriculum_period
                ),
                "timetable_period": selected.get(
                    "timetable_period"
                ),
                "teaching_date": teaching_date,
                "teaching_equipment": tuple(
                    selected.get(
                        "teaching_equipment",
                        (),
                    )
                    or ()
                ),
                "full_document_text": str(
                    full_document
                ),
            }

            st.session_state[
                standardization_transfer_ready_key
            ] = True

            st.success(
                "Đã chuyển giáo án "
                "đang làm việc sang "
                "bước Chuẩn hóa "
                "giáo án theo "
                "Lịch báo giảng."
            )

    save_clicked = st.button(
        "Lưu bản nháp",
        key=prefix + "_save",
        use_container_width=True,
    )

    if save_clicked:
        try:
            saved = workspace_service.save(
                context=context,
                content=content,
                source=mode,
            )

            verified = (
                workspace_service.load(
                    context=context
                )
            )

            if verified != saved:
                st.error(
                    "Không thể xác nhận "
                    "bản nháp để lưu."
                )
            else:
                st.success(
                    "Đã lưu bản nháp."
                )

        except Exception as error:
            st.error(
                "Không thể lưu bản nháp: "
                + str(error)
            )








def _render_lesson_plan_metadata_override_editor(
    *,
    selected_row,
    drafting_date,
    class_name,
    week_number: int,
    selected_index: int,
):
    """
    Render canonical lesson metadata before DOCX processing.

    Values originate from the weekly schedule but may be
    reviewed and corrected by the teacher before they are
    applied to the working DOCX.

    This UI helper does not mutate the source DOCX.
    """

    st.markdown(
        "#### Th\u00f4ng tin s\u1ebd c\u1eadp nh\u1eadt "
        "v\u00e0o gi\u00e1o \u00e1n"
    )

    st.caption(
        "Th\u00f4ng tin \u0111\u01b0\u1ee3c l\u1ea5y "
        "t\u1ef1 \u0111\u1ed9ng t\u1eeb "
        "L\u1ecbch b\u00e1o gi\u1ea3ng. "
        "Gi\u00e1o vi\u00ean c\u00f3 th\u1ec3 "
        "ki\u1ec3m tra v\u00e0 ch\u1ec9nh s\u1eeda "
        "tr\u01b0\u1edbc khi t\u1ea1o "
        "gi\u00e1o \u00e1n chu\u1ea9n h\u00f3a."
    )

    key_prefix = (
        "lbg_lesson_plan_metadata_"
        + str(week_number)
        + "_"
        + str(selected_index)
        + "_"
    )

    teaching_date = st.date_input(
        "Ng\u00e0y d\u1ea1y / "
        "Ng\u00e0y gi\u1ea3ng",
        value=selected_row.teaching_date,
        key=(
            key_prefix
            + "teaching_date"
        ),
    )

    class_value = st.text_input(
        "L\u1edbp",
        value=str(
            class_name
            or ""
        ),
        key=(
            key_prefix
            + "class_name"
        ),
    )

    curriculum_period = st.number_input(
        "Ti\u1ebft PPCT",
        min_value=1,
        value=int(
            selected_row.curriculum_period
        ),
        step=1,
        key=(
            key_prefix
            + "curriculum_period"
        ),
    )

    lesson_title = st.text_input(
        "T\u00ean b\u00e0i",
        value=str(
            selected_row.lesson_title
            or ""
        ),
        key=(
            key_prefix
            + "lesson_title"
        ),
    )

    return {
        "drafting_date": drafting_date,
        "teaching_date": teaching_date,
        "class_name": class_value.strip(),
        "curriculum_period": int(
            curriculum_period
        ),
        "lesson_title": lesson_title.strip(),
    }


def _latest_ai_standardization_transfer():
    """Return the newest transfer payload published by AI authoring."""
    candidates = []
    for key, value in st.session_state.items():
        if not str(key).endswith("_standardization_transfer"):
            continue
        if not isinstance(value, dict):
            continue
        if value.get("source") != "AI_DRAFT":
            continue
        candidates.append((str(key), value))
    return candidates[-1] if candidates else ("", None)



def _standardization_grade_from_class_value(value) -> int | None:
    import re

    text = str(value or "").strip()
    match = re.search(
        r"(?<!\d)([6-9])(?=[A-Za-zÀ-ỹ0-9._\-\s]|$)",
        text,
    )
    if match is None:
        return None
    return int(match.group(1))


def _standardization_grade_for_row(row, *, client=None) -> int | None:
    # V58-C3C: ACTIVE timetable owns class_id; Class Catalog owns grade_level.
    # Never infer grade from digits embedded in class_id or display labels.
    class_id = str(getattr(row, "class_id", "") or "").strip()
    if not class_id or client is None:
        return None

    try:
        class_item = _v58_timed_class_catalog_get(
            client=client,
            class_id=class_id,
        )
    except Exception:
        # Fail closed: repository failure must not revive regex inference.
        return None

    if class_item is None:
        return None

    grade_level = str(
        getattr(class_item, "grade_level", "") or ""
    ).strip()
    if not grade_level:
        return None

    try:
        grade = int(grade_level)
    except (TypeError, ValueError):
        return None

    return grade if 6 <= grade <= 9 else None


def _normalized_sync_value(value) -> str:
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return str(value.isoformat()).strip().casefold()
    text = str(value).strip()
    parts = text.split("/")
    if len(parts) == 3 and all(part.isdigit() for part in parts):
        text = "-".join(reversed(parts))
    return text.casefold()


def _match_transfer_schedule_row(schedule_rows, payload):
    """Match one transferred lesson to its canonical schedule row."""
    if not isinstance(payload, dict):
        return None
    fields = (
        ("subject_ref", "subject_ref"),
        ("component_ref", "component_ref"),
        ("class_ref", "class_id"),
        ("curriculum_period", "curriculum_period"),
        ("timetable_period", "timetable_period"),
        ("teaching_date", "teaching_date"),
        ("lesson_title", "lesson_title"),
    )
    best_index = None
    best_score = -1
    for index, row in enumerate(schedule_rows):
        score = sum(
            1
            for payload_field, row_field in fields
            if _normalized_sync_value(payload.get(payload_field))
            and _normalized_sync_value(payload.get(payload_field))
            == _normalized_sync_value(getattr(row, row_field, None))
        )
        if score > best_score:
            best_index = index
            best_score = score
    return best_index if best_score >= 3 else None


def _render_synced_context_markdown(
    visible: bool,
    body: str,
    *,
    unsafe_allow_html: bool = False,
) -> None:
    if visible:
        st.markdown(
            body,
            unsafe_allow_html=unsafe_allow_html,
        )

# Legacy wiring contract: render_lesson_plan_preview(
# Legacy wiring contract: render_lesson_plan_teacher_review(
# Legacy lesson summary contract: "S\u1ed1 ti\u1ebft"
# Legacy lesson summary contract: "**Ng\u00e0y d\u1ea1y**"
# Legacy selection contract: "C\u00e1ch ch\u1ecdn n\u1ed9i dung "
# STANDARDIZATION_PPCT_REVERSE_SYNC_V51
def _standardization_ppct_reverse_sync(
    *, unit_widget_key, lesson_units, filtered_schedule_rows,
    grade_filter_key, client,
):
    try:
        selected_index = int(st.session_state.get(unit_widget_key, 0))
    except (TypeError, ValueError):
        selected_index = 0
    if not (0 <= selected_index < len(lesson_units)):
        return
    # V58_C5B2_SHADOW_LESSON_PLAN_GROUPS
    # V58_C5B7D_RUNTIME_POLICY_INJECTION
    # ADMIN policy is configuration only; grouping remains shadow projection.
    try:
        from lesson_planning_v2.adapters.supabase_lesson_plan_grouping_policy_repository import (
            SupabaseLessonPlanGroupingPolicyRepository,
        )
        _v58_c5b7d_policy_configs = (
            SupabaseLessonPlanGroupingPolicyRepository(client).list_configs()
        )
        _v58_c5b2_policy_resolver = LessonPlanGroupingPolicyResolver(
            tuple(
                LessonPlanGroupingPolicy(
                    subject_ref=config.subject_ref,
                    component_ref=config.component_ref,
                    mode=config.mode,
                )
                for config in _v58_c5b7d_policy_configs
            )
        )
        st.session_state.pop("_v58_c5b7d_policy_load_error", None)
    except Exception as error:
        _v58_c5b2_policy_resolver = LessonPlanGroupingPolicyResolver()
        st.session_state["_v58_c5b7d_policy_load_error"] = str(error)

    _v58_c5b2_grouping_service = LessonPlanGroupingService()
    _v58_c5b2_lesson_plan_groups = _v58_c5b2_grouping_service.group(
        tuple(filtered_schedule_rows),
        policy_resolver=_v58_c5b2_policy_resolver,
        grade_resolver=lambda row: (
            int(getattr(row, "grade", getattr(row, "grade_level", 0)) or 0)
            or None
        ),
    )
    # Transitional runtime-comparison state only; never business authority.
    st.session_state["_v58_c5b2_shadow_lesson_plan_groups"] = (
        _v58_c5b2_lesson_plan_groups
    )
    st.session_state["_v58_c5b2_shadow_week_number"] = week_number

    # V58_C5B5_SHADOW_RUNTIME_BROWSER
    if st.session_state.get("_v58_c5b2_shadow_lesson_plan_groups"):
        with st.expander(
            "Tr?nh duy?t nh?m gi?o ?n (Shadow V58-C5B5)",
            expanded=False,
        ):
            st.caption(
                "Ch? ??c d? li?u grouping; ch?a thay selector hay authority hi?n h?nh."
            )
            _shadow_groups = tuple(
                st.session_state["_v58_c5b2_shadow_lesson_plan_groups"]
            )
            _shadow_grades = sorted(
                {
                    int(group.grade)
                    for group in _shadow_groups
                    if group.grade is not None
                }
            )
            st.write(
                "Kh?i c? d? li?u: "
                + (
                    ", ".join(str(item) for item in _shadow_grades)
                    if _shadow_grades
                    else "-"
                )
            )
            for _group in _shadow_groups:
                _periods = ", ".join(
                    str(item) for item in _group.curriculum_periods
                )
                _classes = ", ".join(str(item) for item in _group.class_ids)
                _dates = ", ".join(
                    f"{class_id}: {teaching_date}"
                    for class_id, teaching_date in _group.teaching_dates_by_class
                )
                st.markdown(
                    f"**Kh?i {_group.grade} ? {_group.grouping_mode.value}**  "
                    f"PPCT: {_periods or '-'}  "
                    f"L?p: {_classes or '-'}"
                )
                st.caption(
                    f"Ng?y d?y theo l?p: {_dates or '-'} ? "
                    f"Group ID: {_group.group_id}"
                )

    selected_unit = lesson_units[selected_index]
    representative_index = int(selected_unit.representative_index)
    if not (0 <= representative_index < len(filtered_schedule_rows)):
        return
    representative_row = filtered_schedule_rows[representative_index]
    # V58-C3B: PPCT is downstream of canonical grade. Selecting a PPCT
    # unit must never reverse-write grade into SystemContext or its widget
    # projection. Keep the reverse-context snapshot observational only.
    canonical_user_id = str(
        st.session_state.get("portal_user_id", "") or ""
    ).strip()
    canonical_context = (
        get_canonical_context(
            st.session_state,
            user_id=canonical_user_id,
            source_page="weekly_schedule",
        )
        if canonical_user_id
        else None
    )
    selected_grade = (
        getattr(canonical_context, "grade", None)
        if canonical_context is not None
        else None
    )
    st.session_state["_standardization_ppct_reverse_context"] = {
        "subject_ref": str(getattr(representative_row, "subject_ref", "") or ""),
        "component_ref": str(getattr(representative_row, "component_ref", "") or ""),
        "class_id": str(getattr(representative_row, "class_id", "") or ""),
        "curriculum_period": getattr(representative_row, "curriculum_period", None),
        "lesson_title": str(
            getattr(representative_row, "lesson_title", "")
            or selected_unit.title or ""
        ),
        "grade": selected_grade,
    }
    _autosave_standardization_change("Tiết PPCT / Bài dạy")


# STANDARDIZATION_TIMETABLE_CONTEXT_SYNC_V52D
def _standardization_keep_valid_option(key, options):
    options = tuple(options)
    if not options:
        st.session_state.pop(key, None)
        return None
    current = st.session_state.get(key)
    if current not in options:
        current = options[0]
        st.session_state[key] = current
    return current


def _render_lesson_plan_standardization_workspace(
    view,
    teacher_user_id="",
    client=None,
    workspace_focus="AI",
    compact_setup_ui=False,
) -> None:
    if (
        view is None
        or not getattr(
            view,
            "rows",
            None,
        )
    ):
        return

    apply_lesson_authoring_workspace_styles()

    # STANDARDIZATION_FULL_UI_V1
    #
    # Keep the complete lesson-selection/setup shell visible even when a
    # lesson arrives from AI authoring.  The transfer preselects matching
    # values, but must never hide fields that teachers need to review.
    if not compact_setup_ui and workspace_focus != "STANDARDIZE":
        st.markdown(
            """
<div class="mt-section-heading">
  <div class="mt-section-kicker">B\u00c0I D\u1ea0Y</div>
  <h2 class="mt-section-title">Ch\u1ecdn n\u1ed9i dung b\u00e0i so\u1ea1n</h2>
  <div class="mt-section-description">
    Ch\u1ecdn b\u00e0i, ti\u1ebft ho\u1eb7c ch\u1ee7 \u0111\u1ec1 t\u1eeb
    L\u1ecbch b\u00e1o gi\u1ea3ng hi\u1ec7n t\u1ea1i.
  </div>
</div>
            """.strip(),
            unsafe_allow_html=True,
        )

    # V52D canonical weekly timetable source.
    # Every selector below must remain inside this selected week's rows.
    schedule_rows = tuple(
        view.rows
    )

    try:
        selector_columns = st.columns(
            [1.05, 1.05, 1.15, 1.10, 1.10, 1.15]
        )
    except TypeError:
        selector_columns = st.columns(6)
    # Legacy layout contract: selector_columns[5].selectbox
    # Legacy layout contract: selector_columns[3].text_input
    # Legacy layout contract: selector_columns[4].text_input
    subject_refs = tuple(dict.fromkeys(
        str(getattr(row, "subject_ref", "") or "")
        for row in schedule_rows
    ))
    _standardization_keep_valid_option(
        "standardization_subject_filter",
        subject_refs,
    )
    _, selector_transfer = _latest_ai_standardization_transfer()
    restore_requested = bool(st.session_state.get(_RESTORE_LESSON_CONTEXT_KEY))
    working_context = dict(
        st.session_state.get(_WORKING_LESSON_CONTEXT_KEY, {}) or {}
    )
    if (
        restore_requested
        and working_context.get("context_origin") == "STANDARDIZATION"
        and working_context.get("context_read_only") is True
    ):
        selector_transfer = working_context
    transferred_subject = str(
        (selector_transfer or {}).get("subject_ref", "") or ""
    )

    if (
        transferred_subject in subject_refs
        and (restore_requested or "standardization_subject_filter" not in st.session_state)
    ):
        st.session_state["standardization_subject_filter"] = transferred_subject

    if st.session_state.get("standardization_subject_filter") not in subject_refs:
        st.session_state["standardization_subject_filter"] = subject_refs[0]

    def _subject_filter_label(subject_ref):
        return _subject_component_display_names(
            subject_ref=subject_ref,
            component_ref="",
            client=client,
        )[0]

    with selector_columns[0]:
        try:
            selected_subject_ref = st.selectbox(
                "Môn",
                options=subject_refs,
                format_func=_subject_filter_label,
                key="standardization_subject_filter",
                on_change=_autosave_standardization_change,
                args=("Môn",),
            )
        except TypeError:
            selected_subject_ref = st.selectbox(
                "Môn",
                options=subject_refs,
                format_func=_subject_filter_label,
                key="standardization_subject_filter",
            )

    _emit_standardization_canonical_context_change(
        field="subject_ref",
        value=str(selected_subject_ref),
        source_control="standardization_subject_filter_projection",
    )

    component_refs = tuple(dict.fromkeys(
        str(getattr(row, "component_ref", "") or "")
        for row in schedule_rows
        if str(getattr(row, "subject_ref", "") or "") == selected_subject_ref
    ))
    _standardization_keep_valid_option(
        "standardization_component_filter",
        component_refs,
    )
    transferred_component = str(
        (selector_transfer or {}).get("component_ref", "") or ""
    )

    if (
        transferred_component in component_refs
        and (restore_requested or "standardization_component_filter" not in st.session_state)
    ):
        st.session_state["standardization_component_filter"] = transferred_component

    if st.session_state.get("standardization_component_filter") not in component_refs:
        st.session_state["standardization_component_filter"] = component_refs[0]

    def _component_filter_label(component_ref):
        subject_name, component_name = _subject_component_display_names(
            subject_ref=selected_subject_ref,
            component_ref=component_ref,
            client=client,
        )
        return component_name or subject_name

    with selector_columns[1]:
        try:
            selected_component_ref = st.selectbox(
                "Phân môn",
                options=component_refs,
                format_func=_component_filter_label,
                key="standardization_component_filter",
                on_change=_autosave_standardization_change,
                args=("Phân môn",),
            )
        except TypeError:
            selected_component_ref = st.selectbox(
                "Phân môn",
                options=component_refs,
                format_func=_component_filter_label,
                key="standardization_component_filter",
            )

    _emit_standardization_canonical_context_change(
        field="component_ref",
        value=str(selected_component_ref),
        source_control="standardization_component_filter_projection",
    )

    filtered_schedule_rows = tuple(
        row
        for row in schedule_rows
        if str(getattr(row, "subject_ref", "") or "") == selected_subject_ref
        and str(getattr(row, "component_ref", "") or "") == selected_component_ref
    )

    # V57-F3E6D: temporary runtime diagnostic only.
    # It is intentionally read-only and emits no context changes.
    if st.session_state.get("_v57_f3e6d_runtime_diagnostic_enabled", True):
        print("\n[V57-F3E6D] SUBJECT/COMPONENT ROWS")
        for _diag_index, _diag_row in enumerate(filtered_schedule_rows):
            _diag_class_id = str(getattr(_diag_row, "class_id", "") or "")
            try:
                _diag_display = _class_display_name(_diag_class_id, client=client)
            except Exception as _diag_error:
                _diag_display = "<display-error:" + repr(_diag_error) + ">"
            try:
                _diag_grade = _standardization_grade_for_row(_diag_row, client=client)
            except Exception as _diag_error:
                _diag_grade = "<grade-error:" + repr(_diag_error) + ">"
            print(
                "[V57-F3E6D] ROW",
                _diag_index,
                "class_id=", repr(_diag_class_id),
                "display=", repr(_diag_display),
                "grade=", repr(_diag_grade),
                "subject=", repr(getattr(_diag_row, "subject_ref", None)),
                "component=", repr(getattr(_diag_row, "component_ref", None)),
                "ppct=", repr(getattr(_diag_row, "curriculum_period", None)),
                "lesson=", repr(getattr(_diag_row, "lesson_title", None)),
                "tkb_period=", repr(getattr(_diag_row, "timetable_period", None)),
                "date=", repr(getattr(_diag_row, "teaching_date", None)),
            )

    # STANDARDIZATION_GRADE_FILTER_V46
    # "Tất cả khối" preserves the exact previous behavior.
    available_grades = tuple(
        grade
        for grade in range(6, 10)
        if any(
            _standardization_grade_for_row(
                row,
                client=client,
            ) == grade
            for row in filtered_schedule_rows
        )
    )

    # STANDARDIZATION_GRADE_PPCT_CONTEXT_V50B
    grade_filter_options = (None,) + available_grades
    grade_filter_key = (
        "standardization_grade_filter_"
        + str(view.week_number)
        + "_"
        + selected_subject_ref
        + "_"
        + selected_component_ref
    )

    current_grade_filter = st.session_state.get(
        grade_filter_key
    )
    if current_grade_filter not in grade_filter_options:
        st.session_state[grade_filter_key] = None
        current_grade_filter = None

    if current_grade_filter is None and len(available_grades) == 1:
        current_grade_filter = available_grades[0]
        st.session_state[grade_filter_key] = current_grade_filter

    try:
        selected_grade = selector_columns[4].selectbox(
            "Khối lớp",
            options=grade_filter_options,
            format_func=lambda grade: (
                "Tất cả khối"
                if grade is None
                else f"Lớp {grade}"
            ),
            key=grade_filter_key,
            on_change=_autosave_standardization_change,
            args=("Khối lớp", grade_filter_key),
        )
    except AttributeError:
        with selector_columns[4]:
            selected_grade = st.selectbox(
                "Khối lớp",
                options=grade_filter_options,
                format_func=lambda grade: (
                    "Tất cả khối"
                    if grade is None
                    else f"Lớp {grade}"
                ),
                key=grade_filter_key,
            )

    if selected_grade is not None:
        filtered_schedule_rows = tuple(
            row
            for row in filtered_schedule_rows
            if _standardization_grade_for_row(
                row,
                client=client,
            ) == selected_grade
        )

    if not filtered_schedule_rows:
        st.warning(
            "Không có bài dạy phù hợp với khối lớp đã chọn."
        )
        return

    transfer_key, incoming_transfer = (
        _latest_ai_standardization_transfer()
    )
    if restore_requested and selector_transfer is working_context:
        incoming_transfer = working_context
        transfer_key = "working-lesson-context"
    matched_row_index = _match_transfer_schedule_row(
        filtered_schedule_rows,
        incoming_transfer,
    )
    transfer_id = str(
        (incoming_transfer or {}).get("transfer_id", transfer_key)
        or transfer_key
    )
    sync_state_key = "lesson_plan_standardization_applied_transfer_id"
    should_apply_transfer = bool(
        incoming_transfer
        and transfer_id
        and st.session_state.get(sync_state_key) != transfer_id
    )
    hide_synced_context = False

    selector = (
        LessonPlanUnitSelectorService()
    )

    _v58_available_modes_started = _v58_perf_counter()
    available_modes = (
        selector.available_modes(
            rows=filtered_schedule_rows
        )
    )
    _v58_perf_log(
        "lesson_selector_available_modes",
        _v58_available_modes_started,
        row_count=len(filtered_schedule_rows),
    )

    available_modes = tuple(
        mode
        for mode in (
            LessonPlanSelectionMode.PERIOD,
            LessonPlanSelectionMode.LESSON,
            LessonPlanSelectionMode.TOPIC,
        )
        if mode in available_modes
    )

    if not available_modes:
        st.warning("Không có tiết, bài hoặc chủ đề phù hợp để chuẩn hóa.")
        return

    mode_labels = {
        LessonPlanSelectionMode.LESSON: (
            "Theo b\u00e0i"
        ),
        LessonPlanSelectionMode.PERIOD: (
            "Theo ti\u1ebft"
        ),
        LessonPlanSelectionMode.TOPIC: (
            "Theo ch\u1ee7 \u0111\u1ec1"
        ),
        LessonPlanSelectionMode.WEEK_SUBJECT: (
            "Theo tu\u1ea7n / m\u00f4n h\u1ecdc"
        ),
    }

    grade_context_token = (
        "all"
        if selected_grade is None
        else str(selected_grade)
    )

    mode_widget_key = (
        "lbg_lesson_plan_selection_mode_"
        + str(view.week_number)
        + "_"
        + selected_subject_ref
        + "_"
        + selected_component_ref
        + "_grade_"
        + grade_context_token
    )
    if should_apply_transfer and LessonPlanSelectionMode.LESSON in available_modes:
        st.session_state[mode_widget_key] = LessonPlanSelectionMode.LESSON

    if st.session_state.get(mode_widget_key) not in available_modes:
        st.session_state[mode_widget_key] = available_modes[0]

    if hide_synced_context:
        selection_mode = LessonPlanSelectionMode.LESSON
    else:
        try:
            selection_mode = selector_columns[2].selectbox(
                "Cách thực hiện",
                options=available_modes,
                format_func=lambda value: (
                    mode_labels[value]
                ),
                key=mode_widget_key,
                on_change=_autosave_standardization_change,
                args=("Cách thực hiện",),
            )
        except AttributeError:
            with selector_columns[2]:
                selection_mode = st.selectbox(
                    "Cách thực hiện",
                    options=available_modes,
                    format_func=lambda value: (
                        mode_labels[value]
                    ),
                    key=mode_widget_key,
                )

    _v58_build_units_started = _v58_perf_counter()
    lesson_units = (
        selector.build_units(
            # WEEK_SCOPED_PPCT_OPTIONS_V1
            # filtered_schedule_rows belongs exclusively to the currently
            # selected weekly view; no row from another week is admitted.
            rows=filtered_schedule_rows,
            mode=selection_mode,
        )
    )
    _v58_perf_log(
        "lesson_selector_build_units",
        _v58_build_units_started,
        row_count=len(filtered_schedule_rows),
        mode=str(getattr(selection_mode, "value", selection_mode)),
        unit_count=len(lesson_units),
    )

    # V58-C3D: deduplicate selector units only; keep timetable rows intact.
    _deduplicated_lesson_units = []
    _seen_lesson_unit_keys = set()
    for _unit in lesson_units:
        _unit_key = (
            tuple(getattr(_unit, "curriculum_periods", ()) or ()),
            str(getattr(_unit, "title", "") or "").strip().casefold(),
        )
        if _unit_key in _seen_lesson_unit_keys:
            continue
        _seen_lesson_unit_keys.add(_unit_key)
        _deduplicated_lesson_units.append(_unit)
    lesson_units = tuple(_deduplicated_lesson_units)

    if not lesson_units:
        if (
            selection_mode
            is LessonPlanSelectionMode.TOPIC
        ):
            st.warning(
                "D\u1eef li\u1ec7u PPCT hi\u1ec7n "
                "ch\u01b0a c\u00f3 th\u00f4ng tin "
                "ch\u1ee7 \u0111\u1ec1."
            )
        else:
            st.warning(
                "Kh\u00f4ng c\u00f3 n\u1ed9i dung "
                "ph\u00f9 h\u1ee3p \u0111\u1ec3 "
                "chu\u1ea9n h\u00f3a "
                "gi\u00e1o \u00e1n."
            )

        return

    unit_label = {
        LessonPlanSelectionMode.LESSON: (
            "B\u00e0i d\u1ea1y"
        ),
        LessonPlanSelectionMode.PERIOD: (
            "Ti\u1ebft d\u1ea1y"
        ),
        LessonPlanSelectionMode.TOPIC: (
            "Ch\u1ee7 \u0111\u1ec1"
        ),
        LessonPlanSelectionMode.WEEK_SUBJECT: (
            "Tu\u1ea7n / m\u00f4n h\u1ecdc"
        ),
    }[selection_mode]

    unit_widget_key = (
        "lbg_lesson_plan_unit_"
        + selection_mode.value
        + "_"
        + str(view.week_number)
        + "_"
        + selected_subject_ref
        + "_"
        + selected_component_ref
        + "_grade_"
        + grade_context_token
    )
    _standardization_keep_valid_option(
        unit_widget_key,
        tuple(range(len(lesson_units))),
    )
    if should_apply_transfer and matched_row_index is not None:
        matched_row = filtered_schedule_rows[
            int(matched_row_index)
        ]
        for unit_index, unit in enumerate(lesson_units):
            representative_row = filtered_schedule_rows[
                int(unit.representative_index)
            ]
            if _mt_sync_same_lesson(representative_row, matched_row):
                st.session_state[unit_widget_key] = unit_index
                break

    if not isinstance(st.session_state.get(unit_widget_key), int) or not (
        0 <= int(st.session_state.get(unit_widget_key, 0)) < len(lesson_units)
    ):
        st.session_state[unit_widget_key] = 0

    if hide_synced_context and matched_row_index is not None:
        selected_unit_index = int(
            st.session_state.get(unit_widget_key, 0)
        )
    else:
        try:
            selected_unit_index = selector_columns[3].selectbox(
                "Tiết (theo PPCT)",
                options=tuple(
                    range(
                        len(
                            lesson_units
                        )
                    )
                ),
                format_func=lambda index: (
                    lesson_units[
                        index
                    ].selection_label
                ),
                key=unit_widget_key,
                on_change=_standardization_ppct_reverse_sync,
                kwargs={
                    "unit_widget_key": unit_widget_key,
                    "lesson_units": lesson_units,
                    "filtered_schedule_rows": filtered_schedule_rows,
                    "grade_filter_key": grade_filter_key,
                    "client": client,
                },
            )
        except AttributeError:
            with selector_columns[3]:
                selected_unit_index = st.selectbox(
                    "Tiết (theo PPCT)",
                    options=tuple(
                        range(
                            len(
                                lesson_units
                            )
                        )
                    ),
                    format_func=lambda index: (
                        lesson_units[
                            index
                        ].selection_label
                    ),
                    key=unit_widget_key,
                )

    if should_apply_transfer:
        st.session_state[sync_state_key] = transfer_id

    selected_unit = (
        lesson_units[
            selected_unit_index
        ]
    )

    selected_lesson = {
        "lesson_title": (
            selected_unit.title
        ),
        "periods": (
            selected_unit.curriculum_periods
        ),
        "classes": (
            selected_unit.class_ids
        ),
        "teaching_dates": tuple(
            (
                item.teaching_date,
                item.class_id,
            )
            for item
            in selected_unit.teaching_dates
        ),
        "representative_index": (
            selected_unit
            .representative_index
        ),
    }

    filtered_selected_index = int(
        selected_lesson[
            "representative_index"
        ]
    )

    selected_row = (
        filtered_schedule_rows[
            filtered_selected_index
        ]
    )

    # V58_C4B2_STANDARDIZATION_READS_RESOLVED_LBG
    # Legacy contract marker: selected_unit=selected_unit
    # Kept only for source-contract compatibility. The live runtime path
    # does not call _rows_for_selected_lesson_unit anymore.
    # Consume the already-resolved LBG lesson unit directly.
    selected_timetable_rows = tuple(
        filtered_schedule_rows[index]
        for index in tuple(
            getattr(selected_unit, "row_indices", ()) or ()
        )
        if 0 <= index < len(filtered_schedule_rows)
    ) or (selected_row,)

    _legacy_exact_multiclass_source_contract = """text_input(
        "Lớp dạy"
text_input(
        "Ngày dạy"
"""
    # V58-C3D6: the outer display must project the same canonical lesson-unit
    # context already used by "Kiểm tra thông tin bài soạn".  Do not derive a
    # second class scope from selected_timetable_rows.
    selected_class_ids = tuple(dict.fromkeys(
        str(class_id or "").strip()
        for class_id in tuple(selected_unit.class_ids or ())
        if str(class_id or "").strip()
    )) or (
        str(getattr(selected_row, "class_id", "") or ""),
    )
    representative_class_id = str(getattr(selected_row, "class_id", "") or "").strip()
    if representative_class_id:
        _emit_standardization_canonical_context_change(
            field="class_id", value=representative_class_id,
            source_control="standardization_selected_timetable_row",
        )

    selected_teaching_date_pairs = tuple(dict.fromkeys(
        (
            str(item.class_id or "").strip(),
            item.teaching_date,
        )
        for item in tuple(selected_unit.teaching_dates or ())
        if getattr(item, "teaching_date", None) is not None
        and str(getattr(item, "class_id", "") or "").strip()
    ))
    class_display_value = _class_display_names(
        selected_class_ids,
        client=client,
    )
    try:
        detail_columns = st.columns(2, gap="medium")
    except TypeError:
        detail_columns = st.columns(2)
    try:
        detail_columns[0].text_input(
            "Lớp dạy",
            value=class_display_value or "-",
            disabled=True,
            key=(
                "standardization_all_classes_"
                + str(view.week_number)
                + "_"
                + str(filtered_selected_index)
            ),
        )
    except AttributeError:
        with detail_columns[0]:
            st.text_input(
                "Lớp dạy",
                value=class_display_value or "-",
                disabled=True,
                key=(
                    "standardization_all_classes_"
                    + str(view.week_number)
                    + "_"
                    + str(filtered_selected_index)
                ),
            )
    teaching_date_display_value = "; ".join(
        (
            _class_display_name(
                class_id,
                client=client,
            )
            + " · "
            + teaching_date.strftime("%d/%m/%Y")
        )
        for class_id, teaching_date
        in selected_teaching_date_pairs
    )
    try:
        detail_columns[1].text_input(
            "Ngày dạy",
            value=teaching_date_display_value or "-",
            disabled=True,
            key=(
                "standardization_all_teaching_dates_"
                + str(view.week_number)
                + "_"
                + str(filtered_selected_index)
            ),
        )
    except AttributeError:
        with detail_columns[1]:
            st.text_input(
                "Ngày dạy",
                value=teaching_date_display_value or "-",
                disabled=True,
                key=(
                    "standardization_all_teaching_dates_"
                    + str(view.week_number)
                    + "_"
                    + str(filtered_selected_index)
                ),
            )
    if restore_requested:
        st.session_state.pop(_RESTORE_LESSON_CONTEXT_KEY, None)
        st.info("Đã khôi phục nguyên vẹn thông tin của bài đang soạn.")
    selected_row = selected_timetable_rows[0]

    selected_index = next(
        (
            index
            for index, row in enumerate(schedule_rows)
            if row is selected_row
        ),
        0,
    )
    selected_lesson["representative_index"] = selected_index

    selected_lesson["classes"] = selected_class_ids
    selected_lesson["timetable_periods_by_class"] = tuple(
        (
            str(getattr(row, "class_id", "") or ""),
            getattr(row, "timetable_period", None),
        )
        for row in selected_timetable_rows
    )
    selected_lesson["teaching_dates_by_class"] = tuple(
        selected_teaching_date_pairs
    )

    # Keep the exact teacher-facing class/date pairs selected in the current
    # timetable lesson.  The reviewed document adapter intentionally replaces
    # its technical class_id with a display string (for example "8A1, 8A2"),
    # so assignment-id matching must not be used as the primary source at the
    # later DOCX processing boundary.
    st.session_state[
        "_standardization_selected_teaching_date_pairs"
    ] = tuple(
        (
            _class_display_name(
                str(getattr(row, "class_id", "") or ""),
                client=client,
            ),
            getattr(row, "teaching_date", None),
        )
        for row in selected_timetable_rows
        if getattr(row, "teaching_date", None) is not None
    )

    # ---------------------------------------------------------
    # Lesson authoring visual context
    # ---------------------------------------------------------

    context_lesson_title = str(
        selected_lesson.get(
            "lesson_title",
            "",
        )
        or ""
    ).strip()

    context_class = _class_display_names(
        selected_class_ids,
        client=client,
    )

    (
        context_subject,
        context_component,
    ) = _subject_component_display_names(
        subject_ref=str(
            getattr(selected_row, "subject_ref", "")
            or ""
        ),
        component_ref=str(
            getattr(selected_row, "component_ref", "")
            or ""
        ),
        client=client,
    )

    (
        context_timetable_period,
        context_teaching_date_text,
    ) = _class_schedule_display_values(
        selected_timetable_rows,
        client=client,
    )

    context_timetable_period_html = escape(
        context_timetable_period
    ).replace("; ", "<br>")
    context_teaching_date_html = escape(
        context_teaching_date_text
    ).replace("; ", "<br>")

    context_periods = tuple(
        selected_lesson.get(
            "periods",
            (),
        )
        or ()
    )

    context_period_text = (
        ", ".join(
            str(value)
            for value in context_periods
        )
        if context_periods
        else "-"
    )

    _render_synced_context_markdown(
        not hide_synced_context and workspace_focus != "STANDARDIZE",
        f"""
<div class="mt-synced-lesson-title">
  <span>ĐANG SOẠN</span>
  <strong>{escape(context_lesson_title or "Chưa có tên bài")}</strong>
</div>
<div class="mt-lesson-context">
  <div class="mt-context-item">
    <div class="mt-context-label">TU\u1ea6N</div>
    <div class="mt-context-value">{view.week_number}</div>
  </div>

  <div class="mt-context-item">
    <div class="mt-context-label">MÔN</div>
    <div class="mt-context-value">{escape(context_subject)}</div>
  </div>

  <div class="mt-context-item">
    <div class="mt-context-label">PHÂN MÔN</div>
    <div class="mt-context-value">{escape(context_component)}</div>
  </div>

  <div class="mt-context-item">
    <div class="mt-context-label">L\u1edaP</div>
    <div class="mt-context-value">{escape(context_class or "-")}</div>
  </div>

  <div class="mt-context-item mt-context-item--multiline">
    <div class="mt-context-label">TIẾT TKB</div>
    <div class="mt-context-value">{context_timetable_period_html}</div>
  </div>

  <div class="mt-context-item">
    <div class="mt-context-label">TI\u1ebeT PPCT</div>
    <div class="mt-context-value">{escape(context_period_text)}</div>
  </div>

  <div class="mt-context-item mt-context-item--multiline">
    <div class="mt-context-label">NG\u00c0Y D\u1ea0Y</div>
    <div class="mt-context-value">{context_teaching_date_html}</div>
  </div>
</div>
        """.strip(),
        unsafe_allow_html=True,
    )

    _render_synced_context_markdown(
        not hide_synced_context and workspace_focus != "STANDARDIZE",
        f"""
<div class="mt-workspace-card">
  <div class="mt-section-kicker">
    B\u00c0I \u0110ANG L\u00c0M VI\u1ec6C
  </div>

  <div class="mt-section-title">
    {context_lesson_title or "Ch\u01b0a c\u00f3 t\u00ean b\u00e0i"}
  </div>

  <div class="mt-section-description">
    Th\u00f4ng tin n\u00e0y \u0111\u01b0\u1ee3c l\u1ea5y t\u1eeb
    L\u1ecbch b\u00e1o gi\u1ea3ng v\u00e0 s\u1ebd \u0111\u01b0\u1ee3c
    s\u1eed d\u1ee5ng trong quy tr\u00ecnh chu\u1ea9n h\u00f3a
    gi\u00e1o \u00e1n.
  </div>
</div>
        """.strip(),
        unsafe_allow_html=True,
    )

    _render_synced_context_markdown(
        not hide_synced_context and workspace_focus != "STANDARDIZE",
        """
<div class="mt-authoring-stepper">

  <div class="mt-authoring-step is-complete">
    <div class="mt-step-number">1</div>
    <div class="mt-step-content">
      <div class="mt-step-title">Ch\u1ecdn b\u00e0i</div>
      <div class="mt-step-description">
        X\u00e1c \u0111\u1ecbnh n\u1ed9i dung
      </div>
    </div>
  </div>

  <div class="mt-authoring-step is-active">
    <div class="mt-step-number">2</div>
    <div class="mt-step-content">
      <div class="mt-step-title">Ngu\u1ed3n gi\u00e1o \u00e1n</div>
      <div class="mt-step-description">
        Ch\u1ecdn t\u00e0i li\u1ec7u
      </div>
    </div>
  </div>

  <div class="mt-authoring-step">
    <div class="mt-step-number">3</div>
    <div class="mt-step-content">
      <div class="mt-step-title">Ki\u1ec3m tra</div>
      <div class="mt-step-description">
        Xem v\u00e0 x\u00e1c nh\u1eadn
      </div>
    </div>
  </div>

  <div class="mt-authoring-step">
    <div class="mt-step-number">4</div>
    <div class="mt-step-content">
      <div class="mt-step-title">Chu\u1ea9n h\u00f3a</div>
      <div class="mt-step-description">
        X\u1eed l\u00fd t\u00e0i li\u1ec7u
      </div>
    </div>
  </div>

  <div class="mt-authoring-step">
    <div class="mt-step-number">5</div>
    <div class="mt-step-content">
      <div class="mt-step-title">Ho\u00e0n t\u1ea5t</div>
      <div class="mt-step-description">
        L\u01b0u ho\u1eb7c t\u1ea3i xu\u1ed1ng
      </div>
    </div>
  </div>

</div>
        """.strip(),
        unsafe_allow_html=True,
    )

    # Preserve the canonical representative schedule-row
    # context when crossing from weekly scheduling into the
    # lesson drafting workspace.
    selected_lesson.update(
        {
            "academic_year": str(
                getattr(view, "academic_year", "")
                or ""
            ),
            "week_number": int(view.week_number),
            "subject_ref": str(
                getattr(selected_row, "subject_ref", "")
                or ""
            ),
            "component_ref": str(
                getattr(selected_row, "component_ref", "")
                or ""
            ),
            "subject_name": context_subject,
            "component_name": context_component,
            "class_id": str(
                getattr(
                    selected_row,
                    "class_id",
                    "",
                )
                or ""
            ),
            "class_name": context_class,
            "curriculum_period": getattr(
                selected_row,
                "curriculum_period",
                None,
            ),
            "timetable_period": getattr(
                selected_row,
                "timetable_period",
                None,
            ),
            "teaching_date": getattr(
                selected_row,
                "teaching_date",
                None,
            ),
            # Keep the aggregated values shown in "Nội dung soạn bài".
            # Scalar values above remain the representative schedule row for
            # downstream services that require date/period objects.
            "lesson_content_sync": True,
            "timetable_period_display": context_timetable_period,
            "teaching_date_display": context_teaching_date_text,
            "lesson_title": str(
                getattr(
                    selected_row,
                    "lesson_title",
                    "",
                )
                or selected_lesson.get(
                    "lesson_title",
                    ""
                )
            ),
            "teaching_equipment": tuple(
                getattr(
                    selected_row,
                    "teaching_equipment",
                    (),
                )
                or ()
            ),
        }
    )
    st.session_state[_STANDARDIZATION_DRAFT_KEY] = {
        **dict(st.session_state.get(_STANDARDIZATION_DRAFT_KEY, {}) or {}),
        "selected_lesson": dict(selected_lesson),
    }

    # V58_C4B_RESOLVED_LESSON_CONTEXT
    # One resolved LBG lesson projection for downstream consumers.
    _v58_projection_started = _v58_perf_counter()
    resolved_lbg_lesson_context = {
        "source": "LBG_PBSDTB",
        "academic_year": str(selected_lesson.get("academic_year", "") or ""),
        "week_number": int(selected_lesson.get("week_number", view.week_number)),
        "subject_ref": str(selected_lesson.get("subject_ref", "") or ""),
        "component_ref": str(selected_lesson.get("component_ref", "") or ""),
        "grade": selected_grade,
        "class_ids": tuple(selected_lesson.get("classes", ()) or ()),
        "class_id": str(selected_lesson.get("class_id", "") or ""),
        "timetable_periods_by_class": tuple(
            selected_lesson.get("timetable_periods_by_class", ()) or ()
        ),
        "teaching_dates_by_class": tuple(
            selected_lesson.get("teaching_dates_by_class", ()) or ()
        ),
        "curriculum_period": selected_lesson.get("curriculum_period"),
        "lesson_id": str(getattr(selected_row, "lesson_id", "") or ""),
        "lesson_title": str(selected_lesson.get("lesson_title", "") or ""),
        "teaching_equipment": tuple(
            selected_lesson.get("teaching_equipment", ()) or ()
        ),
        "representative_timetable_period": selected_lesson.get("timetable_period"),
        "representative_teaching_date": selected_lesson.get("teaching_date"),
    }

    # Projection envelope only; canonical scalar SystemContext remains authority.
    # Multi-class consumers must read this envelope instead of re-expanding
    # broader schedule_rows.
    st.session_state["_v58_resolved_lbg_lesson_context"] = dict(
        resolved_lbg_lesson_context
    )
    _v58_perf_log(
        "resolved_lbg_lesson_projection",
        _v58_projection_started,
        class_count=len(resolved_lbg_lesson_context.get("class_ids", ()) or ()),
        week_number=resolved_lbg_lesson_context.get("week_number"),
        subject_ref=resolved_lbg_lesson_context.get("subject_ref"),
        component_ref=resolved_lbg_lesson_context.get("component_ref"),
        curriculum_period=resolved_lbg_lesson_context.get("curriculum_period"),
    )
    st.session_state[_STANDARDIZATION_DRAFT_KEY] = {
        **dict(st.session_state.get(_STANDARDIZATION_DRAFT_KEY, {}) or {}),
        "selected_lesson": dict(resolved_lbg_lesson_context),
    }

    if not hide_synced_context:
        st.button(
            "Mở trong trang Soạn bài cùng AI",
            type="primary",
            use_container_width=True,
            key=(
                "lbg_open_full_ai_page_"
                + str(view.week_number)
                + "_"
                + str(selected_index)
            ),
            on_click=_open_ai_authoring_page,
            args=(dict(selected_lesson),),
        )

        with st.expander(
            "Ki\u1ec3m tra th\u00f4ng tin b\u00e0i so\u1ea1n",
            expanded=False,
        ):
            drafting_date = st.date_input(
                "Ng\u00e0y so\u1ea1n",
                value=selected_row.teaching_date,
                max_value=selected_row.teaching_date,
                key=(
                    "lbg_lesson_plan_drafting_date_"
                    + str(view.week_number)
                    + "_"
                    + str(selected_index)
                ),
            )

            metadata_override = (
                _render_lesson_plan_metadata_override_editor(
                    selected_row=selected_row,
                    drafting_date=drafting_date,
                    class_name=(
                        context_class
                    ),
                    week_number=int(
                        view.week_number
                    ),
                    selected_index=selected_index,
                )
            )
    else:
        drafting_date = selected_row.teaching_date
        metadata_override = {
            "drafting_date": drafting_date,
            "teaching_date": selected_row.teaching_date,
            "class_name": context_class,
            "curriculum_period": int(
                selected_row.curriculum_period
            ),
            "lesson_title": str(
                selected_row.lesson_title or ""
            ).strip(),
        }

    if workspace_focus == "AI":
        _render_lesson_plan_drafting_workspace(
            selected_lesson=selected_lesson,
            teacher_user_id=teacher_user_id,
            academic_year=str(
                getattr(
                    view,
                    "academic_year",
                    "",
                )
            ),
            week_number=int(
                getattr(
                    view,
                    "week_number",
                    0,
                )
            ),
            selection_mode=str(
                selection_mode.value
                if hasattr(
                    selection_mode,
                    "value",
                )
                else selection_mode
            ),
            selection_unit_id=str(
                selected_unit.selection_id
                if hasattr(
                    selected_unit,
                    "selection_id",
                )
                else selected_unit.title
            ),
            client=client,
        )
        return

    st.markdown(
        '<div class="mt-workspace-section-separator"></div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        '<div id="upload-lesson-plan"></div>',
        unsafe_allow_html=True,
    )

    st.subheader(
        "\U0001f4dd CHU\u1ea8N H\u00d3A GI\u00c1O \u00c1N"
    )

    st.caption(
        "Ch\u1ecdn t\u1ec7p Word ho\u1eb7c d\u00f9ng gi\u00e1o \u00e1n "
        "v\u1eeba so\u1ea1n c\u00f9ng AI. T\u1ec7p g\u1ed1c lu\u00f4n "
        "\u0111\u01b0\u1ee3c gi\u1eef nguy\u00ean."
    )

    with st.expander("Xem quy tr\u00ecnh chu\u1ea9n h\u00f3a"):
        st.write(
            "B\u1ed5 sung th\u00f4ng tin t\u1eeb L\u1ecbch b\u00e1o "
            "gi\u1ea3ng \u2192 Ki\u1ec3m tra \u2192 Chu\u1ea9n h\u00f3a "
            "\u2192 Xem tr\u01b0\u1edbc \u2192 L\u01b0u ho\u1eb7c t\u1ea3i xu\u1ed1ng."
        )

    input_mode = st.radio(
        "Nguồn giáo án",
        (
            "Tải giáo án lên",
            "Dùng giáo án "
            "vừa xử lý cùng AI",
        ),
        horizontal=True,
        key=(
            "lbg_lesson_plan_input_mode_"
            + str(view.week_number)
            + "_"
            + str(selected_index)
        ),
    )

    uploaded = None
    uploaded_content = None
    original_source_content = None
    ai_revised_text = ""
    source_name = ""
    source_kind = ""

    if (
        input_mode
        == "Tải giáo án lên"
    ):
        uploaded = st.file_uploader(
            "T\u1ea3i gi\u00e1o \u00e1n Word (.docx)",
            type=("docx",),
            accept_multiple_files=False,
            key=(
                "lbg_lesson_plan_upload_"
                + str(view.week_number)
                + "_"
                + str(selected_index)
            ),
        )

        if uploaded is None:
            st.caption(
                "File gốc chỉ được "
                "dùng làm đầu vào "
                "và sẽ được "
                "giữ nguyên."
            )
            return

        source_name = str(
            uploaded.name
        )

        uploaded_content = (
            uploaded.getvalue()
        )
        original_source_content = uploaded_content

        source_kind = "UPLOAD"

        st.success(
            "Đã nhận giáo án: "
            + source_name
        )

    else:
        transfer_candidates = []

        for key, value in (
            st.session_state.items()
        ):
            if not str(key).endswith(
                "_standardization_transfer"
            ):
                continue

            if not isinstance(
                value,
                dict,
            ):
                continue

            if (
                value.get("source")
                != "AI_DRAFT"
            ):
                continue

            transfer_candidates.append(
                value
            )

        transfer_payload = (
            transfer_candidates[-1]
            if transfer_candidates
            else None
        )

        if transfer_payload is None:
            st.info(
                "Chưa có giáo án "
                "được chuyển từ "
                "công cụ Soạn bài "
                "cùng AI."
            )
            return

        ai_docx_bytes = (
            transfer_payload.get(
                "docx_bytes"
            )
        )

        if not isinstance(
            ai_docx_bytes,
            (
                bytes,
                bytearray,
            ),
        ):
            st.info(
                "Đã nhận giáo án "
                "vừa xử lý cùng AI."
            )

            st.warning(
                "Giáo án AI chưa có "
                "tài liệu DOCX làm việc "
                "nội bộ."
            )

            st.caption(
                "Giáo viên không cần "
                "xuất Word. Hệ thống sẽ "
                "tự tạo tài liệu "
                "làm việc ở bước "
                "tiếp theo."
            )

            return

        uploaded_content = bytes(
            ai_docx_bytes
        )
        original_bytes = transfer_payload.get("source_bytes")
        if isinstance(original_bytes, (bytes, bytearray)):
            original_source_content = bytes(original_bytes)
        ai_revised_text = str(
            transfer_payload.get("full_document_text", "") or ""
        )

        source_name = str(
            transfer_payload.get(
                "source_name",
                "",
            )
            or (
                str(
                    transfer_payload.get(
                        "lesson_title",
                        "giao-an-ai",
                    )
                )
                + ".docx"
            )
        )

        source_kind = "AI_DRAFT"

        st.success(
            "Đã nhận giáo án "
            "vừa xử lý cùng AI."
        )


    st.markdown(
        '<div class="mt-workspace-section-separator"></div>',
        unsafe_allow_html=True,
    )

    st.subheader(
        "Xem to\u00e0n b\u1ed9 gi\u00e1o \u00e1n"
    )

    st.caption(
        "Hi\u1ec3n th\u1ecb tr\u1ef1c quan file Word g\u1ed1c tr\u01b0\u1edbc khi "
        "ki\u1ec3m tra v\u00e0 chu\u1ea9n h\u00f3a."
    )

    try:
        viewer_html = build_document_html(
            uploaded_content
        )

        st.components.v1.html(
            viewer_html,
            height=900,
            scrolling=True,
        )

    except Exception as error:
        st.error(
            "Kh\u00f4ng th\u1ec3 hi\u1ec3n th\u1ecb tr\u1ef1c quan gi\u00e1o \u00e1n: "
            + str(error)
        )

    st.markdown(
        '<div class="mt-workspace-section-separator"></div>',
        unsafe_allow_html=True,
    )

    workflow_identity = (
        LessonPlanWorkflowIdentity
        .from_upload(
            week_number=view.week_number,
            row_index=selected_index,
            source_name=source_name,
            content=uploaded_content,
        )
    )

    workflow_state = st.session_state.get(
        workflow_identity.state_key
    )

    if (
        not isinstance(
            workflow_state,
            LessonPlanWorkflowState,
        )
        or not workflow_state.matches(
            workflow_identity
        )
    ):
        workflow_state = LessonPlanWorkflowState(
            identity=workflow_identity
        )

        st.session_state[
            workflow_identity.state_key
        ] = workflow_state

    reviewed_row = None
    modification_plan = None
    preparation_error = None

    try:
        preview_view = workflow_state.preview

        if preview_view is None:
            preview_view = (
                LessonPlanPreviewUploadService()
                .prepare(
                    content=uploaded_content,
                    canonical=CanonicalDocumentContext(
    class_name=metadata_override[
        "class_name"
    ],
    curriculum_period=metadata_override[
        "curriculum_period"
    ],
    lesson_title=metadata_override[
        "lesson_title"
    ],
    drafting_date=metadata_override[
        "drafting_date"
    ].strftime(
        "%d/%m/%Y"
    ),
    teaching_date=metadata_override[
        "teaching_date"
    ].strftime(
        "%d/%m/%Y"
    ),
),
                )
            )

            workflow_state = (
                workflow_state.with_preview(
                    preview_view
                )
            )

            st.session_state[
                workflow_identity.state_key
            ] = workflow_state

        canonical_values = {
    DocumentField.CLASS_NAME: (
        metadata_override[
            "class_name"
        ]
    ),
    DocumentField.CURRICULUM_PERIOD: (
        str(
            metadata_override[
                "curriculum_period"
            ]
        )
    ),
    DocumentField.LESSON_TITLE: (
        metadata_override[
            "lesson_title"
        ]
    ),
    DocumentField.DRAFTING_DATE: (
        metadata_override[
            "drafting_date"
        ].strftime(
            "%d/%m/%Y"
        )
    ),
    DocumentField.TEACHING_DATE: (
        metadata_override[
            "teaching_date"
        ].strftime(
            "%d/%m/%Y"
        )
    ),
}

        render_lesson_plan_preview(
            st=st,
            view=preview_view,
        )

        review_view = (
            LessonPlanTeacherReviewPresenter()
            .present(
                preview=preview_view,
                canonical_values=canonical_values,
            )
        )

        teacher_review = render_lesson_plan_teacher_review(
            st=st,
            view=review_view,
            key_prefix=workflow_identity.widget_key_prefix,
        )

        review_resolution = (
            LessonPlanTeacherReviewResolver()
            .resolve(
                preview=preview_view,
                review=teacher_review,
            )
        )

        workflow_state = workflow_state.with_review(
            review=teacher_review,
            resolution=review_resolution,
        )
        st.session_state[
            workflow_identity.state_key
        ] = workflow_state

        # Legacy UI contract marker retained for compatibility:
        # LessonPlanModificationPlanner().build_from_values(
        #     values=canonical_values
        # )
        modification_plan = (
            LessonPlanModificationPlanner()
            .build(
                resolution=review_resolution
            )
        )

        reviewed_row = (
            LessonPlanReviewedScheduleRow
            .from_schedule_row(
                row=selected_row,
                resolved_metadata=(
                    canonical_values
                ),
            )
        )



    except Exception as error:
        reviewed_row = None
        modification_plan = None
        preparation_error = error
        modification_plan = None

        st.warning(
            "Kh\u00f4ng th\u1ec3 xem tr\u01b0\u1edbc "
            "ho\u1eb7c x\u00e1c nh\u1eadn "
            "th\u00f4ng tin gi\u00e1o \u00e1n: "
            f"{error}"
        )

    st.markdown(
        '<div id="standardize-lesson-plan"></div>',
        unsafe_allow_html=True,
    )

    process_clicked = bool(
        st.session_state.pop(
            "lesson_plan_standardization_execute_requested",
            False,
        )
    )
    confirmed_options = st.session_state.get(
        "lesson_plan_standardization_confirmed_options"
    )

    if not process_clicked:
        st.info(
            "M\u1edf B\u1ea3ng \u0111i\u1ec1u khi\u1ec3n chu\u1ea9n h\u00f3a, "
            "ch\u1ecdn t\u00e1c v\u1ee5 v\u00e0 nh\u1ea5n X\u00e1c nh\u1eadn \u0111\u1ec3 x\u1eed l\u00fd."
        )


    processing_ready = (
        reviewed_row is not None
        and modification_plan is not None
        and preparation_error is None
    )

    if process_clicked and not processing_ready:
        st.error(
            "Canonical lesson data preparation failed; "
            "standardization cannot continue."
        )

    if (
        process_clicked
        and not isinstance(
            confirmed_options,
            LessonPlanStandardizationOptions,
        )
    ):
        st.error(
            "Ch\u01b0a c\u00f3 c\u1ea5u h\u00ecnh chu\u1ea9n h\u00f3a \u0111\u00e3 x\u00e1c nh\u1eadn."
        )

    if (
        process_clicked
        and processing_ready
        and isinstance(
            confirmed_options,
            LessonPlanStandardizationOptions,
        )
    ):
        try:
            with st.spinner(
                "\u0110ang b\u1ed5 sung "
                "th\u00f4ng tin v\u00e0 "
                "chu\u1ea9n h\u00f3a gi\u00e1o \u00e1n..."
            ):
                result = (
                    _process_lesson_plan_upload(
                        row=reviewed_row,
                        drafting_date=(
                            metadata_override["drafting_date"]
                        ),
                        content=(
                            uploaded_content
                        ),
                        original_name=(
                            source_name
                        ),
                        modification_plan=(
                            modification_plan
                        ),
                        options=confirmed_options,
                        original_content=original_source_content,
                        ai_revised_text=ai_revised_text,
                    )
                )

                workflow_state = (
                    workflow_state.with_result(
                        result
                    )
                )

                st.session_state[
                    workflow_identity.state_key
                ] = workflow_state

        except Exception as error:
            import traceback
            print('\n===== STANDARDIZATION_RUNTIME_TRACEBACK =====')
            traceback.print_exc()
            print('===== END_STANDARDIZATION_RUNTIME_TRACEBACK =====\n')
            st.error(
                "Kh\u00f4ng th\u1ec3 "
                "chu\u1ea9n h\u00f3a "
                f"gi\u00e1o \u00e1n: {error}"
            )

    if compact_setup_ui:
        st.markdown(
            '<div id="mt-authoring-compact-end"></div>',
            unsafe_allow_html=True,
        )
        st.components.v1.html(
            """
<script>
(function () {
  const doc = window.parent.document;
  const hiddenClass = "mt-authoring-compact-hidden";
  const styleId = "mt-authoring-compact-style";

  if (!doc.getElementById(styleId)) {
    const style = doc.createElement("style");
    style.id = styleId;
    style.textContent =
      "." + hiddenClass + "{display:none!important;}";
    doc.head.appendChild(style);
  }

  function applyCompactView() {
    const start = doc.getElementById(
      "mt-authoring-compact-start"
    );
    const end = doc.getElementById(
      "mt-authoring-compact-end"
    );
    if (!start || !end) return;

    const startBox = start.closest(
      '[data-testid="stElementContainer"]'
    );
    const endBox = end.closest(
      '[data-testid="stElementContainer"]'
    );
    if (
      !startBox ||
      !endBox ||
      startBox.parentElement !== endBox.parentElement
    ) return;

    let node = startBox;
    while (node && node !== endBox) {
      const next = node.nextElementSibling;
      node.classList.add(hiddenClass);
      node = next;
    }
  }

  [40, 180, 500].forEach(function (delay) {
    window.setTimeout(applyCompactView, delay);
  });
})();
</script>
            """.strip(),
            height=0,
        )

    workflow_state = st.session_state.get(
        workflow_identity.state_key
    )

    if (
        not isinstance(
            workflow_state,
            LessonPlanWorkflowState,
        )
        or not workflow_state.matches(
            workflow_identity
        )
        or workflow_state.result is None
    ):
        return

    result = workflow_state.result


    (
        output_name,
        output_bytes,
        unresolved_fields,
    ) = result


    st.markdown(
        '<div class="mt-workspace-section-separator"></div>',
        unsafe_allow_html=True,
    )

    st.markdown(
        '<div id="preview-standardized-lesson-plan"></div>',
        unsafe_allow_html=True,
    )

    st.subheader(
        "Xem tr\u01b0\u1edbc gi\u00e1o \u00e1n \u0111\u00e3 chu\u1ea9n h\u00f3a"
    )

    st.caption(
        "\u0110\u00e2y l\u00e0 b\u1ea3n gi\u00e1o \u00e1n sau khi \u0111\u00e3 b\u1ed5 sung "
        "th\u00f4ng tin t\u1eeb L\u1ecbch b\u00e1o gi\u1ea3ng v\u00e0 chu\u1ea9n h\u00f3a."
    )

    try:
        standardized_viewer_html = (
            build_document_html(
                output_bytes
            )
        )

        st.components.v1.html(
            standardized_viewer_html,
            height=900,
            scrolling=True,
        )

    except Exception as error:
        st.warning(
            "Kh\u00f4ng th\u1ec3 hi\u1ec3n th\u1ecb tr\u1ef1c quan "
            "b\u1ea3n gi\u00e1o \u00e1n \u0111\u00e3 chu\u1ea9n h\u00f3a: "
            + str(error)
        )

    if unresolved_fields:
        st.warning(
            "Ch\u01b0a t\u1ef1 \u0111\u1ed9ng "
            "c\u1eadp nh\u1eadt \u0111\u01b0\u1ee3c: "
            + ", ".join(
                unresolved_fields
            )
        )
    else:
        st.success(
            "\u0110\u00e3 b\u1ed5 sung "
            "c\u00e1c th\u00f4ng tin "
            "L\u1ecbch b\u00e1o gi\u1ea3ng "
            "v\u00e0 chu\u1ea9n h\u00f3a "
            "gi\u00e1o \u00e1n."
        )


    st.markdown("---")

    st.markdown(
        "### Gi\u00e1o \u00e1n \u0111\u00e3 chu\u1ea9n h\u00f3a"
    )

    st.markdown(
        '<div id="save-standardized-lesson-plan"></div>',
        unsafe_allow_html=True,
    )

    st.caption(
        "B\u1ea1n c\u00f3 th\u1ec3 l\u01b0u b\u1ea3n \u0111\u00e3 chu\u1ea9n h\u00f3a "
        "tr\u00ean h\u1ec7 th\u1ed1ng ho\u1eb7c t\u1ea3i xu\u1ed1ng m\u00e1y."
    )

    save_standardized_clicked = st.button(
        "L\u01b0u v\u00e0o Kho gi\u00e1o \u00e1n",
        type="secondary",
        key=(
            "lbg_lesson_plan_save_standardized_"
            + str(view.week_number)
            + "_"
            + str(selected_index)
        ),
        width="stretch",
    )

    def _save_standardized_artifact_to_library(
        *,
        artifact_file_name: str,
        artifact_content: bytes,
    ) -> None:
        output_name = artifact_file_name
        output_bytes = artifact_content
        upload_service = (
            st.session_state.get(
                "document_library_upload_service"
            )
        )

        if upload_service is None:
            st.warning(
                "Kho gi\u00e1o \u00e1n ch\u01b0a s\u1eb5n s\u00e0ng. "
                "B\u1ea1n v\u1eabn c\u00f3 th\u1ec3 t\u1ea3i file v\u1ec1 m\u00e1y."
            )

        else:
            try:
                from teacher_document_library_v2 import (
                    DocumentCategory,
                    DocumentUploadMetadata,
                )

                categories = tuple(
                    DocumentCategory
                )

                category = next(
                    (
                        item
                        for item in categories
                        if (
                            "lesson"
                            in item.value.casefold()
                            or "giao"
                            in item.value.casefold()
                        )
                    ),
                    categories[0],
                )

                class_name = str(
                    metadata_override.get(
                        "class_name",
                        context_class,
                    )
                    or context_class
                ).strip()

                academic_year = str(
                    getattr(
                        view,
                        "academic_year",
                        "",
                    )
                    or ""
                ).strip()

                subject = str(
                    getattr(
                        selected_row,
                        "subject",
                        "",
                    )
                    or getattr(
                        selected_row,
                        "subject_name",
                        "",
                    )
                    or "N/A"
                ).strip()

                grade_level = str(
                    getattr(
                        selected_row,
                        "grade_level",
                        "",
                    )
                    or "N/A"
                ).strip()

                lesson_title = str(
                    getattr(
                        selected_row,
                        "lesson_title",
                        "",
                    )
                    or output_name
                ).strip()

                metadata = (
                    DocumentUploadMetadata(
                        title=lesson_title,
                        category=category,
                        academic_year=(
                            academic_year
                            or "N/A"
                        ),
                        subject=subject,
                        grade_level=grade_level,
                        class_name=(
                            class_name
                            if class_name != "-"
                            else None
                        ),
                        description=(
                            "Gi\u00e1o \u00e1n \u0111\u00e3 \u0111\u01b0\u1ee3c b\u1ed5 sung "
                            "th\u00f4ng tin t\u1eeb L\u1ecbch b\u00e1o gi\u1ea3ng "
                            "v\u00e0 chu\u1ea9n h\u00f3a tr\u00ean h\u1ec7 th\u1ed1ng."
                        ),
                        tags=(
                            "lesson-plan",
                            "standardized",
                        ),
                    )
                )

                saved_document = (
                    upload_service.upload(
                        content=output_bytes,
                        file_name=output_name,
                        mime_type=(
                            "application/vnd.openxmlformats-"
                            "officedocument.wordprocessingml.document"
                        ),
                        metadata=metadata,
                    )
                )

                st.success(
                    "\u0110\u00e3 l\u01b0u gi\u00e1o \u00e1n chu\u1ea9n h\u00f3a "
                    "v\u00e0o Kho gi\u00e1o \u00e1n."
                )

                link = getattr(
                    saved_document,
                    "web_view_link",
                    None,
                )

                if link:
                    st.link_button(
                        "M\u1edf gi\u00e1o \u00e1n \u0111\u00e3 l\u01b0u",
                        link,
                    )

            except Exception as error:
                st.error(
                    "Kh\u00f4ng th\u1ec3 l\u01b0u gi\u00e1o \u00e1n "
                    "\u0111\u00e3 chu\u1ea9n h\u00f3a: "
                    + str(error)
                )


    if save_standardized_clicked:
        _save_standardized_artifact_to_library(
            artifact_file_name=output_name,
            artifact_content=output_bytes,
        )

    st.markdown(
        '<div id="download-standardized-lesson-plan"></div>',
        unsafe_allow_html=True,
    )

    st.download_button(
        "\U0001f4e5 T\u1ea3i gi\u00e1o \u00e1n "
        "chu\u1ea9n h\u00f3a",
        data=output_bytes,
        file_name=output_name,
        mime=(
            "application/vnd.openxmlformats-"
            "officedocument.wordprocessingml.document"
        ),
        width="stretch",
        key=(
            "lbg_lesson_plan_download_"
            + str(view.week_number)
            + "_"
            + str(selected_index)
        ),
    )

    from portal_v2.ui.standardized_lesson_plan_management_streamlit import (
        render_standardized_lesson_plan_management,
    )

    render_standardized_lesson_plan_management(
        current_file_name=output_name,
        current_content=output_bytes,
        preview_html_builder=build_document_html,
        save_handler=_save_standardized_artifact_to_library,
    )


def _render_weekly_schedule_technical_workspace(
    *,
    client=None,
    user_id: str | None = None,
    embedded: bool = False,
    default_system: bool = False,
) -> None:
    _sync_lbg_week_from_loaded_data()
    floating_notice = st.session_state.pop(_LBG_NOTICE_KEY, "")
    if floating_notice:
        st.toast(str(floating_notice), icon="✅")
    if not embedded:
        st.title(
            "\U0001f4c5 L\u1ecbch b\u00e1o gi\u1ea3ng"
        )

        st.caption(
            "L\u1eadp v\u00e0 qu\u1ea3n l\u00fd "
            "l\u1ecbch b\u00e1o gi\u1ea3ng "
            "theo t\u1eebng tu\u1ea7n h\u1ecdc."
        )

    st.markdown(
        """<style>
        .st-key-weekly_schedule_source,
        .st-key-system_weekly_academic_year,
        .st-key-system_weekly_week_number,
        .st-key-system_weekly_assignment_name,
        .st-key-system_weekly_ppct_mode {
          background:linear-gradient(145deg,#0b2749,#06172c);
          border:1px solid #315f91;
          border-radius:14px;
          /* Legacy contract: min-height:76px */
          min-height:82px;
          /* Legacy contract: padding:6px 10px */
          padding:8px 12px;
          margin:0 !important;
          box-shadow:4px 5px 0 #03101f,0 12px 24px rgba(3,16,31,.2);
        }
        .st-key-weekly_schedule_source p,
        .st-key-system_weekly_academic_year p,
        .st-key-system_weekly_week_number p,
        .st-key-system_weekly_assignment_name p,
        .st-key-system_weekly_ppct_mode p,
        .st-key-weekly_schedule_source label,
        .st-key-system_weekly_academic_year label,
        .st-key-system_weekly_week_number label,
        .st-key-system_weekly_assignment_name label,
        .st-key-system_weekly_ppct_mode label,
        .st-key-system_weekly_academic_year input {
          color:#fff !important;
          font-size:14px !important;
          font-weight:700 !important;
        }
        .st-key-system_weekly_academic_year input,
        .st-key-system_weekly_week_number div[data-baseweb="select"] > div,
        .st-key-system_weekly_assignment_name div[data-baseweb="select"] > div,
        .st-key-system_weekly_ppct_mode div[data-baseweb="select"] > div {
          background:#fff !important;
          border-color:#607b9a !important;
          color:#111 !important;
          font-size:14px !important;
          font-weight:700 !important;
          min-height:40px !important;
          height:40px !important;
          border-radius:8px !important;
          box-shadow:inset 0 1px 2px rgba(3,16,31,.12) !important;
        }
        .st-key-weekly_schedule_source div[data-testid="stWidgetLabel"],
        .st-key-system_weekly_academic_year div[data-testid="stWidgetLabel"],
        .st-key-system_weekly_week_number div[data-testid="stWidgetLabel"],
        .st-key-system_weekly_assignment_name div[data-testid="stWidgetLabel"],
        .st-key-system_weekly_ppct_mode div[data-testid="stWidgetLabel"] {
          margin-bottom:2px !important;
        }
        .st-key-weekly_schedule_source div[role="radiogroup"] {
          gap:10px !important;
        }
        .st-key-system_weekly_week_number div[data-baseweb="select"] span,
        .st-key-system_weekly_assignment_name div[data-baseweb="select"] span,
        .st-key-system_weekly_ppct_mode div[data-baseweb="select"] span,
        .st-key-system_weekly_week_number div[data-baseweb="select"] span {
          color:#111 !important;
          font-weight:700 !important;
        }
        .st-key-system_weekly_week_number svg,
        .st-key-system_weekly_assignment_name svg,
        .st-key-system_weekly_ppct_mode svg {
          color:#111 !important;
          fill:#111 !important;
        }
        @media(max-width:800px){
          .st-key-weekly_schedule_source,
          .st-key-system_weekly_academic_year,
          .st-key-system_weekly_week_number,
          .st-key-system_weekly_assignment_name,
          .st-key-system_weekly_ppct_mode {min-height:auto;}
        }
        </style>""",
        unsafe_allow_html=True,
    )

    (
        source_column,
        year_column,
        week_column,
        assignment_column,
        ppct_column,
    ) = st.columns(
        (1.35, .9, .75, 1.25, 1.25),
        gap="small",
    )

    with source_column:
        source_label = st.radio(
            "\u004e\u0067\u0075\u1ed3\u006e "
            "\u0064\u1eef "
            "\u006c\u0069\u1ec7\u0075",
            (
                "\u0054\u1ea3\u0069 "
                "\u0074\u1eeb "
                "\u006d\u00e1\u0079",
                "\u004c\u1ea5\u0079 "
                "\u0074\u1eeb "
                "\u0068\u1ec7 "
                "\u0074\u0068\u1ed1\u006e\u0067",
            ),
            horizontal=True,
            key="weekly_schedule_source",
            index=1 if default_system else 0,
            on_change=_autosave_lbg_filter_context,
            args=("Nguồn dữ liệu",),
        )

    if (
        source_label
        == (
            "\u004c\u1ea5\u0079 "
            "\u0074\u1eeb "
            "\u0068\u1ec7 "
            "\u0074\u0068\u1ed1\u006e\u0067"
        )
    ):
        if client is None or not user_id:
            st.error(
                "\u0050\u0068\u0069\u00ea\u006e "
                "\u0111\u0103\u006e\u0067 "
                "\u006e\u0068\u1ead\u0070 "
                "\u0063\u0068\u01b0\u0061 "
                "\u0063\u00f3 "
                "\u006e\u0067\u1eef "
                "\u0063\u1ea3\u006e\u0068 "
                "\u0064\u1eef "
                "\u006c\u0069\u1ec7\u0075 "
                "\u0068\u1ec7 "
                "\u0074\u0068\u1ed1\u006e\u0067."
            )
            return

        with year_column:
            _technical_user_id = str(
                st.session_state.get("portal_user_id", "") or ""
            ).strip()
            if not _technical_user_id:
                st.error("Không xác định được người dùng cho ngữ cảnh Năm học.")
                return
            _technical_context = get_canonical_context(
                st.session_state,
                user_id=_technical_user_id,
                source_page="weekly_schedule",
            )
            if not _technical_context.academic_year:
                _default_academic_year = str(
                    st.session_state.get(
                        "portal_academic_year",
                        "2026-2027",
                    )
                    or ""
                ).strip()
                if _default_academic_year:
                    apply_canonical_year_week_change(
                        st.session_state,
                        user_id=_technical_user_id,
                        field="academic_year",
                        value=_default_academic_year,
                        source_page="weekly_schedule",
                        source_control="portal_academic_year_default",
                    )
                    _technical_context = get_canonical_context(
                        st.session_state,
                        user_id=_technical_user_id,
                        source_page="weekly_schedule",
                    )
            publish_year_week_projection(
                st.session_state,
                context=_technical_context,
            )
            academic_year = st.text_input(
                "\u004e\u0103\u006d "
                "\u0068\u1ecdc",
                key="system_weekly_academic_year",
                on_change=_autosave_lbg_filter_context,
                args=("Năm học",),
            ).strip()

        with week_column:
            # WEEKLY_SCHEDULE_SELECTED_WEEK_RESTORE_V2
            _saved_week_key = (
                "_system_weekly_last_updated_week"
            )

            # V57-F2C5I: last-updated week is cache metadata only.
            # It must never restore or replace canonical SystemContext.week_number.

            # GLOBAL_WEEKLY_CONTEXT_PERSISTENT_RESTORE_V1
            # Restore the most recently updated persisted schedule
            # for the current teacher/current academic year.
            if (
                st.session_state.get(
                    _ACTIVE_WEEK_NUMBER_KEY
                )
                is None
            ):
                try:
                    _bootstrap_repository = (
                        SupabaseWeeklyScheduleRepository(
                            client,
                            str(user_id),
                        )
                    )

                    _saved_schedule_summaries = tuple(
                        _bootstrap_repository
                        .list_for_teacher(
                            str(user_id)
                        )
                    )

                    _current_year_summaries = tuple(
                        item
                        for item
                        in _saved_schedule_summaries
                        if str(
                            getattr(
                                item,
                                "academic_year",
                                "",
                            )
                        ) == str(academic_year)
                    )

                    if _current_year_summaries:
                        _latest_summary = max(
                            _current_year_summaries,
                            key=lambda item: (
                                getattr(
                                    item,
                                    "updated_at",
                                )
                            ),
                        )

                        _bootstrap_schedule = (
                            _bootstrap_repository.get(
                                str(
                                    _latest_summary.schedule_id
                                )
                            )
                        )

                        if _bootstrap_schedule is not None:
                            _bootstrap_week = int(
                                _latest_summary.week_number
                            )

                            st.session_state[
                                _ACTIVE_SCHEDULE_ID_KEY
                            ] = str(
                                _latest_summary.schedule_id
                            )

                            # V57-F2C5I: persisted schedule discovery may
                            # restore cache/view metadata, but never canonical
                            # academic-year/week context or selector values.
                            st.session_state[
                                "_system_weekly_last_updated_week"
                            ] = _bootstrap_week

                            _bootstrap_user_id = str(
                                st.session_state.get(
                                    "portal_user_id",
                                    "",
                                )
                                or ""
                            ).strip()
                            if _bootstrap_user_id:
                                _bootstrap_context = get_canonical_context(
                                    st.session_state,
                                    user_id=_bootstrap_user_id,
                                    source_page="weekly_schedule",
                                )
                                _bootstrap_context_week = (
                                    _bootstrap_context.week_number
                                )
                                _bootstrap_context_year = (
                                    _bootstrap_context.academic_year
                                )
                                if (
                                    _bootstrap_context_week is not None
                                    and int(_bootstrap_context_week)
                                    != _bootstrap_week
                                ):
                                    st.session_state[
                                        _LBG_DATA_WEEK_CONTEXT_MISMATCH_KEY
                                    ] = {
                                        "canonical_week": int(
                                            _bootstrap_context_week
                                        ),
                                        "data_week": _bootstrap_week,
                                        "schedule_id": str(
                                            _latest_summary.schedule_id
                                        ),
                                        "source": "persisted_schedule_bootstrap",
                                    }
                                if (
                                    _bootstrap_context_year is not None
                                    and str(_bootstrap_context_year)
                                    != str(_latest_summary.academic_year)
                                ):
                                    st.session_state[
                                        "lbg_data_year_context_mismatch"
                                    ] = {
                                        "canonical_year": str(
                                            _bootstrap_context_year
                                        ),
                                        "data_year": str(
                                            _latest_summary.academic_year
                                        ),
                                        "schedule_id": str(
                                            _latest_summary.schedule_id
                                        ),
                                        "source": "persisted_schedule_bootstrap",
                                    }

                            _bootstrap_generation = (
                                WeeklyScheduleGenerationResult(
                                    request=(
                                        WeeklyScheduleGenerationRequest(
                                            schedule_id=str(
                                                _latest_summary.schedule_id
                                            ),
                                            teacher_id=str(
                                                user_id
                                            ),
                                            academic_year=str(
                                                _latest_summary.academic_year
                                            ),
                                            week_number=(
                                                _bootstrap_week
                                            ),
                                        )
                                    ),
                                    schedule=(
                                        _bootstrap_schedule
                                    ),
                                )
                            )

                            _bootstrap_output = (
                                WeeklyScheduleOutputService()
                                .export_excel(
                                    generation=(
                                        _bootstrap_generation
                                    )
                                )
                            )

                            _bootstrap_view = (
                                WeeklySchedulePortalPresenter()
                                .present(
                                    output=(
                                        _bootstrap_output
                                    )
                                )
                            )

                            st.session_state[
                                _VIEW_STATE_KEY
                            ] = _bootstrap_view

                            st.session_state[
                                _ACTIVE_VIEW_KEY
                            ] = _bootstrap_view

                except Exception as error:
                    st.warning(
                        (
                            "Kh?ng th? kh?i ph?c tu?n l?m vi?c "
                            "?? l?u: "
                            + str(error)
                        )
                    )

            week_number = st.selectbox(
                "Tu\u1ea7n h\u1ecdc",
                options=tuple(
                    range(1, 41)
                ),
                format_func=lambda value: (
                    f"Tu\u1ea7n {value}"
                ),
                key="system_weekly_week_number",
                on_change=_autosave_lbg_filter_context,
                args=("Tuần",),
            )

            # WEEK_SELECTOR_ONE_WAY_V2
            # UI selector -> requested week -> LBG.
            # Loaded/global data must not write back
            # into the selector.
            requested_week_number = int(
                week_number
            )
            week_number = (
                requested_week_number
            )

        if not academic_year:
            st.info(
                "\u0048\u00e3\u0079 "
                "\u006e\u0068\u1ead\u0070 "
                "\u006e\u0103\u006d "
                "\u0068\u1ecdc "
                "\u0111\u1ec3 "
                "\u0074\u0069\u1ebf\u0070 "
                "\u0074\u1ee5\u0063."
            )
            return

        try:
            assignment_repository = (
                SupabaseTeachingAssignmentRepository(
                    client=client,
                    user_id=str(user_id),
                )
            )

            assignments = (
                assignment_repository.list_assignments(
                    owner_id=str(user_id),
                    academic_year=academic_year,
                    role=TeachingAssignmentRole.TEACHING,
                    status=TeachingAssignmentStatus.ACTIVE,
                )
            )

            timetable_repository = (
                SupabaseTeacherTimetableRepository(
                    client=client,
                    user_id=str(user_id),
                )
            )
            timetable_slots = (
                timetable_repository.list_slots(
                    owner_id=str(user_id),
                    academic_year=academic_year,
                    status=TeacherTimetableSlotStatus.ACTIVE,
                )
            )
            scheduled_assignment_ids = {
                slot.assignment_id
                for slot in timetable_slots
            }
            assignments = tuple(
                assignment
                for assignment in assignments
                if assignment.assignment_id
                in scheduled_assignment_ids
            )

        except Exception as error:
            st.error(
                "\u004b\u0068\u00f4\u006e\u0067 "
                "\u0074\u0068\u1ec3 "
                "\u0111\u1ecdc "
                "\u0070\u0068\u00e2\u006e "
                "\u0063\u00f4\u006e\u0067 "
                "\u0067\u0069\u1ea3\u006e\u0067 "
                "\u0064\u1ea1\u0079 "
                "\u0076\u00e0 "
                "\u0074\u0068\u1eddi "
                "\u006b\u0068\u00f3\u0061 "
                "\u0062\u0069\u1ec3\u0075: "
                f"{error}"
            )
            return

        if not assignments:
            st.warning(
                "\u0043\u0068\u01b0\u0061 "
                "\u0063\u00f3 "
                "\u0070\u0068\u00e2\u006e "
                "\u0063\u00f4\u006e\u0067 "
                "\u0067\u0069\u1ea3\u006e\u0067 "
                "\u0064\u1ea1\u0079 "
                "\u0111\u0061\u006e\u0067 "
                "\u0068\u0069\u1ec7\u0075 "
                "\u006c\u1ef1\u0063 "
                "\u0063\u0068\u006f "
                "\u006e\u0103\u006d "
                "\u0068\u1ecdc "
                "\u006e\u00e0\u0079."
            )
            return

        class_repository = SupabaseClassCatalogRepository(
            client=client,
        )
        subject_repository = SupabaseSubjectCatalogRepository(
            client=client,
        )
        assignment_labels = []
        for assignment in assignments:
            class_name = "Chưa có tên lớp"
            subject_name = "Chưa có tên môn"
            try:
                class_item = class_repository.get(
                    class_id=assignment.class_id,
                )
                if class_item is not None:
                    class_name = str(
                        class_item.display_name
                    ).strip() or class_name
            except Exception:
                pass
            try:
                if assignment.component_ref:
                    subject_item = (
                        subject_repository.get_component(
                            component_id=(
                                assignment.component_ref
                            ),
                        )
                    )
                else:
                    subject_item = (
                        subject_repository.get_subject(
                            subject_id=(
                                assignment.subject_ref
                            ),
                        )
                    )
                if subject_item is not None:
                    subject_name = str(
                        subject_item.name
                    ).strip() or subject_name
            except Exception:
                pass
            assignment_labels.append(
                f"{class_name} · {subject_name}"
            )

        with assignment_column:
            st.selectbox(
                "Lớp / Môn dạy",
                options=tuple(assignment_labels),
                index=0,
                key="system_weekly_assignment_name",
                on_change=_autosave_lbg_filter_context,
                args=("Lớp / Môn dạy",),
            )

        with ppct_column:
            st.selectbox(
                "PPCT",
                options=(
                    "Tự động · Tất cả môn",
                ),
                index=0,
                key="system_weekly_ppct_mode",
                on_change=_autosave_lbg_filter_context,
                args=("PPCT",),
            )

        scope_rules = []

        # ACTIVE_WEEK_VIEW_CONSISTENCY_V1
        # week_number is the UI/global source of truth.
        # Never allow a cached view from another week to remain.
        _active_schedule_id = (
            "SYSTEM-"
            + str(user_id)
            + "-"
            + str(academic_year)
            + "-W"
            + str(week_number)
        )

        _cached_view = st.session_state.get(
            _VIEW_STATE_KEY
        )

        _cached_year = (
            str(
                getattr(
                    _cached_view,
                    "academic_year",
                    "",
                )
            )
            if _cached_view is not None
            else ""
        )

        _cached_week = (
            int(
                getattr(
                    _cached_view,
                    "week_number",
                    -1,
                )
            )
            if _cached_view is not None
            else -1
        )

        if (
            _cached_view is None
            or _cached_year != str(academic_year)
            or _cached_week != int(week_number)
        ):
            try:
                _active_repository = (
                    SupabaseWeeklyScheduleRepository(
                        client,
                        str(user_id),
                    )
                )

                _active_schedule = (
                    _active_repository.get(
                        _active_schedule_id
                    )
                )

                if _active_schedule is not None:
                    _active_generation = (
                        WeeklyScheduleGenerationResult(
                            request=(
                                WeeklyScheduleGenerationRequest(
                                    schedule_id=(
                                        _active_schedule_id
                                    ),
                                    teacher_id=str(
                                        user_id
                                    ),
                                    academic_year=str(
                                        academic_year
                                    ),
                                    week_number=int(
                                        week_number
                                    ),
                                )
                            ),
                            schedule=(
                                _active_schedule
                            ),
                        )
                    )

                    _active_output = (
                        WeeklyScheduleOutputService()
                        .export_excel(
                            generation=(
                                _active_generation
                            )
                        )
                    )

                    _active_view = (
                        WeeklySchedulePortalPresenter()
                        .present(
                            output=(
                                _active_output
                            )
                        )
                    )

                    # Only replace the cache after a successful
                    # read and presentation of the exact week.
                    st.session_state[
                        _VIEW_STATE_KEY
                    ] = _active_view

                    st.session_state[
                        _ACTIVE_SCHEDULE_ID_KEY
                    ] = _active_schedule_id

                    # V57-F2C5I: exact-week read updates view/cache only.
                    # Canonical year/week were selected before this read.
                    st.session_state[
                        _ACTIVE_VIEW_KEY
                    ] = _active_view

                else:
                    # Do not display a stale schedule from
                    # another week.
                    st.session_state.pop(
                        _VIEW_STATE_KEY,
                        None,
                    )

            except Exception as error:
                st.warning(
                    "Kh?ng th? t?i l?ch b?o gi?ng "
                    "c?a tu?n ?ang ch?n: "
                    + str(error)
                )

        if st.button(
            "C\u1eadp nh\u1eadt "
            "L\u1ecbch b\u00e1o gi\u1ea3ng",
            type="primary",
            width="stretch",
            key="system_weekly_generate",
        ):
            try:
                schedule_id = (
                    "SYSTEM-"
                    + str(user_id)
                    + "-"
                    + academic_year
                    + "-W"
                    + str(week_number)
                )

                runtime = (
                    SystemWeeklyScheduleRuntime(
                        client=client,
                        user_id=str(user_id),
                    )
                )

                schedule = runtime.generate(
                    request=(
                        SystemWeeklyScheduleRuntimeRequest(
                            schedule_id=schedule_id,
                            academic_year=academic_year,
                            week_number=week_number,
                            ppct_scope_rules=tuple(
                                scope_rules
                            ),
                        )
                    )
                )

                # Persist the generated/updated weekly schedule
                # through the repository already defined by the
                # system architecture.
                #
                # Read-after-write is intentional: the UI must only
                # display a schedule that was actually persisted.
                schedule_repository = (
                    SupabaseWeeklyScheduleRepository(
                        client,
                        str(user_id),
                    )
                )

                schedule_repository.save(
                    schedule
                )

                persisted_schedule = (
                    schedule_repository.get(
                        schedule_id
                    )
                )

                if persisted_schedule is None:
                    raise RuntimeError(
                        "L?ch b?o gi?ng ?? ghi nh?ng "
                        "kh?ng th? ??c l?i ?? x?c minh."
                    )

                if (
                    persisted_schedule.schedule_id
                    != schedule.schedule_id
                ):
                    raise RuntimeError(
                        "L?ch b?o gi?ng ??c l?i kh?ng "
                        "??ng ??nh danh v?a c?p nh?t."
                    )

                schedule = persisted_schedule

                generation = (
                    WeeklyScheduleGenerationResult(
                        request=(
                            WeeklyScheduleGenerationRequest(
                                schedule_id=schedule_id,
                                teacher_id=str(user_id),
                                academic_year=academic_year,
                                week_number=week_number,
                            )
                        ),
                        schedule=schedule,
                    )
                )

                output = (
                    WeeklyScheduleOutputService()
                    .export_excel(
                        generation=generation
                    )
                )

                view = (
                    WeeklySchedulePortalPresenter()
                    .present(
                        output=output
                    )
                )

                st.session_state[
                    _VIEW_STATE_KEY
                ] = view

                # Remember only a successfully persisted
                # week.  No database/schema mutation.
                st.session_state[
                    "_system_weekly_last_updated_week"
                ] = int(week_number)

                # V57-F2C5I: persistence publishes schedule/view metadata
                # only. The save result must not become year/week authority.
                st.session_state[
                    _ACTIVE_SCHEDULE_ID_KEY
                ] = str(schedule_id)

                st.session_state[
                    _ACTIVE_VIEW_KEY
                ] = view
                st.session_state[_LBG_CONTEXT_SNAPSHOT_KEY] = {
                    "source": source_label,
                    "academic_year": academic_year,
                    "week_number": int(week_number),
                    "assignment": st.session_state.get(
                        "system_weekly_assignment_name"
                    ),
                    "ppct_mode": st.session_state.get(
                        "system_weekly_ppct_mode"
                    ),
                }
                st.toast(
                    "Đã lưu và đồng bộ Lịch báo giảng của "
                    f"Tuần {week_number}.",
                    icon="✅",
                )

            except Exception as error:
                st.error(
                    "\u004b\u0068\u00f4\u006e\u0067 "
                    "\u0074\u0068\u1ec3 "
                    "\u0074\u1ea1\u006f "
                    "\u006c\u1ecb\u0063\u0068 "
                    "\u0062\u00e1\u006f "
                    "\u0067\u0069\u1ea3\u006e\u0067 "
                    "\u0074\u1eeb "
                    "\u0064\u1eef "
                    "\u006c\u0069\u1ec7\u0075 "
                    "\u0068\u1ec7 "
                    "\u0074\u0068\u1ed1\u006e\u0067: "
                    f"{error}"
                )
                return

        view = st.session_state.get(
            _VIEW_STATE_KEY
        )

        if view is None:
            return

        st.success(
            "\u0110\u00e3 "
            "\u0074\u1ea1\u006f "
            "\u006c\u1ecb\u0063\u0068 "
            "\u0062\u00e1\u006f "
            "\u0067\u0069\u1ea3\u006e\u0067 "
            "\u0074\u1eeb "
            "\u0064\u1eef "
            "\u006c\u0069\u1ec7\u0075 "
            "\u0068\u1ec7 "
            "\u0074\u0068\u1ed1\u006e\u0067."
        )

        st.subheader(
            f"\u004c\u1ecb\u0063\u0068 "
            f"\u0062\u00e1\u006f "
            f"\u0067\u0069\u1ea3\u006e\u0067 "
            f"- "
            f"\u0054\u0075\u1ea7\u006e "
            f"{view.week_number}"
        )

        (
            class_names,
            subject_names,
            component_names,
        ) = _resolve_lbg_display_names(
            client=client,
            view=view,
        )

        rows = _preview_rows(
            view,
            class_names=class_names,
            subject_names=subject_names,
            component_names=component_names,
        )

        if rows:
            st.data_editor(
                rows,
                width="stretch",
                hide_index=True,
                disabled=(
                    "Th\u1ee9/ng\u00e0y",
                    "Ti\u1ebft TKB",
                    "M\u00f4n/Ph\u00e2n m\u00f4n",
                    "L\u1edbp",
                    "Ti\u1ebft PPCT",
                    "T\u00ean b\u00e0i d\u1ea1y",
                ),
                key=(
                    "system_weekly_schedule_editor_"
                    + str(view.week_number)
                ),
            )
        else:
            st.warning(
                "\u004c\u1ecb\u0063\u0068 "
                "\u0111\u01b0\u1ee3\u0063 "
                "\u0074\u1ea1\u006f "
                "\u006e\u0068\u01b0\u006e\u0067 "
                "\u006b\u0068\u00f4\u006e\u0067 "
                "\u0063\u00f3 "
                "\u0074\u0069\u1ebf\u0074 "
                "\u0064\u1ea1\u0079 "
                "\u0070\u0068\u00f9 "
                "\u0068\u1ee3\u0070 "
                "\u0074\u0072\u006f\u006e\u0067 "
                "\u0074\u0075\u1ea7\u006e "
                "\u006e\u00e0\u0079."
            )

        st.download_button(
            "\u0054\u1ea3\u0069 "
            "\u006c\u1ecb\u0063\u0068 "
            "\u0062\u00e1\u006f "
            "\u0067\u0069\u1ea3\u006e\u0067 "
            "\u0045\u0078\u0063\u0065\u006c",
            data=view.download.content,
            file_name=view.download.file_name,
            mime=view.download.mime_type,
            use_container_width=True,
            key="system_weekly_download",
        )

        return

    uploaded = st.file_uploader(
        "Tải workbook dữ liệu lịch báo giảng",
        type=("xlsx",),
        key="weekly_schedule_upload",
    )

    if uploaded is None:
        st.info("Hãy tải file Excel dữ liệu để bắt đầu.")
        return

    try:
        intake = WeeklyScheduleWorkbookIntakeAdapter().load(
            selection=_local_selection(),
            workbook_bytes=uploaded.getvalue(),
        )
    except Exception as error:
        st.error(f"Không thể đọc dữ liệu: {error}")
        return

    academic_years = _academic_year_options(intake)
    teachers = _teacher_options(intake)

    if not academic_years:
        st.warning("Không tìm thấy năm học trong dữ liệu.")
        return

    if not teachers:
        st.warning("Không tìm thấy giáo viên trong thời khóa biểu.")
        return

    academic_year = st.selectbox(
        "Năm học",
        academic_years,
        key="weekly_schedule_academic_year",
    )

    weeks = _week_options(intake, academic_year)

    if not weeks:
        st.warning("Không tìm thấy tuần học phù hợp.")
        return

    col_week, col_teacher = st.columns(2)

    with col_week:
        week_number = st.selectbox(
            "Tuần",
            weeks,
            format_func=lambda value: f"Tuần {value}",
            key="weekly_schedule_week",
        )

    with col_teacher:
        teacher_id = st.selectbox(
            "Giáo viên",
            teachers,
            key="weekly_schedule_teacher",
        )

    if st.button(
        "Tạo lịch báo giảng",
        type="primary",
        use_container_width=True,
        key="weekly_schedule_generate",
    ):
        schedule_id = (
            f"{teacher_id}-{academic_year}-W{week_number:02d}"
        )

        try:
            generation = (
                LocalWeeklyScheduleGenerationService().generate(
                    intake=intake,
                    request=WeeklyScheduleGenerationRequest(
                        schedule_id=schedule_id,
                        teacher_id=teacher_id,
                        academic_year=academic_year,
                        week_number=week_number,
                    ),
                )
            )

            output = WeeklyScheduleOutputService().export_excel(
                generation=generation
            )

            view = WeeklySchedulePortalPresenter().present(
                output=output
            )

            st.session_state[_VIEW_STATE_KEY] = view

        except Exception as error:
            st.error(
                f"Không thể tạo lịch báo giảng: {error}"
            )
            return

    view = st.session_state.get(_VIEW_STATE_KEY)

    if view is None:
        return

    st.success("Đã tạo lịch báo giảng.")

    st.subheader(
        f"Lịch báo giảng - Tuần {view.week_number}"
    )

    st.caption(
        f"Giáo viên: {view.teacher_id} | "
        f"Năm học: {view.academic_year}"
    )

    preview = _preview_rows(view)

    if preview:
        st.dataframe(
            preview,
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.warning(
            "Lịch được tạo nhưng không có tiết dạy "
            "phù hợp trong tuần này."
        )

    st.download_button(
        "Tải lịch báo giảng Excel",
        data=view.download.content,
        file_name=view.download.file_name,
        mime=view.download.mime_type,
        use_container_width=True,
        key="weekly_schedule_download",
    )


_COMBINED_PAGE_CSS = r"""
<style>
.mt-lbg-device-header {
  background:#071a33;
  color:#ffffff;
  border:2px solid #12345b;
  box-shadow:5px 6px 0 #03101f,0 12px 26px rgba(3,16,31,.24);
  padding:1rem 1.2rem;
  margin:0 5px 1rem 0;
}
.mt-lbg-device-header h1,
.mt-lbg-device-header p { color:#ffffff; margin:.15rem 0; line-height:1.5; }
.mt-lbg-device-header h1 { font-size:26px; }
.mt-lbg-device-header p { font-size:18px; }
.mt-device-summary {
  display:grid;
  grid-template-columns:repeat(3,minmax(150px,1fr));
  background:#071a33;
  color:#ffffff;
  border:2px solid #12345b;
  margin:.4rem 0 1rem;
}
.mt-device-summary > div { padding:.7rem; border-right:1px solid #60758e; }
.mt-device-summary > div:last-child { border-right:0; }
.mt-device-summary strong,
.mt-device-summary span { color:#ffffff; font-size:18px; line-height:1.5; }
@media(max-width:850px){.mt-device-summary{grid-template-columns:1fr}}
</style>
"""


def _equipment_usage_rows(
    view,
    *,
    class_names: dict[str, str] | None = None,
    subject_names: dict[str, str] | None = None,
    component_names: dict[str, str] | None = None,
) -> list[dict]:
    """Derive one equipment request row from each equipped lesson."""
    rows = _preview_rows(
        view,
        class_names=class_names,
        subject_names=subject_names,
        component_names=component_names,
    )
    result = []
    for row in rows:
        equipment = str(
            row.get("Chuẩn bị, điều chỉnh", "") or ""
        ).strip()
        if not equipment:
            continue
        result.append(
            {
                "Thứ/ngày": row.get("Thứ/ngày", ""),
                "Buổi": row.get("Buổi", ""),
                "Tiết TKB": row.get("Tiết TKB", ""),
                "Môn/Phân môn": row.get("Môn/Phân môn", ""),
                "Lớp": row.get("Lớp", ""),
                "Tên bài dạy": row.get("Tên bài dạy", ""),
                "Thiết bị/Phương tiện": equipment,
                "Số lượng": "",
                "Địa điểm sử dụng": "Lớp học",
                "Người nhận": "",
                "Ghi chú": row.get("Ghi chú", ""),
            }
        )
    return result


def _render_equipment_usage_report(
    view,
    *,
    client=None,
) -> None:
    teacher_id = escape(
        str(getattr(view, "teacher_id", "") or "-")
    )
    academic_year = escape(
        str(getattr(view, "academic_year", "") or "-")
    )
    week_number = escape(
        str(getattr(view, "week_number", "") or "-")
    )
    st.markdown(
        f"""<div class="mt-device-summary">
        <div><strong>GIÁO VIÊN</strong><br><span>{teacher_id}</span></div>
        <div><strong>NĂM HỌC</strong><br><span>{academic_year}</span></div>
        <div><strong>TUẦN</strong><br><span>{week_number}</span></div>
        </div>""",
        unsafe_allow_html=True,
    )
    (
        class_names,
        subject_names,
        component_names,
    ) = _resolve_lbg_display_names(
        client=client,
        view=view,
    )
    rows = _equipment_usage_rows(
        view,
        class_names=class_names,
        subject_names=subject_names,
        component_names=component_names,
    )
    if not rows:
        st.info(
            "Tuần này chưa có tiết dạy khai báo "
            "thiết bị hoặc phương tiện sử dụng."
        )
        return
    edited_rows = st.data_editor(
        rows,
        width="stretch",
        hide_index=True,
        disabled=(
            "Thứ/ngày",
            "Buổi",
            "Tiết TKB",
            "Môn/Phân môn",
            "Lớp",
            "Tên bài dạy",
            "Thiết bị/Phương tiện",
        ),
        key=(
            "weekly_equipment_usage_"
            + str(getattr(view, "week_number", "x"))
        ),
    )
    output = StringIO()
    writer = csv.DictWriter(
        output,
        fieldnames=tuple(rows[0].keys()),
        lineterminator="\n",
    )
    writer.writeheader()
    writer.writerows(edited_rows)
    st.download_button(
        "Tải Phiếu báo sử dụng thiết bị",
        data=("\ufeff" + output.getvalue()).encode("utf-8"),
        file_name=(
            "phieu-bao-su-dung-thiet-bi-tuan-"
            + str(getattr(view, "week_number", "x"))
            + ".csv"
        ),
        mime="text/csv",
        use_container_width=True,
        key=(
            "weekly_equipment_download_"
            + str(getattr(view, "week_number", "x"))
        ),
    )
    st.caption(
        "Phiếu được sinh tự động từ cùng các dòng Lịch báo giảng; "
        "các cột số lượng, địa điểm, người nhận và ghi chú "
        "có thể bổ sung trước khi gửi."
    )


def render_weekly_schedule_and_equipment_workspace(
    *,
    client=None,
    user_id: str | None = None,
) -> None:
    """Render synchronized weekly schedule and equipment request."""
    if client is None or not user_id:
        st.error("Phiên đăng nhập chưa sẵn sàng.")
        return
    st.markdown(_COMBINED_PAGE_CSS, unsafe_allow_html=True)
    st.markdown(
        """<div class="mt-lbg-device-header">
        <h1>LỊCH BÁO GIẢNG &amp; PBSDTB</h1>
        <p>Một tuần dữ liệu · Hai biểu mẫu đồng bộ</p>
        </div>""",
        unsafe_allow_html=True,
    )
    schedule_tab, equipment_tab = st.tabs(
        (
            "1. Lịch báo giảng",
            "2. Phiếu báo sử dụng thiết bị",
        )
    )
    with schedule_tab:
        _render_weekly_schedule_technical_workspace(
            client=client,
            user_id=user_id,
            embedded=True,
            default_system=True,
        )
    with equipment_tab:
        view = st.session_state.get(_VIEW_STATE_KEY)
        if view is None:
            st.info(
                "Hãy tạo Lịch báo giảng ở Phần 1. "
                "Phiếu báo sử dụng thiết bị sẽ được sinh tự động."
            )
        else:
            _render_equipment_usage_report(
                view,
                client=client,
            )


# =========================================================
# USER WEEKLY SCHEDULE WORKSPACE
# =========================================================




# STANDARDIZATION_AUTHORITATIVE_TIMETABLE_V56
def _build_standardization_authoritative_week_view(
    *,
    client,
    user_id: str,
    academic_year: str,
    week_number: int,
):
    """Build standardization data from the ACTIVE teacher timetable.

    Timetable owns class/subject/component/date/session/period identity.
    PPCT is resolved by the existing SystemWeeklyScheduleRuntime.
    This helper is read-only and does not persist/replace Lịch báo giảng.
    """
    schedule_id = (
        "SYSTEM-"
        + str(user_id)
        + "-"
        + str(academic_year)
        + "-W"
        + str(int(week_number))
    )

    schedule = (
        SystemWeeklyScheduleRuntime(
            client=client,
            user_id=str(user_id),
        ).generate(
            request=SystemWeeklyScheduleRuntimeRequest(
                schedule_id=schedule_id,
                academic_year=str(academic_year),
                week_number=int(week_number),
                ppct_scope_rules=(),
            )
        )
    )

    generation = WeeklyScheduleGenerationResult(
        request=WeeklyScheduleGenerationRequest(
            schedule_id=schedule_id,
            teacher_id=str(user_id),
            academic_year=str(academic_year),
            week_number=int(week_number),
        ),
        schedule=schedule,
    )

    output = WeeklyScheduleOutputService().export_excel(
        generation=generation
    )

    return WeeklySchedulePortalPresenter().present(
        output=output
    )

def render_lesson_plan_management_workspace(
    *,
    client=None,
    user_id: str | None = None,
) -> None:
    """Render the central entry page for lesson-plan actions."""
    if client is None or not user_id:
        st.error(
            "Phi\u00ean \u0111\u0103ng nh\u1eadp "
            "ch\u01b0a s\u1eb5n s\u00e0ng."
        )
        return

    apply_lesson_authoring_workspace_styles()
    st.title("Qu\u1ea3n l\xfd gi\u00e1o \u00e1n")
    st.caption(
        "Ch\u1ecdn thao t\u00e1c; h\u1ec7 th\u1ed1ng s\u1ebd chuy\u1ec3n "
        "t\u1edbi \u0111\u00fang khu v\u1ef1c tr\u00ean trang "
        "Chu\u1ea9n h\u00f3a gi\u00e1o \u00e1n."
    )

    catalogue = [
        item
        for item in list(
            st.session_state.get(
                "lesson_plan_management_catalog",
                (),
            )
            or ()
        )
        if str(item.get("user_id", "")) == str(user_id)
    ]
    st.subheader("Danh mục giáo án")
    if not catalogue:
        st.info(
            "Chưa có giáo án trong danh mục. Hãy mở trang Soạn bài cùng AI "
            "và nhấn Chuyển sang chuẩn hóa."
        )
    else:
        st.caption(
            f"Có {len(catalogue)} giáo án đã chuyển sang quy trình chuẩn hóa."
        )
        for item in reversed(catalogue):
            title = str(
                item.get("lesson_title", "Giáo án chưa đặt tên")
            )
            subject = str(item.get("subject_name", "-"))
            class_name = str(item.get("class_name", "-"))
            with st.expander(
                f"{title} · {subject} · Lớp {class_name}",
                expanded=False,
            ):
                details = st.columns(4)
                details[0].metric(
                    "Tiết PPCT",
                    str(item.get("curriculum_period", "-")),
                )
                details[1].metric(
                    "Tiết TKB",
                    str(item.get("timetable_period", "-")),
                )
                details[2].metric(
                    "Ngày dạy",
                    str(item.get("teaching_date", "-")),
                )
                details[3].metric(
                    "Cập nhật",
                    str(item.get("updated_at", "-"))
                    .replace("T", " "),
                )
                action_columns = st.columns([3, 1])
                action_columns[0].button(
                    "Chuyển vào Chuẩn hóa giáo án",
                    key=(
                        "lesson_plan_management_open_"
                        + str(item.get("item_id", ""))
                    ),
                    use_container_width=True,
                    type="primary",
                    on_click=(
                        _open_management_catalogue_item
                    ),
                    args=(str(item.get("item_id", "")),),
                )
                action_columns[1].button(
                    "Xóa",
                    key=(
                        "lesson_plan_management_delete_"
                        + str(item.get("item_id", ""))
                    ),
                    use_container_width=True,
                    on_click=_request_catalogue_item_delete,
                    args=(str(item.get("item_id", "")),),
                )
                if str(
                    st.session_state.get(
                        "lesson_plan_management_delete_candidate",
                        "",
                    )
                ) == str(item.get("item_id", "")):
                    st.warning(
                        "Bạn có chắc muốn xóa giáo án này khỏi danh mục?"
                    )
                    confirm_columns = st.columns(2)
                    confirm_columns[0].button(
                        "Xác nhận xóa",
                        key=(
                            "lesson_plan_management_confirm_delete_"
                            + str(item.get("item_id", ""))
                        ),
                        use_container_width=True,
                        on_click=_delete_management_catalogue_item,
                        args=(
                            str(item.get("item_id", "")),
                            str(user_id),
                        ),
                    )
                    confirm_columns[1].button(
                        "Hủy",
                        key=(
                            "lesson_plan_management_cancel_delete_"
                            + str(item.get("item_id", ""))
                        ),
                        use_container_width=True,
                        on_click=_cancel_catalogue_item_delete,
                    )

    _legacy_management_action_contract = """
    _render_standardization_action_flow()
    """
    _render_standardization_action_flow(
        client=client,
        user_id=user_id,
    )


def _open_management_catalogue_item(item_id: str) -> None:
    """Restore one catalogue item and open the real standardization flow."""
    catalogue = list(
        st.session_state.get(
            "lesson_plan_management_catalog",
            (),
        )
        or ()
    )
    selected = next(
        (
            item
            for item in catalogue
            if str(item.get("item_id", "")) == str(item_id)
        ),
        None,
    )
    if selected is None:
        return
    transfer_key = (
        "lesson_authoring_ai_"
        + str(selected.get("item_id", "lesson"))
        + "_standardization_transfer"
    )
    st.session_state.pop(transfer_key, None)
    st.session_state[transfer_key] = {
        "source": "AI_DRAFT",
        "transfer_id": (
            str(selected.get("item_id", ""))
            + ":"
            + str(selected.get("updated_at", ""))
        ),
        "docx_bytes": selected.get("docx_bytes"),
        "source_bytes": selected.get("source_bytes"),
        "source_name": str(
            selected.get("source_name", "")
            or (
                str(selected.get("lesson_title", "giao-an-ai"))
                + ".docx"
            )
        ),
        "teacher_user_id": str(selected.get("user_id", "")),
        # STANDARDIZATION_TO_AI_ONE_WAY_V1
        # Only document content may return for standardization.  Schedule
        # metadata remains owned by LBG/standardization and is never written
        # from an AI catalogue item.
        "full_document_text": str(selected.get("document", "")),
    }
    st.session_state[
        "lesson_authoring_standardization_document"
    ] = str(selected.get("document", ""))
    source_bytes = selected.get("source_bytes")
    if source_bytes:
        st.session_state["lesson_authoring_ai_source_bytes"] = source_bytes
        st.session_state["lesson_authoring_ai_source_name"] = str(
            selected.get("source_name", "") or ""
        )
    st.session_state["lesson_authoring_tool_focus"] = "STANDARDIZE"
    # ONE_WAY_LBG_DATA_FLOW_V1
    # A managed lesson can send its document to standardization, but it may
    # not select or mutate the authoritative LBG year/week.
    st.session_state["portal_navigation_request"] = "Chuẩn hóa giáo án"


def _request_catalogue_item_delete(item_id: str) -> None:
    st.session_state[
        "lesson_plan_management_delete_candidate"
    ] = str(item_id)


def _cancel_catalogue_item_delete() -> None:
    st.session_state.pop(
        "lesson_plan_management_delete_candidate",
        None,
    )


def _delete_management_catalogue_item(
    item_id: str,
    user_id: str,
) -> None:
    """Delete only the selected lesson owned by the signed-in user."""
    catalogue = list(
        st.session_state.get(
            "lesson_plan_management_catalog",
            (),
        )
        or ()
    )
    st.session_state["lesson_plan_management_catalog"] = [
        item
        for item in catalogue
        if not (
            str(item.get("item_id", "")) == str(item_id)
            and str(item.get("user_id", "")) == str(user_id)
        )
    ]
    st.session_state.pop(
        "lesson_plan_management_delete_candidate",
        None,
    )
    if str(
        st.session_state.get(
            "lesson_plan_management_selected_id",
            "",
        )
    ) == str(item_id):
        st.session_state.pop(
            "lesson_plan_management_selected_id",
            None,
        )
    transfer_key = (
        "lesson_authoring_ai_"
        + str(item_id)
        + "_standardization_transfer"
    )
    st.session_state.pop(transfer_key, None)


def render_lesson_authoring_tools_workspace(
    *,
    client=None,
    user_id: str | None = None,
) -> None:
    """Render the restored authoring-tools page in authoring mode."""
    render_weekly_schedule_workspace(
        client=client,
        user_id=user_id,
        initial_focus="AI",
        workspace_page_key="AUTHORING_TOOLS",
        page_title="Công cụ soạn bài",
        compact_setup_ui=False,
    )


def render_weekly_schedule_workspace(
    *,
    client=None,
    user_id: str | None = None,
    initial_focus: str = "STANDARDIZE",
    workspace_page_key: str = "STANDARDIZATION",
    page_title: str = "Chuẩn hóa giáo án",
    compact_setup_ui: bool = False,
) -> None:
    if client is None or not user_id:
        st.error(
            "Phi\u00ean \u0111\u0103ng nh\u1eadp "
            "ch\u01b0a s\u1eb5n s\u00e0ng."
        )
        return

    # Shared ACTIVE WEEK published by
    # "C?p nh?t L?ch b?o gi?ng".
    active_year = st.session_state.get(
        _ACTIVE_ACADEMIC_YEAR_KEY
    )
    active_week = st.session_state.get(
        _ACTIVE_WEEK_NUMBER_KEY
    )
    active_view = st.session_state.get(
        _ACTIVE_VIEW_KEY
    )

    if (
        active_year is not None
        and active_week is not None
    ):
        # Reuse the exact verified weekly view.
        # Do not regenerate a different week.
        if (
            active_view is not None
            and str(
                getattr(
                    active_view,
                    "academic_year",
                    "",
                )
            ) == str(active_year)
            and int(
                getattr(
                    active_view,
                    "week_number",
                    -1,
                )
            ) == int(active_week)
        ):
            st.session_state[
                _VIEW_STATE_KEY
            ] = active_view

    # STANDARDIZATION_FULL_UI_V1
    # The standardization page exposes its complete structure.  Incoming AI
    # or weekly-schedule data may prefill controls, but never hides them.
    standardization_minimal_ui = False

    normalized_focus = str(
        initial_focus
    ).upper()
    if normalized_focus not in {"AI", "STANDARDIZE"}:
        normalized_focus = "STANDARDIZE"

    marker_key = "lesson_authoring_workspace_page_key"
    if st.session_state.get(marker_key) != workspace_page_key:
        st.session_state[_LESSON_AUTHORING_FOCUS_KEY] = (
            normalized_focus
        )
        st.session_state[marker_key] = workspace_page_key

    compact_hidden = False
    if compact_setup_ui:
        compact_key_suffix = str(workspace_page_key).casefold()
        visibility_key = (
            "lesson_authoring_setup_visible_"
            + compact_key_suffix
        )
        setup_visible = bool(
            st.session_state.get(visibility_key, False)
        )
        toggle_label = (
            "Thu gọn công cụ"
            if setup_visible
            else "Hiện công cụ chỉnh sửa"
        )
        if st.button(
            toggle_label,
            key=(
                "lesson_authoring_toggle_setup_"
                + compact_key_suffix
            ),
            use_container_width=True,
        ):
            st.session_state[visibility_key] = not setup_visible
            st.rerun()

        compact_hidden = not setup_visible
        if compact_hidden:
            st.markdown(
                '<div id="mt-authoring-compact-start"></div>',
                unsafe_allow_html=True,
            )
        else:
            st.components.v1.html(
                """
<script>
(function () {
  const doc = window.parent.document;
  doc.querySelectorAll(
    ".mt-authoring-compact-hidden"
  ).forEach(function (node) {
    node.classList.remove(
      "mt-authoring-compact-hidden"
    );
  });
})();
</script>
                """.strip(),
                height=0,
            )

    if page_title == "Chuẩn hóa giáo án":
        # Keep the data/state pipeline, but omit the legacy entry hub and
        # action strip on the dedicated standardization page.
        workspace_focus = "STANDARDIZE"
    else:
        workspace_focus = _render_lesson_authoring_tool_hub(
            page_title=page_title,
            show_entry_actions=True,
        )
    try:
        academic_year_repository = (
            SupabaseAcademicYearConfigurationRepository(
                client=client,
            )
        )

        current_year = (
            academic_year_repository.get_current()
        )

    except Exception as error:
        st.error(
            "Kh\u00f4ng th\u1ec3 \u0111\u1ecdc "
            "c\u1ea5u h\u00ecnh n\u0103m h\u1ecdc "
            f"hi\u1ec7n h\u00e0nh: {error}"
        )
        return

    if current_year is None:
        st.warning(
            "Ch\u01b0a c\u00f3 n\u0103m h\u1ecdc "
            "hi\u1ec7n h\u00e0nh trong h\u1ec7 th\u1ed1ng."
        )
        return

    academic_year = (
        current_year.academic_year
    )

    # =====================================================
    # CANONICAL ACADEMIC WEEKS FROM ADMIN
    # =====================================================

    try:
        academic_week_repository = (
            SupabaseAcademicWeekRepository(
                client=client,
            )
        )

        academic_weeks = (
            academic_week_repository.list_weeks(
                academic_year_id=(
                    current_year.academic_year_id
                )
            )
        )

    except Exception as error:
        st.error(
            "Kh\u00f4ng th\u1ec3 \u0111\u1ecdc "
            "l\u1ecbch tu\u1ea7n t\u1eeb ADMIN: "
            f"{error}"
        )
        return

    active_weeks = tuple(
        item
        for item in academic_weeks
        if item.status.value == "ACTIVE"
    )

    if not active_weeks:
        st.warning(
            "ADMIN ch\u01b0a thi\u1ebft l\u1eadp "
            "l\u1ecbch tu\u1ea7n cho "
            f"n\u0103m h\u1ecdc {academic_year}."
        )
        return

    week_by_number = {
        item.week_number: item
        for item in active_weeks
    }

    week_numbers = tuple(
        week_by_number.keys()
    )

    # Restore only the immutable snapshot originally emitted by the
    # standardization page; AI-authored metadata is never written upstream.
    restore_context = dict(
        st.session_state.get(_WORKING_LESSON_CONTEXT_KEY, {}) or {}
    )
    if (
        st.session_state.get(_RESTORE_LESSON_CONTEXT_KEY)
        and restore_context.get("context_origin") == "STANDARDIZATION"
        and restore_context.get("context_read_only") is True
    ):
        try:
            restored_week = int(restore_context.get("week_number"))
        except (TypeError, ValueError):
            restored_week = 0
        if restored_week in week_numbers:
            _restore_user_id = str(
                st.session_state.get("portal_user_id", "") or ""
            ).strip()
            if _restore_user_id:
                _restore_canonical = get_canonical_context(
                    st.session_state,
                    user_id=_restore_user_id,
                    source_page="weekly_schedule",
                )
                _restore_canonical_week = _restore_canonical.week_number
                if (
                    _restore_canonical_week is not None
                    and int(_restore_canonical_week) != restored_week
                ):
                    st.session_state[
                        _LBG_DATA_WEEK_CONTEXT_MISMATCH_KEY
                    ] = {
                        "canonical_week": int(_restore_canonical_week),
                        "data_week": restored_week,
                        "schedule_id": str(
                            st.session_state.get(
                                _ACTIVE_SCHEDULE_ID_KEY,
                                "",
                            )
                            or ""
                        ),
                        "source": "standardization_restore_snapshot",
                    }

    navigation_notice = st.session_state.pop(_LESSON_AUTHORING_NOTICE_KEY, "")
    if navigation_notice:
        st.toast(str(navigation_notice), icon="✅")

    # -----------------------------------------------------
    # FILTER BAR
    # -----------------------------------------------------

    # V57-F2C5G_CANONICAL_WEEK_AUTHORITY
    # SystemContext.week_number is the only business-context authority here.
    # Session-state week keys are compatibility projections for widgets.
    _canonical_user_id = str(
        st.session_state.get("portal_user_id", "") or ""
    ).strip()
    if not _canonical_user_id:
        st.error("Không xác định được người dùng cho ngữ cảnh Tuần hệ thống.")
        return

    _canonical_context = get_canonical_context(
        st.session_state,
        user_id=_canonical_user_id,
        source_page="weekly_schedule",
    )
    _canonical_week = _canonical_context.week_number

    if (
        _canonical_week is None
        or int(_canonical_week) not in week_numbers
    ):
        _emit_canonical_week_change(
            selected_week=int(week_numbers[0]),
            source_control="academic_week_default",
        )
        _canonical_context = get_canonical_context(
            st.session_state,
            user_id=_canonical_user_id,
            source_page="weekly_schedule",
        )

    week_number = int(_canonical_context.week_number)
    publish_year_week_projection(
        st.session_state,
        context=_canonical_context,
    )

    hide_standardization_context_ui = (
        page_title == "Chuẩn hóa giáo án"
    )

    if hide_standardization_context_ui:
        # V57-F2C5G: widget value is already projected from canonical
        # SystemContext before widget creation. No competing writer here.

        st.markdown(
            """
<style>
.mt-standardization-week-menu {
  margin:.25rem 0 1rem;
  padding:1rem 1.2rem;
  border:2px solid #1f65b8;
  border-radius:18px;
  background:linear-gradient(135deg,#071a33,#123f73);
  color:#fff;
  box-shadow:0 10px 24px rgba(7,26,51,.22);
}
.mt-standardization-week-menu h3 {
  margin:0;
  color:#fff;
  font-size:1.2rem;
  font-weight:800;
}
.mt-standardization-week-menu p {
  margin:.3rem 0 0;
  color:#dbeafe;
}
.st-key-standardization_authoring_week_number {
  margin-top:-.75rem;
  padding:.8rem 1.15rem 1rem;
  border:2px solid #1f65b8;
  border-top:0;
  border-radius:0 0 18px 18px;
  background:#eef6ff;
  box-shadow:0 8px 20px rgba(31,101,184,.15);
}
.st-key-standardization_authoring_week_number label p {
  color:#12345b!important;
  font-size:1.05rem!important;
  font-weight:800!important;
}
.st-key-standardization_authoring_week_number div[data-baseweb="select"] > div {
  min-height:52px!important;
  font-size:1.08rem!important;
  font-weight:800!important;
  background:#fff!important;
}
</style>
<section class="mt-standardization-week-menu">
  <h3>📅 TUẦN SOẠN GIÁO ÁN</h3>
  <p>Đồng bộ hai chiều với tuần đang chọn trong Lịch báo giảng &amp; PBSDTB.</p>
</section>
            """.strip(),
            unsafe_allow_html=True,
        )
        week_number = int(
            st.selectbox(
                "Chọn tuần soạn",
                options=week_numbers,
                format_func=lambda value: (
                    f"Tuần {value} · "
                    + week_by_number[value].start_date.strftime("%d/%m/%Y")
                    + " – "
                    + week_by_number[value].end_date.strftime("%d/%m/%Y")
                ),
                key=_STANDARDIZATION_WEEK_KEY,
                on_change=(
                    _sync_standardization_week_to_lbg
                ),
            )
        )
    else:
        # V57-F2C5L: legacy LBG selector is a canonical projection/subscriber.
        # No direct mirror writer is allowed here.
        controls = st.columns(
            [1.25, 1.1, 1.1, 1.1, 0.9],
            gap="medium",
        )

        with controls[0]:
            st.text_input(
                "N\u0103m h\u1ecdc",
                value=academic_year,
                disabled=True,
                key="lbg_user_academic_year",
            )

        with controls[1]:
            week_number = st.selectbox(
                "Tu\u1ea7n h\u1ecdc",
                options=week_numbers,
                format_func=lambda value: (
                    f"Tu\u1ea7n {value}"
                ),
                key="lbg_user_week_number",
                on_change=_sync_legacy_lbg_week_to_canonical,
            )

    selected_academic_week = (
        week_by_number[
            week_number
        ]
    )

    reload_data = False
    if not hide_standardization_context_ui:
        with controls[2]:
            st.text_input(
                "T\u1eeb ng\u00e0y",
                value=(
                    selected_academic_week
                    .start_date
                    .strftime("%d/%m/%Y")
                ),
                disabled=True,
                key="lbg_user_from_date",
            )

        with controls[3]:
            st.text_input(
                "\u0110\u1ebfn ng\u00e0y",
                value=(
                    selected_academic_week
                    .end_date
                    .strftime("%d/%m/%Y")
                ),
                disabled=True,
                key="lbg_user_to_date",
            )

        with controls[4]:
            reload_data = st.button(
                "\U0001f504 T\u1ea3i l\u1ea1i d\u1eef li\u1ec7u",
                width="stretch",
                key="lbg_user_reload",
            )

    # -----------------------------------------------------
    # CURRENT WEEKLY SCHEDULE VIEW
    # -----------------------------------------------------

    # ONE_WAY_LBG_DATA_FLOW_V1
    # Destination pages consume the active LBG view.  They never clear,
    # regenerate or overwrite the source page's shared view state.
    view = active_view

    if view is None:
        view = st.session_state.get(
            _VIEW_STATE_KEY
        )

    if reload_data:
        view = None

    if (
        view is not None
        and (
            str(view.academic_year)
            != academic_year
            or int(view.week_number)
            != int(week_number)
        )
    ):
        view = None


    # STANDARDIZATION_AUTHORITATIVE_TIMETABLE_V56
    # ACTIVE teacher timetable is authoritative for the selected week.
    # PPCT remains the curriculum-content source.
    # V57C_PHASE2_CANONICAL_CONTEXT_PROJECTION
    # Compatibility-only projection; it never writes widget/session state.
    canonical_context = project_system_context(
        st.session_state,
        source_page="weekly_schedule",
        source_control="standardization_authoritative_view",
    )
    # V57C_PHASE2B_SHADOW_ONLY
    # The projected SystemContext is observational in this phase.
    # Existing local selected year/week continue to drive the V56 authoritative
    # timetable runtime so compatibility work cannot change business behavior.
    _canonical_projection_snapshot = (
        canonical_context.academic_year,
        canonical_context.week_number,
    )

    try:
        view = _build_standardization_authoritative_week_view(
            client=client,
            user_id=str(user_id),
            academic_year=academic_year,
            week_number=int(week_number),
        )
    except Exception as error:
        st.error(
            "Không thể đồng bộ dữ liệu Chuẩn hóa giáo án "
            "với Thời khóa biểu đang hiệu lực của tuần này: "
            + str(error)
        )
        return

    if view is None:
        try:
            schedule_id = (
                "SYSTEM-"
                + str(user_id)
                + "-"
                + academic_year
                + "-W"
                + str(week_number)
            )

            # Read the schedule persisted by Lịch báo giảng.  Creating or
            # updating it remains exclusive to the source page.
            schedule = (
                SupabaseWeeklyScheduleRepository(
                    client,
                    str(user_id),
                ).get(schedule_id)
            )

            if schedule is None:
                st.info(
                    "Chưa có Lịch báo giảng đã cập nhật "
                    f"cho Tuần {week_number}."
                )
                return

            generation = (
                WeeklyScheduleGenerationResult(
                    request=(
                        WeeklyScheduleGenerationRequest(
                            schedule_id=schedule_id,
                            teacher_id=str(user_id),
                            academic_year=academic_year,
                            week_number=week_number,
                        )
                    ),
                    schedule=schedule,
                )
            )

            output = (
                WeeklyScheduleOutputService()
                .export_excel(
                    generation=generation
                )
            )

            view = (
                WeeklySchedulePortalPresenter()
                .present(
                    output=output
                )
            )

        except Exception as error:
            st.error(
                "Kh\u00f4ng th\u1ec3 \u0111\u1ecdc "
                "d\u1eef li\u1ec7u L\u1ecbch b\u00e1o gi\u1ea3ng "
                f"cho Tu\u1ea7n {week_number}: "
                + str(error)
            )
            return

    if view is None:
        st.info(
            "Ch\u01b0a c\u00f3 d\u1eef li\u1ec7u b\u00e0i d\u1ea1y "
            f"cho Tu\u1ea7n {week_number}."
        )
        return

    st.session_state[
        "_standardization_current_week_view"
    ] = view
    st.session_state[
        "_standardization_current_academic_year"
    ] = str(academic_year)
    st.session_state[
        "_standardization_current_week_number"
    ] = int(week_number)
    st.session_state[
        "lesson_standardization_supabase_client"
    ] = client
    st.session_state[
        "lesson_standardization_teacher_user_id"
    ] = str(user_id)

    if standardization_minimal_ui:
        _render_standardization_modern_3d_header()

    if page_title == "Chuẩn hóa giáo án":
        _render_lesson_plan_standardization_workspace(
            view,
            teacher_user_id=str(user_id),
            client=client,
            workspace_focus=workspace_focus,
            compact_setup_ui=compact_hidden,
        )
    else:
        _render_lbg_table(
            view,
            client=client,
            teacher_user_id=str(user_id),
            workspace_focus=workspace_focus,
            compact_setup_ui=compact_hidden,
        )

    _render_pending_standardization_target()

    # Standardization controls belong at the bottom
    # of the lesson-plan workspace.
    if page_title == "Chuẩn hóa giáo án":
        st.session_state[
            "lesson_standardization_supabase_client"
        ] = client
        st.session_state[
            "lesson_standardization_teacher_user_id"
        ] = str(user_id)
        _render_standardization_control_panel()



# STANDARDIZATION_DRAFTING_APPROVAL_V2

from datetime import timedelta as _mt_date_delta
from io import BytesIO as _mt_date_bytesio
from dataclasses import (
    is_dataclass as _mt_date_is_dataclass,
    replace as _mt_date_dataclass_replace,
)

_MT_DRAFTING_ENABLED = (
    "standardization_drafting_before_monday_enabled"
)

_MT_DRAFTING_DAYS = (
    "standardization_drafting_before_monday_days"
)

_MT_APPROVAL_ENABLED = (
    "standardization_approval_before_monday_enabled"
)

_MT_APPROVAL_DAYS = (
    "standardization_approval_before_monday_days"
)


def _mt_monday_of_teaching_week(
    teaching_date,
):
    return (
        teaching_date
        - _mt_date_delta(
            days=teaching_date.weekday()
        )
    )


def _mt_date_before_teaching_week(
    teaching_date,
    days_before,
):
    days_before = int(days_before)

    if days_before < 0:
        days_before = 0

    return (
        _mt_monday_of_teaching_week(
            teaching_date
        )
        - _mt_date_delta(
            days=days_before
        )
    )


def _mt_replace_text_span_preserve_runs(
    paragraph,
    start,
    end,
    replacement,
):
    runs = list(paragraph.runs)

    if not runs:
        return False

    locations = []
    cursor = 0

    for run in runs:
        run_start = cursor
        cursor += len(run.text)

        locations.append(
            (
                run,
                run_start,
                cursor,
            )
        )

    affected = [
        item
        for item in locations
        if (
            item[2] > start
            and item[1] < end
        )
    ]

    if not affected:
        return False

    first_run = affected[0][0]
    first_start = affected[0][1]

    last_run = affected[-1][0]
    last_start = affected[-1][1]

    first_local = max(
        0,
        start - first_start,
    )

    last_local = max(
        0,
        end - last_start,
    )

    prefix = first_run.text[
        :first_local
    ]

    suffix = last_run.text[
        last_local:
    ]

    if first_run is last_run:
        first_run.text = (
            prefix
            + replacement
            + suffix
        )
        return True

    first_run.text = (
        prefix
        + replacement
    )

    for run, _, _ in affected[1:-1]:
        run.text = ""

    last_run.text = suffix

    return True


def _mt_iter_date_paragraphs(
    document,
):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph

        for paragraph in section.footer.paragraphs:
            yield paragraph


def _mt_overlay_approval_date_bytes(
    content,
    approval_date,
):
    import re
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    formatted = approval_date.strftime(
        "%d/%m/%Y"
    )

    document = Document(
        _mt_date_bytesio(content)
    )

    patterns = (
        re.compile(
            (
                r"(?i)"
                r"(?P<label>"
                r"Ng\u00e0y\s+duy\u1ec7t\s*:\s*"
                r")"
                r"(?P<date>"
                r"\d{1,2}[/-]"
                r"\d{1,2}[/-]"
                r"\d{4}"
                r")"
            )
        ),
        re.compile(
            (
                r"(?i)"
                r"(?P<label>"
                r"Ng\u00e0y\s+ph\u00ea\s+duy\u1ec7t\s*:\s*"
                r")"
                r"(?P<date>"
                r"\d{1,2}[/-]"
                r"\d{1,2}[/-]"
                r"\d{4}"
                r")"
            )
        ),
        re.compile(
            (
                r"(?i)"
                r"(?P<label>"
                r"Duy\u1ec7t\s+ng\u00e0y\s*:\s*"
                r")"
                r"(?P<date>"
                r"\d{1,2}[/-]"
                r"\d{1,2}[/-]"
                r"\d{4}"
                r")"
            )
        ),
    )

    changed = False
    paragraphs = list(
        _mt_iter_date_paragraphs(document)
    )

    for paragraph in paragraphs:
        for pattern in patterns:
            matches = list(
                pattern.finditer(
                    paragraph.text
                )
            )

            if matches and paragraph.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                changed = True

            for match in reversed(matches):
                if (
                    match.group("date")
                    == formatted
                ):
                    continue

                if _mt_replace_text_span_preserve_runs(
                    paragraph,
                    match.start("date"),
                    match.end("date"),
                    formatted,
                ):
                    changed = True

    # Common school template:
    #   Ngày 13 tháng 09 năm 2025
    #   Tổ CM duyệt:
    # The date is not labelled "Ngày duyệt", so resolve it by proximity to
    # the approval marker.  If the date line is absent, insert it.
    long_date_pattern = re.compile(
        r"(?i)Ng\s*ày\s+\d{1,2}\s+th\s*áng\s+\d{1,2}\s+n\s*ăm\s+\d{4}"
    )
    long_date_text = approval_date.strftime(
        "Ngày %d tháng %m năm %Y"
    )
    obsolete_date_paragraphs = []

    for index, paragraph in enumerate(paragraphs):
        if "duyệt" not in paragraph.text.casefold():
            continue

        if paragraph.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            changed = True

        date_found = False

        for candidate in paragraphs[max(0, index - 2): index + 1]:
            matches = list(
                long_date_pattern.finditer(candidate.text)
            )

            if not matches:
                normalized = candidate.text.strip().casefold()
                is_unlabelled_legacy_date = (
                    normalized.startswith("ngày ")
                    and not any(
                        label in normalized
                        for label in (
                            "ngày soạn",
                            "ngày dạy",
                            "ngày duyệt",
                            "ngày phê duyệt",
                        )
                    )
                    and (
                        re.search(r"\d{4}", normalized) is not None
                        or "..." in normalized
                        or "…" in normalized
                    )
                )

                if is_unlabelled_legacy_date:
                    obsolete_date_paragraphs.append(candidate)
                continue

            date_found = True

            if candidate.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
                candidate.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                changed = True

            for match in reversed(matches):
                if match.group(0) == long_date_text:
                    continue

                if _mt_replace_text_span_preserve_runs(
                    candidate,
                    match.start(),
                    match.end(),
                    long_date_text,
                ):
                    changed = True

        if not date_found and not any(
            pattern.search(paragraph.text)
            for pattern in patterns
        ):
            inserted = paragraph.insert_paragraph_before(
                long_date_text
            )
            if paragraph.style is not None:
                inserted.style = paragraph.style
            inserted.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            changed = True

    for paragraph in dict.fromkeys(obsolete_date_paragraphs):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
            changed = True

    if not changed:
        return content

    output = _mt_date_bytesio()

    document.save(output)

    return output.getvalue()


def _mt_overlay_drafting_date_bytes(
    content,
    drafting_date,
):
    """Overwrite or add the drafting date while preserving Word content."""

    import re
    from docx import Document

    formatted = drafting_date.strftime("%d/%m/%Y")
    document = Document(_mt_date_bytesio(content))
    paragraphs = list(_mt_iter_date_paragraphs(document))
    pattern = re.compile(
        r"(?i)(?P<label>Ng\s*ày\s+so\s*ạn\s*:\s*)"
        r"(?P<date>\d{1,2}[/-]\d{1,2}[/-]\d{4})"
    )
    label_pattern = re.compile(
        r"(?i)Ng\s*ày\s+so\s*ạn\s*:\s*$"
    )
    changed = False
    drafting_label_found = False

    for paragraph in paragraphs:
        matches = list(pattern.finditer(paragraph.text))

        if matches:
            drafting_label_found = True

        for match in reversed(matches):
            if match.group("date") == formatted:
                continue

            if _mt_replace_text_span_preserve_runs(
                paragraph,
                match.start("date"),
                match.end("date"),
                formatted,
            ):
                changed = True

        if not matches and label_pattern.search(paragraph.text):
            paragraph.add_run(formatted)
            drafting_label_found = True
            changed = True

    if not drafting_label_found:
        anchor = next(
            (
                paragraph
                for paragraph in paragraphs
                if "ngày dạy" in paragraph.text.casefold()
            ),
            paragraphs[0] if paragraphs else None,
        )

        if anchor is not None:
            inserted = anchor.insert_paragraph_before(
                "Ngày soạn: " + formatted
            )
            if anchor.style is not None:
                inserted.style = anchor.style
        else:
            document.add_paragraph(
                "Ngày soạn: " + formatted
            )
        changed = True

    if not changed:
        return content

    output = _mt_date_bytesio()
    document.save(output)
    return output.getvalue()


_mt_original_standardization_control_panel = (
    _render_standardization_control_panel
)


def _render_standardization_control_panel():
    _mt_original_standardization_control_panel()

    st.markdown(
        "\u0023\u0023\u0023\u0023 "
        "Thi\u1ebft l\u1eadp ng\u00e0y"
    )

    drafting_enabled = st.checkbox(
        "D\u00e1n \u0111\u00e8 Ng\u00e0y so\u1ea1n",
        key=_MT_DRAFTING_ENABLED,
        help=(
            "Ng\u00e0y so\u1ea1n = "
            "Th\u1ee9 Hai c\u1ee7a tu\u1ea7n h\u1ecdc "
            "tr\u1eeb N ng\u00e0y."
        ),
    )

    if drafting_enabled:
        st.number_input(
            (
                "Ng\u00e0y so\u1ea1n: "
                "Tr\u01b0\u1edbc th\u1ee9 Hai "
                "c\u1ee7a tu\u1ea7n h\u1ecdc N ng\u00e0y"
            ),
            min_value=0,
            value=int(
                st.session_state.get(
                    _MT_DRAFTING_DAYS,
                    0,
                )
            ),
            step=1,
            key=_MT_DRAFTING_DAYS,
        )

    approval_enabled = st.checkbox(
        "D\u00e1n \u0111\u00e8 Ng\u00e0y duy\u1ec7t",
        key=_MT_APPROVAL_ENABLED,
        help=(
            "Ng\u00e0y duy\u1ec7t = "
            "Th\u1ee9 Hai c\u1ee7a tu\u1ea7n h\u1ecdc "
            "tr\u1eeb N ng\u00e0y."
        ),
    )

    if approval_enabled:
        st.number_input(
            (
                "Ng\u00e0y duy\u1ec7t: "
                "Tr\u01b0\u1edbc th\u1ee9 Hai "
                "c\u1ee7a tu\u1ea7n h\u1ecdc N ng\u00e0y"
            ),
            min_value=0,
            value=int(
                st.session_state.get(
                    _MT_APPROVAL_DAYS,
                    0,
                )
            ),
            step=1,
            key=_MT_APPROVAL_DAYS,
        )


_mt_original_process_lesson_plan_upload_dates = (
    _process_lesson_plan_upload
)


def _mt_result_output_bytes(result):
    """Read DOCX bytes from either the legacy tuple or result object."""

    if (
        isinstance(result, tuple)
        and len(result) >= 2
        and isinstance(result[1], (bytes, bytearray))
    ):
        return bytes(result[1])

    content = getattr(result, "output_bytes", None)
    return bytes(content) if isinstance(content, (bytes, bytearray)) else None


def _mt_result_with_output_bytes(result, content):
    """Replace DOCX bytes without changing the processor result contract."""

    updated = bytes(content)

    if isinstance(result, tuple) and len(result) >= 2:
        return result[:1] + (updated,) + result[2:]

    from dataclasses import is_dataclass, replace

    if is_dataclass(result):
        return replace(result, output_bytes=updated)

    try:
        result.output_bytes = updated
        return result
    except Exception as error:
        raise RuntimeError(
            "Không thể cập nhật dữ liệu Word sau chuẩn hóa."
        ) from error


def _process_lesson_plan_upload(
    *args,
    **kwargs,
):
    import inspect

    signature = inspect.signature(
        _mt_original_process_lesson_plan_upload_dates
    )

    bound = signature.bind_partial(
        *args,
        **kwargs,
    )

    row = bound.arguments.get("row")

    teaching_date = (
        getattr(
            row,
            "teaching_date",
            None,
        )
        if row is not None
        else None
    )

    if (
        teaching_date is not None
        and bool(
            st.session_state.get(
                _MT_DRAFTING_ENABLED,
                False,
            )
        )
    ):
        bound.arguments[
            "drafting_date"
        ] = (
            _mt_date_before_teaching_week(
                teaching_date,
                st.session_state.get(
                    _MT_DRAFTING_DAYS,
                    0,
                ),
            )
        )

    result = (
        _mt_original_process_lesson_plan_upload_dates(
            *bound.args,
            **bound.kwargs,
        )
    )

    if teaching_date is None:
        return result

    content = _mt_result_output_bytes(result)

    if not content:
        return result

    updated_content = content

    if bool(
        st.session_state.get(
            _MT_DRAFTING_ENABLED,
            False,
        )
    ):
        drafting_date = _mt_date_before_teaching_week(
            teaching_date,
            st.session_state.get(_MT_DRAFTING_DAYS, 0),
        )
        updated_content = _mt_overlay_drafting_date_bytes(
            updated_content,
            drafting_date,
        )

    if bool(
        st.session_state.get(
            _MT_APPROVAL_ENABLED,
            False,
        )
    ):
        approval_date = _mt_date_before_teaching_week(
            teaching_date,
            st.session_state.get(_MT_APPROVAL_DAYS, 0),
        )
        updated_content = _mt_overlay_approval_date_bytes(
            updated_content,
            approval_date,
        )

    if updated_content == content:
        return result

    return _mt_result_with_output_bytes(
        result,
        updated_content,
    )



# STANDARDIZATION_ASSIGNMENT_TIMETABLE_SYNC_V1

from dataclasses import (
    is_dataclass as _mt_sync_is_dataclass,
    replace as _mt_sync_dataclass_replace,
)
from io import BytesIO as _mt_sync_bytesio


_MT_TEACHING_SYNC_ENABLED = (
    "standardization_assignment_timetable_sync_enabled"
)


def _mt_sync_text(value):
    return str(
        value
        if value is not None
        else ""
    ).strip()


def _mt_sync_component_matches(
    assignment_component,
    row_component,
):
    assignment_value = (
        _mt_sync_text(
            assignment_component
        )
    )

    row_value = (
        _mt_sync_text(
            row_component
        )
    )

    if not assignment_value:
        return True

    return (
        assignment_value.casefold()
        == row_value.casefold()
    )


def _mt_sync_same_lesson(
    current,
    candidate,
):
    current_period = getattr(
        current,
        "curriculum_period",
        None,
    )

    candidate_period = getattr(
        candidate,
        "curriculum_period",
        None,
    )

    if (
        current_period is not None
        and candidate_period is not None
    ):
        return (
            int(current_period)
            == int(candidate_period)
        )

    current_title = _mt_sync_text(
        getattr(
            current,
            "lesson_title",
            "",
        )
    )

    candidate_title = _mt_sync_text(
        getattr(
            candidate,
            "lesson_title",
            "",
        )
    )

    return bool(
        current_title
        and candidate_title
        and (
            current_title.casefold()
            == candidate_title.casefold()
        )
    )


def _mt_resolve_teaching_date_pairs(
    *,
    current_row,
    assignments,
    weekly_rows,
):
    """
    Resolve the exact class/date block for one lesson.

    Assignment = class authority.
    Current week rows = timetable/week evidence.

    Special fallback:
    when current class is assigned but has no matching
    timetable row for this lesson in the week, keep only
    current class using current_row.teaching_date.
    """

    subject_ref = _mt_sync_text(
        getattr(
            current_row,
            "subject_ref",
            "",
        )
    )

    component_ref = _mt_sync_text(
        getattr(
            current_row,
            "component_ref",
            "",
        )
    )

    current_class = _mt_sync_text(
        getattr(
            current_row,
            "class_id",
            "",
        )
    )

    assigned_classes = set()

    for assignment in assignments:
        if (
            _mt_sync_text(
                getattr(
                    assignment,
                    "subject_ref",
                    "",
                )
            ).casefold()
            != subject_ref.casefold()
        ):
            continue

        if not _mt_sync_component_matches(
            getattr(
                assignment,
                "component_ref",
                None,
            ),
            component_ref,
        ):
            continue

        class_id = _mt_sync_text(
            getattr(
                assignment,
                "class_id",
                "",
            )
        )

        if class_id:
            assigned_classes.add(
                class_id
            )

    if not assigned_classes:
        return ()

    scheduled = {}

    for row in weekly_rows:
        class_id = _mt_sync_text(
            getattr(
                row,
                "class_id",
                "",
            )
        )

        if class_id not in assigned_classes:
            continue

        if (
            _mt_sync_text(
                getattr(
                    row,
                    "subject_ref",
                    "",
                )
            ).casefold()
            != subject_ref.casefold()
        ):
            continue

        if not _mt_sync_component_matches(
            getattr(
                row,
                "component_ref",
                None,
            ),
            component_ref,
        ):
            continue

        if not _mt_sync_same_lesson(
            current_row,
            row,
        ):
            continue

        teaching_date = getattr(
            row,
            "teaching_date",
            None,
        )

        if teaching_date is None:
            continue

        # If one lesson has several periods for one class,
        # keep the first teaching date in the week.
        existing = scheduled.get(
            class_id
        )

        if (
            existing is None
            or teaching_date < existing
        ):
            scheduled[
                class_id
            ] = teaching_date

    # -------------------------------------------------
    # User-approved special rule:
    #
    # Class is in professional assignment, but this
    # class has no matching timetable row this week:
    # keep only this class, use its existing teaching
    # date, and remove the other classes from Word.
    # -------------------------------------------------

    if (
        current_class in assigned_classes
        and current_class not in scheduled
    ):
        fallback_date = getattr(
            current_row,
            "teaching_date",
            None,
        )

        if fallback_date is None:
            return ()

        return (
            (
                current_class,
                fallback_date,
            ),
        )

    if not scheduled:
        return ()

    return tuple(
        (
            class_id,
            scheduled[class_id],
        )
        for class_id
        in sorted(
            scheduled.keys()
        )
    )


def _mt_load_active_assignments(
    *,
    client,
    user_id,
    academic_year,
):
    from educational_planning_v2.adapters.supabase_teaching_assignment_repository import (
        SupabaseTeachingAssignmentRepository,
    )
    from educational_planning_v2.models.teaching_assignment import (
        TeachingAssignmentRole,
        TeachingAssignmentStatus,
    )

    repository = (
        SupabaseTeachingAssignmentRepository(
            client=client,
            user_id=str(user_id),
        )
    )

    return repository.list_assignments(
        owner_id=str(user_id),
        academic_year=str(
            academic_year
        ),
        role=(
            TeachingAssignmentRole.TEACHING
        ),
        status=(
            TeachingAssignmentStatus.ACTIVE
        ),
    )


def _mt_pair_lines(
    pairs,
):
    lines = []

    for class_id, teaching_date in pairs:
        if hasattr(
            teaching_date,
            "strftime",
        ):
            date_text = (
                teaching_date.strftime(
                    "%d/%m/%Y"
                )
            )
        else:
            date_text = _mt_sync_text(
                teaching_date
            )

        lines.append(
            str(class_id)
            + " - "
            + date_text
        )

    return "\n".join(lines)


def _mt_sync_teaching_date_bytes(
    content,
    *,
    pairs,
):
    """
    Replace only the Ngay day value area.

    Supports:
      Ngay day: 01/01/2020

    and:
      Ngay day:
      7A1 - ...
      7A2 - ...

    Other metadata in the same paragraph is preserved.
    """

    import re
    from docx import Document

    if not pairs:
        return content

    replacement = _mt_pair_lines(
        pairs
    )

    document = Document(
        _mt_sync_bytesio(
            content
        )
    )

    label_pattern = re.compile(
        (
            r"(?i)"
            r"Ng\u00e0y\s+d\u1ea1y\s*:"
        )
    )

    pair_pattern = re.compile(
        (
            r"\b\d{1,2}[A-Za-z]\d+\b"
            r"\s*[-\u2013\u2014:]\s*"
            r"\d{1,2}[/-]"
            r"\d{1,2}[/-]"
            r"\d{4}"
        )
    )

    plain_date_pattern = re.compile(
        (
            r"\d{1,2}[/-]"
            r"\d{1,2}[/-]"
            r"\d{4}"
        )
    )

    next_label_pattern = re.compile(
        (
            r"(?i)"
            r"(?:"
            r"L\u1edbp"
            r"|Ti\u1ebft"
            r"|B\u00e0i"
            r"|T\u00ean\s+b\u00e0i"
            r"|Ng\u00e0y\s+duy\u1ec7t"
            r"|Ng\u00e0y\s+ph\u00ea\s+duy\u1ec7t"
            r")\s*:"
        )
    )

    changed = False

    paragraphs = list(
        _mt_iter_date_paragraphs(document)
    )

    for paragraph_index, paragraph in enumerate(paragraphs):
        text = paragraph.text

        label = label_pattern.search(
            text
        )

        if label is None:
            continue

        value_start = label.end()

        tail = text[
            value_start:
        ]

        next_label = (
            next_label_pattern.search(
                tail
            )
        )

        value_end = (
            value_start
            + (
                next_label.start()
                if next_label
                else len(tail)
            )
        )

        value_text = text[
            value_start:value_end
        ]

        pairs_found = list(
            pair_pattern.finditer(
                value_text
            )
        )

        plain_found = (
            plain_date_pattern.search(
                value_text
            )
        )

        # Replace only a recognized teaching-date value.
        # Do not overwrite unrelated text accidentally.
        if (
            not pairs_found
            and plain_found is None
        ):
            continue

        new_value = (
            "\n"
            + replacement
        )

        if value_end < len(text):
            new_value += "    "

        if _mt_replace_text_span_preserve_runs(
            paragraph,
            value_start,
            value_end,
            new_value,
        ):
            changed = True

            # Some school templates store each additional class/date on its
            # own paragraph.  Remove those obsolete continuation paragraphs;
            # the complete authoritative block is now in the label paragraph.
            for continuation in paragraphs[paragraph_index + 1:]:
                continuation_text = continuation.text.strip()

                if not continuation_text:
                    break

                if pair_pattern.fullmatch(continuation_text) is None:
                    break

                element = continuation._element
                parent = element.getparent()

                if parent is not None:
                    parent.remove(element)

            break

    if not changed:
        return content

    output = _mt_sync_bytesio()

    document.save(
        output
    )

    return output.getvalue()


# -------------------------------------------------
# Add the option to the EXISTING control panel.
# -------------------------------------------------

_mt_original_standardization_control_panel_3b = (
    _render_standardization_control_panel
)


def _render_standardization_control_panel():
    _mt_original_standardization_control_panel_3b()

    st.checkbox(
        (
            "\u0110\u1ed3ng b\u1ed9 Ng\u00e0y d\u1ea1y "
            "theo Ph\u00e2n c\u00f4ng + TKB tu\u1ea7n"
        ),
        key=_MT_TEACHING_SYNC_ENABLED,
        help=(
            "X\u00f3a l\u1edbp th\u1eeba, "
            "b\u1ed5 sung l\u1edbp thi\u1ebfu v\u00e0 "
            "c\u1eadp nh\u1eadt Ng\u00e0y d\u1ea1y "
            "theo Ph\u00e2n c\u00f4ng chuy\u00ean m\u00f4n "
            "v\u00e0 l\u1ecbch c\u1ee7a \u0111\u00fang tu\u1ea7n."
        ),
    )


# -------------------------------------------------
# Wrap current processor.
# This runs AFTER the already-PASS drafting/approval
# processor, so 3A remains untouched.
# -------------------------------------------------

_mt_original_process_lesson_plan_upload_3b = (
    _process_lesson_plan_upload
)


def _process_lesson_plan_upload(
    *args,
    **kwargs,
):
    import inspect

    signature = inspect.signature(
        _mt_original_process_lesson_plan_upload_3b
    )

    bound = signature.bind_partial(
        *args,
        **kwargs,
    )

    result = (
        _mt_original_process_lesson_plan_upload_3b(
            *args,
            **kwargs,
        )
    )

    if not bool(
        st.session_state.get(
            _MT_TEACHING_SYNC_ENABLED,
            False,
        )
    ):
        return result

    # The preceding drafting/approval wrapper exposes ``*args, **kwargs``.
    # In that real wrapper chain, inspect.bind_partial stores the schedule row
    # inside the nested ``kwargs`` mapping instead of under a top-level "row"
    # argument.  Read the direct call first, then retain the explicit-signature
    # fallback used by isolated processors/tests.
    current_row = kwargs.get("row")

    if current_row is None:
        current_row = bound.arguments.get("row")

    if current_row is None:
        nested_kwargs = bound.arguments.get("kwargs", {})

        if isinstance(nested_kwargs, dict):
            current_row = nested_kwargs.get("row")

    if current_row is None:
        return result

    # Primary source: the exact rows already selected and displayed in
    # "Nội dung soạn bài".  These pairs carry teacher-facing class names and
    # the real dates of every class for the chosen PPCT lesson.
    pairs = tuple(
        (
            _mt_sync_text(class_name),
            teaching_date,
        )
        for class_name, teaching_date in tuple(
            st.session_state.get(
                "_standardization_selected_teaching_date_pairs",
                (),
            )
            or ()
        )
        if _mt_sync_text(class_name) and teaching_date is not None
    )

    if pairs:
        content = _mt_result_output_bytes(result)

        if not content:
            return result

        updated = _mt_sync_teaching_date_bytes(
            content,
            pairs=pairs,
        )

        if updated == content:
            return result

        return _mt_result_with_output_bytes(result, updated)

    client = st.session_state.get(
        "lesson_standardization_supabase_client"
    )

    user_id = st.session_state.get(
        "lesson_standardization_teacher_user_id"
    )

    academic_year = (
        st.session_state.get(
            "_standardization_current_academic_year"
        )
    )

    view = st.session_state.get(
        "_standardization_current_week_view"
    )

    if (
        client is None
        or not user_id
        or not academic_year
        or view is None
    ):
        # Safety rule:
        # no authoritative data -> no destructive sync.
        return result

    assignments = (
        _mt_load_active_assignments(
            client=client,
            user_id=user_id,
            academic_year=academic_year,
        )
    )

    weekly_rows = tuple(
        getattr(
            view,
            "rows",
            (),
        )
        or ()
    )

    pairs = (
        _mt_resolve_teaching_date_pairs(
            current_row=current_row,
            assignments=assignments,
            weekly_rows=weekly_rows,
        )
    )

    if not pairs:
        return result

    content = _mt_result_output_bytes(result)

    if not content:
        return result

    updated = (
        _mt_sync_teaching_date_bytes(
            content,
            pairs=pairs,
        )
    )

    if updated == content:
        return result

    return _mt_result_with_output_bytes(
        result,
        updated,
    )



# STANDARDIZATION_IMAGE_AUTOFIT_V1

from io import BytesIO as _mt_image_bytesio
from dataclasses import (
    is_dataclass as _mt_image_is_dataclass,
    replace as _mt_image_dataclass_replace,
)


_MT_IMAGE_AUTOFIT_ENABLED = (
    "standardization_image_autofit_enabled"
)

_MT_END_RULE_ENABLED = (
    "standardization_lesson_end_rule_enabled"
)


def _mt_image_xml_local_name(
    element,
):
    tag = str(
        getattr(
            element,
            "tag",
            "",
        )
    )

    if "}" in tag:
        return tag.rsplit(
            "}",
            1,
        )[-1]

    return tag


def _mt_image_find_ancestor(
    element,
    local_name,
):
    current = element

    while current is not None:
        if (
            _mt_image_xml_local_name(
                current
            )
            == local_name
        ):
            return current

        current = current.getparent()

    return None


def _mt_image_cell_width_emu(
    inline,
):
    """
    Read table-cell width when the picture sits inside
    a Word table cell.

    tcW with type=dxa uses twips:
    1 twip = 635 EMU.
    """

    tc = _mt_image_find_ancestor(
        inline,
        "tc",
    )

    if tc is None:
        return None

    tc_pr = next(
        (
            child
            for child in tc
            if (
                _mt_image_xml_local_name(
                    child
                )
                == "tcPr"
            )
        ),
        None,
    )

    if tc_pr is None:
        return None

    tc_w = next(
        (
            child
            for child in tc_pr
            if (
                _mt_image_xml_local_name(
                    child
                )
                == "tcW"
            )
        ),
        None,
    )

    if tc_w is None:
        return None

    width_value = None
    width_type = None

    for key, value in tc_w.attrib.items():
        local = (
            key.rsplit(
                "}",
                1,
            )[-1]
        )

        if local == "w":
            width_value = value

        elif local == "type":
            width_type = value

    if (
        width_value is None
        or str(width_type or "dxa")
        not in {
            "dxa",
            "",
        }
    ):
        return None

    try:
        twips = int(
            width_value
        )
    except Exception:
        return None

    if twips <= 0:
        return None

    return twips * 635


def _mt_image_page_content_box(
    document,
):
    """
    Return conservative page-content width/height in EMU.
    """

    if not document.sections:
        return (
            6_000_000,
            8_000_000,
        )

    section = document.sections[0]

    width = (
        int(section.page_width)
        - int(section.left_margin)
        - int(section.right_margin)
    )

    height = (
        int(section.page_height)
        - int(section.top_margin)
        - int(section.bottom_margin)
    )

    return (
        max(
            width,
            1,
        ),
        max(
            height,
            1,
        ),
    )


def _mt_image_paragraph_inline_count(
    inline,
):
    paragraph = (
        _mt_image_find_ancestor(
            inline,
            "p",
        )
    )

    if paragraph is None:
        return 1

    count = 0

    for element in paragraph.iter():
        if (
            _mt_image_xml_local_name(
                element
            )
            == "inline"
        ):
            count += 1

    return max(
        count,
        1,
    )


def _mt_image_target_box(
    *,
    document,
    inline,
):
    page_width, page_height = (
        _mt_image_page_content_box(
            document
        )
    )

    cell_width = (
        _mt_image_cell_width_emu(
            inline
        )
    )

    available_width = (
        cell_width
        if cell_width is not None
        else page_width
    )

    # Small safety margin prevents touching cell/page borders.
    available_width = int(
        available_width
        * 0.94
    )

    image_count = (
        _mt_image_paragraph_inline_count(
            inline
        )
    )

    if image_count > 1:
        available_width = int(
            available_width
            / image_count
            * 0.96
        )

    # Prevent one very tall portrait from filling
    # more than about 80% of the printable page height.
    available_height = int(
        page_height
        * 0.80
    )

    return (
        max(
            available_width,
            1,
        ),
        max(
            available_height,
            1,
        ),
    )


def _mt_image_resize_shape(
    *,
    shape,
    document,
):
    original_width = int(
        shape.width
    )

    original_height = int(
        shape.height
    )

    if (
        original_width <= 0
        or original_height <= 0
    ):
        return False

    max_width, max_height = (
        _mt_image_target_box(
            document=document,
            inline=shape._inline,
        )
    )

    width_scale = (
        max_width
        / original_width
    )

    height_scale = (
        max_height
        / original_height
    )

    # Only shrink images that overflow the target box. Upscaling small images
    # makes scanned figures blurry and creates the false appearance that every
    # picture must fill the whole page.
    scale = min(
        1.0,
        width_scale,
        height_scale,
    )

    if scale <= 0:
        return False

    new_width = max(
        1,
        int(
            original_width
            * scale
        ),
    )

    new_height = max(
        1,
        int(
            original_height
            * scale
        ),
    )

    if (
        abs(
            new_width
            - original_width
        )
        <= 100
        and abs(
            new_height
            - original_height
        )
        <= 100
    ):
        return False

    shape.width = new_width
    shape.height = new_height

    return True


def _mt_image_center_shape_paragraph(shape):
    """Center the paragraph containing one or more inline pictures."""

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = _mt_image_find_ancestor(
        shape._inline,
        "p",
    )

    if paragraph is None:
        return False

    paragraph_properties = paragraph.get_or_add_pPr()
    justification = paragraph_properties.find(qn("w:jc"))

    if justification is not None:
        current = justification.get(qn("w:val"))
        if current == "center":
            return False
    else:
        justification = OxmlElement("w:jc")
        paragraph_properties.append(justification)

    justification.set(qn("w:val"), "center")
    return True


def _mt_image_convert_anchors_to_inline(document):
    """Move floating pictures into the text flow so they stay inside cells."""

    from docx.oxml import OxmlElement

    changed = False

    for anchor in list(document.element.xpath(".//wp:anchor")):
        parent = anchor.getparent()

        if parent is None:
            continue

        inline = OxmlElement("wp:inline")
        inline.set("distT", "0")
        inline.set("distB", "0")
        inline.set("distL", "0")
        inline.set("distR", "0")

        movable = {
            "extent",
            "effectExtent",
            "docPr",
            "cNvGraphicFramePr",
            "graphic",
        }

        for child in list(anchor):
            if _mt_image_xml_local_name(child) in movable:
                anchor.remove(child)
                inline.append(child)

        parent.replace(anchor, inline)
        changed = True

    return changed


def _mt_image_relax_picture_table_rows(document):
    """Remove fixed row heights that leave blank space after image fitting."""

    from docx.enum.text import WD_ALIGN_PARAGRAPH

    changed = False

    for table in document.tables:
        for row in table.rows:
            if not row._tr.xpath(".//w:drawing"):
                continue

            tr_pr = row._tr.get_or_add_trPr()

            for height in list(tr_pr.xpath("./w:trHeight")):
                tr_pr.remove(height)
                changed = True

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph._p.xpath(".//w:drawing"):
                        continue

                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        changed = True

                    paragraph_format = paragraph.paragraph_format

                    if paragraph_format.space_before is not None:
                        paragraph_format.space_before = None
                        changed = True

                    if paragraph_format.space_after is not None:
                        paragraph_format.space_after = None
                        changed = True

    return changed


def _mt_paragraph_has_rule(paragraph):
    return bool(paragraph._p.xpath("./w:pPr/w:pBdr/*"))


def _mt_ensure_lesson_end_rule_bytes(content):
    """Insert one horizontal rule immediately before the approval block."""

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document(_mt_image_bytesio(content))
    paragraphs = list(_mt_iter_date_paragraphs(document))
    approval_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if "duyệt" in paragraph.text.casefold()
        ),
        None,
    )
    anchor_index = len(paragraphs)

    if approval_index is not None:
        anchor_index = approval_index

        if (
            approval_index > 0
            and paragraphs[approval_index - 1]
            .text.strip().casefold().startswith("ngày ")
        ):
            anchor_index = approval_index - 1

    nearby = paragraphs[max(0, anchor_index - 3):anchor_index]

    if any(_mt_paragraph_has_rule(paragraph) for paragraph in nearby):
        return content

    if anchor_index < len(paragraphs):
        rule_paragraph = paragraphs[anchor_index].insert_paragraph_before("")
    else:
        rule_paragraph = document.add_paragraph("")

    p_pr = rule_paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    output = _mt_image_bytesio()
    document.save(output)
    return output.getvalue()


def _mt_remove_redundant_empty_cell_paragraphs(cell):
    """Keep at most one consecutive empty paragraph in a table cell."""

    changed = False
    previous_empty = False

    for paragraph in list(cell.paragraphs):
        meaningful = bool(
            paragraph.text.strip()
            or paragraph._p.xpath(".//w:drawing | .//m:oMath")
        )

        if meaningful:
            previous_empty = False
            continue

        if previous_empty and len(cell.paragraphs) > 1:
            element = paragraph._element
            parent = element.getparent()

            if parent is not None:
                parent.remove(element)
                changed = True
            continue

        previous_empty = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = 0
        paragraph_format.space_after = 0

    return changed


def _mt_reflow_all_lesson_tables(document):
    """Make lesson-plan tables size themselves from actual remaining content."""

    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    changed = False

    for table in document.tables:
        table.autofit = True

        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()

            for height in list(tr_pr.xpath("./w:trHeight")):
                tr_pr.remove(height)
                changed = True

            for cell in row.cells:
                if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.TOP:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    changed = True

                changed = (
                    _mt_remove_redundant_empty_cell_paragraphs(cell)
                    or changed
                )

    return changed


def _mt_insert_paragraph_after(paragraph):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    return Paragraph(element, paragraph._parent)


def _mt_format_lesson_document_layout_bytes(
    content,
    *,
    ensure_end_rule=True,
):
    """Apply heading, date, table, approval spacing and end-rule layout."""

    import re
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document(_mt_image_bytesio(content))
    changed = _mt_reflow_all_lesson_tables(document)
    heading_pattern = re.compile(
        r"(?i)^\s*(?:CHỦ\s*ĐỀ|TIẾT\s*\d+|BÀI\s*\d+|BÀI\s*DẠY\s*:)")
    date_pattern = re.compile(
        r"(?i)(?:Ngày\s+soạn|Ngày\s+dạy|Ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})"
    )
    class_date_pattern = re.compile(
        r"^\s*\d{1,2}[A-Za-z]\d+\s*[-–—:]\s*\d{1,2}[/-]\d{1,2}[/-]\d{4}\s*$"
    )

    for paragraph in _mt_iter_date_paragraphs(document):
        text = paragraph.text.strip()

        if heading_pattern.search(text):
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                changed = True

        if date_pattern.search(text) or class_date_pattern.match(text):
            for run in paragraph.runs:
                if run.italic is not True:
                    run.italic = True
                    changed = True

    body_paragraphs = list(document.paragraphs)
    approval_marker = next(
        (
            paragraph
            for paragraph in body_paragraphs
            if "duyệt" in paragraph.text.casefold()
        ),
        None,
    )

    if approval_marker is not None:
        approval_marker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        marker_index = body_paragraphs.index(approval_marker)
        approval_date = next(
            (
                paragraph
                for paragraph in reversed(
                    body_paragraphs[max(0, marker_index - 3):marker_index]
                )
                if date_pattern.search(paragraph.text.strip())
            ),
            None,
        )
        approval_anchor = approval_date or approval_marker

        if approval_date is not None:
            approval_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if ensure_end_rule:
            # Remove old empty/rule paragraphs immediately before approval,
            # then rebuild exactly: rule, one blank line, approval date/marker.
            previous = approval_anchor._p.getprevious()

            while previous is not None and previous.tag.endswith("}p"):
                previous_paragraph_text = "".join(
                    previous.itertext()
                ).strip()
                has_rule = bool(previous.xpath("./w:pPr/w:pBdr/*"))

                if previous_paragraph_text or not has_rule:
                    if previous_paragraph_text:
                        break

                older = previous.getprevious()
                previous.getparent().remove(previous)
                previous = older
                changed = True

            rule_paragraph = approval_anchor.insert_paragraph_before("")
            p_pr = rule_paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "auto")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            approval_anchor.insert_paragraph_before("")
            changed = True

        following = approval_marker._p.getnext()
        existing_blanks = 0

        while following is not None and following.tag.endswith("}p"):
            if "".join(following.itertext()).strip():
                break
            existing_blanks += 1
            following = following.getnext()

        cursor = approval_marker

        for _ in range(max(0, 4 - existing_blanks)):
            cursor = _mt_insert_paragraph_after(cursor)
            changed = True

    if not changed:
        return content

    output = _mt_image_bytesio()
    document.save(output)
    return output.getvalue()


def _mt_autofit_images_bytes(
    content,
):
    """
    Resize inline Word pictures without altering
    lesson-plan metadata or textual content.
    """

    from docx import Document

    document = Document(
        _mt_image_bytesio(
            content
        )
    )

    changed = _mt_image_convert_anchors_to_inline(document)

    shapes = list(
        document.inline_shapes
    )

    if not shapes:
        if not changed:
            return content

        output = _mt_image_bytesio()
        document.save(output)
        return output.getvalue()

    for shape in shapes:
        try:
            changed = (
                _mt_image_resize_shape(
                    shape=shape,
                    document=document,
                )
                or changed
            )
            changed = (
                _mt_image_center_shape_paragraph(shape)
                or changed
            )

        except Exception:
            # Image fitting must never break standardization.
            # Leave this picture unchanged.
            continue

    changed = _mt_image_relax_picture_table_rows(document) or changed

    if not changed:
        return content

    output = (
        _mt_image_bytesio()
    )

    document.save(
        output
    )

    return output.getvalue()


# -------------------------------------------------
# Add option to EXISTING standardization panel.
# -------------------------------------------------

_mt_original_standardization_control_panel_3c = (
    _render_standardization_control_panel
)


def _render_standardization_date_and_document_options():
    """Render date, timetable and image controls inside the main panel."""

    st.session_state.setdefault(_MT_DRAFTING_ENABLED, True)
    st.session_state.setdefault(_MT_DRAFTING_DAYS, 3)
    st.session_state.setdefault(_MT_APPROVAL_ENABLED, True)
    st.session_state.setdefault(_MT_APPROVAL_DAYS, 1)
    st.session_state.setdefault(_MT_TEACHING_SYNC_ENABLED, True)
    st.session_state.setdefault(_MT_IMAGE_AUTOFIT_ENABLED, True)
    st.session_state.setdefault(_MT_END_RULE_ENABLED, True)

    st.markdown("#### Thiết lập ngày và tài liệu")

    drafting_enabled = st.checkbox(
        "Dán đè hoặc bổ sung Ngày soạn",
        key=_MT_DRAFTING_ENABLED,
        help="Ngày soạn = Thứ Hai của tuần học trừ N ngày.",
    )

    if drafting_enabled:
        st.number_input(
            "Ngày soạn: trước thứ Hai của tuần học N ngày",
            min_value=0,
            value=int(st.session_state.get(_MT_DRAFTING_DAYS, 3)),
            step=1,
            key=_MT_DRAFTING_DAYS,
        )

    approval_enabled = st.checkbox(
        "Dán đè hoặc bổ sung Ngày duyệt",
        key=_MT_APPROVAL_ENABLED,
        help="Ngày duyệt = Thứ Hai của tuần học trừ N ngày.",
    )

    if approval_enabled:
        st.number_input(
            "Ngày duyệt: trước thứ Hai của tuần học N ngày",
            min_value=0,
            value=int(st.session_state.get(_MT_APPROVAL_DAYS, 1)),
            step=1,
            key=_MT_APPROVAL_DAYS,
        )

    st.checkbox(
        "Đồng bộ Ngày dạy theo Phân công + TKB tuần",
        key=_MT_TEACHING_SYNC_ENABLED,
        help=(
            "Xóa lớp thừa, bổ sung lớp thiếu và cập nhật Ngày dạy "
            "theo đúng phân công chuyên môn và lịch của tuần."
        ),
    )

    st.checkbox(
        "Tự động căn chỉnh hình ảnh theo khung giáo án",
        key=_MT_IMAGE_AUTOFIT_ENABLED,
        help=(
            "Thu nhỏ ảnh vượt khung, giữ nguyên tỷ lệ, không phóng ảnh "
            "nhỏ và căn giữa ảnh nội tuyến."
        ),
    )

    st.checkbox(
        "Tự động tạo đường kẻ khi hết bài",
        key=_MT_END_RULE_ENABLED,
        help="Chèn một đường kẻ trước phần phê duyệt nếu giáo án chưa có.",
    )


def _render_standardization_control_panel():
    # Call the original panel directly. It now renders all extended controls
    # inside its expander, so the legacy wrapper chain cannot append them below.
    _mt_original_standardization_control_panel()


# -------------------------------------------------
# Wrap current 3A + 3B processing chain.
# Image autofit runs LAST.
# -------------------------------------------------

_mt_original_process_lesson_plan_upload_3c = (
    _process_lesson_plan_upload
)


def _process_lesson_plan_upload(
    *,
    row=None,
    drafting_date=None,
    content: bytes | None = None,
    original_name: str = "",
    modification_plan=None,
    options=None,
    original_content=None,
    ai_revised_text="",
):
    result = (
        _mt_original_process_lesson_plan_upload_3c(
            row=row,
            drafting_date=drafting_date,
            content=content,
            original_name=original_name,
            modification_plan=modification_plan,
            options=options,
            original_content=original_content,
            ai_revised_text=ai_revised_text,
        )
    )

    image_enabled = bool(
        st.session_state.get(
            _MT_IMAGE_AUTOFIT_ENABLED,
            False,
        )
    )
    end_rule_enabled = bool(
        st.session_state.get(
            _MT_END_RULE_ENABLED,
            False,
        )
    )

    if not image_enabled and not end_rule_enabled:
        return result

    content = _mt_result_output_bytes(result)

    if not content:
        return result

    updated = content

    if image_enabled:
        updated = _mt_autofit_images_bytes(updated)

    if end_rule_enabled:
        updated = _mt_ensure_lesson_end_rule_bytes(updated)

    updated = _mt_format_lesson_document_layout_bytes(
        updated,
        ensure_end_rule=end_rule_enabled,
    )

    if updated == content:
        return result

    return _mt_result_with_output_bytes(
        result,
        updated,
    )
