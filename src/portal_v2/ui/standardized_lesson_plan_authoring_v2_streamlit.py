from __future__ import annotations
from datetime import datetime, timezone

import re
from html import escape
from typing import Any, Callable, Mapping

import streamlit as st

from lesson_planning_v2.services.lesson_plan_smart_up_resolver import (
    SmartUpContext,
    resolve_from_catalog,
)
from teacher_document_library_v2.adapters.supabase_document_repository import (
    SupabaseTeacherDocumentRepository,
)
from teacher_document_library_v2.services import TeacherDocumentCatalog
from teacher_document_library_v2.storage import StoredDocumentFile
from lesson_planning_v2.services.lesson_plan_local_file_search import (
    find_local_lesson_plans,
    read_local_lesson_plan,
)


GROUP_CONTEXT_KEY = "lesson_plan_group_context_v2"
ORIGINAL_DOCUMENT_KEY = "g1b_v2_original_document"
STANDARDIZED_DOCUMENT_KEY = "g1b_v2_standardized_document"
AUDIT_RESULT_KEY = "g1b_v2_standardization_audit_result"
AUDIT_FIELD_EVIDENCE_KEY = "g1b_v2_canonical_field_evidence"
TEACHER_VERIFICATION_KEY = "g1b_v2_teacher_canonical_verification"
AI_REVISION_KEY = "g1b_v2_ai_standardization_revision"
AI_TASK_MONITOR_KEY = "g1b_v2_ai_task_monitor"
AI_TASK_EVIDENCE_KEY = "g1b_v2_ai_task_evidence"


_AI_TASKS = (
    ("CONFIG", "Đọc và khóa cấu hình ADMIN"),
    ("PAGE", "Khổ giấy và lề trang"),
    ("FONT", "Font, cỡ chữ và màu chữ"),
    ("SPACING", "Giãn dòng và giãn chữ"),
    ("TABLE", "Khung trang, bảng và biểu"),
    ("ROW", "Kiểm tra tách hàng của bảng"),
    ("FORMULA_FONT", "Font công thức Toán"),
    ("FORMULA_VALUE", "Giá trị công thức không đổi"),
    ("INTEGRITY", "Nội dung, hình ảnh và cấu trúc"),
    ("GATE", "Cổng tuân thủ cuối cùng"),
    ("RELEASE", "Mở quyền Lưu, Tải xuống và Gộp"),
)


def _monitor_state(*, phase="idle", active=0, checks=None, details=None, message=""):
    return {
        "phase": str(phase),
        "active": int(active),
        "checks": dict(checks or {}),
        "details": dict(details or {}),
        "message": str(message or ""),
    }


def _compliance_monitor_state(compliance):
    raw_checks = tuple(compliance.get("checks") or ()) if isinstance(compliance, Mapping) else ()
    checks = {}
    groups = {
        "CONFIG": ("ACTIVE_CONFIGURATION_SNAPSHOT",),
        "PAGE": ("PAGE_SIZE", "PAGE_MARGINS"),
        "FONT": ("BODY_FONT", "FONT_COLOR"),
        "SPACING": ("CHARACTER_SPACING", "LINE_SPACING"),
        "TABLE": ("PAGE_BORDER", "TABLE_REPEAT_HEADER"),
        "ROW": ("TABLE_ROW_SPLIT",),
        "FORMULA_FONT": ("OLE_FORMULA_FONT",),
        "FORMULA_VALUE": ("FORMULA_VALUE_INTEGRITY",),
        "INTEGRITY": ("CONTENT_INTEGRITY", "MEDIA_INTEGRITY"),
    }
    by_code = {
        str(item.get("code")): str(item.get("status") or "UNVERIFIED").upper()
        for item in raw_checks if isinstance(item, Mapping)
    }
    evidence_by_code = {
        str(item.get("code")): dict(item)
        for item in raw_checks if isinstance(item, Mapping)
    }
    for task_code, evidence_codes in groups.items():
        statuses = [by_code[code] for code in evidence_codes if code in by_code]
        if "FAIL" in statuses:
            checks[task_code] = "blocked"
        elif "REVIEW_REQUIRED" in statuses:
            checks[task_code] = "review"
        elif statuses and all(value == "PASS" for value in statuses):
            checks[task_code] = "pass"
        else:
            checks[task_code] = "unverified"
    if "OLE_FORMULA_FONT" not in by_code and by_code.get("FORMULA_VALUE_INTEGRITY") == "PASS":
        checks["FORMULA_FONT"] = "pass"
    final_status = str(compliance.get("status") or "UNVERIFIED").upper() if isinstance(compliance, Mapping) else "UNVERIFIED"
    checks["GATE"] = "pass" if final_status == "PASS" else ("review" if final_status == "REVIEW_REQUIRED" else "blocked")
    checks["RELEASE"] = "pass" if final_status == "PASS" else "blocked"
    return _monitor_state(
        phase="complete",
        active=len(_AI_TASKS),
        checks=checks,
        details={"CONFIG": evidence_by_code.get("ACTIVE_CONFIGURATION_SNAPSHOT", {})},
        message="Kết luận: " + final_status,
    )


def _render_ai_task_monitor(target, state):
    state = state if isinstance(state, Mapping) else _monitor_state()
    phase = str(state.get("phase") or "idle")
    active = int(state.get("active") or 0)
    checks = dict(state.get("checks") or {})
    rows = []
    completed = 0
    for index, (code, label) in enumerate(_AI_TASKS):
        status = checks.get(code)
        if not status:
            if phase == "running" and index < active:
                status = "pass"
            elif phase == "running" and index == active:
                status = "running"
            else:
                status = "queued"
        if status == "pass":
            completed += 1
        icon = {"pass": "✓", "running": "●", "blocked": "!", "review": "?", "unverified": "?"}.get(status, "○")
        rows.append(
            '<div class="g1b-task g1b-' + escape(status) + '"><span>'
            + icon + '</span><div><b>' + escape(label) + '</b><small>'
            + {"pass": "Đạt", "running": "Đang thực hiện", "blocked": "Không đạt", "review": "Cần kiểm tra", "unverified": "Chưa xác minh"}.get(status, "Chờ xử lý")
            + '</small></div></div>'
        )
    percent = round(completed * 100 / len(_AI_TASKS))
    message = escape(str(state.get("message") or "Sẵn sàng theo dõi lần chuẩn hóa tiếp theo."))
    target.markdown(
        '<style>.block-container{max-width:100%;padding-left:1rem;padding-right:1rem}.g1b-ai-monitor{position:relative;width:calc(300% + 2rem);height:132px;overflow-x:auto;overflow-y:hidden;background:linear-gradient(145deg,#fff,#eef4ff);border:1px solid #d6e2f3;border-radius:12px;padding:9px 11px;margin:6px 0 12px;box-shadow:4px 6px 14px rgba(26,55,95,.13)}.g1b-ai-head{display:flex;align-items:center;gap:12px;height:25px}.g1b-ai-head h4{margin:0;color:#173a67;font-size:15px;white-space:nowrap}.g1b-ai-head small{color:#58708d;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}.g1b-progress{height:5px;background:#dfe8f5;border-radius:8px;margin:5px 0 7px}.g1b-progress i{display:block;height:100%;background:linear-gradient(90deg,#2374e1,#21b879);border-radius:8px;transition:width .35s ease}.g1b-task-flow{display:grid;grid-template-columns:repeat(11,minmax(92px,1fr));gap:5px;min-width:1050px}.g1b-task{position:relative;display:flex;align-items:center;gap:5px;height:68px;padding:6px;border:1px solid #dce5f0;border-radius:9px;color:#52647a;background:#f4f7fb;overflow:hidden}.g1b-task span{flex:0 0 19px;width:19px;height:19px;border-radius:50%;text-align:center;font-weight:700;z-index:1}.g1b-task div{z-index:1;min-width:0}.g1b-task b{display:block;font-size:10px;line-height:1.15}.g1b-task small{font-size:9px}.g1b-pass{background:linear-gradient(135deg,#e8fbf0,#cef2df);border-color:#a9dfc1}.g1b-pass span{background:#bcebcf;color:#08783e}.g1b-running{color:#073c73;border-color:#65aef1;background:linear-gradient(100deg,#d9efff 0%,#65b8ff 35%,#c9f5ff 50%,#478eea 70%,#d9efff 100%);background-size:260% 100%;animation:g1bwater 1.25s linear infinite}.g1b-running span{background:#fff;color:#1262c5;animation:g1bpulse .9s infinite}.g1b-blocked{background:#ffe7e7;border-color:#efb0b0}.g1b-blocked span{background:#ffcaca;color:#bc2525}.g1b-review,.g1b-unverified{background:#fff5d9;border-color:#edd797}.g1b-review span,.g1b-unverified span{background:#ffe9a9;color:#926700}.g1b-blocked::after,.g1b-review::after,.g1b-unverified::after{content:"";position:absolute;right:7px;top:16px;width:3px;height:38px;border-radius:3px;transform:rotate(34deg);opacity:.9}.g1b-blocked::after{background:#d93636}.g1b-review::after{background:#d49a00}.g1b-unverified::after{background:#8090a4}@keyframes g1bwater{0%{background-position:100% 0}100%{background-position:-160% 0}}@keyframes g1bpulse{50%{transform:scale(1.18)}}@media(max-width:900px){.g1b-ai-monitor{width:calc(300% + 1rem)}}</style>'
        '<section class="g1b-ai-monitor"><div class="g1b-ai-head"><h4>Tiến trình AI</h4><small>' + message + '</small></div><div class="g1b-progress"><i style="width:' + str(percent) + '%"></i></div><div class="g1b-task-flow">' + ''.join(rows) + '</div></section>',
        unsafe_allow_html=True,
    )


def _apply_real_progress_event(target, event):
    if not isinstance(event, Mapping):
        return
    code = str(event.get("code") or "").upper()
    valid_codes = [item[0] for item in _AI_TASKS]
    if code not in valid_codes:
        return
    status = str(event.get("status") or "running").lower()
    if status not in {"queued", "running", "pass", "blocked", "review", "unverified"}:
        status = "unverified"
    previous = st.session_state.get(AI_TASK_MONITOR_KEY)
    checks = dict(previous.get("checks") or {}) if isinstance(previous, Mapping) else {}
    details = dict(previous.get("details") or {}) if isinstance(previous, Mapping) else {}
    checks[code] = status
    state = _monitor_state(
        phase="running" if status == "running" else "complete" if code == "RELEASE" else "running",
        active=valid_codes.index(code),
        checks=checks,
        details=details,
        message=str(event.get("message") or "Pipeline đang xử lý tài liệu."),
    )
    st.session_state[AI_TASK_MONITOR_KEY] = state
    _render_ai_task_monitor(target, state)


def _render_admin_configuration_diagnostic(state):
    details = dict(state.get("details") or {}) if isinstance(state, Mapping) else {}
    evidence = details.get("CONFIG")
    if not isinstance(evidence, Mapping) or not evidence:
        return
    status = str(evidence.get("status") or "UNVERIFIED").upper()
    status_display = {
        "PASS": "ĐẠT",
        "FAIL": "KHÔNG ĐẠT",
        "BLOCKED": "BỊ CHẶN",
        "WARNING": "CẢNH BÁO",
        "REVIEW": "CẦN KIỂM TRA",
        "UNVERIFIED": "CHƯA XÁC MINH",
    }.get(status, status)
    expected = evidence.get("expected")
    actual = evidence.get("actual")
    with st.container(key="g1b_report_card_1"):
        with st.expander(
            "Chi tiết nhiệm vụ 1: Cấu hình ADMIN — " + status_display,
            expanded=status != "PASS",
        ):
            st.markdown("**Yêu cầu bắt buộc:** " + str(expected or "Bản chụp cấu hình ACTIVE bất biến"))
            if isinstance(actual, Mapping):
                st.write({
                    "Mã cấu hình toàn hệ thống": actual.get("global_profile_id"),
                    "Mã phiên bản ACTIVE": actual.get("global_version_id"),
                    "Môn áp dụng": actual.get("subject_ref"),
                    "Phân môn áp dụng": actual.get("component_ref"),
                    "Mã kiểm tra cấu hình": actual.get("configuration_hash"),
                    "Các trường bị khóa": actual.get("locked_paths"),
                })
            else:
                st.code("ACTIVE_CONFIGURATION_SNAPSHOT = MISSING", language="text")
            if status == "PASS":
                st.success("Pipeline đã nhận và khóa đúng cấu hình ADMIN đang có hiệu lực cho lần chuẩn hóa này.")
            else:
                missing = []
                if not isinstance(actual, Mapping):
                    missing.append("snapshot cấu hình")
                else:
                    if not actual.get("global_version_id"):
                        missing.append("global_version_id của phiên bản ACTIVE")
                    if not actual.get("configuration_hash"):
                        missing.append("configuration_hash")
                st.error("Nguyên nhân: thiếu " + ", ".join(missing or ["bằng chứng cấu hình hợp lệ"]) + ".")
                st.info(
                    "Giải pháp: ADMIN phải Publish/Activate cấu hình toàn hệ thống; "
                    "sau đó mở lại trang Chuẩn hóa để runtime nạp lại cấu hình trước khi chạy."
                )


def _text(value: Any, fallback: str = "—") -> str:
    normalized = str(value or "").strip()
    return normalized or fallback


def _smart_up_filename_identity(value: str) -> str:
    import re

    name = str(value or "").strip().casefold()
    if name.endswith(".docx"):
        name = name[:-5]
    name = re.sub(r"[\s_-]+", ".", name)
    name = re.sub(r"\.+", ".", name).strip(".")
    name = re.sub(
        r"\btuan0*(\d+)\b",
        lambda match: f"tuan{int(match.group(1)):02d}",
        name,
    )
    name = re.sub(
        r"\bbai0*(\d+)\b",
        lambda match: f"bai{int(match.group(1)):02d}",
        name,
    )
    return name


def _ppct_preferred_file_names(context, *, preferred_file_name):
    periods=tuple(int(x) for x in context.get('curriculum_periods', ()) if str(x).strip().isdigit() and int(x)>0)
    if not periods:
        return ()
    m=re.match(r'^KHBD\.([A-Z0-9]+?)(\d+)\.(?:TUAN\d+|BAI\d+|\d+)\.DOCX$', str(preferred_file_name or '').strip().upper())
    if m is None:
        return ()
    code,grade=m.groups()
    return tuple(f'KHBD.{code}{grade}.{period:03d}.docx' for period in periods)

def _matching_drive_files(
    files: tuple[StoredDocumentFile, ...],
    *,
    preferred_file_name: str,
    legacy_file_name: str,
    aliases: tuple[str, ...],
    additional_preferred_file_names: tuple[str, ...] = (),
) -> tuple[tuple[StoredDocumentFile, str], ...]:
    targets: list[tuple[str, str]] = []
    if preferred_file_name:
        targets.append(
            (
                _smart_up_filename_identity(preferred_file_name),
                "DRIVE_PREFERRED_FILENAME",
            )
        )
    if legacy_file_name:
        targets.append(
            (
                _smart_up_filename_identity(legacy_file_name),
                "DRIVE_LEGACY_FILENAME",
            )
        )
    for additional_name in additional_preferred_file_names:
        if str(additional_name or '').strip():
            targets.append((_smart_up_filename_identity(additional_name), 'DRIVE_PPCT_PREFERRED_FILENAME'))
    for alias in aliases:
        if str(alias or "").strip():
            targets.append(
                (
                    _smart_up_filename_identity(alias),
                    "DRIVE_ALIAS_FILENAME",
                )
            )

    matches: list[tuple[StoredDocumentFile, str]] = []
    for item in files:
        if not str(item.file_name or "").casefold().endswith(".docx"):
            continue
        identity = _smart_up_filename_identity(item.file_name)
        for target, reason in targets:
            if target and identity == target:
                matches.append((item, reason))
                break
    return tuple(matches)


def selected_group_context(
    session_state: Mapping[str, Any],
) -> dict[str, Any] | None:
    value = session_state.get(GROUP_CONTEXT_KEY)
    if not isinstance(value, Mapping):
        return None
    group_id = _text(value.get("group_id"), fallback="")
    if not group_id:
        return None
    return dict(value)


def _apply_styles() -> None:
    st.markdown(
        """
        <style>
        [data-testid="stMainBlockContainer"] {
            max-width:100%; padding:1rem 1.35rem 4rem;
        }
        .g1b-hero {
            background:#fff; border:1px solid #dbe3ee; border-radius:18px;
            padding:1rem 1.2rem; box-shadow:0 10px 28px rgba(15,23,42,.06);
        }
        .g1b-hero h1 { margin:0; color:#0f2747; font-size:1.65rem; }
        .g1b-hero p { margin:.4rem 0 0; color:#607086; }
        .g1b-viewer {
            background:#fff; border:1px solid #d5deea; border-radius:16px;
            padding:.75rem; min-height:1050px;
        }

        /* G1B_UI_P2A_APPROVED_STANDARDIZED_WORKSPACE */
        .g1b-hero {
            background:linear-gradient(180deg,rgba(255,255,255,.99),rgba(248,250,253,.98));
            border:1px solid #d8e1ec;
            border-radius:22px;
            padding:1.25rem 1.35rem 1.15rem;
            box-shadow:inset 0 1px 0 rgba(255,255,255,.95),0 2px 4px rgba(11,31,58,.10),0 16px 34px rgba(11,31,58,.10);
            margin-bottom:1rem;
        }
        .g1b-hero h1 { margin:0; color:#071a33; font-size:1.72rem; font-weight:800; letter-spacing:-.02em; }
        .g1b-hero p { margin:.5rem 0 0; color:#5e6f84; font-size:.98rem; }

        .g1b-context-grid {
            display:grid;
            grid-template-columns:repeat(4,minmax(0,1fr));
            gap:1rem;
            margin:.2rem 0 1rem;
        }
        .g1b-context-card {
            position:relative;
            overflow:hidden;
            min-height:112px;
            padding:1rem 1.05rem;
            border-radius:16px;
            color:#fff;
            background:radial-gradient(circle at 20% 0%,rgba(47,111,237,.24),transparent 40%),linear-gradient(145deg,#153b67 0%,#0b1f3a 54%,#061426 100%);
            border:1px solid #061426;
            box-shadow:inset 0 1px 0 rgba(255,255,255,.18),inset 0 -1px 0 rgba(0,0,0,.28),0 3px 0 #04101f,0 10px 22px rgba(6,20,38,.24);
        }
        .g1b-context-card::after {
            content:"";
            position:absolute;
            inset:0;
            pointer-events:none;
            background:linear-gradient(120deg,rgba(255,255,255,.08),transparent 34%);
        }
        .g1b-context-label { position:relative; z-index:1; color:#c7d6e8; font-size:.88rem; font-weight:600; margin-bottom:.55rem; }
        .g1b-context-value { position:relative; z-index:1; color:#fff; font-size:1.14rem; font-weight:800; line-height:1.25; word-break:break-word; }
        .g1b-context-meta {
            background:#fff;
            border:1px solid #d8e1ec;
            border-radius:12px;
            padding:.75rem .95rem;
            margin:.15rem 0 1rem;
            color:#31445b;
            box-shadow:0 4px 12px rgba(11,31,58,.06);
        }

        .stButton > button,
        .stDownloadButton > button,
        button[kind="primary"],
        button[kind="secondary"],
        [data-testid="stBaseButton-primary"],
        [data-testid="stBaseButton-secondary"] {
            background:linear-gradient(180deg,#173e69 0%,#0b1f3a 56%,#061426 100%) !important;
            color:#ffffff !important;
            border:1px solid #04101f !important;
            border-top-color:#365b7d !important;
            border-radius:12px !important;
            box-shadow:inset 0 1px 0 rgba(255,255,255,.20),inset 0 -1px 0 rgba(0,0,0,.28),0 3px 0 #04101f,0 8px 17px rgba(6,20,38,.24) !important;
            font-weight:750 !important;
            min-height:2.72rem !important;
        }
        .stButton > button:hover,
        .stDownloadButton > button:hover,
        button[kind="primary"]:hover,
        button[kind="secondary"]:hover,
        [data-testid="stBaseButton-primary"]:hover,
        [data-testid="stBaseButton-secondary"]:hover {
            background:linear-gradient(180deg,#1b4a7b 0%,#102a43 54%,#07182d 100%) !important;
            transform:translateY(-1px);
            box-shadow:inset 0 1px 0 rgba(255,255,255,.23),0 4px 0 #04101f,0 12px 24px rgba(6,20,38,.28) !important;
        }
        .stButton > button:active,
        .stDownloadButton > button:active,
        button[kind="primary"]:active,
        button[kind="secondary"]:active {
            transform:translateY(2px);
            box-shadow:inset 0 2px 4px rgba(0,0,0,.20),0 1px 0 #04101f,0 4px 10px rgba(6,20,38,.20) !important;
        }
        .stButton > button:disabled,
        .stDownloadButton > button:disabled,
        button[kind="primary"]:disabled,
        button[kind="secondary"]:disabled {
            background:linear-gradient(180deg,#d9e1eb,#cbd5e1) !important;
            color:#718096 !important;
            border-color:#b8c4d1 !important;
            box-shadow:inset 0 1px 0 rgba(255,255,255,.85) !important;
            transform:none !important;
        }

        [data-baseweb="tab-list"] { border-bottom:1px solid #d8e1ec; gap:.25rem; }
        [data-baseweb="tab"] { min-height:3rem; padding:.7rem 1rem; border-radius:10px 10px 0 0; color:#42566e; font-weight:700; }
        [data-baseweb="tab"][aria-selected="true"] { color:#0b1f3a !important; background:#ffffff; box-shadow:0 -2px 10px rgba(11,31,58,.06); }

        .g1b-viewer {
            background:#eef3f8 !important;
            border:1px solid #d8e1ec !important;
            border-radius:16px !important;
            padding:1rem !important;
            box-shadow:inset 0 1px 0 rgba(255,255,255,.9),0 8px 22px rgba(11,31,58,.08) !important;
        }

        @media (max-width:1100px) {
            .g1b-context-grid { grid-template-columns:repeat(2,minmax(0,1fr)); }
        }
        @media (max-width:700px) {
            .g1b-context-grid { grid-template-columns:1fr; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def _render_group_header(context: Mapping[str, Any]) -> None:
    periods = ", ".join(
        str(item) for item in context.get("curriculum_periods", ())
    ) or "\u2014"
    group_label = _text(
        context.get("canonical_group_name"),
        fallback=(
            f"Kh\u1ed1i {_text(context.get('grade'))} \u00b7 "
            f"Tu\u1ea7n {_text(context.get('week_number'))}"
        ),
    )
    subject_value = _text(
        context.get("subject_display") or context.get("subject_ref")
    )
    grade_value = _text(context.get("grade"))

    st.markdown(
        '<section class="g1b-hero">'
        '<h1>So\u1ea1n b\u00e0i c\u00f9ng chu\u1ea9n gi\u00e1o \u00e1n</h1>'
        '<p>Kh\u00f4ng gian l\u00e0m vi\u1ec7c c\u1ee7a nh\u00f3m gi\u00e1o \u00e1n '
        '\u0111\u00e3 ch\u1ecdn t\u1eeb trang So\u1ea1n b\u00e0i theo tu\u1ea7n.</p>'
        '</section>',
        unsafe_allow_html=True,
    )

    cards = (
        ("Nh\u00f3m gi\u00e1o \u00e1n", group_label),
        ("M\u00f4n", subject_value),
        ("Kh\u1ed1i", grade_value),
        ("Ti\u1ebft PPCT", periods),
    )
    card_html = ['<section class="g1b-context-grid">']
    for label, value in cards:
        card_html.append(
            '<article class="g1b-context-card">'
            f'<div class="g1b-context-label">{label}</div>'
            f'<div class="g1b-context-value">{value}</div>'
            '</article>'
        )
    card_html.append('</section>')
    st.markdown("".join(card_html), unsafe_allow_html=True)

    meta_lines = [f"<strong>B\u00e0i:</strong> {_text(context.get('lesson_title'))}"]
    for item in tuple(context.get("occurrences", ()) or ()):
        if isinstance(item, Mapping):
            class_display = item.get("class_display") or item.get("class_id")
            meta_lines.append(
                f"L\u1edbp {_text(class_display)} \u00b7 "
                f"Ng\u00e0y d\u1ea1y: {_text(item.get('teaching_date'))} \u00b7 "
                f"Ti\u1ebft TKB: {_text(item.get('timetable_period'))}"
            )
    st.markdown(
        '<div class="g1b-context-meta">' + "<br>".join(meta_lines) + "</div>",
        unsafe_allow_html=True,
    )

def _render_document(
    *,
    content: bytes | None,
    preview_html_builder: Callable[[bytes], str] | None,
    empty_message: str,
) -> None:
    if not content:
        st.info(empty_message)
        return
    if preview_html_builder is None:
        st.warning("Trình xem giáo án chưa được kết nối.")
        return
    st.components.v1.html(
        preview_html_builder(content),
        height=1050,
        scrolling=True,
    )



def _docx_text_for_ai(content: bytes) -> str:
    from io import BytesIO
    from docx import Document

    document = Document(BytesIO(bytes(content)))
    blocks: list[str] = []
    blocks.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
    for table in document.tables:
        for row in table.rows:
            value = " | ".join(cell.text.strip() for cell in row.cells)
            if value.strip(" |"):
                blocks.append(value)
    return "\n".join(blocks)


def _apply_ai_revision_bytes(source_content: bytes, revised_text: str) -> tuple[bytes, object]:
    from pathlib import Path
    import os
    import tempfile
    from document_standardization.lesson_plan_ai_revision_overlay import LessonPlanAiRevisionOverlay

    source_path = output_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as source_file:
            source_path = source_file.name
            source_file.write(bytes(source_content))
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as output_file:
            output_path = output_file.name
        result = LessonPlanAiRevisionOverlay().apply(
            source=Path(source_path), output=Path(output_path), revised_text=revised_text
        )
        return Path(output_path).read_bytes(), result
    finally:
        for path in (source_path, output_path):
            if path and os.path.exists(path):
                os.unlink(path)


def render_standardized_lesson_plan_authoring_v2(
    *,
    client: Any,
    user_id: str,
    preview_html_builder: Callable[[bytes], str] | None = None,
    standardize_handler: Callable[..., tuple[str, bytes]] | None = None,
    ai_handler: Callable[..., str] | None = None,
    ai_status: str = "AI chưa được cấu hình",
    smart_up_content_loader: Callable[[Any], bytes] | None = None,
    save_handler: Callable[..., Any] | None = None,
    back_handler: Callable[[], Any] | None = None,
) -> None:
    # SMART_UP_RUNTIME_STORAGE_BINDING
    if smart_up_content_loader is None:
        runtime_storage = st.session_state.get("document_library_storage")
        if runtime_storage is not None and callable(
            getattr(runtime_storage, "download", None)
        ):
            smart_up_content_loader = lambda document: runtime_storage.download(
                document.storage_file_id
            )

    _apply_styles()

    # V14B6I_RC3B_REPORT_CARDS_ONLY
    st.markdown(
        """
<style>
.st-key-g1b_report_card_1 div[data-testid="stExpander"] details,
.st-key-g1b_report_card_2 div[data-testid="stExpander"] details,
.st-key-g1b_report_card_3 div[data-testid="stExpander"] details,
.st-key-g1b_report_card_4 div[data-testid="stExpander"] details,
.st-key-g1b_report_card_5 div[data-testid="stExpander"] details,
.st-key-g1b_report_card_6 div[data-testid="stExpander"] details {
    border-radius: 12px !important;
    border: 1px solid rgba(29, 73, 112, .55) !important;
    background: #ffffff !important;
    overflow: hidden !important;
    box-shadow:
        0 4px 0 rgba(7, 22, 39, .28),
        0 9px 18px rgba(4, 25, 48, .14) !important;
    margin: .18rem 0 .48rem 0 !important;
}

.st-key-g1b_report_card_1 div[data-testid="stExpander"] summary,
.st-key-g1b_report_card_2 div[data-testid="stExpander"] summary,
.st-key-g1b_report_card_3 div[data-testid="stExpander"] summary,
.st-key-g1b_report_card_4 div[data-testid="stExpander"] summary,
.st-key-g1b_report_card_5 div[data-testid="stExpander"] summary,
.st-key-g1b_report_card_6 div[data-testid="stExpander"] summary {
    min-height: 45px !important;
    border-radius: 10px !important;
    padding: .48rem .78rem !important;
    color: #ffffff !important;
    font-weight: 760 !important;
    letter-spacing: .003em !important;
    box-shadow:
        inset 0 1px 0 rgba(255,255,255,.16),
        inset 0 -1px 0 rgba(0,0,0,.22) !important;
    transition:
        transform .14s ease,
        filter .14s ease,
        box-shadow .14s ease !important;
}

.st-key-g1b_report_card_1 div[data-testid="stExpander"] summary {
    background: linear-gradient(145deg,#050b12 0%,#0a1726 58%,#10243a 100%) !important;
}

.st-key-g1b_report_card_2 div[data-testid="stExpander"] summary {
    background: linear-gradient(145deg,#081829 0%,#102d4d 58%,#17456e 100%) !important;
}

.st-key-g1b_report_card_3 div[data-testid="stExpander"] summary {
    background: linear-gradient(145deg,#0c2742 0%,#17466c 58%,#1f5a86 100%) !important;
}

.st-key-g1b_report_card_4 div[data-testid="stExpander"] summary {
    background: linear-gradient(145deg,#123b5d 0%,#1d5a84 58%,#2874a4 100%) !important;
}

.st-key-g1b_report_card_5 div[data-testid="stExpander"] summary {
    background: linear-gradient(145deg,#1b5278 0%,#2877a6 58%,#3b91bd 100%) !important;
}

.st-key-g1b_report_card_6 div[data-testid="stExpander"] summary {
    background: linear-gradient(145deg,#317da5 0%,#4b9fc5 58%,#69b9d8 100%) !important;
}

.st-key-g1b_report_card_1 div[data-testid="stExpander"] summary *,
.st-key-g1b_report_card_2 div[data-testid="stExpander"] summary *,
.st-key-g1b_report_card_3 div[data-testid="stExpander"] summary *,
.st-key-g1b_report_card_4 div[data-testid="stExpander"] summary *,
.st-key-g1b_report_card_5 div[data-testid="stExpander"] summary *,
.st-key-g1b_report_card_6 div[data-testid="stExpander"] summary * {
    color: #ffffff !important;
}

.st-key-g1b_report_card_1 div[data-testid="stExpander"] summary:hover,
.st-key-g1b_report_card_2 div[data-testid="stExpander"] summary:hover,
.st-key-g1b_report_card_3 div[data-testid="stExpander"] summary:hover,
.st-key-g1b_report_card_4 div[data-testid="stExpander"] summary:hover,
.st-key-g1b_report_card_5 div[data-testid="stExpander"] summary:hover,
.st-key-g1b_report_card_6 div[data-testid="stExpander"] summary:hover {
    filter: brightness(1.07);
    transform: translateY(-1px);
}

.st-key-g1b_report_card_1 div[data-testid="stExpander"] details[open],
.st-key-g1b_report_card_2 div[data-testid="stExpander"] details[open],
.st-key-g1b_report_card_3 div[data-testid="stExpander"] details[open],
.st-key-g1b_report_card_4 div[data-testid="stExpander"] details[open],
.st-key-g1b_report_card_5 div[data-testid="stExpander"] details[open],
.st-key-g1b_report_card_6 div[data-testid="stExpander"] details[open] {
    background: linear-gradient(180deg,#ffffff 0%,#f7fbff 100%) !important;
}
</style>
        """,
        unsafe_allow_html=True,
    )

    context = selected_group_context(st.session_state)
    if context is None:
        st.warning(
            "Chưa có nhóm giáo án được chọn. "
            "Hãy mở trang Soạn bài theo tuần và chọn một nhóm giáo án."
        )
        return
    _render_group_header(context)
    group_id = str(context["group_id"])
    st.markdown("**Tìm và tải giáo án của nhóm**")
    st.caption(
        "Chọn đúng tệp DOCX trên máy. Tệp sẽ được gắn riêng với "
        "nhóm giáo án đang mở và hiển thị ngay ở tab xem trước."
    )
    preferred_file_name = _text(context.get("preferred_file_name"), fallback="")
    expected_file_name = _text(context.get("canonical_file_name"), fallback="")
    displayed_file_name = preferred_file_name or expected_file_name
    ppct_preferred_file_names = _ppct_preferred_file_names(context, preferred_file_name=preferred_file_name)
    if displayed_file_name:
        st.info(f"Tên giáo án cần tìm: {displayed_file_name}")

    # Smart Up discovery is read-only and scoped to the current user.
    smart_up_context = SmartUpContext(
        expected_file_name=expected_file_name,
        preferred_file_name=preferred_file_name,
        subject_ref=_text(context.get("subject_ref"), fallback=""),
        component_ref=_text(context.get("component_ref"), fallback=""),
        grade=_text(context.get("grade"), fallback=""),
        week_number=context.get("week_number"),
        lesson_id=_text(context.get("lesson_id"), fallback=""),
        lesson_title=_text(context.get("lesson_title"), fallback=""),
        curriculum_periods=tuple(
            int(item)
            for item in context.get("curriculum_periods", ())
            if str(item).strip().isdigit()
        ),
        # G1B_H5D_DEFAULT_FALLBACK_PPCT
        # PPCT preferred names are discovery aliases only.
        aliases=tuple(dict.fromkeys(
            tuple(context.get("legacy_file_names", ()) or ())
            + tuple(ppct_preferred_file_names or ())
        )),
    )

    smart_up_resolution = None
    if st.button(
        "Up giáo án",
        key=f"g1b_v2_smart_up_{group_id}",
        type="primary",
        use_container_width=True,
        help=(
            "Ưu tiên kho Google Drive do ADMIN đăng ký; "
            "sau đó kho đã đăng ký, máy cục bộ và tải thủ công."
        ),
    ):
        blocked_by_ambiguity = False
        runtime_storage = st.session_state.get("document_library_storage")
        drive_folder_id = _text(
            st.session_state.get(
                "admin_lesson_plan_google_drive_folder_id"
            ),
            fallback="",
        )

        # G1B_H4D3_DRIVE_FIRST_SMART_UP
        if (
            drive_folder_id
            and runtime_storage is not None
            and callable(getattr(runtime_storage, "list_folder_tree", None))
            and callable(getattr(runtime_storage, "download", None))
        ):
            try:
                drive_files = runtime_storage.list_folder_tree(
                    drive_folder_id,
                    recursive=True,
                    mime_type=None,
                )
                # G1B_H5C_R2_RUNTIME_DIAGNOSTIC
                drive_docx_names = tuple(
                    str(getattr(item, 'file_name', '') or '')
                    for item in tuple(drive_files or ())
                    if str(getattr(item, 'file_name', '') or '').lower().endswith('.docx')
                )
                with st.container(key="g1b_report_card_2"):
                    with st.expander('Chẩn đoán Smart Up PPCT (tạm thời)', expanded=True):
                        st.write('Tên tệp ưu tiên:', preferred_file_name or '—')
                        st.write('Tên tệp dự kiến:', expected_file_name or '—')
                        st.write('Tiết PPCT của nhóm:', tuple(context.get('curriculum_periods', ()) or ()))
                        st.write('Các tên tệp ưu tiên theo PPCT:', ppct_preferred_file_names or ('—',))
                        st.write('Mã thư mục Google Drive:', drive_folder_id or '—')
                        st.write('Tổng số tệp trong Google Drive:', len(tuple(drive_files or ())))
                        st.write('Số tệp giáo án DOCX:', len(drive_docx_names))
                        st.write(
                            'Các tệp KHBD phù hợp:',
                            tuple(name for name in drive_docx_names if name.upper().startswith('KHBD.'))[:100],
                        )
                drive_matches = _matching_drive_files(
                    tuple(drive_files or ()),
                    preferred_file_name=preferred_file_name,
                    legacy_file_name=expected_file_name,
                    aliases=tuple(
                        context.get("legacy_file_names", ()) or ()
                    ),
                    additional_preferred_file_names=ppct_preferred_file_names,
                )
                if len(drive_matches) == 1:
                    item, reason = drive_matches[0]
                    content = runtime_storage.download(item.file_id)
                    if not isinstance(content, bytes) or not content:
                        raise ValueError(
                            "Google Drive returned empty DOCX content"
                        )
                    st.session_state[ORIGINAL_DOCUMENT_KEY] = {
                        "file_name": item.file_name,
                        "content": content,
                        "group_id": group_id,
                        "source": "GOOGLE_DRIVE_SMART_UP",
                        "storage_provider": item.provider,
                        "storage_file_id": item.file_id,
                        "web_view_link": item.web_view_link,
                        "match_reason": reason,
                    }
                    st.session_state.pop(
                        STANDARDIZED_DOCUMENT_KEY,
                        None,
                    )
                    st.success(
                        "Đã tìm thấy giáo án trong kho Google Drive "
                        "và nạp vào khu vực xem trước."
                    )
                    st.rerun()
                elif len(drive_matches) > 1:
                    blocked_by_ambiguity = True
                    st.warning(
                        "Tìm thấy nhiều giáo án Google Drive có cùng tên "
                        "phù hợp. Hệ thống dừng để tránh tải nhầm."
                    )
            except Exception as error:
                st.session_state[
                    f"g1b_v2_drive_smart_up_error_{group_id}"
                ] = str(error)
                st.info(
                    "Smart Up chưa đọc được kho Google Drive đã cấu hình; "
                    "hệ thống sẽ thử nguồn dự phòng."
                )

        if not blocked_by_ambiguity:
            try:
                smart_up_catalog = TeacherDocumentCatalog(
                    SupabaseTeacherDocumentRepository(
                        client=client,
                        user_id=str(user_id),
                    )
                )
                smart_up_resolution = resolve_from_catalog(
                    smart_up_catalog,
                    smart_up_context,
                )
                # G1B_H5F_FINAL_PPCT_CATALOG_FALLBACK
                if (
                    getattr(smart_up_resolution, 'status', None) == 'NOT_FOUND'
                    and ppct_preferred_file_names
                ):
                    try:
                        _h5f_docs = tuple(catalog.search())
                        _h5f_targets = {
                            _normalize_lesson_plan_file_identity(name)
                            for name in ppct_preferred_file_names
                            if str(name or '').strip()
                        }
                        _h5f_matches = tuple(
                            doc
                            for doc in _h5f_docs
                            if _normalize_lesson_plan_file_identity(
                                getattr(doc, 'file_name', '')
                            ) in _h5f_targets
                        )
                        if len(_h5f_matches) == 1:
                            from lesson_planning_v2.services.lesson_plan_smart_up_resolver import resolve_documents as _h5f_resolve_documents
                            smart_up_resolution = _h5f_resolve_documents(
                                _h5f_matches,
                                smart_up_context,
                            )
                        elif len(_h5f_matches) > 1:
                            st.warning(
                                'Có nhiều giáo án PPCT trùng tên trong kho đã đăng ký; hệ thống không tự chọn.'
                            )
                    except Exception:
                        pass
                st.session_state[
                    f"g1b_v2_smart_up_resolution_{group_id}"
                ] = smart_up_resolution
                if (
                    smart_up_resolution is not None
                    and smart_up_resolution.status == "MULTIPLE"
                ):
                    blocked_by_ambiguity = True
            except Exception as error:
                st.session_state[
                    f"g1b_v2_smart_up_error_{group_id}"
                ] = str(error)

        if (
            not blocked_by_ambiguity
            and (
                smart_up_resolution is None
                or smart_up_resolution.status == "NOT_FOUND"
            )
        ):
            try:
                local_matches = find_local_lesson_plans(
                    preferred_file_name=preferred_file_name,
                    legacy_file_name=expected_file_name,
                    aliases=tuple(dict.fromkeys(
                        tuple(context.get("legacy_file_names", ()) or ())
                        + tuple(ppct_preferred_file_names or ())
                    )),
                )
                if len(local_matches) == 1:
                    item = local_matches[0]
                    st.session_state[ORIGINAL_DOCUMENT_KEY] = {
                        "file_name": item.path.name,
                        "content": read_local_lesson_plan(item),
                        "group_id": group_id,
                        "source": "LOCAL_SMART_UP",
                        "local_path": str(item.path),
                        "match_reason": item.match_reason,
                    }
                    st.session_state.pop(
                        STANDARDIZED_DOCUMENT_KEY,
                        None,
                    )
                    st.success(
                        "Đã tìm thấy và tải giáo án trên máy: "
                        f"{item.path.name}"
                    )
                    st.rerun()
                elif len(local_matches) > 1:
                    blocked_by_ambiguity = True
                    st.warning(
                        "Tìm thấy nhiều giáo án trên máy có cùng tên "
                        "phù hợp. Hệ thống dừng để tránh tải nhầm."
                    )
            except (OSError, PermissionError, ValueError):
                pass

    if smart_up_resolution is None:
        smart_up_resolution = st.session_state.get(
            f"g1b_v2_smart_up_resolution_{group_id}"
        )

    if smart_up_resolution is not None:
        if smart_up_resolution.status == "NOT_FOUND":
            st.info(
                "Chưa tìm thấy giáo án phù hợp trong kho đã đăng ký. "
                "Hãy chọn tệp DOCX thủ công."
            )
        elif smart_up_resolution.status == "FOUND" and smart_up_resolution.best:
            best = smart_up_resolution.best
            st.success(
                "Smart Up đã tìm thấy: "
                f"{best.document.file_name} · {best.match_reason}"
            )
            st.caption(
                "Đã xác định đúng tài liệu. "
                "Bước tiếp theo sẽ nạp nội dung read-only vào Preview."
            )
            # SMART_UP_AUTO_PREVIEW_LOAD
            if smart_up_content_loader is not None:
                try:
                    smart_up_content = smart_up_content_loader(best.document)
                    if isinstance(smart_up_content, bytes) and smart_up_content:
                        st.session_state[ORIGINAL_DOCUMENT_KEY] = {
                            "file_name": best.document.file_name,
                            "content": smart_up_content,
                            "group_id": group_id,
                            "source": "SMART_UP",
                            "storage_provider": best.document.storage_provider,
                            "storage_file_id": best.document.storage_file_id,
                            "match_reason": best.match_reason,
                        }
                        st.session_state.pop(STANDARDIZED_DOCUMENT_KEY, None)
                        st.success("Đã nạp giáo án vào khu vực xem trước.")
                    else:
                        st.warning("Đã tìm thấy tài liệu nhưng nội dung tải về đang rỗng.")
                except Exception as error:
                    st.warning(
                        "Đã tìm thấy giáo án nhưng chưa thể nạp nội dung tự động. "
                        "Bạn vẫn có thể chọn tệp DOCX thủ công."
                    )
                    st.session_state[f"g1b_v2_smart_up_load_error_{group_id}"] = str(error)
        elif smart_up_resolution.status == "MULTIPLE":
            names = [
                item.document.file_name
                for item in smart_up_resolution.candidates[:5]
            ]
            st.warning(
                "Tìm thấy nhiều giáo án có độ phù hợp tương đương: "
                + " · ".join(names)
                + ". Hệ thống chưa tự chọn để tránh tải nhầm."
            )

    uploaded = st.file_uploader(
        "Tìm và tải giáo án từ máy (.docx)",
        type=("docx",),
        accept_multiple_files=False,
        key=f"g1b_v2_upload_{group_id}",
        help="Chỉ nhận tệp Word định dạng DOCX cho nhóm đang mở.",
    )
    if uploaded is not None:
        uploaded_content = uploaded.getvalue()
        if not uploaded_content:
            st.error("Tệp giáo án đang rỗng. Hãy chọn lại tệp DOCX.")
        else:
            if (
                displayed_file_name
                and uploaded.name.casefold() != displayed_file_name.casefold()
            ):
                st.warning(
                    "Tên tệp chưa theo tên chuẩn của nhóm. "
                    f"Tên chuẩn gợi ý: {displayed_file_name}. "
                    "Hệ thống vẫn nhận tệp đã chọn để không làm gián đoạn công việc."
                )
            st.session_state[ORIGINAL_DOCUMENT_KEY] = {
                "file_name": uploaded.name,
                "content": uploaded_content,
                "group_id": group_id,
            }
            st.success(f"Đã tải {uploaded.name} cho nhóm giáo án đang mở.")
    original = st.session_state.get(ORIGINAL_DOCUMENT_KEY)
    if (
        isinstance(original, Mapping)
        and str(original.get("group_id", "")) != group_id
    ):
        original = None
        st.session_state.pop(ORIGINAL_DOCUMENT_KEY, None)
        st.session_state.pop(STANDARDIZED_DOCUMENT_KEY, None)
    original_content = (
        bytes(original.get("content", b""))
        if isinstance(original, Mapping)
        else b""
    )
    # G1B_V2_AI_STANDARDIZATION_ASSISTANT
    working_content = original_content
    if original_content:
        from hashlib import sha256
        source_hash = sha256(original_content).hexdigest()
        revision = st.session_state.get(AI_REVISION_KEY)
        if isinstance(revision, Mapping) and revision.get("group_id") == group_id and revision.get("source_hash") == source_hash:
            applied_content = bytes(revision.get("applied_content", b""))
            if applied_content:
                working_content = applied_content
        with st.expander("AI hỗ trợ kiểm tra và đề xuất chuẩn hóa", expanded=False):
            if callable(ai_handler):
                st.success("Trạng thái AI: " + str(ai_status))
            else:
                st.warning(str(ai_status or "AI chưa được cấu hình"))
            ai_request = st.text_area(
                "Yêu cầu dành cho AI",
                value="Kiểm tra nội dung giáo án, phát hiện điểm chưa đầy đủ và đề xuất bản chỉnh sửa. Giữ nguyên dữ liệu lịch, lớp, tiết PPCT và ngày dạy.",
                key=f"g1b_v2_ai_request_{group_id}",
                height=100,
            )
            if st.button("AI kiểm tra và đề xuất", key=f"g1b_v2_ai_review_{group_id}", disabled=not callable(ai_handler)):
                try:
                    with st.spinner("AI đang kiểm tra giáo án..."):
                        revised_text = ai_handler(request=ai_request, document=_docx_text_for_ai(original_content), context=context)
                    if not str(revised_text or "").strip():
                        raise RuntimeError("AI không trả về nội dung đề xuất.")
                    st.session_state[AI_REVISION_KEY] = {
                        "group_id": group_id, "source_hash": source_hash,
                        "revised_text": str(revised_text), "applied_content": b"",
                    }
                    st.rerun()
                except Exception as error:
                    st.error("AI chưa xử lý được giáo án: " + str(error))
            revision = st.session_state.get(AI_REVISION_KEY)
            if isinstance(revision, Mapping) and revision.get("group_id") == group_id and revision.get("source_hash") == source_hash:
                revised_text = st.text_area(
                    "Bản đề xuất của AI - giáo viên có thể chỉnh trước khi áp dụng",
                    value=str(revision.get("revised_text", "")),
                    key=f"g1b_v2_ai_revision_text_{group_id}", height=300,
                )
                apply_col, reset_col = st.columns(2)
                if apply_col.button("Áp dụng gợi ý AI vào bản làm việc", key=f"g1b_v2_ai_apply_{group_id}"):
                    try:
                        applied, overlay = _apply_ai_revision_bytes(original_content, revised_text)
                        revision = dict(revision)
                        revision.update({"revised_text": revised_text, "applied_content": applied})
                        st.session_state[AI_REVISION_KEY] = revision
                        warnings = tuple(getattr(overlay, "warnings", ()) or ())
                        if warnings:
                            st.warning(" ".join(warnings))
                        else:
                            st.success("Đã tạo bản làm việc từ gợi ý AI; file gốc vẫn được giữ nguyên.")
                        st.rerun()
                    except Exception as error:
                        st.error("Không thể áp dụng gợi ý AI an toàn: " + str(error))
                if reset_col.button("Bỏ gợi ý AI", key=f"g1b_v2_ai_reset_{group_id}"):
                    st.session_state.pop(AI_REVISION_KEY, None)
                    st.rerun()

    actions = st.columns(3)
    can_standardize = bool(working_content)
    # G1B_P6A_V2_USER_APPROVAL_TOGGLE
    st.session_state.setdefault(
        "g1b_v2_include_approval_block",
        True,
    )
    st.checkbox(
        "Sinh ph\u1ea7n T\u1ed5 CM duy\u1ec7t",
        key="g1b_v2_include_approval_block",
        help=(
            "USER ch\u1ec9 ch\u1ecdn c\u00f3 ho\u1eb7c kh\u00f4ng sinh kh\u1ed1i ph\u00ea duy\u1ec7t. "
            "Nh\u00e3n, c\u0103n l\u1ec1 v\u00e0 ch\u00ednh s\u00e1ch ng\u00e0y do c\u1ea5u h\u00ecnh "
            "ADMIN \u0111ang c\u00f3 hi\u1ec7u l\u1ef1c quy\u1ebft \u0111\u1ecbnh."
        ),
    )

    standardize_clicked = actions[0].button(
        "Chu\u1ea9n h\u00f3a",
        type="primary",
        disabled=not can_standardize,
        use_container_width=True,
        key=f"g1b_v2_standardize_{group_id}",
    )
    # G1B_V13B_INLINE_REAL_AI_TASK_MONITOR
    task_monitor_slot = actions[0].empty()
    st.session_state.setdefault(AI_TASK_MONITOR_KEY, _monitor_state())
    _render_ai_task_monitor(
        task_monitor_slot,
        st.session_state.get(AI_TASK_MONITOR_KEY),
    )
    # G1B_V14A_ADMIN_CONFIGURATION_DIAGNOSTIC
    _render_admin_configuration_diagnostic(
        st.session_state.get(AI_TASK_MONITOR_KEY)
    )
    if standardize_clicked:
        # V14B3_CLEAR_STALE_CANONICAL_EVIDENCE
        st.session_state.pop("_g1b_v2_pipeline_evidence", None)
        st.session_state.pop("_g1b_v2_runtime_compliance_diagnostic", None)
        st.session_state.pop(AUDIT_FIELD_EVIDENCE_KEY, None)
        st.session_state.pop(TEACHER_VERIFICATION_KEY, None)
        st.session_state.pop(AUDIT_RESULT_KEY, None)
        st.session_state[AI_TASK_MONITOR_KEY] = _monitor_state(
            phase="running",
            active=0,
            message="Đang đọc cấu hình ADMIN đang có hiệu lực.",
        )
        _render_ai_task_monitor(
            task_monitor_slot,
            st.session_state[AI_TASK_MONITOR_KEY],
        )
        if standardize_handler is None:
            st.session_state[AI_TASK_MONITOR_KEY] = _monitor_state(
                phase="complete",
                checks={"CONFIG": "blocked", "GATE": "blocked", "RELEASE": "blocked"},
                message="Không tìm thấy bộ chuẩn hóa đã kết nối.",
            )
            _render_ai_task_monitor(task_monitor_slot, st.session_state[AI_TASK_MONITOR_KEY])
            st.error(
                "B\u1ed9 chu\u1ea9n h\u00f3a ch\u01b0a \u0111\u01b0\u1ee3c "
                "k\u1ebft n\u1ed1i v\u1edbi trang n\u00e0y."
            )
        else:
            try:
                st.session_state[AI_TASK_MONITOR_KEY] = _monitor_state(
                    phase="running",
                    active=-1,
                    message="Đang chuyển tài liệu vào pipeline chuẩn hóa.",
                )
                _render_ai_task_monitor(task_monitor_slot, st.session_state[AI_TASK_MONITOR_KEY])
                with st.spinner(
                    "\u0110ang chu\u1ea9n h\u00f3a gi\u00e1o \u00e1n..."
                ):
                    import inspect

                    handler_arguments = dict(
                        file_name=str(
                            original.get("file_name", "giao-an.docx")
                        ),
                        content=working_content,
                        group_context=context,
                    )
                    handler_signature = inspect.signature(standardize_handler)
                    if (
                        "progress_callback" in handler_signature.parameters
                        or any(
                            parameter.kind == inspect.Parameter.VAR_KEYWORD
                            for parameter in handler_signature.parameters.values()
                        )
                    ):
                        handler_arguments["progress_callback"] = lambda event: _apply_real_progress_event(
                            task_monitor_slot, event
                        )
                    name, content = standardize_handler(**handler_arguments)
                standardized_bytes = bytes(content)
                if not standardized_bytes:
                    raise RuntimeError(
                        "Standardizer returned empty DOCX content."
                    )
                st.session_state[STANDARDIZED_DOCUMENT_KEY] = {
                    "file_name": str(name or "giao-an-da-chuan.docx"),
                    "content": standardized_bytes,
                    "group_id": context["group_id"],
                }
                # G1B_ENGLISH_PILOT01_A5H_FULL_AUDIT_RUNTIME
                try:
                    from document_standardization.lesson_plan_standardization_audit_evidence_adapter import (
                        build_full_audit_evidence,
                    )
                    from document_standardization.lesson_plan_standardization_audit_gate import (
                        LessonPlanStandardizationAuditGate,
                    )

                    audit_gate = LessonPlanStandardizationAuditGate()
                    pipeline_evidence = st.session_state.get(
                        "_g1b_v2_pipeline_evidence"
                    )
                    evidence_bundle = build_full_audit_evidence(
                        group_context=context,
                        standardized_content=standardized_bytes,
                        pipeline_evidence=(
                            pipeline_evidence
                            if isinstance(pipeline_evidence, Mapping)
                            else None
                        ),
                    )
                    if evidence_bundle.ready:
                        audit_result = audit_gate.evaluate(
                            original_content=original_content,
                            standardized_content=standardized_bytes,
                            canonical_context=evidence_bundle.canonical_context,
                            validated_analysis=evidence_bundle.validated_analysis,
                            context_result=evidence_bundle.context_result,
                            standardization_report=evidence_bundle.standardization_report,
                        )
                    else:
                        audit_result = audit_gate.evaluate_artifact_only(
                            original_content=original_content,
                            standardized_content=standardized_bytes,
                        )
                    # V14A: formatting evidence remains authoritative even when
                    # another canonical audit component is not ready.
                    standardization_report = evidence_bundle.standardization_report
                    compliance = (
                        standardization_report.get("compliance", {})
                        if isinstance(standardization_report, Mapping)
                        else {}
                    )
                    compliance_status = str(
                        compliance.get("status") or "UNVERIFIED"
                    ).upper()
                    # V14B6F_A7_PERSIST_RUNTIME_COMPLIANCE_EVIDENCE
                    # Diagnostic-only snapshot. It is read after Streamlit reruns.
                    st.session_state["_g1b_v2_runtime_compliance_diagnostic"] = (
                        dict(compliance) if isinstance(compliance, Mapping) else {}
                    )
                    if not compliance:
                        st.session_state[AI_TASK_MONITOR_KEY] = _monitor_state(
                            phase="complete",
                            active=len(_AI_TASKS),
                            checks={"CONFIG": "unverified", "GATE": "blocked", "RELEASE": "blocked"},
                            details={
                                "CONFIG": {
                                    "code": "ACTIVE_CONFIGURATION_SNAPSHOT",
                                    "status": "UNVERIFIED",
                                    "expected": "immutable ACTIVE snapshot",
                                    "actual": None,
                                }
                            },
                            message="MISSING_STANDARDIZATION_REPORT: pipeline không trả về báo cáo cấu hình.",
                        )
                    if compliance_status != "PASS":
                        audit_result = {
                            "status": "fail",
                            "trust_score": 0,
                            "message": (
                                "ADMIN Configuration Enforcement Gate: "
                                + compliance_status
                            ),
                            "evidence": tuple(compliance.get("checks") or ()),
                        }
                    st.session_state[AI_TASK_MONITOR_KEY] = _compliance_monitor_state(
                        compliance
                    ) if compliance else st.session_state[AI_TASK_MONITOR_KEY]
                    _render_ai_task_monitor(
                        task_monitor_slot,
                        st.session_state[AI_TASK_MONITOR_KEY],
                    )
                    st.session_state[AUDIT_RESULT_KEY] = audit_result
                    # V14B3_STORE_FIELD_LEVEL_CANONICAL_EVIDENCE
                    canonical_field_rows = {}
                    for validated_item in tuple(
                        getattr(
                            evidence_bundle.validated_analysis,
                            "proposals",
                            (),
                        )
                        or ()
                    ):
                        proposal = getattr(validated_item, "proposal", None)
                        field_value = getattr(proposal, "field", None)
                        field_key = getattr(field_value, "value", field_value)
                        field_key = str(field_key or "").strip()
                        status_raw = getattr(validated_item, "status", None)
                        status_text = getattr(status_raw, "value", status_raw)
                        canonical_field_rows[field_key] = {
                            "expected": getattr(
                                validated_item,
                                "canonical_value",
                                None,
                            ),
                            "found": getattr(
                                validated_item,
                                "found_value",
                                getattr(proposal, "value", None),
                            ),
                            "status": str(
                                status_text or "unverified"
                            ).lower(),
                        }
                    st.session_state[
                        AUDIT_FIELD_EVIDENCE_KEY
                    ] = canonical_field_rows
                except Exception as audit_error:
                    # V14B3_CLEAR_FIELD_EVIDENCE_ON_AUDIT_ERROR
                    st.session_state.pop(AUDIT_FIELD_EVIDENCE_KEY, None)
                    st.session_state[AI_TASK_MONITOR_KEY] = _monitor_state(
                        phase="complete",
                        checks={"GATE": "blocked", "RELEASE": "blocked"},
                        message="Không thể xác minh bằng chứng: " + str(audit_error),
                    )
                    _render_ai_task_monitor(
                        task_monitor_slot,
                        st.session_state[AI_TASK_MONITOR_KEY],
                    )
                    st.session_state[AUDIT_RESULT_KEY] = {
                        "status": "unverified",
                        "trust_score": 0,
                        "message": str(audit_error),
                    }
                final_monitor = st.session_state.get(AI_TASK_MONITOR_KEY, {})
                release_status = dict(final_monitor.get("checks") or {}).get("RELEASE") if isinstance(final_monitor, Mapping) else None
                st.session_state[f"g1b_v2_standardize_notice_{group_id}"] = {
                    "level": "success" if release_status == "pass" else "error",
                    "message": (
                        "Chuẩn hóa và kiểm duyệt giáo án đã đạt PASS."
                        if release_status == "pass"
                        else "Pipeline đã tạo DOCX nhưng cổng kiểm duyệt chưa đạt; hệ thống đã khóa Lưu/Tải/Gộp."
                    ),
                }
                st.rerun()
            except Exception as error:
                st.session_state[AI_TASK_MONITOR_KEY] = _monitor_state(
                    phase="complete",
                    checks={"GATE": "blocked", "RELEASE": "blocked"},
                    message="Chuẩn hóa bị dừng: " + str(error),
                )
                _render_ai_task_monitor(
                    task_monitor_slot,
                    st.session_state[AI_TASK_MONITOR_KEY],
                )
                st.error(
                    "Chu\u1ea9n h\u00f3a gi\u00e1o \u00e1n ch\u01b0a "
                    "th\u00e0nh c\u00f4ng: " + str(error)
                )

    standardize_notice = st.session_state.pop(
        f"g1b_v2_standardize_notice_{group_id}",
        "",
    )
    if standardize_notice:
        if isinstance(standardize_notice, Mapping):
            notice_message = str(standardize_notice.get("message") or "")
            if standardize_notice.get("level") == "success":
                st.success(notice_message)
            else:
                st.error(notice_message)
        else:
            st.success(str(standardize_notice))

    # V14B6F_A7_PERSISTED_RUNTIME_COMPLIANCE_DIAGNOSTIC
    runtime_compliance = st.session_state.get(
        "_g1b_v2_runtime_compliance_diagnostic", {}
    )
    if isinstance(runtime_compliance, Mapping) and runtime_compliance:
        runtime_compliance_status = str(
            runtime_compliance.get("status") or "UNVERIFIED"
        ).upper()
        with st.container(key="g1b_report_card_3"):
            with st.expander(
                "\u0043h\u1ea9n \u0111o\u00e1n t\u1ea1m th\u1eddi: Kiểm tra tuân thủ cấu hình ADMIN khi vận hành",
                expanded=True,
            ):
                st.caption(
                    "\u0043h\u1ec9 \u0111\u1ecdc b\u1eb1ng ch\u1ee9ng runtime c\u1ee7a l\u1ea7n chu\u1ea9n h\u00f3a hi\u1ec7n t\u1ea1i. "
                    "Kh\u1ed1i n\u00e0y kh\u00f4ng thay \u0111\u1ed5i PASS/FAIL hay quy\u1ec1n L\u01b0u/T\u1ea3i/G\u1ed9p."
                )
                runtime_compliance_status_display = {
                    "PASS": "ĐẠT",
                    "FAIL": "KHÔNG ĐẠT",
                    "BLOCKED": "BỊ CHẶN",
                    "WARNING": "CẢNH BÁO",
                    "REVIEW": "CẦN KIỂM TRA",
                    "UNVERIFIED": "CHƯA XÁC MINH",
                }.get(runtime_compliance_status, runtime_compliance_status)
                st.write(
                    "**Kết quả kiểm tra tuân thủ:**",
                    runtime_compliance_status_display,
                )
                runtime_checks = tuple(runtime_compliance.get("checks") or ())
                if runtime_checks:
                    for runtime_check in runtime_checks:
                        if not isinstance(runtime_check, Mapping):
                            continue
                        runtime_code = str(runtime_check.get("code") or "UNKNOWN")
                        runtime_status = str(
                            runtime_check.get("status") or "UNVERIFIED"
                        ).upper()
                        runtime_status_display = {
                            "PASS": "ĐẠT",
                            "FAIL": "KHÔNG ĐẠT",
                            "BLOCKED": "BỊ CHẶN",
                            "WARNING": "CẢNH BÁO",
                            "REVIEW": "CẦN KIỂM TRA",
                            "UNVERIFIED": "CHƯA XÁC MINH",
                        }.get(runtime_status, runtime_status)

                        runtime_code_display = {
                            "ACTIVE_CONFIGURATION_SNAPSHOT":
                                "Ảnh chụp cấu hình ADMIN đang có hiệu lực",
                        }.get(runtime_code, runtime_code)

                        st.markdown(
                            f"**{runtime_code_display} — {runtime_status_display}**"
                        )
                        runtime_expected = runtime_check.get("expected")
                        runtime_expected_display = {
                            "immutable ACTIVE snapshot":
                                "Bản chụp cấu hình ACTIVE bất biến",
                        }.get(runtime_expected, runtime_expected)
                        st.write(
                            "**Yêu cầu:**",
                            runtime_expected_display,
                        )
                        runtime_actual = runtime_check.get("actual")
                        if isinstance(runtime_actual, Mapping):
                            st.write("**Dữ liệu thực tế:**")
                            st.write({
                                "Mã cấu hình toàn hệ thống":
                                    runtime_actual.get("global_profile_id"),
                                "Mã phiên bản ACTIVE":
                                    runtime_actual.get("global_version_id"),
                                "Mã hồ sơ cấu hình môn học":
                                    runtime_actual.get("subject_profile_id"),
                                "Mã phiên bản cấu hình môn học":
                                    runtime_actual.get("subject_version_id"),
                                "Môn áp dụng":
                                    runtime_actual.get("subject_ref"),
                                "Phân môn áp dụng":
                                    runtime_actual.get("component_ref"),
                                "Mã kiểm tra cấu hình":
                                    runtime_actual.get("configuration_hash"),
                                "Các trường bị khóa":
                                    runtime_actual.get("locked_paths"),
                            })
                        else:
                            st.write(
                                "**Dữ liệu thực tế:**",
                                runtime_actual,
                            )
                else:
                    st.info("Chưa có danh sách tiêu chí kiểm tra tuân thủ cấu hình ADMIN.")

    # G1B_ENGLISH_PILOT01_A5H_AUDIT_STATUS_UI
    audit_blocks_save = False
    audit_result = st.session_state.get(AUDIT_RESULT_KEY)
    if audit_result is not None:
        raw_status = getattr(audit_result, "status", None)
        status_value = getattr(raw_status, "value", raw_status)
        trust_score = getattr(audit_result, "trust_score", None)
        if isinstance(audit_result, Mapping):
            status_value = audit_result.get("status", status_value)
            trust_score = audit_result.get("trust_score", trust_score)
        normalized_status = str(status_value or "unverified").lower()
        canonical_field_rows = st.session_state.get(AUDIT_FIELD_EVIDENCE_KEY, {})
        if not isinstance(canonical_field_rows, Mapping):
            canonical_field_rows = {}
        teacher_verification = st.session_state.get(TEACHER_VERIFICATION_KEY, {})
        if not isinstance(teacher_verification, Mapping):
            teacher_verification = {}

        standardized_snapshot = st.session_state.get(STANDARDIZED_DOCUMENT_KEY)
        standardized_snapshot_content = (
            bytes(standardized_snapshot.get("content", b""))
            if isinstance(standardized_snapshot, Mapping)
            else b""
        )
        verification_scope = {
            "group_id": str(context.get("group_id") or ""),
            "output_sha256": (
                sha256(standardized_snapshot_content).hexdigest()
                if standardized_snapshot_content
                else ""
            ),
        }
        required_field_keys = (
            "class_name",
            "curriculum_period",
            "lesson_title",
            "drafting_date",
            "teaching_date",
        )
        effective_field_status = {}
        for _field_key in required_field_keys:
            _row = canonical_field_rows.get(_field_key, {})
            if not isinstance(_row, Mapping):
                _row = {}
            _auto_status = str(_row.get("status") or "unverified").lower()
            _verified = teacher_verification.get(_field_key, {})
            _teacher_ok = (
                isinstance(_verified, Mapping)
                and bool(_verified.get("confirmed"))
                and str(_verified.get("group_id") or "") == verification_scope["group_id"]
                and str(_verified.get("output_sha256") or "") == verification_scope["output_sha256"]
                and str(_verified.get("expected_snapshot") or "") == str(_row.get("expected") or "")
                and str(_verified.get("found_snapshot") or "") == str(_row.get("found") or "")
                and str(_verified.get("teacher_value") or "").strip() != ""
            )
            effective_field_status[_field_key] = (
                "teacher_verified"
                if _auto_status in ("conflict", "unverified") and _teacher_ok
                else _auto_status
            )
        canonical_pass_100 = bool(canonical_field_rows) and all(
            effective_field_status.get(_field_key) in ("accepted", "teacher_verified")
            for _field_key in required_field_keys
        )

        final_monitor_for_release = st.session_state.get(AI_TASK_MONITOR_KEY, {})
        final_checks_for_release = (
            dict(final_monitor_for_release.get("checks") or {})
            if isinstance(final_monitor_for_release, Mapping)
            else {}
        )
        admin_enforcement_pass = (
            str(final_checks_for_release.get("GATE") or "").lower() == "pass"
            and str(final_checks_for_release.get("RELEASE") or "").lower() == "pass"
        )
        # V14B6F_A8_R2_LOCKED_TEACHER_AUTHORITY_RELEASE_POLICY
        # ADMIN Compliance remains an independent supervision result.
        release_allowed = canonical_pass_100
        audit_blocks_save = not release_allowed
        score_text = str(int(trust_score)) + "%" if isinstance(trust_score, (int, float)) else "N/A"
        # V14B3_ADMIN_FORMAT_CONCLUSION
        st.success(
            "\u2713 Chu\u1ea9n h\u00f3a \u0111\u1ecbnh d\u1ea1ng theo c\u1ea5u h\u00ecnh ADMIN: "
            "TH\u00c0NH C\u00d4NG"
        )
        if canonical_pass_100:
            st.success(
                "\u2713 Ki\u1ec3m duy\u1ec7t d\u1eef li\u1ec7u canonical: PASS 100% - "
                "\u0111\u00e3 x\u00e1c minh \u0111\u1ea7y 5/5 tr\u01b0\u1eddng b\u1eaft bu\u1ed9c."
            )
        elif normalized_status == "warning":
            st.warning(
                "\u26a0 Ki\u1ec3m duy\u1ec7t d\u1eef li\u1ec7u canonical: C\u1ea2NH B\u00c1O - "
                "c\u00f2n tr\u01b0\u1eddng ch\u01b0a \u0111\u1ee7 b\u1eb1ng ch\u1ee9ng x\u00e1c minh. "
                "M\u1ee9c \u0111\u1ed9 bao ph\u1ee7 b\u1eb1ng ch\u1ee9ng: " + score_text
            )
        elif normalized_status == "fail":
            st.error(
                "\u2717 Ki\u1ec3m duy\u1ec7t d\u1eef li\u1ec7u canonical: FAIL - "
                "ph\u00e1t hi\u1ec7n xung \u0111\u1ed9t ho\u1eb7c l\u1ed7i c\u1ea7n x\u1eed l\u00fd."
            )
        else:
            st.warning(
                "? Ki\u1ec3m duy\u1ec7t d\u1eef li\u1ec7u canonical: CH\u01afA X\u00c1C MINH - "
                "ch\u01b0a \u0111\u1ee7 b\u1eb1ng ch\u1ee9ng \u0111\u1ec3 k\u1ebft lu\u1eadn."
            )
        # V14B3_CANONICAL_FIELD_REPORT
        canonical_field_rows = st.session_state.get(
            AUDIT_FIELD_EVIDENCE_KEY,
            {},
        )
        if isinstance(canonical_field_rows, Mapping):
            field_labels = {
                "class_name": "L\u1edbp",
                "curriculum_period": "Ti\u1ebft PPCT",
                "lesson_title": "T\u00ean b\u00e0i",
                "drafting_date": "Ng\u00e0y so\u1ea1n",
                "teaching_date": "Ng\u00e0y d\u1ea1y",
            }
            status_labels = {
                "accepted": "\u0110\u00c3 X\u00c1C MINH",
                "conflict": "XUNG \u0110\u1ed8T",
                "unverified": "CH\u01afA X\u00c1C MINH",
            }
            with st.container(key="g1b_report_card_4"):
                with st.expander(
                    "\u0110\u1ed1i chi\u1ebfu 5 tr\u01b0\u1eddng d\u1eef li\u1ec7u canonical",
                    expanded=(not canonical_pass_100),
                ):
                    st.caption(
                        "M\u1ed7i tr\u01b0\u1eddng ch\u1ec9 \u0111\u01b0\u1ee3c ch\u1ea5p nh\u1eadn khi gi\u00e1 tr\u1ecb "
                        "t\u00ecm th\u1ea5y trong DOCX kh\u1edbp v\u1edbi gi\u00e1 tr\u1ecb h\u1ec7 th\u1ed1ng "
                        "y\u00eau c\u1ea7u. CH\u01afA X\u00c1C MINH kh\u00f4ng \u0111\u01b0\u1ee3c t\u00ednh l\u00e0 PASS."
                    )
                    for field_key in (
                        "class_name",
                        "curriculum_period",
                        "lesson_title",
                        "drafting_date",
                        "teaching_date",
                    ):
                        row = canonical_field_rows.get(field_key, {})
                        if not isinstance(row, Mapping):
                            row = {}
                        expected = row.get("expected")
                        found = row.get("found")
                        row_status = str(
                            row.get("status") or "unverified"
                        ).lower()
                        expected_text = (
                            str(expected)
                            if expected not in (None, "")
                            else "Ch\u01b0a c\u00f3 gi\u00e1 tr\u1ecb canonical"
                        )
                        found_text = (
                            str(found)
                            if found not in (None, "")
                            else "Ch\u01b0a t\u00ecm th\u1ea5y \u0111\u1ee7 b\u1eb1ng ch\u1ee9ng trong DOCX"
                        )
                        effective_status = effective_field_status.get(field_key, row_status)
                        verified_record = teacher_verification.get(field_key, {})
                        teacher_value = (
                            str(verified_record.get("teacher_value") or "")
                            if isinstance(verified_record, Mapping)
                            else ""
                        )
                        effective_label = (
                            "\u0110\u00c3 X\u00c1C MINH B\u1edeI GI\u00c1O VI\u00caN"
                            if effective_status == "teacher_verified"
                            else ("\u0110\u00c3 X\u00c1C MINH T\u1ef0 \u0110\u1ed8NG" if row_status == "accepted" else status_labels.get(row_status, "CH\u01afA X\u00c1C MINH"))
                        )
                        st.markdown(
                            "**" + field_labels[field_key] + "**  \n"
                            + "- Gi\u00e1 tr\u1ecb h\u1ec7 th\u1ed1ng y\u00eau c\u1ea7u: `" + expected_text + "`  \n"
                            + "- Gi\u00e1 tr\u1ecb t\u00ecm th\u1ea5y trong DOCX: `" + found_text + "`  \n"
                            + "- Tr\u1ea1ng th\u00e1i: **" + effective_label + "**"
                        )
                        if effective_status == "teacher_verified" and teacher_value:
                            st.caption("Gi\u00e1 tr\u1ecb gi\u00e1o vi\u00ean x\u00e1c nh\u1eadn: " + teacher_value)
                        if row_status in ("conflict", "unverified"):
                            st.caption("Y\u00eau c\u1ea7u gi\u00e1o vi\u00ean x\u00e1c minh tr\u01b0\u1eddng n\u00e0y.")
                            input_key = f"g1b_v2_teacher_value_{group_id}_{field_key}"
                            proposed_teacher_value = st.text_input(
                                "Gi\u00e1 tr\u1ecb gi\u00e1o vi\u00ean x\u00e1c nh\u1eadn",
                                value=(teacher_value or expected_text),
                                key=input_key,
                            )
                            if st.button(
                                "X\u00e1c nh\u1eadn l\u00e0 \u0111\u00fang",
                                key=f"g1b_v2_teacher_confirm_{group_id}_{field_key}",
                                disabled=not str(proposed_teacher_value or "").strip(),
                            ):
                                updated = dict(teacher_verification)
                                updated[field_key] = {
                                    "confirmed": True,
                                    "teacher_value": str(proposed_teacher_value).strip(),
                                    "expected_snapshot": expected,
                                    "found_snapshot": found,
                                    "group_id": verification_scope["group_id"],
                                    "output_sha256": verification_scope["output_sha256"],
                                    "verified_by": str(user_id or ""),
                                    "verified_at": datetime.now(timezone.utc).isoformat(),
                                }
                                st.session_state[TEACHER_VERIFICATION_KEY] = updated
                                st.rerun()
                            # V14B6F_A8_R2_TEACHER_CONFIRMATION_IS_FINAL
                            # No revoke control: teacher confirmation is final business verification.

        # G1B_ENGLISH_PILOT01_A5J2A_AUDIT_EXPLAINABILITY
        audit_evidence = getattr(audit_result, "evidence", ())
        if isinstance(audit_result, Mapping):
            audit_evidence = audit_result.get("evidence", audit_evidence)
        audit_evidence = tuple(audit_evidence or ())

        if audit_evidence:
            with st.container(key="g1b_report_card_5"):
                with st.expander("B\u1eb1ng ch\u1ee9ng k\u1ef9 thu\u1eadt", expanded=False):
                    st.caption(
                        "B\u1ea3ng n\u00e0y ghi l\u1ea1i b\u1eb1ng ch\u1ee9ng \u0111\u1ed9c l\u1eadp "
                        "\u0111\u00e3 d\u00f9ng \u0111\u1ec3 k\u1ebft lu\u1eadn. FAIL/XUNG \u0110\u1ed8T "
                        "l\u00e0 l\u1ed7i c\u1ea7n x\u1eed l\u00fd; CH\u01afA X\u00c1C MINH l\u00e0 "
                        "ch\u01b0a \u0111\u1ee7 b\u1eb1ng ch\u1ee9ng."
                    )
                    for item in audit_evidence:
                        item_code = getattr(item, "code", "")
                        item_status_raw = getattr(item, "status", None)
                        item_status = getattr(item_status_raw, "value", item_status_raw)
                        item_message = getattr(item, "message", "")
                        item_values = getattr(item, "evidence", ())
                        if isinstance(item, Mapping):
                            item_code = item.get("code", item_code)
                            item_status = item.get("status", item_status)
                            item_message = item.get("message", item_message)
                            item_values = item.get("evidence", item_values)

                        item_status_text = str(item_status or "unverified").upper()
                        item_code_text = str(item_code or "AUDIT_EVIDENCE")
                        technical_messages = {
                            "SOURCE_DOCX_READABLE": "\u0110\u1ecdc \u0111\u01b0\u1ee3c gi\u00e1o \u00e1n g\u1ed1c \u0111\u1ed9c l\u1eadp.",
                            "OUTPUT_DOCX_READABLE": "\u0110\u1ecdc \u0111\u01b0\u1ee3c gi\u00e1o \u00e1n sau chu\u1ea9n h\u00f3a \u0111\u1ed9c l\u1eadp.",
                            "ARTIFACT_HASH_EVIDENCE": "\u0110\u00e3 ghi nh\u1eadn m\u00e3 b\u0103m c\u1ee7a t\u1ec7p g\u1ed1c v\u00e0 t\u1ec7p k\u1ebft qu\u1ea3.",
                            "CANONICAL_EVIDENCE_NOT_WIRED": "Ch\u01b0a nh\u1eadn \u0111\u1ee7 b\u1eb1ng ch\u1ee9ng canonical t\u1eeb pipeline.",
                        }
                        technical_message = technical_messages.get(
                            item_code_text,
                            "B\u1eb1ng ch\u1ee9ng k\u1ef9 thu\u1eadt cho ti\u00eau ch\u00ed "
                            + item_code_text
                            + ".",
                        )
                        technical_status_labels = {
                            "PASS": "\u0110\u1ea0T",
                            "WARNING": "C\u1ea2NH B\u00c1O",
                            "FAIL": "KH\u00d4NG \u0110\u1ea0T",
                            "UNVERIFIED": "CH\u01afA X\u00c1C MINH",
                        }
                        st.markdown(
                            "**"
                            + technical_status_labels.get(
                                item_status_text,
                                item_status_text,
                            )
                            + " | "
                            + item_code_text
                            + "** - "
                            + technical_message
                        )
                        for value in tuple(item_values or ()):
                            st.caption(str(value))
        else:
            st.caption(
                "Chi tiet kiem duyet: chua co bang chung theo tung tieu chi "
                "de hien thi."
            )
    standardized = st.session_state.get(STANDARDIZED_DOCUMENT_KEY)
    standardized_content = (
        bytes(standardized.get("content", b""))
        if isinstance(standardized, Mapping)
        else b""
    )
    # G1B_13H1R4B5J_TOP_SAVE_RUNTIME_WIRING
    if actions[1].button(
        "Lưu hệ thống",
        disabled=(save_handler is None or not standardized_content or audit_blocks_save),
        use_container_width=True,
    ):
        if save_handler is not None:
            try:
                save_handler(
                    artifact_file_name=(
                        str(
                            standardized.get(
                                "file_name",
                                "giao-an-da-chuan.docx",
                            )
                        )
                        if isinstance(standardized, Mapping)
                        else "giao-an-da-chuan.docx"
                    ),
                    artifact_content=standardized_content,
                )
            except Exception as error:
                st.error(
                    "Lưu hệ thống chưa thành công: " + str(error)
                )
    actions[2].download_button(
        "Tải xuống",
        data=standardized_content,
        file_name=(
            str(standardized.get("file_name", "giao-an-da-chuan.docx"))
            if isinstance(standardized, Mapping)
            else "giao-an-da-chuan.docx"
        ),
        disabled=(not standardized_content or audit_blocks_save),
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        use_container_width=True,
    )
    if original_content and standardize_handler is None:
        st.caption(
            "Bản nền V2 đã nhận đúng nhóm và tệp. "
            "Bộ chuẩn hóa do ADMIN cấu hình sẽ được "
            "nối ở bước tiếp theo."
        )
    original_tab, standardized_tab = st.tabs(
        ("Xem trước giáo án gốc", "Xem giáo án đã chuẩn")
    )
    with original_tab:
        _render_document(
            content=original_content,
            preview_html_builder=preview_html_builder,
            empty_message="Chưa tải giáo án gốc.",
        )
    with standardized_tab:
        trace_state = st.session_state.get(AI_TASK_MONITOR_KEY)
        if isinstance(trace_state, Mapping):
            trace_message = str(trace_state.get("message") or "").strip()
            if trace_message:
                st.info("Quá trình tạo bản xem: " + trace_message)
            trace_checks = dict(trace_state.get("checks") or {})
            if trace_checks:
                with st.container(key="g1b_report_card_6"):
                    with st.expander("Xem nhật ký các công đoạn đã tác động đến giáo án", expanded=False):
                        for task_code, task_label in _AI_TASKS:
                            task_status = str(trace_checks.get(task_code) or "queued")
                            st.caption(task_label + ": " + {
                                "pass": "Đạt",
                                "running": "Đang thực hiện",
                                "blocked": "Không đạt",
                                "review": "Cần kiểm tra",
                                "unverified": "Chưa xác minh",
                            }.get(task_status, "Chờ xử lý"))
        _render_document(
            content=standardized_content,
            preview_html_builder=preview_html_builder,
            empty_message="Chưa có giáo án đã chuẩn hóa.",
        )
    # G1B_13H1R4B_STANDARDIZED_LIST_WIRING
    if standardized_content:
        if not audit_blocks_save:
            from portal_v2.ui.standardized_lesson_plan_management_streamlit import (
                render_standardized_lesson_plan_management,
            )

            render_standardized_lesson_plan_management(
                current_file_name=(
                    str(
                        standardized.get(
                            "file_name",
                            "standardized-lesson-plan.docx",
                        )
                    )
                    if isinstance(standardized, Mapping)
                    else "standardized-lesson-plan.docx"
                ),
                current_content=standardized_content,
                preview_html_builder=preview_html_builder,
                save_handler=save_handler if not audit_blocks_save else None,
            )
        else:
            st.error(
                "Tệp chưa đạt ADMIN Configuration Enforcement Gate nên chưa được "
                "đưa vào danh sách Lưu/Tải/Gộp."
            )

    # G1B_13H1R4B4J_SAVE_AND_BACK_NAV
    st.markdown("---")
    if back_handler is not None:
        if st.button(
            '← Quay lại Soạn bài theo tuần',
            key="g1b_v2_back_to_weekly_schedule",
            use_container_width=True,
        ):
            back_handler()
            st.rerun()
