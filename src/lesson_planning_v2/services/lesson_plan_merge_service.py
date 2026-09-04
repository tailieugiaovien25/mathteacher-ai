from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from io import BytesIO
import re
import unicodedata
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


@dataclass(frozen=True)
class LessonPlanMergeSource:
    source_id: str
    file_name: str
    content: bytes


@dataclass(frozen=True)
class LessonPlanMergeResult:
    content: bytes
    source_ids: tuple[str, ...]
    source_file_names: tuple[str, ...]


class LessonPlanMergeError(ValueError):
    pass


class LessonPlanMergeService:
    """Merge standardized lesson-plan DOCX artifacts conservatively.

    The service is intentionally independent from Streamlit and storage.
    Source order is preserved exactly as supplied.

    Supported relationship-bearing body content:
    - images
    - external hyperlinks

    Unknown embedded relationship types are rejected rather than silently
    producing a damaged DOCX.
    """

    _APPROVAL_MARKERS = (
        "to cm duyet",
        "to chuyen mon duyet",
    )

    _DATE_LINE_RE = re.compile(
        r"(ngay|thang|nam).*(ngay|thang|nam)",
        re.IGNORECASE,
    )

    _REL_ATTRS = (
        qn("r:id"),
        qn("r:embed"),
        qn("r:link"),
    )

    def merge(
        self,
        sources: Sequence[LessonPlanMergeSource],
        *,
        include_approval: bool = True,
    ) -> LessonPlanMergeResult:
        ordered = tuple(sources)

        if not ordered:
            raise LessonPlanMergeError(
                "At least one standardized lesson plan is required."
            )

        self._validate_sources(ordered)

        documents = tuple(
            self._load_document(item)
            for item in ordered
        )

        approval_ranges = tuple(
            self._approval_range(document)
            for document in documents
        )

        final_range = approval_ranges[-1]

        if include_approval and final_range is None:
            raise LessonPlanMergeError(
                "The final lesson plan does not contain an approval block."
            )

        base = Document()

        # Remove the initial empty paragraph created by python-docx.
        body = base.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

        for source_index, (source, document) in enumerate(
            zip(ordered, documents)
        ):
            approval_range = approval_ranges[source_index]

            if source_index < len(ordered) - 1:
                block_elements = self._body_elements_without_approval(
                    document,
                    approval_range,
                )
            else:
                block_elements = self._body_elements_without_approval(
                    document,
                    approval_range,
                )

            for element in block_elements:
                cloned = self._clone_with_relationships(
                    source_document=document,
                    target_document=base,
                    element=element,
                )
                self._append_before_sectpr(base, cloned)

        # USER decides whether the merged artifact needs approval.
        if include_approval:
            final_document = documents[-1]
            assert final_range is not None
            for element in self._approval_elements(
                final_document,
                final_range,
            ):
                cloned = self._clone_with_relationships(
                    source_document=final_document,
                    target_document=base,
                    element=element,
                )
                self._append_before_sectpr(base, cloned)

        self._collapse_trailing_empty_paragraphs(base)

        output = BytesIO()
        base.save(output)
        content = output.getvalue()

        # Structural self-validation before returning bytes.
        check = Document(BytesIO(content))

        marker_count = sum(
            1
            for text in self._paragraph_texts(check)
            if self._is_approval_marker(text)
        )

        expected_marker_count = 1 if include_approval else 0
        if marker_count != expected_marker_count:
            raise LessonPlanMergeError(
                "Merged document approval marker count does not match "
                "the USER approval choice."
            )

        return LessonPlanMergeResult(
            content=content,
            source_ids=tuple(
                item.source_id
                for item in ordered
            ),
            source_file_names=tuple(
                item.file_name
                for item in ordered
            ),
        )

    def _validate_sources(
        self,
        sources: Sequence[LessonPlanMergeSource],
    ) -> None:
        seen_ids: set[str] = set()

        for item in sources:
            source_id = str(item.source_id).strip()
            file_name = str(item.file_name).strip()

            if not source_id:
                raise LessonPlanMergeError(
                    "Every merge source requires a stable source_id."
                )

            if source_id in seen_ids:
                raise LessonPlanMergeError(
                    "Duplicate source_id is not allowed: "
                    + source_id
                )

            seen_ids.add(source_id)

            if not file_name.casefold().endswith(".docx"):
                raise LessonPlanMergeError(
                    "Only DOCX standardized lesson plans can be merged."
                )

            if not isinstance(item.content, bytes) or not item.content:
                raise LessonPlanMergeError(
                    "Every merge source requires non-empty DOCX bytes."
                )

    def _load_document(
        self,
        source: LessonPlanMergeSource,
    ) -> _Document:
        try:
            return Document(BytesIO(source.content))
        except Exception as error:
            raise LessonPlanMergeError(
                "Invalid DOCX source: "
                + source.file_name
            ) from error

    def _body_elements(
        self,
        document: _Document,
    ) -> list:
        return [
            child
            for child in document.element.body
            if child.tag != qn("w:sectPr")
        ]

    def _body_elements_without_approval(
        self,
        document: _Document,
        approval_range: tuple[int, int] | None,
    ) -> list:
        elements = self._body_elements(document)

        if approval_range is None:
            return elements

        start, end = approval_range

        return [
            element
            for index, element in enumerate(elements)
            if not (start <= index < end)
        ]

    def _approval_elements(
        self,
        document: _Document,
        approval_range: tuple[int, int],
    ) -> list:
        start, end = approval_range
        elements = self._body_elements(document)
        return elements[start:end]

    def _approval_range(
        self,
        document: _Document,
    ) -> tuple[int, int] | None:
        elements = self._body_elements(document)

        marker_index: int | None = None

        for index, element in enumerate(elements):
            if element.tag != qn("w:p"):
                continue

            text = self._element_text(element)

            if self._is_approval_marker(text):
                marker_index = index

        if marker_index is None:
            return None

        # Approval must be near the end. This prevents accidental removal
        # if the phrase appears inside ordinary lesson content.
        remaining = len(elements) - marker_index

        if remaining > max(20, len(elements) // 3):
            raise LessonPlanMergeError(
                "Approval marker was found too far from the document end."
            )

        start = marker_index

        # Include the nearby administrative date line when present.
        for index in range(
            marker_index - 1,
            max(-1, marker_index - 4),
            -1,
        ):
            if index < 0:
                break

            candidate = elements[index]

            if candidate.tag != qn("w:p"):
                break

            text = self._element_text(candidate).strip()

            if not text:
                start = index
                continue

            normalized = self._normalize(text)

            if self._looks_like_date_line(normalized):
                start = index
                break

            break

        # Standardized approval/signature area is expected to be the
        # document tail. Keeping the tail together avoids orphan remnants.
        return start, len(elements)

    def _looks_like_date_line(self, normalized: str) -> bool:
        return bool(
            self._DATE_LINE_RE.search(normalized)
            or (
                "ngay" in normalized
                and "thang" in normalized
            )
        )

    def _is_approval_marker(self, text: str) -> bool:
        normalized = self._normalize(text)
        return any(
            marker in normalized
            for marker in self._APPROVAL_MARKERS
        )

    def _normalize(self, value: str) -> str:
        normalized = unicodedata.normalize("NFD", value)
        without_marks = "".join(
            char
            for char in normalized
            if unicodedata.category(char) != "Mn"
        )
        return " ".join(
            without_marks.casefold().split()
        )

    def _element_text(self, element) -> str:
        return "".join(
            node.text or ""
            for node in element.iter()
            if node.tag == qn("w:t")
        )

    def _paragraph_texts(
        self,
        document: _Document,
    ) -> Iterable[str]:
        for paragraph in document.paragraphs:
            yield paragraph.text

    def _clone_with_relationships(
        self,
        *,
        source_document: _Document,
        target_document: _Document,
        element,
    ):
        cloned = deepcopy(element)

        relationship_map: dict[str, str] = {}

        for node in cloned.iter():
            for attribute in self._REL_ATTRS:
                old_rid = node.get(attribute)

                if not old_rid:
                    continue

                if old_rid not in relationship_map:
                    relationship_map[old_rid] = (
                        self._copy_relationship(
                            source_document=source_document,
                            target_document=target_document,
                            relationship_id=old_rid,
                        )
                    )

                node.set(
                    attribute,
                    relationship_map[old_rid],
                )

        return cloned

    def _copy_relationship(
        self,
        *,
        source_document: _Document,
        target_document: _Document,
        relationship_id: str,
    ) -> str:
        relationship = source_document.part.rels.get(
            relationship_id
        )

        if relationship is None:
            raise LessonPlanMergeError(
                "Missing DOCX relationship: "
                + relationship_id
            )

        if relationship.reltype == RT.IMAGE:
            # G1B_13H1R4B5J_RAW_OPC_IMAGE_PART_COPY
            from docx.opc.part import Part

            image_part = relationship.target_part
            source_partname = str(image_part.partname)
            suffix = (
                "." + source_partname.rsplit(".", 1)[-1]
                if "." in source_partname
                else ".bin"
            )
            package = target_document.part.package
            target_partname = package.next_partname(
                "/word/media/image%d" + suffix
            )
            copied_part = Part(
                target_partname,
                image_part.content_type,
                image_part.blob,
                package,
            )
            return target_document.part.relate_to(
                copied_part,
                RT.IMAGE,
            )

        if (
            relationship.reltype
            == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
            and not relationship.is_external
        ):
            # G1B_13H1R4B5U1_RAW_OPC_OLE_PART_COPY
            from docx.opc.part import Part

            embedded_part = relationship.target_part
            source_partname = str(embedded_part.partname)
            suffix = (
                "." + source_partname.rsplit(".", 1)[-1]
                if "." in source_partname
                else ".bin"
            )
            package = target_document.part.package
            target_partname = package.next_partname(
                "/word/embeddings/oleObject%d" + suffix
            )
            copied_part = Part(
                target_partname,
                embedded_part.content_type,
                embedded_part.blob,
                package,
            )
            return target_document.part.relate_to(
                copied_part,
                relationship.reltype,
            )

        if (
            relationship.reltype == RT.HYPERLINK
            and relationship.is_external
        ):
            return target_document.part.relate_to(
                relationship.target_ref,
                RT.HYPERLINK,
                is_external=True,
            )

        raise LessonPlanMergeError(
            "Unsupported embedded DOCX relationship type: "
            + relationship.reltype
        )

    def _append_before_sectpr(
        self,
        document: _Document,
        element,
    ) -> None:
        body = document.element.body
        sectpr = body.sectPr

        if sectpr is None:
            body.append(element)
        else:
            sectpr.addprevious(element)

    def _collapse_trailing_empty_paragraphs(
        self,
        document: _Document,
    ) -> None:
        body = document.element.body

        while True:
            non_section = [
                child
                for child in body
                if child.tag != qn("w:sectPr")
            ]

            if len(non_section) <= 1:
                return

            last = non_section[-1]

            if last.tag != qn("w:p"):
                return

            if self._element_text(last).strip():
                return

            # Do not delete a paragraph containing drawings/equations.
            if last.xpath(".//w:drawing | .//m:oMath | .//m:oMathPara"):
                return

            body.remove(last)
