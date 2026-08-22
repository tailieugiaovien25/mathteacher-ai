"""Application helpers for the complete lesson-plan V1 workspace.

This module owns lesson-plan workspace identity and state coordination.
It is storage-neutral and contains no Streamlit or Supabase dependency.
"""

from __future__ import annotations

from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
import re
from typing import Any

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from lesson_planning_v2.services.lesson_plan_draft_workspace_service import (
    LessonPlanDraftWorkspaceService,
)
from lesson_planning_v2.workspace_draft import (
    LessonPlanWorkspaceDraft,
)


@dataclass(frozen=True)
class LessonPlanWorkspaceContext:
    teacher_user_id: str
    academic_year: str
    week_number: int
    subject_ref: str
    selection_mode: str
    selection_unit_id: str
    class_or_grade_ref: str = ""
    lesson_id: str = ""
    title: str = ""

    def __post_init__(self) -> None:
        required = (
            ("teacher_user_id", self.teacher_user_id),
            ("academic_year", self.academic_year),
            ("subject_ref", self.subject_ref),
            ("selection_mode", self.selection_mode),
            ("selection_unit_id", self.selection_unit_id),
        )

        for field_name, value in required:
            if not str(value).strip():
                raise ValueError(
                    f"{field_name} must not be empty"
                )

        if (
            not isinstance(self.week_number, int)
            or isinstance(self.week_number, bool)
            or self.week_number <= 0
        ):
            raise ValueError(
                "week_number must be a positive integer"
            )

    @property
    def draft_id(self) -> str:
        """
        Stable identity for exactly one teacher-owned
        lesson-plan drafting unit.
        """
        raw = "|".join(
            (
                self.teacher_user_id.strip(),
                self.academic_year.strip(),
                str(self.week_number),
                self.subject_ref.strip(),
                self.selection_mode.strip(),
                self.selection_unit_id.strip(),
            )
        )

        digest = sha256(
            raw.encode("utf-8")
        ).hexdigest()[:24]

        return f"lbg-draft-{digest}"

    @property
    def widget_prefix(self) -> str:
        return (
            "lbg_workspace_"
            + self.draft_id
        )


@dataclass(frozen=True)
class LessonPlanWorkspaceContent:
    objectives_text: str = ""
    materials_text: str = ""
    teaching_process_text: str = ""
    full_document_text: str = ""

    @classmethod
    def from_draft(
        cls,
        draft: LessonPlanWorkspaceDraft | None,
    ) -> "LessonPlanWorkspaceContent":
        if draft is None:
            return cls()

        full_document_text = str(
            getattr(
                draft,
                "full_document_text",
                "",
            )
            or ""
        ).strip()

        if not full_document_text:
            parts = []

            if draft.objectives_text.strip():
                parts.extend(
                    (
                        "I. M?C TI?U",
                        draft.objectives_text.strip(),
                    )
                )

            if draft.materials_text.strip():
                parts.extend(
                    (
                        "II. THI?T B? V? H?C LI?U",
                        draft.materials_text.strip(),
                    )
                )

            if draft.teaching_process_text.strip():
                parts.extend(
                    (
                        "III. TI?N TR?NH D?Y H?C",
                        draft.teaching_process_text.strip(),
                    )
                )

            full_document_text = "\n\n".join(
                parts
            ).strip()

        return cls(
            objectives_text=draft.objectives_text,
            materials_text=draft.materials_text,
            teaching_process_text=(
                draft.teaching_process_text
            ),
            full_document_text=(
                full_document_text
            ),
        )


class LessonPlanWorkspaceV1Service:
    """Coordinate load/save operations for one drafting unit."""

    def __init__(
        self,
        *,
        draft_service: LessonPlanDraftWorkspaceService,
    ) -> None:
        if draft_service is None:
            raise ValueError(
                "draft_service must not be None"
            )

        self._draft_service = draft_service

    def load(
        self,
        *,
        context: LessonPlanWorkspaceContext,
    ) -> LessonPlanWorkspaceDraft | None:
        return self._draft_service.get_draft(
            draft_id=context.draft_id,
            teacher_user_id=(
                context.teacher_user_id
            ),
        )

    def save(
        self,
        *,
        context: LessonPlanWorkspaceContext,
        content: LessonPlanWorkspaceContent,
        source: str = "editor",
        metadata: dict[str, Any] | None = None,
    ) -> LessonPlanWorkspaceDraft:
        normalized_metadata = dict(
            metadata or {}
        )

        normalized_metadata[
            "workspace_source"
        ] = str(source).strip() or "editor"

        draft = LessonPlanWorkspaceDraft(
            draft_id=context.draft_id,
            teacher_user_id=(
                context.teacher_user_id
            ),
            academic_year=(
                context.academic_year
            ),
            week_number=(
                context.week_number
            ),
            subject_ref=(
                context.subject_ref
            ),
            selection_mode=(
                context.selection_mode
            ),
            selection_unit_id=(
                context.selection_unit_id
            ),
            objectives_text=(
                content.objectives_text
            ),
            materials_text=(
                content.materials_text
            ),
            teaching_process_text=(
                content.full_document_text
                or content.teaching_process_text
            ),
            class_or_grade_ref=(
                context.class_or_grade_ref
                or None
            ),
            lesson_id=(
                context.lesson_id
                or context.selection_unit_id
            ),
            title=context.title,
            status="DRAFT",
            metadata=normalized_metadata,
        )

        return self._draft_service.save_draft(
            draft
        )




class LessonPlanFullDocumentDocxAdapter:
    """
    Convert the whole-document working text into an
    internal DOCX representation.

    This DOCX is an implementation detail used to pass
    the AI-edited lesson plan into the existing
    standardization pipeline.

    It is not a teacher-facing export action.
    """

    def build_bytes(
        self,
        full_document_text: str,
    ) -> bytes:
        value = str(
            full_document_text or ""
        ).strip()

        if not value:
            raise ValueError(
                "full_document_text must not be empty"
            )

        stream = BytesIO()

        document = Document()

        blocks = value.splitlines()

        for line in blocks:
            text_line = line.rstrip()

            if not text_line:
                document.add_paragraph("")
                continue

            document.add_paragraph(
                text_line
            )

        document.save(
            stream
        )

        return stream.getvalue()


class LessonPlanDocxWholeDocumentImporter:
    """
    DOCX -> editable whole-document text importer.

    Paragraphs and tables are emitted in their original
    document order.

    This importer does not split the lesson plan into
    fixed I/II/III sections.
    """

    def import_bytes(
        self,
        content: bytes,
    ) -> str:
        if not isinstance(
            content,
            (bytes, bytearray),
        ):
            raise TypeError(
                "content must be bytes"
            )

        if not content:
            raise ValueError(
                "content must not be empty"
            )

        document = Document(
            BytesIO(
                bytes(content)
            )
        )

        blocks: list[str] = []

        for child in (
            document.element.body.iterchildren()
        ):
            if isinstance(
                child,
                CT_P,
            ):
                paragraph = Paragraph(
                    child,
                    document,
                )

                value = (
                    paragraph.text.strip()
                )

                if value:
                    blocks.append(
                        value
                    )

                continue

            if isinstance(
                child,
                CT_Tbl,
            ):
                table = Table(
                    child,
                    document,
                )

                table_lines: list[str] = []

                for row in table.rows:
                    values = [
                        cell.text.strip()
                        for cell in row.cells
                    ]

                    if any(values):
                        table_lines.append(
                            " | ".join(
                                values
                            )
                        )

                if table_lines:
                    blocks.append(
                        "\n".join(
                            table_lines
                        )
                    )

        return "\n\n".join(
            blocks
        ).strip()


class LessonPlanDocxTextImporter:
    """
    Conservative DOCX -> editable-text importer.

    It does not mutate the source file.  It extracts text
    and maps obvious I/II/III sections when present.
    """

    _OBJECTIVES = (
        "I. MỤC TIÊU",
        "I. MUC TIEU",
        "MỤC TIÊU",
    )

    _MATERIALS = (
        "II. THIẾT BỊ VÀ HỌC LIỆU",
        "II. THIET BI VA HOC LIEU",
        "THIẾT BỊ VÀ HỌC LIỆU",
    )

    _PROCESS = (
        "III. TIẾN TRÌNH DẠY HỌC",
        "III. TIEN TRINH DAY HOC",
        "TIẾN TRÌNH DẠY HỌC",
    )

    def import_bytes(
        self,
        content: bytes,
    ) -> LessonPlanWorkspaceContent:
        if not isinstance(
            content,
            (bytes, bytearray),
        ):
            raise TypeError(
                "content must be bytes"
            )

        if not content:
            raise ValueError(
                "content must not be empty"
            )

        document = Document(
            BytesIO(bytes(content))
        )

        lines: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                lines.append(text)

        for table in document.tables:
            for row in table.rows:
                values = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if values:
                    lines.append(
                        " | ".join(values)
                    )

        return self._split_sections(
            lines
        )

    def _split_sections(
        self,
        lines: list[str],
    ) -> LessonPlanWorkspaceContent:
        sections = {
            "objectives": [],
            "materials": [],
            "process": [],
        }

        current = None

        for raw_line in lines:
            line = raw_line.strip()
            upper = self._normalize_heading(
                line
            )

            if self._matches(
                upper,
                self._OBJECTIVES,
            ):
                current = "objectives"
                continue

            if self._matches(
                upper,
                self._MATERIALS,
            ):
                current = "materials"
                continue

            if self._matches(
                upper,
                self._PROCESS,
            ):
                current = "process"
                continue

            if current is not None:
                sections[current].append(
                    line
                )

        # If the DOCX has no recognizable headings,
        # preserve all readable content in process rather
        # than silently throwing information away.
        if not any(
            sections.values()
        ):
            sections["process"] = list(
                lines
            )

        return LessonPlanWorkspaceContent(
            objectives_text="\n".join(
                sections["objectives"]
            ).strip(),
            materials_text="\n".join(
                sections["materials"]
            ).strip(),
            teaching_process_text="\n".join(
                sections["process"]
            ).strip(),
        )

    @staticmethod
    def _normalize_heading(
        value: str,
    ) -> str:
        return re.sub(
            r"\s+",
            " ",
            value.strip().upper(),
        )

    @classmethod
    def _matches(
        cls,
        actual: str,
        candidates: tuple[str, ...],
    ) -> bool:
        normalized_actual = (
            cls._normalize_heading(
                actual
            )
        )

        return any(
            normalized_actual
            == cls._normalize_heading(
                candidate
            )
            for candidate in candidates
        )


class LessonPlanSimpleDocxExporter:
    """
    V1 editor DOCX exporter.

    The professional weekly renderer remains the canonical
    weekly-output path.  This exporter exists for one
    teacher-edited lesson drafting unit.
    """

    def export(
        self,
        *,
        context: LessonPlanWorkspaceContext,
        content: LessonPlanWorkspaceContent,
    ) -> bytes:
        document = Document()

        section = document.sections[0]
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

        normal = document.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(14)

        title = document.add_paragraph()
        title.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = title.add_run(
            context.title
            or "KẾ HOẠCH BÀI DẠY"
        )
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        info = (
            f"Năm học: {context.academic_year}"
            f" | Tuần: {context.week_number}"
        )

        document.add_paragraph(
            info
        )

        if context.class_or_grade_ref:
            document.add_paragraph(
                "Lớp/Khối: "
                + context.class_or_grade_ref
            )

        self._add_section(
            document,
            "I. MỤC TIÊU",
            content.objectives_text,
        )

        self._add_section(
            document,
            "II. THIẾT BỊ VÀ HỌC LIỆU",
            content.materials_text,
        )

        self._add_section(
            document,
            "III. TIẾN TRÌNH DẠY HỌC",
            content.teaching_process_text,
        )

        output = BytesIO()
        document.save(output)

        return output.getvalue()

    @staticmethod
    def _add_section(
        document,
        heading: str,
        text: str,
    ) -> None:
        paragraph = document.add_paragraph()

        run = paragraph.add_run(
            heading
        )
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

        for value in (
            text.splitlines()
            if text.strip()
            else ("",)
        ):
            body = document.add_paragraph(
                value
            )

            for run in body.runs:
                run.font.name = (
                    "Times New Roman"
                )
                run.font.size = Pt(14)

@dataclass(frozen=True)
class LessonPlanLibraryItem:
    document_id: str
    title: str
    file_name: str
    mime_type: str
    storage_file_id: str
    web_view_link: str | None = None
    academic_year: str = ""
    subject: str = ""
    grade_level: str = ""
    class_name: str = ""


class LessonPlanLibrarySourceService:
    """
    Read-only application boundary between the teacher
    document catalog and the lesson-plan editor.
    """

    DOCX_MIME = (
        "application/vnd.openxmlformats-"
        "officedocument.wordprocessingml.document"
    )

    def __init__(
        self,
        *,
        repository,
        storage=None,
    ) -> None:
        if repository is None:
            raise ValueError(
                "repository must not be None"
            )

        self._repository = repository
        self._storage = storage

    def list_docx(
        self,
        *,
        academic_year: str = "",
        subject: str = "",
    ) -> tuple[LessonPlanLibraryItem, ...]:
        documents = (
            self._repository.list_all()
        )

        result = []

        for document in documents:
            mime_type = str(
                getattr(
                    document,
                    "mime_type",
                    "",
                )
            )

            file_name = str(
                getattr(
                    document,
                    "file_name",
                    "",
                )
            )

            if (
                mime_type != self.DOCX_MIME
                and not file_name
                .casefold()
                .endswith(".docx")
            ):
                continue

            document_year = str(
                getattr(
                    document,
                    "academic_year",
                    "",
                )
            )

            document_subject = str(
                getattr(
                    document,
                    "subject",
                    "",
                )
            )

            if (
                academic_year.strip()
                and document_year
                != academic_year.strip()
            ):
                continue

            if (
                subject.strip()
                and document_subject.casefold()
                != subject.strip().casefold()
            ):
                continue

            result.append(
                LessonPlanLibraryItem(
                    document_id=str(
                        document.document_id
                    ),
                    title=str(
                        document.title
                    ),
                    file_name=file_name,
                    mime_type=mime_type,
                    storage_file_id=str(
                        document.storage_file_id
                    ),
                    web_view_link=(
                        document.web_view_link
                    ),
                    academic_year=(
                        document_year
                    ),
                    subject=(
                        document_subject
                    ),
                    grade_level=str(
                        getattr(
                            document,
                            "grade_level",
                            "",
                        )
                    ),
                    class_name=str(
                        getattr(
                            document,
                            "class_name",
                            "",
                        )
                        or ""
                    ),
                )
            )

        return tuple(
            sorted(
                result,
                key=lambda item: (
                    item.title.casefold(),
                    item.file_name.casefold(),
                ),
            )
        )

    def load_bytes(
        self,
        item: LessonPlanLibraryItem,
    ) -> bytes:
        if self._storage is None:
            raise RuntimeError(
                "Kho file ch?a ???c k?t n?i "
                "cho thao t?c ??c."
            )

        loader = getattr(
            self._storage,
            "download",
            None,
        )

        if not callable(loader):
            raise RuntimeError(
                "Storage hi?n t?i kh?ng h? tr? download."
            )

        content = loader(
            item.storage_file_id
        )

        if not isinstance(
            content,
            bytes,
        ) or not content:
            raise RuntimeError(
                "Kh?ng ??c ???c n?i dung gi?o ?n."
            )

        return content
