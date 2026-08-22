from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from lesson_planning_v2.weekly_lesson_plan_docx_presentation import (
    WeeklyLessonPlanDocxPresentationProfile,
)

from lesson_planning_v2.weekly_lesson_plan_docx_layout import (
    WeeklyLessonPlanDocxLayoutProfile,
)

from lesson_planning_v2.weekly_lesson_plan_word_document import (
    WeeklyLessonPlanWordDocument,
    WeeklyLessonPlanWordSection,
)


class WeeklyLessonPlanDocxRenderer:
    def render(
        self,
        *,
        document: WeeklyLessonPlanWordDocument,
        output_path: str | Path,
        presentation_profile: WeeklyLessonPlanDocxPresentationProfile | None = None,
        layout_profile: WeeklyLessonPlanDocxLayoutProfile | None = None,
    ) -> Path:
        path = Path(output_path)

        path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        docx = Document()

        effective_layout = (
            layout_profile
            if layout_profile is not None
            else WeeklyLessonPlanDocxLayoutProfile.default()
        )

        self._configure_document(
            docx,
            effective_layout,
        )

        self._render_header(
            docx,
            document,
            presentation_profile,
            effective_layout,
        )

        for section in document.sections:
            self._render_section(
                docx,
                section,
                presentation_profile,
                effective_layout,
            )

            if (
                presentation_profile is not None
                and presentation_profile.page_break_between_sections
                and section is not document.sections[-1]
            ):
                docx.add_page_break()

        if document.approval is not None:
            self._render_approval(
                docx,
                document.approval.approver_role,
                presentation_profile,
                effective_layout,
            )

        self._apply_pagination_rules(
            docx
        )

        docx.save(path)

        return path

    @staticmethod
    def _configure_document(
        docx: DocumentObject,
        layout_profile: WeeklyLessonPlanDocxLayoutProfile,
    ) -> None:
        normal = docx.styles["Normal"]

        normal.font.name = layout_profile.body_font
        normal.font.size = Pt(
            layout_profile.body_size
        )

        paragraph_format = (
            normal.paragraph_format
        )

        paragraph_format.line_spacing = (
            layout_profile.line_spacing
        )
        paragraph_format.space_before = Pt(
            layout_profile.space_before_pt
        )
        paragraph_format.space_after = Pt(
            layout_profile.space_after_pt
        )

        for section in docx.sections:
            section.top_margin = Cm(
                layout_profile.top_margin_cm
            )
            section.bottom_margin = Cm(
                layout_profile.bottom_margin_cm
            )
            section.left_margin = Cm(
                layout_profile.left_margin_cm
            )
            section.right_margin = Cm(
                layout_profile.right_margin_cm
            )

    @staticmethod
    def _render_header(
        docx: DocumentObject,
        document: WeeklyLessonPlanWordDocument,
        presentation_profile: WeeklyLessonPlanDocxPresentationProfile | None = None,
        layout_profile: WeeklyLessonPlanDocxLayoutProfile | None = None,
    ) -> None:
        header = document.header

        title = docx.add_paragraph()

        title.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        title.paragraph_format.keep_with_next = (
            True
        )

        run = title.add_run(
            "GIÁO ÁN TUẦN "
            f"{header.week_number}"
        )

        run.bold = True

        if layout_profile is not None:
            run.font.name = (
                layout_profile.body_font
            )
            run.font.size = Pt(
                layout_profile.title_size
            )

        table = docx.add_table(
            rows=2,
            cols=2,
        )

        table.autofit = True

        values = (
            (
                "Giáo viên: "
                f"{header.teacher_id}",
                "Môn học: "
                f"{header.subject_ref}",
            ),
            (
                "Lớp / Khối: "
                f"{header.scope_label}",
                "Năm học: "
                f"{header.academic_year}",
            ),
        )

        for row_index, row_values in enumerate(
            values
        ):
            for column_index, value in enumerate(
                row_values
            ):
                cell = table.cell(
                    row_index,
                    column_index,
                )

                paragraph = (
                    cell.paragraphs[0]
                )

                paragraph.text = value

                if layout_profile is not None:
                    for cell_run in paragraph.runs:
                        cell_run.font.name = (
                            layout_profile.body_font
                        )
                        cell_run.font.size = Pt(
                            layout_profile.body_size
                        )

    @classmethod
    def _render_section(
        cls,
        docx: DocumentObject,
        section: WeeklyLessonPlanWordSection,
        presentation_profile: WeeklyLessonPlanDocxPresentationProfile | None = None,
        layout_profile: WeeklyLessonPlanDocxLayoutProfile | None = None,
    ) -> None:
        heading = docx.add_paragraph()

        heading.paragraph_format.keep_with_next = (
            True
        )

        run = heading.add_run(
            (
                f"TIẾT "
                f"{section.period_number}: "
                f"{section.title}"
            )
        )

        run.bold = True

        if layout_profile is not None:
            run.font.name = (
                layout_profile.body_font
            )
            run.font.size = Pt(
                layout_profile.heading_size
            )

        table = docx.add_table(
            rows=3,
            cols=2,
        )

        table.autofit = True

        component_text = (
            "Phân môn: "
            f"{section.component_ref}"
            if section.component_ref is not None
            else ""
        )

        values = (
            (
                "Tiết PPCT: "
                f"{section.curriculum_period}",
                "Lớp: "
                f"{section.class_id}",
            ),
            (
                "Ngày soạn: "
                f"{cls._format_date(section.preparation_date)}",
                "Ngày dạy: "
                f"{cls._format_date(section.teaching_date)}",
            ),
            (
                component_text,
                "",
            ),
        )

        for row_index, row_values in enumerate(
            values
        ):
            for column_index, value in enumerate(
                row_values
            ):
                cell = table.cell(
                    row_index,
                    column_index,
                )

                paragraph = (
                    cell.paragraphs[0]
                )

                paragraph.text = value

                if layout_profile is not None:
                    for cell_run in paragraph.runs:
                        cell_run.font.name = (
                            layout_profile.body_font
                        )
                        cell_run.font.size = Pt(
                            layout_profile.body_size
                        )

        cls._render_content(
            docx,
            section,
            presentation_profile,
            layout_profile,
        )

    @staticmethod
    def _render_content(
        docx: DocumentObject,
        section: WeeklyLessonPlanWordSection,
        presentation_profile: WeeklyLessonPlanDocxPresentationProfile | None = None,
        layout_profile: WeeklyLessonPlanDocxLayoutProfile | None = None,
    ) -> None:
        labels = (
            (
                "objectives",
                "I. M\u1ee5c ti\u00eau",
            ),
            (
                "materials",
                (
                    "II. Thi\u1ebft b\u1ecb "
                    "v\u00e0 h\u1ecdc li\u1ec7u"
                ),
            ),
            (
                "teaching_process",
                (
                    "III. Ti\u1ebfn tr\u00ecnh "
                    "d\u1ea1y h\u1ecdc"
                ),
            ),
        )

        for key, label in labels:
            if key not in section.content:
                continue

            heading = docx.add_paragraph()

            heading.paragraph_format.keep_with_next = True

            run = heading.add_run(label)
            run.bold = True
            if layout_profile is not None:
                run.font.name = layout_profile.body_font
                run.font.size = Pt(
                    layout_profile.heading_size
                )

            docx.add_paragraph(
                str(
                    section.content[key]
                )
            )

    @staticmethod
    def _render_approval(
        docx: DocumentObject,
        approver_role: str,
        presentation_profile: WeeklyLessonPlanDocxPresentationProfile | None = None,
        layout_profile: WeeklyLessonPlanDocxLayoutProfile | None = None,
    ) -> None:
        blank_lines = (
            presentation_profile.approval_blank_lines
            if presentation_profile is not None
            else 1
        )

        for _ in range(blank_lines):
            docx.add_paragraph("")

        paragraph = docx.add_paragraph()

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )

        run = paragraph.add_run(
            approver_role
        )

        run.bold = True
        if layout_profile is not None:
            run.font.name = layout_profile.body_font
            run.font.size = Pt(
                layout_profile.heading_size
            )

    @staticmethod
    def _format_date(value) -> str:
        return value.strftime(
            "%d/%m/%Y"
        )


    @staticmethod
    def _apply_pagination_rules(
        docx: DocumentObject,
    ) -> None:
        """
        Keep structural Word blocks together across page breaks.

        Metadata rows form one continuous header chain and
        I/II/III content headings must remain with the
        paragraph that follows them.
        """

        # Period/header metadata is rendered in tables.
        # Every populated metadata paragraph continues the
        # structural chain into the following paragraph/block.
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            (
                                paragraph
                                .paragraph_format
                                .keep_with_next
                            ) = True

        content_heading_prefixes = (
            "I. M?c ti?u",
            "II. Thi?t b? v? h?c li?u",
            "III. Ti?n tr?nh d?y h?c",
        )

        for paragraph in docx.paragraphs:
            normalized = (
                paragraph.text.strip()
            )

            if any(
                normalized.startswith(prefix)
                for prefix
                in content_heading_prefixes
            ):
                (
                    paragraph
                    .paragraph_format
                    .keep_with_next
                ) = True
