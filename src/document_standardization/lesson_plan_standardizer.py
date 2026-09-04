"""Standardize lesson-plan DOCX formatting without changing its content."""

from __future__ import annotations

import hashlib
import json
import os
import re
import tempfile
import zipfile
from collections import Counter
from copy import deepcopy
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree


@dataclass(frozen=True)
class DocumentInventory:
    paragraphs: int
    tables: int
    table_cells: int
    inline_shapes: int
    sections: int
    omml_equations: int
    equation_paragraphs: int
    ole_objects: int
    drawings: int
    images: int
    fonts: dict[str, int] = field(default_factory=dict)
    body_text_sha256: str = ""
    header_footer_text_sha256: str = ""
    formula_value_sha256: str = ""
    media_fingerprints: tuple[str, ...] = ()
    embedded_fingerprints: tuple[str, ...] = ()
    external_relationships: tuple[str, ...] = ()


@dataclass(frozen=True)
class LessonPlanStandardizationOptions:
    """User-confirmed operations for one standardization run."""

    preserve_original_maximum: bool = False
    sync_context: bool = True
    normalize_font: bool = True
    normalize_equations: bool = True
    normalize_tables: bool = True
    normalize_page_layout: bool = True
    normalize_spacing: bool = True
    normalize_header_footer: bool = True

    @property
    def has_selected_operation(self) -> bool:
        return any(
            (
                self.sync_context,
                self.normalize_font,
                self.normalize_equations,
                self.normalize_tables,
                self.normalize_page_layout,
                self.normalize_spacing,
                self.normalize_header_footer,
            )
        )


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _package_counts(path: Path) -> dict[str, int]:
    counts = Counter()
    with zipfile.ZipFile(path) as archive:
        counts["images"] = sum(name.startswith("word/media/") for name in archive.namelist())
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            xml = archive.read(name)
            counts["omml_equations"] += xml.count(b"<m:oMath>")
            counts["equation_paragraphs"] += xml.count(b"<m:oMathPara")
            counts["ole_objects"] += xml.count(b"<o:OLEObject")
            counts["drawings"] += xml.count(b"<w:drawing")
    return dict(counts)


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _text_sha256(values) -> str:
    payload = json.dumps(
        list(values),
        ensure_ascii=False,
        separators=(",", ":"),
    ).encode("utf-8")

    return _sha256_bytes(payload)


def _formula_value_fingerprint(path: Path) -> str:
    """Hash formula tokens only; formatting nodes are deliberately excluded."""
    math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    values = []
    with zipfile.ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(archive.read(name))
            tokens = root.xpath(".//m:oMath//m:t/text() | .//m:oMathPara//m:t/text()", namespaces={"m": math_ns})
            if tokens:
                values.append((name, tokens))
    return _text_sha256(values)


def _body_text_fingerprint(document) -> str:
    values = []

    for index, paragraph in enumerate(
        document.paragraphs
    ):
        values.append(
            (
                "paragraph",
                index,
                paragraph.text,
            )
        )

    for table_index, table in enumerate(
        document.tables
    ):
        for row_index, row in enumerate(
            table.rows
        ):
            for cell_index, cell in enumerate(
                row.cells
            ):
                values.append(
                    (
                        "cell",
                        table_index,
                        row_index,
                        cell_index,
                        cell.text,
                    )
                )

    return _text_sha256(values)


def _header_footer_text_fingerprint(
    document,
) -> str:
    values = []

    for section_index, section in enumerate(
        document.sections
    ):
        containers = (
            (
                "header",
                section.header,
            ),
            (
                "first_page_header",
                section.first_page_header,
            ),
            (
                "even_page_header",
                section.even_page_header,
            ),
            (
                "footer",
                section.footer,
            ),
            (
                "first_page_footer",
                section.first_page_footer,
            ),
            (
                "even_page_footer",
                section.even_page_footer,
            ),
        )

        for kind, container in containers:
            values.append(
                (
                    section_index,
                    kind,
                    tuple(
                        paragraph.text
                        for paragraph
                        in container.paragraphs
                    ),
                )
            )

    return _text_sha256(values)


def _package_fingerprints(
    path: Path,
) -> dict[str, tuple[str, ...]]:
    media = []
    embedded = []
    external_relationships = []

    with zipfile.ZipFile(path) as archive:
        names = set(
            archive.namelist()
        )

        for name in sorted(names):
            if name.startswith(
                "word/media/"
            ):
                media.append(
                    (
                        name
                        + ":"
                        + _sha256_bytes(
                            archive.read(name)
                        )
                    )
                )

            elif name.startswith(
                "word/embeddings/"
            ):
                embedded.append(
                    (
                        name
                        + ":"
                        + _sha256_bytes(
                            archive.read(name)
                        )
                    )
                )

        relationship_names = sorted(
            name
            for name in names
            if (
                name.endswith(".rels")
                and (
                    name.startswith(
                        "word/"
                    )
                    or name.startswith(
                        "_rels/"
                    )
                )
            )
        )

        rel_ns = (
            "http://schemas.openxmlformats.org/"
            "package/2006/relationships"
        )

        for name in relationship_names:
            root = etree.fromstring(
                archive.read(name)
            )

            for relationship in root:
                if (
                    relationship.tag
                    != f"{{{rel_ns}}}Relationship"
                ):
                    continue

                if (
                    relationship.get(
                        "TargetMode"
                    )
                    != "External"
                ):
                    continue

                external_relationships.append(
                    "|".join(
                        (
                            name,
                            relationship.get(
                                "Type",
                                "",
                            ),
                            relationship.get(
                                "Target",
                                "",
                            ),
                        )
                    )
                )

    return {
        "media": tuple(media),
        "embedded": tuple(embedded),
        "external_relationships": tuple(
            sorted(
                external_relationships
            )
        ),
    }


def inventory(path: Path) -> DocumentInventory:
    document = Document(path)

    fonts = Counter()
    table_cells = 0

    paragraphs = list(
        document.paragraphs
    )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_cells += 1
                paragraphs.extend(
                    cell.paragraphs
                )

    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.font.name:
                fonts[
                    run.font.name
                ] += 1

    package = _package_counts(
        path
    )

    fingerprints = (
        _package_fingerprints(
            path
        )
    )

    return DocumentInventory(
        paragraphs=len(
            paragraphs
        ),
        tables=len(
            document.tables
        ),
        table_cells=table_cells,
        inline_shapes=len(
            document.inline_shapes
        ),
        sections=len(
            document.sections
        ),
        omml_equations=package.get(
            "omml_equations",
            0,
        ),
        equation_paragraphs=package.get(
            "equation_paragraphs",
            0,
        ),
        ole_objects=package.get(
            "ole_objects",
            0,
        ),
        drawings=package.get(
            "drawings",
            0,
        ),
        images=package.get(
            "images",
            0,
        ),
        fonts=dict(
            sorted(
                fonts.items()
            )
        ),
        body_text_sha256=(
            _body_text_fingerprint(
                document
            )
        ),
        header_footer_text_sha256=(
            _header_footer_text_fingerprint(
                document
            )
        ),
        formula_value_sha256=_formula_value_fingerprint(path),
        media_fingerprints=(
            fingerprints[
                "media"
            ]
        ),
        embedded_fingerprints=(
            fingerprints[
                "embedded"
            ]
        ),
        external_relationships=(
            fingerprints[
                "external_relationships"
            ]
        ),
    )


class LessonPlanWordStandardizer:
    """Apply deterministic formatting rules to a copy of a DOCX file."""

    def __init__(self, profile: dict[str, object]):
        self.profile = deepcopy(profile)

    @classmethod
    def from_json(cls, path: Path) -> "LessonPlanWordStandardizer":
        return cls(json.loads(path.read_text(encoding="utf-8")))

    def standardize(
        self,
        source: Path,
        output: Path,
        report_path: Path,
        *,
        options: LessonPlanStandardizationOptions | None = None,
        progress_callback: Callable[[dict[str, Any]], None] | None = None,
    ) -> dict[str, object]:
        def emit(code: str, status: str, message: str) -> None:
            if progress_callback is None:
                return
            try:
                progress_callback({"code": code, "status": status, "message": message})
            except Exception:
                # Monitoring must never mutate or interrupt DOCX processing.
                pass

        emit("CONFIG", "running", "Äang xÃ¡c nháº­n snapshot cáº¥u hÃ¬nh ADMIN.")
        source = source.resolve()
        output = output.resolve()
        if source == output:
            raise ValueError("Không được ghi đè tệp Word gốc.")
        if source.suffix.lower() != ".docx" or output.suffix.lower() != ".docx":
            raise ValueError("CÃ´ng cá»¥ V1 chá»‰ xá»­ lÃ½ tá»‡p .docx.")

        options = options or LessonPlanStandardizationOptions()
        emit("CONFIG", "pass", "ÄÃ£ nháº­n cáº¥u hÃ¬nh Ã¡p dá»¥ng cho láº§n cháº¡y nÃ y.")
        before = inventory(source)
        source_hash = _sha256(source)
        document = Document(source)
        changes = Counter()
        if options.normalize_page_layout:
            emit("PAGE", "running", "Äang Ã¡p dá»¥ng khá»• giáº¥y, lá» vÃ  khung trang.")
            self._normalize_sections(document, changes)
        if options.normalize_font:
            emit("FONT", "running", "Äang Ã¡p dá»¥ng font, cá»¡ chá»¯ vÃ  mÃ u chá»¯.")
            self._normalize_styles(document)
        if options.normalize_font or options.normalize_spacing:
            emit("SPACING", "running", "Äang Ã¡p dá»¥ng giÃ£n dÃ²ng vÃ  giÃ£n chá»¯.")
            self._normalize_paragraphs(
                document,
                changes,
                normalize_font=options.normalize_font,
                normalize_spacing=options.normalize_spacing,
            )
        # Canonical output must not retain review highlights or cell shading,
        # even when the UI font/table switches are disabled.
        self._normalize_global_text_color(
            document,
            changes,
        )
        if options.normalize_tables:
            emit("TABLE", "running", "Äang Ä‘á»‹nh dáº¡ng báº£ng, biá»ƒu vÃ  hÃ ng báº£ng.")
            emit("ROW", "running", "Äang thiáº¿t láº­p quy táº¯c khÃ´ng tÃ¡ch hÃ ng.")
            self._normalize_tables(document, changes)
        if options.normalize_header_footer:
            self._normalize_headers_and_footers(document, changes)

        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        if (
            options.normalize_equations
            and self.profile.get("equations", {}).get("mode") == "force_times"
        ):
            emit("FORMULA_FONT", "running", "Äang chuyá»ƒn font cÃ´ng thá»©c ToÃ¡n an toÃ n.")
            changes["omml_runs_forced_to_times"] = self._force_omml_font(
                output,
                self.profile["equations"].get("text_font", "Times New Roman"),
                self.profile["body"]["size_pt"],
            )
        emit("FORMULA_VALUE", "running", "Äang Ä‘á»‘i chiáº¿u giÃ¡ trá»‹ cÃ´ng thá»©c trÆ°á»›c vÃ  sau.")
        emit("INTEGRITY", "running", "Äang Ä‘á»‘i chiáº¿u ná»™i dung vÃ  hÃ¬nh áº£nh.")
        after = inventory(output)
        self._validate_integrity(before, after)
        emit("GATE", "running", "Äang cháº¡y cá»•ng tuÃ¢n thá»§ cuá»‘i cÃ¹ng.")
        compliance = self._evaluate_format_compliance(
            document_path=output, before=before, after=after
        )
        status_by_code = {
            str(item.get("code")): str(item.get("status") or "FAIL").upper()
            for item in tuple(compliance.get("checks") or ())
        }
        progress_groups = {
            "PAGE": ("PAGE_SIZE", "PAGE_MARGINS", "PAGE_BORDER"),
            "FONT": ("BODY_FONT", "FONT_COLOR"),
            "SPACING": ("CHARACTER_SPACING", "LINE_SPACING"),
            "TABLE": ("TABLE_REPEAT_HEADER",),
            "ROW": ("TABLE_ROW_SPLIT",),
            "FORMULA_VALUE": ("FORMULA_VALUE_INTEGRITY",),
            "INTEGRITY": ("CONTENT_INTEGRITY", "MEDIA_INTEGRITY"),
        }
        for task_code, evidence_codes in progress_groups.items():
            statuses = [status_by_code.get(code, "FAIL") for code in evidence_codes]
            task_status = "pass" if all(value == "PASS" for value in statuses) else "blocked"
            emit(task_code, task_status, "ÄÃ£ kiá»ƒm chá»©ng báº±ng DOCX Ä‘áº§u ra.")
        formula_font_status = status_by_code.get("OLE_FORMULA_FONT", "PASS")
        emit("FORMULA_FONT", "review" if formula_font_status == "REVIEW_REQUIRED" else "pass", "ÄÃ£ kiá»ƒm tra kiá»ƒu cÃ´ng thá»©c hiá»‡n cÃ³.")
        final_status = str(compliance.get("status") or "BLOCKED").upper()
        emit("GATE", "pass" if final_status == "PASS" else ("review" if final_status == "REVIEW_REQUIRED" else "blocked"), "Káº¿t luáº­n cá»•ng tuÃ¢n thá»§: " + final_status)
        emit("RELEASE", "pass" if final_status == "PASS" else "blocked", "Quyá»n LÆ°u/Táº£i/Gá»™p Ä‘Æ°á»£c quyáº¿t Ä‘á»‹nh tá»« káº¿t luáº­n tuÃ¢n thá»§.")

        report = {
            "result": "completed" if compliance["status"] == "PASS" else "blocked",
            "source": str(source), "output": str(output), "source_sha256": source_hash,
            "source_preserved": _sha256(source) == source_hash,
            "profile_name": self.profile.get("profile_name", "lesson-plan-default"),
            "selected_options": asdict(options),
            "before": asdict(before), "after": asdict(after), "changes": dict(changes),
            "equations": {
                "mode": self.profile.get("equations", {}).get("mode", "safe"),
                "plain_text_font": self.profile.get("equations", {}).get("text_font", "Times New Roman"),
                "omml_preserved": before.omml_equations,
                "omml_runs_forced_to_times": changes.get("omml_runs_forced_to_times", 0),
                "legacy_or_embedded_requires_review": before.ole_objects,
            },
            "warnings": self._warnings(before),
            "compliance": compliance,
            "configuration_snapshot": dict(
                self.profile.get("_admin_configuration_snapshot", {}) or {}
            ),
        }
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        return report

    def _normalize_sections(self, document, changes):
        page = self.profile["page"]
        sizes = {"A3": (29.7, 42.0), "A4": (21.0, 29.7), "A5": (14.8, 21.0)}
        paper_size = str(page.get("paper_size", "A4")).upper()
        width_cm, height_cm = sizes.get(paper_size, sizes["A4"])
        landscape = str(page.get("orientation", "portrait")).lower() == "landscape"
        for section in document.sections:
            section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
            section_width_cm, section_height_cm = (height_cm, width_cm) if landscape else (width_cm, height_cm)
            section.page_width, section.page_height = Cm(section_width_cm), Cm(section_height_cm)
            section.left_margin = Cm(page["margin_left_cm"])
            section.right_margin = Cm(page["margin_right_cm"])
            section.top_margin = Cm(page["margin_top_cm"])
            section.bottom_margin = Cm(page["margin_bottom_cm"])
            self._set_page_border(section, page)
            changes["sections_normalized"] += 1

    @staticmethod
    def _set_page_border(section, page):
        sect_pr = section._sectPr
        existing = sect_pr.find(qn("w:pgBorders"))
        if existing is not None:
            sect_pr.remove(existing)
        if not page.get("border_enabled", False):
            return
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement("w:" + edge)
            border.set(qn("w:val"), str(page.get("border_style", "single")))
            border.set(qn("w:sz"), str(max(2, round(float(page.get("border_width_pt", 0.5)) * 8))))
            border.set(qn("w:space"), "24")
            border.set(qn("w:color"), str(page.get("border_color_rgb", "000000")))
            borders.append(border)
        sect_pr.append(borders)

    def _normalize_styles(self, document):
        body = self.profile["body"]
        style = document.styles["Normal"]
        style.font.name = body["font"]
        style.font.size = Pt(body["size_pt"])
        if body.get("color_rgb"):
            style.font.color.rgb = RGBColor.from_string(str(body["color_rgb"]))
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), body["font"])

    @staticmethod
    def _clear_header_or_footer(container):
        element = container._element
        for child in list(element):
            element.remove(child)
        element.append(OxmlElement("w:p"))

    def _normalize_headers_and_footers(
        self,
        document,
        changes,
    ):
        profile = self.profile.get(
            "header_footer",
            {},
        )

        # Preservation-first contract:
        # existing headers and footers belong to
        # the teacher's source document and must
        # remain untouched unless destructive
        # replacement is explicitly requested.
        if not profile.get(
            "remove_existing",
            False,
        ):
            changes[
                "headers_footers_preserved"
            ] += 1
            return

        body = self.profile["body"]

        for section in document.sections:
            headers = (
                section.header,
                section.first_page_header,
                section.even_page_header,
            )

            footers = (
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            )

            for header in headers:
                header.is_linked_to_previous = False
                self._clear_header_or_footer(
                    header
                )
                changes[
                    "headers_cleared"
                ] += 1

            for footer in footers:
                footer.is_linked_to_previous = False
                self._clear_header_or_footer(
                    footer
                )
                changes[
                    "footers_cleared"
                ] += 1


            if profile.get(
                "page_number",
                True,
            ):
                for footer in footers:
                    paragraph = (
                        footer.paragraphs[0]
                    )

                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                    )

                    run = paragraph.add_run()
                    run.font.name = body["font"]
                    run.font.size = Pt(
                        body["size_pt"]
                    )

                    fonts = (
                        run._element
                        .get_or_add_rPr()
                        .rFonts
                    )

                    for key in (
                        "ascii",
                        "hAnsi",
                        "eastAsia",
                        "cs",
                    ):
                        fonts.set(
                            qn(f"w:{key}"),
                            body["font"],
                        )

                    field_begin = OxmlElement(
                        "w:fldChar"
                    )
                    field_begin.set(
                        qn("w:fldCharType"),
                        "begin",
                    )

                    instruction = OxmlElement(
                        "w:instrText"
                    )
                    instruction.set(
                        qn("xml:space"),
                        "preserve",
                    )
                    instruction.text = " PAGE "

                    field_end = OxmlElement(
                        "w:fldChar"
                    )
                    field_end.set(
                        qn("w:fldCharType"),
                        "end",
                    )

                    run._r.extend(
                        (
                            field_begin,
                            instruction,
                            field_end,
                        )
                    )

                    changes[
                        "automatic_page_numbers_added"
                    ] += 1

        settings = document.settings._element

        update_fields = settings.find(
            qn("w:updateFields")
        )

        if update_fields is None:
            update_fields = OxmlElement(
                "w:updateFields"
            )
            settings.append(update_fields)

        update_fields.set(
            qn("w:val"),
            "true",
        )


    def _paragraph_kind(self, text: str) -> str:
        stripped = " ".join(text.split())
        if re.match(r"^(BÃ€I|CHá»¦ Äá»€)\b", stripped, re.I): return "title"
        if re.match(r"^[IVX]+\.\s+", stripped): return "heading_1"
        if re.match(r"^Hoáº¡t Ä‘á»™ng\s+\d+", stripped, re.I): return "activity"
        if re.match(r"^\d+\.\s+", stripped): return "heading_2"
        return "body"

    def _all_paragraphs(self, document):
        yield from document.paragraphs
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    def _normalize_paragraphs(
        self,
        document,
        changes,
        *,
        normalize_font=True,
        normalize_spacing=True,
    ):
        body = self.profile["body"]
        for paragraph in self._all_paragraphs(document):
            kind = self._paragraph_kind(paragraph.text)
            in_table = paragraph._p.getparent().tag == qn("w:tc")
            size = self.profile["table"]["size_pt"] if in_table else body["size_pt"]
            if normalize_font:
                for run in paragraph.runs:
                    run.font.name = body["font"]
                    run.font.size = Pt(size)
                    if body.get("color_rgb"):
                        run.font.color.rgb = RGBColor.from_string(str(body["color_rgb"]))
                        color = run._element.get_or_add_rPr().find(qn("w:color"))
                        if color is not None:
                            for attribute in ("themeColor", "themeTint", "themeShade"):
                                color.attrib.pop(qn("w:" + attribute), None)
                    fonts = run._element.get_or_add_rPr().rFonts
                    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                        fonts.set(qn(f"w:{key}"), body["font"])
                    spacing = run._element.get_or_add_rPr().find(qn("w:spacing"))
                    if spacing is None:
                        spacing = OxmlElement("w:spacing")
                        run._element.get_or_add_rPr().append(spacing)
                    spacing.set(qn("w:val"), str(round(float(body.get("character_spacing_pt", 0.0)) * 20)))
            fmt = paragraph.paragraph_format
            if normalize_spacing:
                fmt.line_spacing = body["line_spacing"]
                if not in_table:
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    if kind == "body" and paragraph.text.strip():
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    elif kind == "title":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if normalize_font and kind in {
                "title",
                "heading_1",
                "heading_2",
                "activity",
            }:
                for run in paragraph.runs:
                    run.bold = True
                    if kind == "title":
                        run.font.size = Pt(
                            self.profile["title"]["size_pt"]
                        )
            changes[f"paragraph_{kind}"] += 1

    @staticmethod
    def _set_table_layout(table, mode: str) -> None:
        table_properties = table._tbl.tblPr
        layout = table_properties.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            table_properties.append(layout)
        layout.set(qn("w:type"), mode)

    @staticmethod
    def _set_table_preferred_width_percent(
        table,
        value: int = 5000,
    ) -> None:
        table_properties = table._tbl.tblPr
        table_width = table_properties.first_child_found_in("w:tblW")
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            table_properties.append(table_width)
        table_width.set(qn("w:w"), str(value))
        table_width.set(qn("w:type"), "pct")

    @staticmethod
    def _table_unique_cells(row):
        result = []
        seen = set()
        for cell in row.cells:
            marker = id(cell._tc)
            if marker in seen:
                continue
            seen.add(marker)
            result.append(cell)
        return result

    @staticmethod
    def _cell_content_demand(cell) -> float:
        text = "\n".join(
            paragraph.text
            for paragraph in cell.paragraphs
        )
        compact = " ".join(text.split())
        paragraph_count = max(
            1,
            sum(
                1
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            ),
        )
        long_token_penalty = sum(
            max(0, len(token) - 14)
            for token in compact.split()
        )
        drawing_count = len(
            cell._tc.xpath(
                './/*[local-name()="drawing" or local-name()="pict"]'
            )
        )
        equation_count = len(
            cell._tc.xpath(
                './/*[local-name()="oMath" or local-name()="oMathPara"]'
            )
        )
        return (
            len(compact)
            + (paragraph_count * 12)
            + (long_token_penalty * 1.5)
            + (drawing_count * 120)
            + (equation_count * 80)
        )

    def _two_column_balance_ratio(self, table):
        if len(table.columns) != 2 or len(table.rows) < 3:
            return None

        left_demand = 0.0
        right_demand = 0.0
        usable_rows = 0

        for row_index, row in enumerate(table.rows):
            unique_cells = self._table_unique_cells(row)
            if len(unique_cells) != 2:
                continue
            if row_index == 0:
                continue

            left_demand += self._cell_content_demand(unique_cells[0])
            right_demand += self._cell_content_demand(unique_cells[1])
            usable_rows += 1

        if usable_rows < 2:
            return None

        total = left_demand + right_demand
        if total <= 0:
            return 0.50

        candidates = [
            value / 100.0
            for value in range(36, 65, 2)
        ]
        best_ratio = 0.50
        best_score = None

        for left_ratio in candidates:
            right_ratio = 1.0 - left_ratio
            wrapped_cost = (
                left_demand / left_ratio
                + right_demand / right_ratio
            )
            balance_penalty = (
                total
                * abs(left_ratio - 0.50)
                * 0.08
            )
            score = wrapped_cost + balance_penalty

            if best_score is None or score < best_score:
                best_score = score
                best_ratio = left_ratio

        return best_ratio

    @staticmethod
    def _apply_fixed_widths(
        table,
        widths,
    ) -> None:
        grid_columns = list(
            table._tbl.tblGrid.gridCol_lst
        )

        for column, width in zip(
            grid_columns,
            widths,
        ):
            column.set(
                qn("w:w"),
                str(width),
            )

        table.autofit = False
        LessonPlanWordStandardizer._set_table_layout(
            table,
            "fixed",
        )

        table_width = (
            table._tbl.tblPr
            .first_child_found_in("w:tblW")
        )

        if table_width is not None and widths:
            table_width.set(
                qn("w:w"),
                str(sum(widths)),
            )
            table_width.set(
                qn("w:type"),
                "dxa",
            )

    @staticmethod
    def _apply_content_autofit(table) -> None:
        LessonPlanWordStandardizer._set_table_preferred_width_percent(
            table,
            5000,
        )
        LessonPlanWordStandardizer._set_table_layout(
            table,
            "autofit",
        )
        table.autofit = True

        seen_cells = set()
        for row in table.rows:
            for cell in row.cells:
                marker = id(cell._tc)
                if marker in seen_cells:
                    continue
                seen_cells.add(marker)

                cell_width = (
                    cell._tc
                    .get_or_add_tcPr()
                    .get_or_add_tcW()
                )

                cell_width.set(
                    qn("w:w"),
                    "0",
                )
                cell_width.set(
                    qn("w:type"),
                    "auto",
                )

    def _normalize_global_text_color(
        self,
        document,
        changes,
    ):
        seen_cells = set()

        def normalize_paragraph(paragraph):
            p_pr = paragraph._p.get_or_add_pPr()
            paragraph_shading = p_pr.find(qn("w:shd"))
            if paragraph_shading is not None:
                p_pr.remove(paragraph_shading)
                changes["paragraph_shading_removed"] += 1

            for run in paragraph.runs:
                properties = (
                    run._element
                    .get_or_add_rPr()
                )

                color = properties.find(
                    qn("w:color")
                )
                if color is None:
                    color = OxmlElement(
                        "w:color"
                    )
                    properties.append(color)

                old_color = (
                    color.get(qn("w:val"))
                    or ""
                ).upper()

                if old_color not in {
                    "",
                    "000000",
                    "AUTO",
                }:
                    changes[
                        "text_colors_normalized"
                    ] += 1

                color.set(
                    qn("w:val"),
                    "000000",
                )

                for attr in (
                    "themeColor",
                    "themeTint",
                    "themeShade",
                ):
                    color.attrib.pop(
                        qn(f"w:{attr}"),
                        None,
                    )

                highlight = properties.find(
                    qn("w:highlight")
                )
                if highlight is not None:
                    properties.remove(
                        highlight
                    )
                    changes[
                        "text_highlights_removed"
                    ] += 1

                run_shading = properties.find(
                    qn("w:shd")
                )
                if run_shading is not None:
                    properties.remove(
                        run_shading
                    )
                    changes[
                        "run_shading_removed"
                    ] += 1

        def normalize_table(table):
            for row in table.rows:
                for cell in row.cells:
                    marker = id(cell._tc)
                    if marker in seen_cells:
                        continue
                    seen_cells.add(marker)

                    tc_pr = cell._tc.get_or_add_tcPr()
                    cell_shading = tc_pr.find(qn("w:shd"))
                    if cell_shading is not None:
                        tc_pr.remove(cell_shading)
                        changes["cell_shading_removed"] += 1

                    for paragraph in cell.paragraphs:
                        normalize_paragraph(
                            paragraph
                        )

                    for nested in cell.tables:
                        normalize_table(nested)

        for paragraph in document.paragraphs:
            normalize_paragraph(paragraph)

        for table in document.tables:
            normalize_table(table)

        for section in document.sections:
            containers = (
                section.header,
                section.first_page_header,
                section.even_page_header,
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            )

            for container in containers:
                for paragraph in container.paragraphs:
                    normalize_paragraph(
                        paragraph
                    )

                for table in container.tables:
                    normalize_table(table)

    @staticmethod
    def _normalize_cell_borders(cell, table_profile, changes) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        existing = tc_pr.find(qn("w:tcBorders"))
        if existing is not None:
            tc_pr.remove(existing)
            changes["cell_border_overrides_removed"] += 1

        borders = OxmlElement("w:tcBorders")
        size = str(
            max(
                2,
                round(
                    float(table_profile.get("border_width_pt", 0.5))
                    * 8
                ),
            )
        )
        color = str(table_profile.get("border_color_rgb", "000000"))
        style = str(table_profile.get("border_style", "single"))
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement("w:" + edge)
            border.set(qn("w:val"), style)
            border.set(qn("w:sz"), size)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color)
            borders.append(border)
        tc_pr.append(borders)

    @staticmethod
    def _fit_cell_drawings(cell, width_dxa: int, changes) -> None:
        if width_dxa <= 0:
            return

        available_emu = max(635, (width_dxa - 720) * 635)
        drawings = cell._tc.xpath(
            './/*[local-name()="drawing"]'
        )
        for drawing in drawings:
            # Remove inherited picture cropping before calculating the safe
            # in-cell size. This covers both inline and anchored DrawingML.
            for source_rect in drawing.xpath(
                './/*[local-name()="srcRect"]'
            ):
                for edge in ("l", "t", "r", "b"):
                    if source_rect.get(edge) is not None:
                        source_rect.set(edge, "0")

            for offset in drawing.xpath(
                './/*[local-name()="positionH" or '
                'local-name()="positionV"]/'
                '*[local-name()="posOffset"]'
            ):
                offset.text = "0"

            word_extents = drawing.xpath(
                './/*[local-name()="extent" and '
                'namespace-uri()=' 
                '"http://schemas.openxmlformats.org/'
                'drawingml/2006/wordprocessingDrawing"]'
            )
            if not word_extents:
                continue
            extent = word_extents[0]
            try:
                current_cx = int(extent.get("cx") or 0)
                current_cy = int(extent.get("cy") or 0)
            except (TypeError, ValueError):
                continue
            if current_cx <= available_emu or current_cx <= 0:
                continue

            ratio = available_emu / current_cx
            target_cx = int(round(current_cx * ratio))
            target_cy = max(1, int(round(current_cy * ratio)))
            extent.set("cx", str(target_cx))
            extent.set("cy", str(target_cy))

            for drawing_extent in drawing.xpath(
                './/*[local-name()="xfrm"]/'
                '*[local-name()="ext"]'
            ):
                drawing_extent.set("cx", str(target_cx))
                drawing_extent.set("cy", str(target_cy))

            changes["cell_images_scaled_to_fit"] += 1

    def _normalize_tables(
        self,
        document,
        changes,
    ):
        table_profile = self.profile["table"]
        section = document.sections[0]

        usable_dxa = round(
            (
                section.page_width
                - section.left_margin
                - section.right_margin
            )
            / 635
        )

        for table in document.tables:
            self._set_table_borders(table, table_profile)
            grid_columns = list(
                table._tbl.tblGrid.gridCol_lst
            )

            original_widths = [
                int(
                    column.get(qn("w:w"))
                    or 0
                )
                for column in grid_columns
            ]

            original_total = sum(
                original_widths
            )

            balance_ratio = (
                self._two_column_balance_ratio(
                    table
                )
            )

            if (
                balance_ratio is not None
                and len(original_widths) == 2
                and usable_dxa > 0
            ):
                left_width = round(
                    usable_dxa
                    * balance_ratio
                )

                widths = [
                    left_width,
                    usable_dxa - left_width,
                ]

                self._apply_fixed_widths(
                    table,
                    widths,
                )

                changes[
                    "two_column_tables_balanced"
                ] += 1

                changes[
                    "two_column_balance_ratio_basis_points"
                ] += round(
                    balance_ratio
                    * 10000
                )

            elif (
                original_total > usable_dxa
                and original_total
            ):
                self._apply_content_autofit(
                    table
                )
                widths = []

                changes[
                    "tables_content_autofit"
                ] += 1

            else:
                widths = original_widths

                self._apply_fixed_widths(
                    table,
                    widths,
                )

            for row_index, row in enumerate(
                table.rows
            ):
                row_pr = (
                    row._tr.get_or_add_trPr()
                )

                existing_no_split = row_pr.find(
                    qn("w:cantSplit")
                )

                if existing_no_split is not None:
                    row_pr.remove(
                        existing_no_split
                    )

                if (
                    row_index == 0
                    and len(table.rows) > 1
                ):
                    row_pr.append(
                        OxmlElement(
                            "w:cantSplit"
                        )
                    )

                    header = OxmlElement(
                        "w:tblHeader"
                    )
                    header.set(
                        qn("w:val"),
                        "true",
                    )
                    row_pr.append(header)

                unique_cells = (
                    self._table_unique_cells(
                        row
                    )
                )

                for cell_index, cell in enumerate(
                    unique_cells
                ):
                    cell.vertical_alignment = (
                        WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    )

                    effective_width = 0
                    if (
                        widths
                        and cell_index < len(widths)
                    ):
                        effective_width = widths[cell_index]
                        cell_width = (
                            cell._tc
                            .get_or_add_tcPr()
                            .get_or_add_tcW()
                        )

                        cell_width.set(
                            qn("w:w"),
                            str(effective_width),
                        )

                        cell_width.set(
                            qn("w:type"),
                            "dxa",
                        )
                    elif unique_cells and usable_dxa > 0:
                        effective_width = round(
                            usable_dxa / len(unique_cells)
                        )

                    self._normalize_cell_borders(
                        cell,
                        table_profile,
                        changes,
                    )
                    self._fit_cell_drawings(
                        cell,
                        effective_width,
                        changes,
                    )

                    if (
                        row_index == 0
                        and len(table.rows) > 1
                    ):
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = (
                                WD_ALIGN_PARAGRAPH.CENTER
                            )

                            paragraph.paragraph_format.keep_with_next = True

                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(
                                    table_profile[
                                        "size_pt"
                                    ]
                                )

            changes[
                "tables_normalized"
            ] += 1

    @staticmethod
    def _set_table_borders(table, profile):
        tbl_pr = table._tbl.tblPr
        existing = tbl_pr.find(qn("w:tblBorders"))
        if existing is not None:
            tbl_pr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement("w:" + edge)
            border.set(qn("w:val"), str(profile.get("border_style", "single")))
            border.set(qn("w:sz"), str(max(2, round(float(profile.get("border_width_pt", 0.5)) * 8))))
            border.set(qn("w:color"), str(profile.get("border_color_rgb", "000000")))
            borders.append(border)
        tbl_pr.append(borders)

    @staticmethod
    def _force_omml_font(path: Path, font: str, size_pt: int) -> int:
        """Set editable OMML runs to normal text font while preserving math structures."""
        math_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        word_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        forced = 0
        handle, temporary_name = tempfile.mkstemp(suffix=".docx", dir=path.parent)
        os.close(handle)
        try:
            with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
                temporary_name, "w", zipfile.ZIP_DEFLATED
            ) as target:
                for item in source.infolist():
                    data = source.read(item.filename)
                    if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                        root = etree.fromstring(data)
                        if item.filename == "word/settings.xml":
                            for node in root.xpath(".//m:mathPr/m:mathFont", namespaces={"m": math_ns}):
                                node.set(f"{{{math_ns}}}val", font)
                        for math_run in root.xpath(".//m:r", namespaces={"m": math_ns}):
                            math_properties = math_run.find(f"{{{math_ns}}}rPr")
                            if math_properties is None:
                                math_properties = etree.Element(f"{{{math_ns}}}rPr")
                                math_run.insert(0, math_properties)
                            if math_properties.find(f"{{{math_ns}}}nor") is None:
                                math_properties.append(etree.Element(f"{{{math_ns}}}nor"))
                            word_properties = math_run.find(f"{{{word_ns}}}rPr")
                            if word_properties is None:
                                word_properties = etree.Element(f"{{{word_ns}}}rPr")
                                math_run.insert(1, word_properties)
                            fonts = word_properties.find(f"{{{word_ns}}}rFonts")
                            if fonts is None:
                                fonts = etree.Element(f"{{{word_ns}}}rFonts")
                                word_properties.insert(0, fonts)
                            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                                fonts.set(f"{{{word_ns}}}{key}", font)
                            for tag in ("sz", "szCs"):
                                size = word_properties.find(f"{{{word_ns}}}{tag}")
                                if size is None:
                                    size = etree.Element(f"{{{word_ns}}}{tag}")
                                    word_properties.append(size)
                                size.set(f"{{{word_ns}}}val", str(round(size_pt * 2)))
                            forced += 1
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                    target.writestr(item, data)
            os.replace(temporary_name, path)
        finally:
            if os.path.exists(temporary_name):
                os.unlink(temporary_name)
        return forced

    def _validate_integrity(
        self,
        before,
        after,
    ):
        count_fields = (
            "paragraphs",
            "tables",
            "table_cells",
            "inline_shapes",
            "sections",
            "omml_equations",
            "equation_paragraphs",
            "ole_objects",
            "drawings",
            "images",
        )

        for name in count_fields:
            if (
                getattr(before, name)
                != getattr(after, name)
            ):
                raise ValueError(
                    "Document integrity failure: "
                    + name
                    + " count changed."
                )

        if before.body_text_sha256 != after.body_text_sha256:
            raise ValueError("Document integrity failure: body or table text changed.")
        if before.formula_value_sha256 != after.formula_value_sha256:
            raise ValueError("Document integrity failure: mathematical formula value changed.")
        for name in (
            "media_fingerprints", "embedded_fingerprints", "external_relationships"
        ):
            if getattr(before, name) != getattr(after, name):
                raise ValueError("Document integrity failure: " + name + " changed.")
        remove_existing = bool(
            self.profile.get("header_footer", {}).get("remove_existing", False)
        )
        if not remove_existing and before.header_footer_text_sha256 != after.header_footer_text_sha256:
            raise ValueError("Document integrity failure: header/footer content changed.")

    def _evaluate_format_compliance(self, *, document_path: Path, before, after):
        document = Document(document_path)
        page = self.profile.get("page", {})
        body = self.profile.get("body", {})
        table_profile = self.profile.get("table", {})
        snapshot = dict(self.profile.get("_admin_configuration_snapshot", {}) or {})
        checks = []

        def add(code, passed, expected, actual, status=None):
            checks.append({
                "code": code,
                "status": status or ("PASS" if passed else "FAIL"),
                "expected": expected,
                "actual": actual,
            })

        snapshot_ok = bool(snapshot.get("configuration_hash")) and bool(
            snapshot.get("global_version_id")
        )
        add("ACTIVE_CONFIGURATION_SNAPSHOT", snapshot_ok, "immutable ACTIVE snapshot", snapshot or None)

        sizes = {"A3": (29.7, 42.0), "A4": (21.0, 29.7), "A5": (14.8, 21.0)}
        expected_size = sizes.get(str(page.get("paper_size", "A4")).upper(), sizes["A4"])
        actual_sizes = [(round(s.page_width.cm, 1), round(s.page_height.cm, 1)) for s in document.sections]
        add("PAGE_SIZE", all(size == expected_size for size in actual_sizes), expected_size, actual_sizes)
        expected_margins = tuple(float(page.get(k, 0)) for k in (
            "margin_left_cm", "margin_right_cm", "margin_top_cm", "margin_bottom_cm"
        ))
        actual_margins = [tuple(round(getattr(s, name).cm, 2) for name in (
            "left_margin", "right_margin", "top_margin", "bottom_margin"
        )) for s in document.sections]
        add("PAGE_MARGINS", all(all(abs(a-b) < .03 for a,b in zip(row, expected_margins)) for row in actual_margins), expected_margins, actual_margins)

        expected_font = str(body.get("font") or "")
        expected_color = str(body.get("color_rgb") or "").upper()
        expected_spacing = str(round(float(body.get("character_spacing_pt", 0.0)) * 20))
        bad_fonts, bad_colors, bad_spacing = [], [], []
        for paragraph in self._all_paragraphs(document):
            for run in paragraph.runs:
                if not run.text:
                    continue
                if expected_font and run.font.name != expected_font:
                    bad_fonts.append(run.font.name or "INHERITED")
                if expected_color:
                    color = run._element.get_or_add_rPr().find(qn("w:color"))
                    value = color.get(qn("w:val")) if color is not None else None
                    theme = color.get(qn("w:themeColor")) if color is not None else None
                    if str(value or "").upper() != expected_color or theme:
                        bad_colors.append({"value": value, "theme": theme})
                spacing = run._element.get_or_add_rPr().find(qn("w:spacing"))
                value = spacing.get(qn("w:val")) if spacing is not None else None
                if value != expected_spacing:
                    bad_spacing.append(value)
        add("BODY_FONT", not bad_fonts, expected_font, bad_fonts[:20])
        add("FONT_COLOR", not bad_colors, expected_color or "not enforced", bad_colors[:20])
        add("CHARACTER_SPACING", not bad_spacing, expected_spacing, bad_spacing[:20])

        expected_line = float(body.get("line_spacing", 1.0))
        bad_lines = []
        for paragraph in self._all_paragraphs(document):
            if paragraph.text and paragraph.paragraph_format.line_spacing != expected_line:
                bad_lines.append(str(paragraph.paragraph_format.line_spacing))
        add("LINE_SPACING", not bad_lines, expected_line, bad_lines[:20])

        border_expected = bool(page.get("border_enabled", False))
        border_actual = [s._sectPr.find(qn("w:pgBorders")) is not None for s in document.sections]
        add("PAGE_BORDER", all(value == border_expected for value in border_actual), border_expected, border_actual)

        repeat_expected = bool(table_profile.get("repeat_header", True))
        split_allowed = bool(table_profile.get("allow_row_split", True))
        missing_headers, splittable_rows = [], []
        for table_index, table in enumerate(document.tables):
            if repeat_expected and len(table.rows) > 1:
                if table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is None:
                    missing_headers.append(table_index)
            if not split_allowed:
                for row_index, row in enumerate(table.rows):
                    if row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is None:
                        splittable_rows.append((table_index, row_index))
        add("TABLE_REPEAT_HEADER", not missing_headers, repeat_expected, missing_headers)
        add("TABLE_ROW_SPLIT", not splittable_rows, split_allowed, splittable_rows[:30])

        add("CONTENT_INTEGRITY", before.body_text_sha256 == after.body_text_sha256, before.body_text_sha256, after.body_text_sha256)
        add("MEDIA_INTEGRITY", before.media_fingerprints == after.media_fingerprints, len(before.media_fingerprints), len(after.media_fingerprints))
        add("FORMULA_VALUE_INTEGRITY", before.formula_value_sha256 == after.formula_value_sha256, before.formula_value_sha256, after.formula_value_sha256)
        if before.ole_objects:
            add("OLE_FORMULA_FONT", False, "verified Times New Roman", f"{before.ole_objects} embedded OLE object(s)", status="REVIEW_REQUIRED")

        statuses = {item["status"] for item in checks}
        status = "BLOCKED" if "FAIL" in statuses else ("REVIEW_REQUIRED" if "REVIEW_REQUIRED" in statuses else "PASS")
        return {"status": status, "checks": checks, "blocking": status != "PASS"}

    @staticmethod
    def _warnings(before):
        warnings = []
        if before.omml_equations:
            warnings.append(
                f"Giá»¯ nguyÃªn cáº¥u trÃºc cá»§a {before.omml_equations} cÃ´ng thá»©c Word; cáº§n kiá»ƒm tra trá»±c quan font vÃ  kÃ½ hiá»‡u."
            )
        if before.ole_objects:
            warnings.append(f"CÃ³ {before.ole_objects} Ä‘á»‘i tÆ°á»£ng OLE/MathType/Equation cÅ© cáº§n kiá»ƒm tra thá»§ cÃ´ng.")
        return warnings
