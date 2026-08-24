"""Conservatively apply AI text revisions to an original DOCX copy."""

from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class AiRevisionOverlayResult:
    changed_paragraphs: int = 0
    inserted_paragraphs: int = 0
    preserved_ambiguous_blocks: int = 0
    warnings: tuple[str, ...] = ()


class LessonPlanAiRevisionOverlay:
    """Apply only safely mappable AI paragraph changes to the source copy."""

    @staticmethod
    def _editable_lines(value: str) -> list[str]:
        return [
            line.strip()
            for line in str(value or "").splitlines()
            if line.strip() and " | " not in line
        ]

    @staticmethod
    def _source_paragraphs(document) -> list[Paragraph]:
        return [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

    @staticmethod
    def _replace_text(paragraph: Paragraph, value: str) -> None:
        runs = list(paragraph.runs)
        if runs:
            runs[0].text = value
            for run in runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(value)

    @staticmethod
    def _insert_after(paragraph: Paragraph, value: str) -> Paragraph:
        element = OxmlElement("w:p")
        paragraph._p.addnext(element)
        inserted = Paragraph(element, paragraph._parent)
        if paragraph._p.pPr is not None:
            inserted._p.insert(0, deepcopy(paragraph._p.pPr))
        inserted.add_run(value)
        return inserted

    def apply(
        self,
        *,
        source: Path,
        output: Path,
        revised_text: str,
    ) -> AiRevisionOverlayResult:
        document = Document(source)
        paragraphs = self._source_paragraphs(document)
        original = [paragraph.text.strip() for paragraph in paragraphs]
        revised = self._editable_lines(revised_text)

        if not revised or revised == original:
            document.save(output)
            return AiRevisionOverlayResult()

        matcher = SequenceMatcher(a=original, b=revised, autojunk=False)
        changed = 0
        inserted_count = 0
        ambiguous = 0

        for tag, i1, i2, j1, j2 in reversed(matcher.get_opcodes()):
            if tag == "equal":
                continue
            old_count = i2 - i1
            new_count = j2 - j1
            if tag == "replace" and old_count == new_count:
                for offset in range(old_count):
                    self._replace_text(
                        paragraphs[i1 + offset],
                        revised[j1 + offset],
                    )
                    changed += 1
                continue
            if tag == "insert" and i1 > 0:
                anchor = paragraphs[i1 - 1]
                for value in revised[j1:j2]:
                    anchor = self._insert_after(anchor, value)
                    inserted_count += 1
                continue
            ambiguous += max(old_count, new_count, 1)

        warnings = ()
        if ambiguous:
            warnings = (
                "Một số thay đổi AI không ánh xạ an toàn vào Word gốc; "
                "hệ thống đã giữ nguyên các phần đó để giáo viên kiểm tra.",
            )

        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        return AiRevisionOverlayResult(
            changed_paragraphs=changed,
            inserted_paragraphs=inserted_count,
            preserved_ambiguous_blocks=ambiguous,
            warnings=warnings,
        )
