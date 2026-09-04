from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document

from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)
from educational_planning_v2.models import TeachingSession

@dataclass(frozen=True)
class ContextApplicationResult:
    applied_fields: tuple[str, ...]
    unresolved_fields: tuple[str, ...]


# G1B_13H1R4B5S5_REHYDRATE_STALE_TEACHING_SESSION
def _g1b_rehydrate_teaching_session(session):
    if isinstance(session, TeachingSession):
        return session

    candidates = (
        getattr(session, "value", None),
        getattr(session, "name", None),
        session,
    )
    aliases = {
        "SANG": TeachingSession.MORNING,
        "BUOI_SANG": TeachingSession.MORNING,
        "MORNING": TeachingSession.MORNING,
        "CHIEU": TeachingSession.AFTERNOON,
        "BUOI_CHIEU": TeachingSession.AFTERNOON,
        "AFTERNOON": TeachingSession.AFTERNOON,
    }

    for candidate in candidates:
        if candidate is None:
            continue
        normalized = (
            str(candidate)
            .strip()
            .upper()
            .removeprefix("TEACHINGSESSION.")
        )
        if normalized in aliases:
            return aliases[normalized]
        try:
            return TeachingSession(normalized)
        except (TypeError, ValueError):
            continue

    raise TypeError("session must be TeachingSession")


# G1B_13H1R4B5S4_REHYDRATE_STALE_CONTEXT
def _g1b_rehydrate_scheduled_lesson_context(context):
    if isinstance(context, ScheduledLessonContext):
        return context

    actual_type = type(context)
    same_canonical_symbol = (
        actual_type.__module__ == ScheduledLessonContext.__module__
        and actual_type.__name__ == ScheduledLessonContext.__name__
    )
    if not same_canonical_symbol:
        raise TypeError(
            "context must be ScheduledLessonContext"
            + " [actual="
            + actual_type.__module__
            + "."
            + actual_type.__name__
            + ", actual_type_id="
            + str(id(actual_type))
            + ", expected="
            + ScheduledLessonContext.__module__
            + "."
            + ScheduledLessonContext.__name__
            + ", expected_type_id="
            + str(id(ScheduledLessonContext))
            + "]"
        )

    field_names = tuple(ScheduledLessonContext.__dataclass_fields__.keys())
    try:
        payload = {name: getattr(context, name) for name in field_names}
        payload["session"] = _g1b_rehydrate_teaching_session(
            payload["session"]
        )
        return ScheduledLessonContext(**payload)
    except Exception as error:
        raise TypeError(
            "context must be ScheduledLessonContext"
            + " [reload-safe rehydrate failed: "
            + type(error).__name__
            + ": "
            + str(error)
            + "]"
        ) from error


class LessonPlanDocumentContextApplier:
    """
    Apply scheduled teaching metadata to an existing
    lesson-plan DOCX without owning formatting rules.
    """

    _FIELD_PATTERNS = {
        "drafting_date": (
            r"^\s*ngày\s+soạn\s*:",
        ),
        "teaching_date": (
            r"^\s*ngày\s+(?:dạy|giảng)\s*:",
        ),
        "class_id": (
            r"^\s*lớp\s*:",
        ),
        "curriculum_period": (
            r"^\s*(?:tiết|ppct|tiết\s+ppct)\s*:",
        ),
        "lesson_title": (
            r"^\s*(?:tên\s+bài|bài)\s*:",
        ),
    }

    def apply(
            self,
            source: Path,
            output: Path,
            context: ScheduledLessonContext,
        ) -> ContextApplicationResult:
            """
            Compatibility entry point.

            Public API remains unchanged, but canonical
            metadata mutation is delegated to the
            preservation-first metadata overlay engine.
            """
            source = source.resolve()
            output = output.resolve()

            if source == output:
                raise ValueError(
                    "Kh\u00f4ng \u0111\u01b0\u1ee3c ghi \u0111\u00e8 t\u1ec7p Word g\u1ed1c."
                )

            if (
                source.suffix.lower() != ".docx"
                or output.suffix.lower() != ".docx"
            ):
                raise ValueError(
                    "Context applier ch\u1ec9 x\u1eed l\u00fd t\u1ec7p .docx."
                )

            # G1B_13H1R4B5S4_APPLY_RELOAD_SAFE_CONTEXT
            context = _g1b_rehydrate_scheduled_lesson_context(context)

            from document_standardization.lesson_plan_metadata import (
                LessonPlanMetadata,
            )
            from document_standardization.lesson_plan_metadata_overlay import (
                LessonPlanMetadataOverlay,
            )

            document = Document(source)

            metadata = LessonPlanMetadata(
                drafting_date=context.drafting_date,
                teaching_date=context.teaching_date,
                class_name=context.class_id,
                curriculum_period=(
                    context.curriculum_period
                ),
                lesson_title=context.lesson_title,
            )

            overlay_result = (
                LessonPlanMetadataOverlay()
                .apply(
                    document=document,
                    metadata=metadata,
                )
            )

            output.parent.mkdir(
                parents=True,
                exist_ok=True,
            )

            document.save(output)

            field_name_map = {
                "drafting_date": "drafting_date",
                "teaching_date": "teaching_date",
                "class_name": "class_id",
                "curriculum_period": (
                    "curriculum_period"
                ),
                "lesson_title": "lesson_title",
            }

            unresolved_metadata = {
                field.value
                for field
                in overlay_result.unresolved_fields
            }

            requested = (
                metadata.overlay_values()
            )

            applied = []

            unresolved = []

            for metadata_name in requested:
                public_name = field_name_map.get(
                    metadata_name,
                    metadata_name,
                )

                if (
                    metadata_name
                    in unresolved_metadata
                ):
                    unresolved.append(
                        public_name
                    )
                else:
                    applied.append(
                        public_name
                    )

            return ContextApplicationResult(
                applied_fields=tuple(applied),
                unresolved_fields=tuple(unresolved),
            )

    @staticmethod
    def _context_values(
        context: ScheduledLessonContext,
    ) -> dict[str, str]:
        drafting_date = (
            context.drafting_date.strftime(
                "%d/%m/%Y"
            )
            if context.drafting_date
            else ""
        )

        return {
            "drafting_date": drafting_date,
            "teaching_date": (
                context.teaching_date.strftime(
                    "%d/%m/%Y"
                )
            ),
            "class_id": context.class_id,
            "curriculum_period": str(
                context.curriculum_period
            ),
            "lesson_title": context.lesson_title,
        }

    def _apply_field(
        self,
        document,
        field_name: str,
        value: str,
        *,
        context: ScheduledLessonContext,
    ) -> bool:
        if not value:
            return False

        for paragraph in self._all_paragraphs(
            document
        ):
            original_text = paragraph.text

            updated_text = (
                self._replace_field_in_text(
                    original_text,
                    field_name=field_name,
                    value=value,
                    context=context,
                )
            )

            if updated_text is None:
                continue

            if updated_text != original_text:
                self._replace_paragraph_text(
                    paragraph,
                    updated_text,
                )

            return True

        if self._apply_table_label_value(
            document,
            field_name=field_name,
            value=value,
        ):
            return True

        return False

    @classmethod
    def _apply_table_label_value(
        cls,
        document,
        *,
        field_name: str,
        value: str,
    ) -> bool:
        """
        Support common lesson-plan tables where
        the metadata label and value occupy
        separate cells in the same row.

        Example:
            | Lớp       | 8A1        |
            | Tiết PPCT | 1          |
            | Tên bài   | Bài cũ     |
        """
        patterns = cls._FIELD_PATTERNS.get(
            field_name,
            (),
        )

        for table in document.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue

                label = row.cells[0].text.strip()

                if not cls._matches_table_label(
                    label,
                    patterns,
                ):
                    continue

                value_cell = row.cells[1]

                if value_cell.paragraphs:
                    paragraph = (
                        value_cell.paragraphs[0]
                    )
                    cls._replace_paragraph_text(
                        paragraph,
                        value,
                    )

                    for extra_paragraph in (
                        value_cell.paragraphs[1:]
                    ):
                        cls._replace_paragraph_text(
                            extra_paragraph,
                            "",
                        )
                else:
                    value_cell.text = value

                return True

        return False

    @staticmethod
    def _matches_table_label(
        label: str,
        patterns,
    ) -> bool:
        normalized = label.strip()

        for pattern in patterns:
            table_pattern = pattern

            if table_pattern.endswith(
                r"\s*:"
            ):
                table_pattern = (
                    table_pattern[:-4]
                )
            elif table_pattern.endswith(":"):
                table_pattern = (
                    table_pattern[:-1]
                )

            if re.match(
                table_pattern + r"\s*$",
                normalized,
                flags=re.IGNORECASE,
            ):
                return True

        return False

    @classmethod
    def _replace_field_in_text(
        cls,
        text: str,
        *,
        field_name: str,
        value: str,
        context: ScheduledLessonContext,
    ) -> str | None:
        if field_name == "drafting_date":
            return cls._replace_simple_value(
                text,
                pattern=(
                    r"(?P<prefix>"
                    r"(?:ng\u00e0y\s+so\u1ea1n|date\s+of\s+planning)"
                    r"\s*:?\s*)"
                    r"(?:\d{1,2}/\d{1,2}/\d{2,4}|[.\u2026\u00b7_\-\s]+)"
                ),
                value=value,
            )

        if field_name == "teaching_date":
            return cls._replace_simple_value(
                text,
                pattern=(
                    r"(?P<prefix>"
                    r"(?:ng\u00e0y\s+(?:d\u1ea1y|gi\u1ea3ng)|date\s+of\s+teaching)"
                    r"\s*:?\s*)"
                    r"(?:\d{1,2}/\d{1,2}/\d{2,4}|[.\u2026\u00b7_\-\s]+)"
                ),
                value=value,
            )

        if field_name == "class_id":
            return cls._replace_simple_value(
                text,
                pattern=(
                    r"(?P<prefix>"
                    r"l\u1edbp\s*:?\s*)"
                    r"[A-Za-z0-9._-]+"
                ),
                value=value,
            )

        if field_name == "curriculum_period":
            explicit = (
                cls._replace_simple_value(
                    text,
                    pattern=(
                        r"(?P<prefix>"
                        r"ti\u1ebft\s+ppct"
                        r"\s*:\s*)"
                        r"\d+"
                    ),
                    value=value,
                )
            )

            if explicit is not None:
                return explicit

            explicit = (
                cls._replace_simple_value(
                    text,
                    pattern=(
                        r"(?P<prefix>"
                        r"ppct\s*:\s*)"
                        r"\d+"
                    ),
                    value=value,
                )
            )

            if explicit is not None:
                return explicit

            explicit = (
                cls._replace_simple_value(
                    text,
                    pattern=(
                        r"(?P<prefix>"
                        r"^\s*ti\u1ebft"
                        r"\s*:\s*)"
                        r"\d+"
                    ),
                    value=value,
                )
            )

            if explicit is not None:
                return explicit

            return cls._replace_period_heading(
                text,
                curriculum_period=(
                    context.curriculum_period
                ),
                period_in_lesson=(
                    context.period_in_lesson
                ),
            )

        if field_name == "lesson_title":
            explicit = re.match(
                (
                    r"(?P<prefix>"
                    r"^\s*(?:"
                    r"t\u00ean\s+b\u00e0i"
                    r"|b\u00e0i"
                    r")\s*:\s*)"
                    r"(?P<old>.*)$"
                ),
                text,
                flags=re.IGNORECASE,
            )

            if explicit is not None:
                return (
                    explicit.group("prefix")
                    + value
                )

            return cls._replace_lesson_heading(
                text,
                lesson_title=value,
            )

        return None

    @staticmethod
    def _replace_simple_value(
        text: str,
        *,
        pattern: str,
        value: str,
    ) -> str | None:
        match = re.search(
            pattern,
            text,
            flags=re.IGNORECASE,
        )

        if match is None:
            return None

        return (
            text[:match.start()]
            + match.group("prefix")
            + value
            + text[match.end():]
        )

    @staticmethod
    def _replace_period_heading(
        text: str,
        *,
        curriculum_period: int,
        period_in_lesson: int,
    ) -> str | None:
        match = re.match(
            (
                r"(?P<prefix>"
                r"^\s*(?:ti\u1ebft|period)\s+)"
                r"(?P<periods>"
                r"\d+(?:\s*(?:,|\+)\s*\d+)*"
                r")"
                r"(?P<separator>"
                r"\s*[.\-:]?)"
            ),
            text,
            flags=re.IGNORECASE,
        )

        if match is None:
            return None

        existing_periods = re.findall(
            r"\d+",
            match.group("periods"),
        )

        count = max(
            1,
            len(existing_periods),
        )

        start_period = (
            curriculum_period
            - period_in_lesson
            + 1
        )

        if start_period <= 0:
            start_period = (
                curriculum_period
            )

        original_periods = (
            match.group("periods")
        )

        period_joiner = (
            " + "
            if "+" in original_periods
            else ","
        )

        replacement_periods = (
            period_joiner.join(
                str(
                    start_period + offset
                )
                for offset in range(count)
            )
        )

        return (
            text[:match.start()]
            + match.group("prefix")
            + replacement_periods
            + match.group("separator")
            + text[match.end():]
        )

    @staticmethod
    def _replace_lesson_heading(
        text: str,
        *,
        lesson_title: str,
    ) -> str | None:
        patterns = (
            # Standalone section heading:
            # ?7: Ten bai
            (
                r"(?P<prefix>"
                r"^\s*\u00a7\s*\d+"
                r"\s*:\s*)"
                r"(?P<title>.*?)"
                r"(?P<suffix>"
                r"\s*\(\s*\d+\s+"
                r"ti\u1ebft\s*\)\s*"
                r")?$"
            ),

            # Tiet 1. BAI 1. Ten bai
            (
                r"(?P<prefix>"
                r"^\s*ti\u1ebft\s+"
                r"\d+(?:\s*,\s*\d+)*"
                r"\s*[.\-:]\s*"
                r"b\u00e0i\s+\d+"
                r"\s*[.\-:]\s*)"
                r"(?P<title>.*?)"
                r"(?P<suffix>"
                r"\s*\(\s*\d+\s+"
                r"ti\u1ebft\s*\)\s*"
                r")?$"
            ),

            # TIET 4 - ?4: Ten bai
            (
                r"(?P<prefix>"
                r"^\s*ti\u1ebft\s+"
                r"\d+(?:\s*,\s*\d+)*"
                r"\s*[.\-:]\s*"
                r"\u00a7\s*\d+"
                r"\s*:\s*)"
                r"(?P<title>.*?)"
                r"(?P<suffix>"
                r"\s*\(\s*\d+\s+"
                r"ti\u1ebft\s*\)\s*"
                r")?$"
            ),
        )

        for pattern in patterns:
            match = re.match(
                pattern,
                text,
                flags=re.IGNORECASE,
            )

            if match is None:
                continue

            suffix = (
                match.group("suffix")
                or ""
            )

            return (
                match.group("prefix")
                + lesson_title
                + suffix
            )

        return None

    @staticmethod
    def _all_paragraphs(document):
        yield from document.paragraphs

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    @staticmethod
    def _replace_paragraph_text(
        paragraph,
        value: str,
    ) -> None:
        if paragraph.runs:
            paragraph.runs[0].text = value

            for run in paragraph.runs[1:]:
                run.text = ""

            return

        paragraph.add_run(value)


# MULTICLASS_TEACHING_DATE_OVERLAY_V1
#
# Compatibility layer for teacher lesson plans whose
# teaching-date field contains several class/date pairs:
#
#   Ng?y d?y:
#   7A1 - 02/10/2025
#   7A2 - 30/09/2025
#
# Only the date belonging to context.class_id is replaced.
# Other classes, document structure and unrelated content
# remain unchanged.

def _mt_replace_span_preserving_runs(
    paragraph,
    start: int,
    end: int,
    replacement: str,
) -> bool:
    runs = list(paragraph.runs)

    if not runs:
        return False

    positions = []
    cursor = 0

    for run in runs:
        run_start = cursor
        cursor += len(run.text)
        positions.append(
            (run, run_start, cursor)
        )

    affected = [
        item
        for item in positions
        if item[2] > start
        and item[1] < end
    ]

    if not affected:
        return False

    first_run, first_start, _ = affected[0]
    last_run, last_start, _ = affected[-1]

    first_local = max(
        0,
        start - first_start,
    )

    last_local = max(
        0,
        end - last_start,
    )

    prefix = first_run.text[
        :first_local
    ]

    suffix = last_run.text[
        last_local:
    ]

    if first_run is last_run:
        first_run.text = (
            prefix
            + replacement
            + suffix
        )
        return True

    first_run.text = (
        prefix
        + replacement
    )

    for run, _, _ in affected[1:-1]:
        run.text = ""

    last_run.text = suffix

    return True


def _mt_iter_all_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _mt_overlay_multiclass_teaching_date(
    output,
    context,
) -> bool:
    import re
    from docx import Document

    class_id = str(
        getattr(
            context,
            "class_id",
            "",
        )
        or ""
    ).strip()

    teaching_date = getattr(
        context,
        "teaching_date",
        None,
    )

    if (
        not class_id
        or teaching_date is None
    ):
        return False

    formatted_date = (
        teaching_date.strftime(
            "%d/%m/%Y"
        )
        if hasattr(
            teaching_date,
            "strftime",
        )
        else str(teaching_date)
    )

    document = Document(output)

    # Accept "-", en dash, em dash and ":".
    # Only replace the date attached to the selected class.
    pattern = re.compile(
        r"(?<![A-Za-z0-9])"
        + re.escape(class_id)
        + r"\s*[-\u2013\u2014:]\s*"
        + r"(?P<date>"
          r"\d{1,2}[/-]\d{1,2}[/-]\d{4}"
          r")",
        flags=re.IGNORECASE,
    )

    changed = False

    for paragraph in _mt_iter_all_paragraphs(
        document
    ):
        text = paragraph.text

        matches = list(
            pattern.finditer(text)
        )

        # Work backwards so earlier character
        # positions are not shifted.
        for match in reversed(matches):
            date_start = match.start(
                "date"
            )
            date_end = match.end(
                "date"
            )

            current_date = match.group(
                "date"
            )

            if current_date == formatted_date:
                continue

            if _mt_replace_span_preserving_runs(
                paragraph,
                date_start,
                date_end,
                formatted_date,
            ):
                changed = True

    if changed:
        document.save(output)

    return changed


_mt_original_context_apply = (
    LessonPlanDocumentContextApplier.apply
)


def _mt_context_apply_with_multiclass_date(
    self,
    source,
    output,
    context,
):
    result = _mt_original_context_apply(
        self,
        source,
        output,
        context,
    )

    _mt_overlay_multiclass_teaching_date(
        output,
        context,
    )

    return result


LessonPlanDocumentContextApplier.apply = (
    _mt_context_apply_with_multiclass_date
)


