from __future__ import annotations

from dataclasses import dataclass
from enum import Enum
import re
import unicodedata

from docx.document import Document as DocxDocument
from docx.table import Table, _Cell


class MetadataField(str, Enum):
    SCHOOL_NAME = "school_name"
    TEACHER_NAME = "teacher_name"
    SUBJECT_NAME = "subject_name"
    CLASS_NAME = "class_name"
    LESSON_TITLE = "lesson_title"
    CURRICULUM_PERIOD = "curriculum_period"
    DRAFTING_DATE = "drafting_date"
    TEACHING_DATE = "teaching_date"


class MetadataLocationKind(str, Enum):
    PARAGRAPH = "paragraph"
    TABLE_CELL = "table_cell"
    TABLE_PAIRED_CELL = "table_paired_cell"


class MetadataMatchStrategy(str, Enum):
    INLINE_EXPLICIT = "inline_explicit"
    INLINE_SPACE = "inline_space"
    PAIRED_CELL = "paired_cell"
    LESSON_HEADING = "lesson_heading"


@dataclass(
    frozen=True,
    slots=True,
)
class MetadataLocation:
    field: MetadataField
    kind: MetadataLocationKind
    text: str
    value_text: str
    strategy: MetadataMatchStrategy
    confidence: float

    # Label/source coordinates.
    paragraph_index: int | None = None

    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    cell_paragraph_index: int | None = None

    # Actual value target coordinates.
    #
    # For inline metadata these are normally the
    # same paragraph/cell as the label.
    #
    # For paired-cell metadata these point to the
    # neighbouring cell that contains the value.
    value_paragraph_index: int | None = None

    value_table_index: int | None = None
    value_row_index: int | None = None
    value_cell_index: int | None = None
    value_cell_paragraph_index: int | None = None

    @property
    def is_high_confidence(self) -> bool:
        return self.confidence >= 0.90


@dataclass(
    frozen=True,
    slots=True,
)
class _Candidate:
    field: MetadataField
    value_text: str
    strategy: MetadataMatchStrategy
    confidence: float
    label_only: bool = False


def _fold_text(
    value: str,
) -> str:
    """
    Accent-insensitive matching for Vietnamese.

    The original DOCX text is never modified.

    The source code intentionally uses ASCII-only
    labels so PowerShell/terminal encoding cannot
    corrupt the matching contract.
    """

    value = str(
        value
    ).strip().lower()

    value = unicodedata.normalize(
        "NFD",
        value,
    )

    value = "".join(
        character
        for character in value
        if unicodedata.category(
            character
        ) != "Mn"
    )

    # Vietnamese D with stroke does not decompose.
    value = value.replace(
        "\u0111",
        "d",
    )

    value = re.sub(
        r"\s+",
        " ",
        value,
    )

    return value


_LABELS = (
    (
        MetadataField.DRAFTING_DATE,
        "ngay soan",
    ),
    (
        MetadataField.TEACHING_DATE,
        "ngay day",
    ),
    (
        MetadataField.TEACHING_DATE,
        "ngay giang",
    ),
    (
        MetadataField.CURRICULUM_PERIOD,
        "tiet ppct",
    ),
    (
        MetadataField.CURRICULUM_PERIOD,
        "ppct",
    ),
    (
        MetadataField.CURRICULUM_PERIOD,
        "tiet",
    ),
    (
        MetadataField.CLASS_NAME,
        "lop",
    ),
    (
        MetadataField.SUBJECT_NAME,
        "mon",
    ),
    (
        MetadataField.TEACHER_NAME,
        "giao vien",
    ),
    (
        MetadataField.TEACHER_NAME,
        "gv",
    ),
    (
        MetadataField.SCHOOL_NAME,
        "truong",
    ),
    (
        MetadataField.LESSON_TITLE,
        "ten bai",
    ),
    (
        MetadataField.LESSON_TITLE,
        "bai",
    ),
)


_DATE_PATTERN = re.compile(
    r"^\d{1,2}\s*[/.-]\s*"
    r"\d{1,2}\s*[/.-]\s*"
    r"\d{2,4}\b"
)

_CLASS_PATTERN = re.compile(
    r"^[0-9]{1,2}"
    r"[A-Za-z]"
    r"[0-9A-Za-z]*$"
)

_PERIOD_PATTERN = re.compile(
    r"^\d+\b"
)


def _looks_like_value(
    field: MetadataField,
    value: str,
) -> bool:

    value = value.strip()

    if not value:
        return False

    folded = _fold_text(
        value
    )

    if field in {
        MetadataField.DRAFTING_DATE,
        MetadataField.TEACHING_DATE,
    }:
        return bool(
            _DATE_PATTERN.match(
                value
            )
        )

    if (
        field
        == MetadataField.CURRICULUM_PERIOD
    ):
        return bool(
            _PERIOD_PATTERN.match(
                value
            )
        )

    if field == MetadataField.CLASS_NAME:
        compact = re.sub(
            r"\s+",
            "",
            value,
        )

        return bool(
            _CLASS_PATTERN.match(
                compact
            )
        )

    if field == MetadataField.SUBJECT_NAME:
        # Avoid generic prose such as:
        # "Mon hoc hom nay..."
        return not folded.startswith(
            "hoc "
        )

    return True


def _extract_value_from_original(
    *,
    original: str,
    label: str,
) -> tuple[
    str,
    MetadataMatchStrategy,
]:

    stripped = original.strip()

    if ":" in stripped:
        return (
            stripped.split(
                ":",
                1,
            )[1].strip(),
            MetadataMatchStrategy
            .INLINE_EXPLICIT,
        )

    words = stripped.split()

    label_word_count = len(
        label.split()
    )

    value = " ".join(
        words[
            label_word_count:
        ]
    ).strip()

    if value.startswith("-"):
        value = value[
            1:
        ].strip()

        return (
            value,
            MetadataMatchStrategy
            .INLINE_EXPLICIT,
        )

    return (
        value,
        MetadataMatchStrategy
        .INLINE_SPACE,
    )


def _extract_candidate(
    text: str,
) -> _Candidate | None:

    stripped = str(
        text
    ).strip()

    if not stripped:
        return None

    folded = _fold_text(
        stripped
    )

    # -----------------------------------------------------
    # Special lesson heading recognition.
    #
    # Detect:
    #   Bai 7. ...
    #   Bai 12: ...
    #
    # Do NOT detect:
    #   Bai tap 1
    #   Bai hoc hom nay
    # -----------------------------------------------------

    lesson_heading = re.match(
        r"^bai\s+"
        r"(?:\d+|[ivxlcdm]+)"
        r"(?:\s*[.:\-]|\s+)"
        r".+",
        folded,
    )

    if lesson_heading is not None:
        return _Candidate(
            field=(
                MetadataField
                .LESSON_TITLE
            ),
            value_text=stripped,
            strategy=(
                MetadataMatchStrategy
                .LESSON_HEADING
            ),
            confidence=0.95,
        )

    # -----------------------------------------------------
    # Named metadata labels.
    # -----------------------------------------------------

    for field, label in _LABELS:

        if folded == label:
            return _Candidate(
                field=field,
                value_text="",
                strategy=(
                    MetadataMatchStrategy
                    .INLINE_EXPLICIT
                ),
                confidence=1.0,
                label_only=True,
            )

        if folded in {
            label + ":",
            label + " -",
            label + "-",
        }:
            return _Candidate(
                field=field,
                value_text="",
                strategy=(
                    MetadataMatchStrategy
                    .INLINE_EXPLICIT
                ),
                confidence=1.0,
                label_only=True,
            )

        if not folded.startswith(
            label
        ):
            continue

        remainder = folded[
            len(label):
        ]

        # Prevent matching:
        # lopHoc...
        # tietHoc...
        if (
            remainder
            and not (
                remainder[0].isspace()
                or remainder[0]
                in ":-"
            )
        ):
            continue

        value, strategy = (
            _extract_value_from_original(
                original=stripped,
                label=label,
            )
        )

        # The short lesson-title alias "bai"
        # is intentionally explicit-only.
        #
        # Accept:
        #   Bai: Phan so
        #   Bai - Phan so
        #
        # Reject:
        #   Bai Phan so
        #   Bai tap 1
        #   Bai hoc hom nay
        #
        # Numbered headings such as
        # "BAI 7. PHAN SO" are handled by
        # LESSON_HEADING above.
        if (
            field
            == MetadataField.LESSON_TITLE
            and label == "bai"
            and strategy
            == MetadataMatchStrategy.INLINE_SPACE
        ):
            continue

        if not value:
            return _Candidate(
                field=field,
                value_text="",
                strategy=strategy,
                confidence=1.0,
                label_only=True,
            )

        if not _looks_like_value(
            field,
            value,
        ):
            continue

        confidence = (
            1.0
            if strategy
            == MetadataMatchStrategy
            .INLINE_EXPLICIT
            else 0.92
        )

        return _Candidate(
            field=field,
            value_text=value,
            strategy=strategy,
            confidence=confidence,
        )

    return None




# =========================================================
# Multi-candidate paragraph parser
# =========================================================

_COMPOSITE_LABEL_PATTERN = re.compile(
    "(?i)"
    "(?<!\\w)"
    "("
    "ng\\u00e0y\\s+so\\u1ea1n"
    "|ng\\u00e0y\\s+d\\u1ea1y"
    "|ng\\u00e0y\\s+gi\\u1ea3ng"
    "|l\\u1edbp"
    "|ti\\u1ebft"
    ")"
    "\\s*(?=:|-|\\d)"
)



# G1B_P6B_ENGLISH_STRUCTURAL_METADATA_RECOGNITION
_ENGLISH_DATE_VALUE = (
    r"\d{1,2}\s*[/.-]\s*"
    r"\d{1,2}\s*[/.-]\s*"
    r"\d{2,4}"
)

def _extract_english_structural_candidates(original: str) -> tuple[_Candidate, ...]:
    stripped = str(original).strip()
    if not stripped:
        return ()

    match = re.match(
        rf"(?i)^(?:date\s+of\s+planning|date\s+of\s+preparation|preparation\s+date)"
        rf"\s*:\s*(?P<value>{_ENGLISH_DATE_VALUE})\s*$",
        stripped,
    )
    if match:
        return (_Candidate(
            field=MetadataField.DRAFTING_DATE,
            value_text=match.group("value"),
            strategy=MetadataMatchStrategy.INLINE_EXPLICIT,
            confidence=1.0,
        ),)

    match = re.match(
        rf"(?i)^(?:date\s+of\s+teaching|teaching\s+date)"
        rf"(?:\s+(?P<class>[0-9]{{1,2}}[A-Za-z][0-9A-Za-z]*))?"
        rf"\s*:\s*(?P<date>{_ENGLISH_DATE_VALUE})\s*$",
        stripped,
    )
    if match:
        result = [_Candidate(
            field=MetadataField.TEACHING_DATE,
            value_text=match.group("date"),
            strategy=MetadataMatchStrategy.INLINE_EXPLICIT,
            confidence=1.0,
        )]
        if match.group("class"):
            result.append(_Candidate(
                field=MetadataField.CLASS_NAME,
                value_text=match.group("class"),
                strategy=MetadataMatchStrategy.INLINE_EXPLICIT,
                confidence=1.0,
            ))
        return tuple(result)

    match = re.match(
        rf"(?is)^(?P<class>[0-9]{{1,2}}[A-Za-z][0-9A-Za-z]*)\s*:\s*"
        rf"(?P<date>{_ENGLISH_DATE_VALUE})\s*"
        rf"period\s+(?P<period>\d+)\s*"
        rf"(?P<title>\S.+?)\s*$",
        stripped,
    )
    if match:
        return (
            _Candidate(field=MetadataField.CLASS_NAME, value_text=match.group("class"),
                       strategy=MetadataMatchStrategy.INLINE_EXPLICIT, confidence=1.0),
            _Candidate(field=MetadataField.TEACHING_DATE, value_text=match.group("date"),
                       strategy=MetadataMatchStrategy.INLINE_EXPLICIT, confidence=1.0),
            _Candidate(field=MetadataField.CURRICULUM_PERIOD, value_text=match.group("period"),
                       strategy=MetadataMatchStrategy.INLINE_EXPLICIT, confidence=1.0),
            _Candidate(field=MetadataField.LESSON_TITLE, value_text=match.group("title"),
                       strategy=MetadataMatchStrategy.INLINE_EXPLICIT, confidence=1.0),
        )
    return ()

def _extract_candidates(
    original: str,
) -> tuple[_Candidate, ...]:
    """
    Extract zero, one or many metadata candidates
    from one paragraph or table-cell paragraph.

    This is a read-only parser.

    Compatibility contract:
        ordinary single-field content continues
        through _extract_candidate().

    Composite contract:
        one paragraph may produce multiple fields.
    """

    if not isinstance(
        original,
        str,
    ):
        return ()

    stripped = original.strip()

    if not stripped:
        return ()

    english_structural = _extract_english_structural_candidates(stripped)
    if english_structural:
        return english_structural

    matches = list(
        _COMPOSITE_LABEL_PATTERN.finditer(
            stripped
        )
    )

    folded = _fold_text(
        stripped
    )

    has_numbered_heading = bool(
        re.search(
            r"\bbai\s+"
            r"\d+\s*"
            r"[.:\-]?\s*"
            r".+$",
            folded,
        )
    )


    # -----------------------------------------------------
    # Locate the numbered lesson heading in ORIGINAL text.
    #
    # We cannot use indexes from folded text because
    # whitespace normalization may change offsets.
    # -----------------------------------------------------

    heading_start = None

    if has_numbered_heading:

        for position in range(
            len(stripped)
        ):
            suffix = stripped[
                position:
            ]

            suffix_folded = _fold_text(
                suffix
            )

            if re.match(
                r"^bai\s+"
                r"\d+\s*"
                r"[.:\-]?\s*"
                r".+$",
                suffix_folded,
            ):
                heading_start = position
                break


    # -----------------------------------------------------
    # Preserve the legacy path for an ordinary
    # single-field paragraph.
    #
    # Important:
    # do NOT take this shortcut when a numbered lesson
    # heading is also present, e.g.
    #
    #     Tiet 9. BAI 7. ...
    # -----------------------------------------------------

    if (
        len(matches) <= 1
        and heading_start is None
    ):
        candidate = _extract_candidate(
            stripped
        )

        if candidate is None:
            return ()

        return (
            candidate,
        )


    candidates: list[
        _Candidate
    ] = []


    # -----------------------------------------------------
    # Parse every ordinary metadata fragment.
    # -----------------------------------------------------

    for index, match in enumerate(
        matches
    ):
        start = match.start()

        if index + 1 < len(matches):
            end = matches[
                index + 1
            ].start()

        elif (
            heading_start is not None
            and heading_start > start
        ):
            end = heading_start

        else:
            end = len(
                stripped
            )

        fragment = (
            stripped[
                start:end
            ]
            .strip()
            .rstrip("-")
            .strip()
        )

        if not fragment:
            continue

        candidate = (
            _extract_candidate(
                fragment
            )
        )

        if candidate is None:
            continue

        # ---------------------------------------------
        # Composite period headings:
        #
        #     Ti?t 9. B?I 7. ...
        #
        # The punctuation after the period number is
        # presentation syntax and must survive overlay.
        #
        # Store only the numeric metadata span:
        #
        #     9
        #
        # rather than:
        #
        #     9.
        # ---------------------------------------------

        if (
            candidate.field
            == MetadataField.CURRICULUM_PERIOD
        ):
            period_match = re.fullmatch(
                r"\s*(\d+)\s*[.\-:]?\s*",
                candidate.value_text,
            )

            if period_match is not None:
                candidate = _Candidate(
                    field=candidate.field,
                    value_text=(
                        period_match.group(1)
                    ),
                    strategy=candidate.strategy,
                    confidence=(
                        candidate.confidence
                    ),
                    label_only=(
                        candidate.label_only
                    ),
                )

        candidates.append(
            candidate
        )


    # -----------------------------------------------------
    # Parse numbered lesson heading independently.
    #
    # Example:
    #
    #     Tiet 9. BAI 7. PHAN SO
    #
    # produces:
    #     curriculum_period
    #     lesson_title
    # -----------------------------------------------------

    if heading_start is not None:

        heading_text = stripped[
            heading_start:
        ].strip()

        heading_candidate = (
            _extract_candidate(
                heading_text
            )
        )

        if (
            heading_candidate
            is not None
            and heading_candidate.field
            == MetadataField.LESSON_TITLE
        ):
            title_match = re.match(
                (
                    r"^\s*"
                    r"B\u00e0i\s+"
                    r"(?:\d+|[IVXLCDMivxlcdm]+)"
                    r"\s*[.:\-]?\s*"
                    r"(?P<title>.+?)"
                    r"\s*$"
                ),
                heading_text,
                flags=re.IGNORECASE,
            )

            if title_match is not None:
                heading_candidate = _Candidate(
                    field=(
                        heading_candidate.field
                    ),
                    value_text=(
                        title_match.group(
                            "title"
                        )
                    ),
                    strategy=(
                        heading_candidate.strategy
                    ),
                    confidence=(
                        heading_candidate.confidence
                    ),
                    label_only=False,
                )

            candidates.append(
                heading_candidate
            )


    # -----------------------------------------------------
    # One source paragraph should not produce duplicate
    # values for the same canonical metadata field.
    # -----------------------------------------------------

    unique: list[
        _Candidate
    ] = []

    seen_fields: set[
        MetadataField
    ] = set()

    for candidate in candidates:

        if (
            candidate.field
            in seen_fields
        ):
            continue

        seen_fields.add(
            candidate.field
        )

        unique.append(
            candidate
        )


    if unique:
        return tuple(
            unique
        )


    # -----------------------------------------------------
    # Last compatibility fallback.
    # -----------------------------------------------------

    candidate = _extract_candidate(
        stripped
    )

    if candidate is None:
        return ()

    return (
        candidate,
    )

class LessonPlanMetadataLocator:
    """
    Preservation-safe metadata locator.

    This component is READ ONLY.

    It identifies:
      - inline paragraph metadata
      - inline table-cell metadata
      - paired label/value cells
      - numbered lesson headings

    No paragraph, run, table cell, relationship,
    header, footer or XML node is mutated.
    """

    def locate(
        self,
        document: DocxDocument,
    ) -> tuple[
        MetadataLocation,
        ...,
    ]:

        locations: list[
            MetadataLocation
        ] = []

        self._locate_body_paragraphs(
            document=document,
            locations=locations,
        )

        for table_index, table in enumerate(
            document.tables
        ):
            self._locate_table(
                table=table,
                table_index=table_index,
                locations=locations,
            )

        return tuple(
            locations
        )

    def _locate_body_paragraphs(
        self,
        *,
        document: DocxDocument,
        locations: list[
            MetadataLocation
        ],
    ) -> None:

        for paragraph_index, paragraph in enumerate(
            document.paragraphs
        ):
            candidates = (
                _extract_candidates(
                    paragraph.text
                )
            )

            for candidate in candidates:

                locations.append(
                    MetadataLocation(
                        field=candidate.field,
                        kind=(
                            MetadataLocationKind
                            .PARAGRAPH
                        ),
                        text=paragraph.text,
                        value_text=(
                            candidate.value_text
                        ),
                        strategy=(
                            candidate.strategy
                        ),
                        confidence=(
                            candidate.confidence
                        ),
                        paragraph_index=(
                            paragraph_index
                        ),
                        value_paragraph_index=(
                            paragraph_index
                        ),
                    )
                )

    def _locate_table(
        self,
        *,
        table: Table,
        table_index: int,
        locations: list[
            MetadataLocation
        ],
    ) -> None:

        for row_index, row in enumerate(
            table.rows
        ):

            unique_cells: list[
                tuple[
                    int,
                    _Cell,
                ]
            ] = []

            visited_cells: set[int] = set()

            for cell_index, cell in enumerate(
                row.cells
            ):
                identity = id(
                    cell._tc
                )

                if identity in visited_cells:
                    continue

                visited_cells.add(
                    identity
                )

                unique_cells.append(
                    (
                        cell_index,
                        cell,
                    )
                )

            for position, (
                cell_index,
                cell,
            ) in enumerate(
                unique_cells
            ):

                for paragraph_index, paragraph in enumerate(
                    cell.paragraphs
                ):
                    candidates = (
                        _extract_candidates(
                            paragraph.text
                        )
                    )

                    if not candidates:
                        continue

                    for candidate in candidates:

                        # ---------------------------------
                        # Inline metadata.
                        #
                        # One paragraph/cell may contain:
                        #
                        # Ngay soan ... Ngay day ... Lop ...
                        #
                        # Every candidate becomes its own
                        # MetadataLocation.
                        # ---------------------------------

                        if not candidate.label_only:
                            locations.append(
                                MetadataLocation(
                                    field=(
                                        candidate.field
                                    ),
                                    kind=(
                                        MetadataLocationKind
                                        .TABLE_CELL
                                    ),
                                    text=(
                                        paragraph.text
                                    ),
                                    value_text=(
                                        candidate
                                        .value_text
                                    ),
                                    strategy=(
                                        candidate.strategy
                                    ),
                                    confidence=(
                                        candidate
                                        .confidence
                                    ),
                                    table_index=(
                                        table_index
                                    ),
                                    row_index=(
                                        row_index
                                    ),
                                    cell_index=(
                                        cell_index
                                    ),
                                    cell_paragraph_index=(
                                        paragraph_index
                                    ),
                                    value_table_index=(
                                        table_index
                                    ),
                                    value_row_index=(
                                        row_index
                                    ),
                                    value_cell_index=(
                                        cell_index
                                    ),
                                    value_cell_paragraph_index=(
                                        paragraph_index
                                    ),
                                )
                            )

                            continue

                        # ---------------------------------
                        # Label-only metadata:
                        #
                        # Ngay day: | 15/09/2026
                        #
                        # Preserve the established paired
                        # cell contract.
                        # ---------------------------------

                        next_position = (
                            position + 1
                        )

                        if (
                            next_position
                            >= len(
                                unique_cells
                            )
                        ):
                            locations.append(
                                MetadataLocation(
                                    field=(
                                        candidate.field
                                    ),
                                    kind=(
                                        MetadataLocationKind
                                        .TABLE_CELL
                                    ),
                                    text=(
                                        paragraph.text
                                    ),
                                    value_text="",
                                    strategy=(
                                        candidate.strategy
                                    ),
                                    confidence=0.75,
                                    table_index=(
                                        table_index
                                    ),
                                    row_index=(
                                        row_index
                                    ),
                                    cell_index=(
                                        cell_index
                                    ),
                                    cell_paragraph_index=(
                                        paragraph_index
                                    ),
                                )
                            )

                            continue

                        (
                            target_cell_index,
                            target_cell,
                        ) = unique_cells[
                            next_position
                        ]

                        target_text = (
                            target_cell.text
                            .strip()
                        )

                        # Do not pair with a cell that
                        # itself contains metadata labels.
                        target_candidates = (
                            _extract_candidates(
                                target_text
                            )
                        )

                        if (
                            not target_text
                            or target_candidates
                            or not _looks_like_value(
                                candidate.field,
                                target_text,
                            )
                        ):
                            locations.append(
                                MetadataLocation(
                                    field=(
                                        candidate.field
                                    ),
                                    kind=(
                                        MetadataLocationKind
                                        .TABLE_CELL
                                    ),
                                    text=(
                                        paragraph.text
                                    ),
                                    value_text="",
                                    strategy=(
                                        candidate.strategy
                                    ),
                                    confidence=0.75,
                                    table_index=(
                                        table_index
                                    ),
                                    row_index=(
                                        row_index
                                    ),
                                    cell_index=(
                                        cell_index
                                    ),
                                    cell_paragraph_index=(
                                        paragraph_index
                                    ),
                                )
                            )

                            continue

                        locations.append(
                            MetadataLocation(
                                field=(
                                    candidate.field
                                ),
                                kind=(
                                    MetadataLocationKind
                                    .TABLE_PAIRED_CELL
                                ),
                                text=(
                                    paragraph.text
                                ),
                                value_text=(
                                    target_text
                                ),
                                strategy=(
                                    MetadataMatchStrategy
                                    .PAIRED_CELL
                                ),
                                confidence=0.99,
                                table_index=(
                                    table_index
                                ),
                                row_index=(
                                    row_index
                                ),
                                cell_index=(
                                    cell_index
                                ),
                                cell_paragraph_index=(
                                    paragraph_index
                                ),
                                value_table_index=(
                                    table_index
                                ),
                                value_row_index=(
                                    row_index
                                ),
                                value_cell_index=(
                                    target_cell_index
                                ),
                                value_cell_paragraph_index=0,
                            )
                        )
