from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
import re
from types import SimpleNamespace
from typing import Any, Mapping

from docx import Document

from document_intelligence.contracts import DocumentField
from document_intelligence.validation import (
    CanonicalDocumentContext,
    ValidationStatus,
)


@dataclass(frozen=True)
class FullAuditEvidenceBundle:
    canonical_context: CanonicalDocumentContext
    validated_analysis: Any
    context_result: Any
    standardization_report: Any
    ready: bool
    missing_evidence: tuple[str, ...]


def _first(context: Mapping[str, Any], occurrence: Mapping[str, Any], *names: str):
    for source in (context, occurrence):
        for name in names:
            value = source.get(name)
            if value not in (None, ""):
                return value
    return None


def _date_text(value: Any) -> str | None:
    if value in (None, ""):
        return None
    if isinstance(value, datetime):
        value = value.date()
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    text = str(value).strip()
    if not text:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(text, fmt).strftime("%d/%m/%Y")
        except ValueError:
            pass
    return text


def build_canonical_context(group_context: Mapping[str, Any] | None) -> CanonicalDocumentContext:
    context = dict(group_context or {})
    occurrences = tuple(context.get("occurrences", ()) or ())
    occurrence = next(
        (dict(item) for item in occurrences if isinstance(item, Mapping)),
        {},
    )

    period = _first(context, occurrence, "curriculum_period", "period", "ppct_period")
    if period in (None, ""):
        periods = tuple(context.get("curriculum_periods", ()) or ())
        period = periods[0] if periods else None
    try:
        period = int(period) if period not in (None, "") else None
    except (TypeError, ValueError):
        period = None

    class_name = _first(
        context,
        occurrence,
        "class_name",
        "class_display",
        "class",
        "class_id",
        "class_ref",
    )
    lesson_title = _first(context, occurrence, "lesson_title", "title")
    drafting_date = _date_text(_first(context, occurrence, "drafting_date"))
    teaching_date = _date_text(
        _first(context, occurrence, "teaching_date", "date", "lesson_date")
    )

    return CanonicalDocumentContext(
        class_name=str(class_name).strip() if class_name not in (None, "") else None,
        curriculum_period=period,
        lesson_title=str(lesson_title).strip() if lesson_title not in (None, "") else None,
        drafting_date=drafting_date,
        teaching_date=teaching_date,
    )


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(bytes(content)))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _norm(value: str) -> str:
    return " ".join(str(value or "").casefold().split())


def _date_candidates(value: str) -> tuple[str, ...]:
    text = str(value or "").strip()
    if not text:
        return ()
    result = [text]
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"):
        try:
            parsed = datetime.strptime(text, fmt)
            result.extend(
                (
                    parsed.strftime("%d/%m/%Y"),
                    parsed.strftime("%d-%m-%Y"),
                    parsed.strftime("%Y-%m-%d"),
                    parsed.strftime("%d/%m/%y"),
                )
            )
            break
        except ValueError:
            pass
    return tuple(dict.fromkeys(result))


# V14B3_RECOGNIZED_CANONICAL_VALUE
def _recognized_field_value(
    text: str,
    field: DocumentField,
    canonical_value: str,
) -> str | None:
    canonical_text = str(canonical_value or "").strip()
    if not canonical_text:
        return None

    if field is DocumentField.CURRICULUM_PERIOD:
        patterns = (
            r"(?i)\bperiod\s*[:\-]?\s*(\d+)\b",
            r"(?i)\bti[e\u1ebf]t\s*[:\-]?\s*(\d+)\b",
            r"(?i)\bppct\s*[:\-]?\s*(\d+)\b",
        )
        matches = []
        for pattern in patterns:
            for match in re.finditer(pattern, text):
                matches.append(match.group(1))
        if not matches:
            return None
        try:
            wanted = int(canonical_text)
            for value in matches:
                if int(value) == wanted:
                    return value
        except (TypeError, ValueError):
            pass
        return matches[0]

    if field in (DocumentField.DRAFTING_DATE, DocumentField.TEACHING_DATE):
        for candidate in _date_candidates(canonical_text):
            match = re.search(re.escape(candidate), text, flags=re.IGNORECASE)
            if match is not None:
                return match.group(0)

        if field is DocumentField.DRAFTING_DATE:
            labels = (
                "date\\s+of\\s+planning",
                "date\\s+of\\s+preparation",
                "ng\\u00e0y\\s+so\\u1ea1n",
            )
        else:
            labels = (
                "date\\s+of\\s+teaching",
                "teaching\\s+date",
                "ng\\u00e0y\\s+d\\u1ea1y",
            )
        pattern = (
            "(?i)(?:" + "|".join(labels) + ")"
            "\\s*[:\\-]?\\s*"
            "([0-3]?\\d[\\/\\-][01]?\\d[\\/\\-](?:\\d{2}|\\d{4}))"
        )
        match = re.search(pattern, text)
        return match.group(1) if match is not None else None

    exact = re.search(re.escape(canonical_text), text, flags=re.IGNORECASE)
    if exact is not None:
        return exact.group(0)

    if field is DocumentField.CLASS_NAME:
        match = re.search(
            "(?im)^\\s*(?:class|l\\u1edbp)\\s*[:\\-]?\\s*"
            "([0-9]{1,2}[A-Za-z][0-9A-Za-z\\-]*)\\s*$",
            text,
        )
        return match.group(1) if match is not None else None

    if field is DocumentField.LESSON_TITLE:
        unit_match = re.search(
            "(?im)^\\s*(UNIT\\s+\\d+\\s*:\\s*.+?)\\s*$",
            text,
        )
        lesson_match = re.search(
            "(?im)^\\s*Lesson\\s+\\d+\\s*:\\s*(.+?)\\s*$",
            text,
        )
        if unit_match is not None and lesson_match is not None:
            return (
                unit_match.group(1).strip()
                + " - "
                + lesson_match.group(1).strip()
            )
        match = re.search(
            "(?im)^\\s*(?:lesson|b\\u00e0i)\\s*[:\\-]\\s*(.+?)\\s*$",
            text,
        )
        return match.group(1).strip() if match is not None else None

    return None


def _recognized_matches_canonical(
    *,
    text: str,
    field: DocumentField,
    canonical_value: str,
    found_value: str,
) -> bool:
    if field in (DocumentField.DRAFTING_DATE, DocumentField.TEACHING_DATE):
        return _field_present(text, field, canonical_value)

    if field is DocumentField.CURRICULUM_PERIOD:
        try:
            return int(str(canonical_value).strip()) == int(str(found_value).strip())
        except (TypeError, ValueError):
            return False

    return _norm(str(canonical_value)) == _norm(str(found_value))


def _field_present(text: str, field: DocumentField, canonical_value: str) -> bool:
    normalized_text = _norm(text)
    if field is DocumentField.CURRICULUM_PERIOD:
        try:
            period = int(str(canonical_value).strip())
        except (TypeError, ValueError):
            return False
        patterns = (
            rf"(?i)\bperiod\s*[:\-]?\s*{period}\b",
            rf"(?i)\bti[eáº¿]t\s*[:\-]?\s*{period}\b",
            rf"(?i)\bppct\s*[:\-]?\s*{period}\b",
        )
        return any(re.search(pattern, text) is not None for pattern in patterns)

    if field in (DocumentField.DRAFTING_DATE, DocumentField.TEACHING_DATE):
        return any(_norm(item) in normalized_text for item in _date_candidates(canonical_value))

    needle = _norm(canonical_value)
    return bool(needle and needle in normalized_text)


def build_validated_output_evidence(
    *,
    standardized_content: bytes,
    canonical_context: CanonicalDocumentContext,
):
    text = _docx_text(standardized_content)
    proposals = []

    for field in (
        DocumentField.CLASS_NAME,
        DocumentField.CURRICULUM_PERIOD,
        DocumentField.LESSON_TITLE,
        DocumentField.DRAFTING_DATE,
        DocumentField.TEACHING_DATE,
    ):
        canonical_value = canonical_context.value_for(field)
        status = ValidationStatus.UNVERIFIED
        found_value = None
        if canonical_value not in (None, ""):
            found_value = _recognized_field_value(
                text,
                field,
                str(canonical_value),
            )
            if found_value not in (None, ""):
                if _recognized_matches_canonical(
                    text=text,
                    field=field,
                    canonical_value=str(canonical_value),
                    found_value=str(found_value),
                ):
                    status = ValidationStatus.ACCEPTED
                else:
                    status = ValidationStatus.CONFLICT

        proposals.append(
            SimpleNamespace(
                proposal=SimpleNamespace(
                    field=field,
                    value=found_value,
                ),
                status=status,
                canonical_value=canonical_value,
                found_value=found_value,
            )
        )

    return SimpleNamespace(proposals=tuple(proposals))


def build_full_audit_evidence(
    *,
    group_context: Mapping[str, Any] | None,
    standardized_content: bytes,
    pipeline_evidence: Mapping[str, Any] | None,
) -> FullAuditEvidenceBundle:
    pipeline = dict(pipeline_evidence or {})
    resolved_canonical_context = pipeline.get("resolved_canonical_context")
    canonical = build_canonical_context(
        resolved_canonical_context
        if isinstance(resolved_canonical_context, Mapping)
        else group_context
    )
    validated = build_validated_output_evidence(
        standardized_content=standardized_content,
        canonical_context=canonical,
    )

    context_result = pipeline.get("context_result")
    standardization_report = pipeline.get("standardization_report")

    missing = []
    if context_result is None:
        missing.append("context_result")
    if standardization_report is None:
        missing.append("standardization_report")

    canonical_values = (
        canonical.class_name,
        canonical.curriculum_period,
        canonical.lesson_title,
        canonical.drafting_date,
        canonical.teaching_date,
    )
    if any(value in (None, "") for value in canonical_values):
        missing.append("canonical_context")

    return FullAuditEvidenceBundle(
        canonical_context=canonical,
        validated_analysis=validated,
        context_result=context_result,
        standardization_report=standardization_report,
        ready=not missing,
        missing_evidence=tuple(missing),
    )
