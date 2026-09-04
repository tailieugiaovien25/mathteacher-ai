from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from enum import Enum
from typing import Mapping


class PeriodScopeStatus(str, Enum):
    SUCCESS = "SUCCESS"
    PARTIAL = "PARTIAL"
    FAILED = "FAILED"


@dataclass(frozen=True)
class PeriodTeachingOccurrence:
    curriculum_period: int
    teaching_date: date
    class_id: str = ""
    class_display: str = ""
    timetable_period: int | None = None


@dataclass(frozen=True)
class PeriodScopeWarning:
    code: str
    curriculum_period: int | None
    message: str


@dataclass(frozen=True)
class PeriodScopeResult:
    status: PeriodScopeStatus
    selected_periods: tuple[int, ...]
    occurrences_by_period: tuple[tuple[int, tuple[PeriodTeachingOccurrence, ...]], ...]
    document_periods_outside_scope: tuple[int, ...]
    missing_document_periods: tuple[int, ...]
    warnings: tuple[PeriodScopeWarning, ...]

    def occurrence_map(self) -> dict[int, tuple[PeriodTeachingOccurrence, ...]]:
        return dict(self.occurrences_by_period)


def _positive_period(value) -> int | None:
    try:
        result = int(value)
    except (TypeError, ValueError):
        return None
    return result if result > 0 else None


def _date_value(value) -> date | None:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value or "").strip()
    for pattern in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(text, pattern).date()
        except ValueError:
            continue
    return None


def resolve_multi_period_scope(
    *,
    group_context: Mapping,
    document_periods,
) -> PeriodScopeResult:
    if not isinstance(group_context, Mapping):
        raise TypeError("group_context must be a mapping")

    selected = tuple(sorted({
        period
        for value in tuple(group_context.get("curriculum_periods", ()) or ())
        if (period := _positive_period(value)) is not None
    }))
    present = tuple(sorted({
        period
        for value in tuple(document_periods or ())
        if (period := _positive_period(value)) is not None
    }))

    warnings: list[PeriodScopeWarning] = []
    grouped: dict[int, list[PeriodTeachingOccurrence]] = {
        period: [] for period in selected
    }

    for raw in tuple(group_context.get("occurrences", ()) or ()):
        if not isinstance(raw, Mapping):
            warnings.append(PeriodScopeWarning(
                code="INVALID_OCCURRENCE",
                curriculum_period=None,
                message="Occurrence is not a mapping and was ignored.",
            ))
            continue
        period = _positive_period(raw.get("curriculum_period"))
        teaching_date = _date_value(raw.get("teaching_date"))
        if period is None or teaching_date is None:
            warnings.append(PeriodScopeWarning(
                code="INVALID_OCCURRENCE",
                curriculum_period=period,
                message="Occurrence has no valid curriculum period or teaching date.",
            ))
            continue
        if period not in grouped:
            warnings.append(PeriodScopeWarning(
                code="OCCURRENCE_OUTSIDE_SELECTED_SCOPE",
                curriculum_period=period,
                message=f"Period {period} occurrence is outside the selected week.",
            ))
            continue
        timetable_period = _positive_period(raw.get("timetable_period"))
        item = PeriodTeachingOccurrence(
            curriculum_period=period,
            teaching_date=teaching_date,
            class_id=str(raw.get("class_id") or ""),
            class_display=str(raw.get("class_display") or ""),
            timetable_period=timetable_period,
        )
        if item not in grouped[period]:
            grouped[period].append(item)

    for period in selected:
        grouped[period].sort(key=lambda item: (
            item.teaching_date,
            item.class_display,
            item.class_id,
            item.timetable_period or 0,
        ))
        if not grouped[period]:
            warnings.append(PeriodScopeWarning(
                code="MISSING_PERIOD_OCCURRENCE",
                curriculum_period=period,
                message=f"Period {period} has no valid teaching occurrence.",
            ))

    outside = tuple(period for period in present if period not in selected)
    for period in outside:
        warnings.append(PeriodScopeWarning(
            code="DOCUMENT_PERIOD_OUTSIDE_SELECTED_SCOPE",
            curriculum_period=period,
            message=f"Period {period} is outside the selected week and must remain unchanged.",
        ))

    missing = tuple(period for period in selected if period not in present)
    for period in missing:
        warnings.append(PeriodScopeWarning(
            code="SELECTED_PERIOD_NOT_FOUND_IN_DOCUMENT",
            curriculum_period=period,
            message=f"Period {period} was selected but was not found in the document.",
        ))

    usable = tuple(
        period for period in selected
        if period in present and grouped.get(period)
    )
    if not selected or not usable:
        status = PeriodScopeStatus.FAILED
    elif warnings:
        status = PeriodScopeStatus.PARTIAL
    else:
        status = PeriodScopeStatus.SUCCESS

    return PeriodScopeResult(
        status=status,
        selected_periods=selected,
        occurrences_by_period=tuple(
            (period, tuple(grouped[period])) for period in selected
        ),
        document_periods_outside_scope=outside,
        missing_document_periods=missing,
        warnings=tuple(warnings),
    )



# G1B_ENGLISH_PATCH03C2_SCOPED_PERIOD_DATE_OVERLAY
@dataclass(frozen=True)
class ScopedPeriodDateOverlayResult:
    content: bytes
    applied_periods: tuple[int, ...]
    warnings: tuple[PeriodScopeWarning, ...]


def _replace_paragraph_text_preserve_runs(paragraph, value: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(value)
        return
    runs[0].text = value
    for run in runs[1:]:
        run.text = ""


def _table_period(table) -> int | None:
    import re
    for row in table.rows:
        visited = set()
        for cell in row.cells:
            identity = id(cell._tc)
            if identity in visited:
                continue
            visited.add(identity)
            for paragraph in cell.paragraphs:
                match = re.search(r"(?i)\bperiod\s+(\d+)\b", paragraph.text)
                if match:
                    return int(match.group(1))
    return None


def _copy_table_paragraph_text(*, source_table, output_table) -> None:
    source_paragraphs = []
    output_paragraphs = []
    for table, target in ((source_table, source_paragraphs), (output_table, output_paragraphs)):
        visited = set()
        for row in table.rows:
            for cell in row.cells:
                identity = id(cell._tc)
                if identity in visited:
                    continue
                visited.add(identity)
                target.extend(cell.paragraphs)
    if len(source_paragraphs) != len(output_paragraphs):
        raise ValueError("HEADER_TABLE_PARAGRAPH_SHAPE_MISMATCH")
    for source_paragraph, output_paragraph in zip(source_paragraphs, output_paragraphs):
        if output_paragraph.text != source_paragraph.text:
            _replace_paragraph_text_preserve_runs(output_paragraph, source_paragraph.text)


def _replace_english_date_field(paragraph, *, label: str, value: date) -> bool:
    import re
    pattern = re.compile(
        rf"(?i)(?P<prefix>\b{re.escape(label)}\s*:\s*)"
        r"(?P<value>(?:\d{1,2}[ \t]*[/.-][ \t]*\d{1,2}[ \t]*[/.-][ \t]*\d{2,4})|[.\u2026\u00b7_\- \t]+)"
    )
    match = pattern.search(paragraph.text)
    if not match:
        return False
    replacement = value.strftime("%d/%m/%Y")
    updated = paragraph.text[:match.start("value")] + replacement + paragraph.text[match.end("value"):]
    _replace_paragraph_text_preserve_runs(paragraph, updated)
    return True


def _remove_english_planning_field(paragraph) -> bool:
    import re

    pattern = re.compile(
        r"(?i)\bDate\s+of\s+planning\s*:\s*"
        r"(?:(?:\d{1,2}[ \t]*[/.-][ \t]*\d{1,2}[ \t]*[/.-][ \t]*\d{2,4})|[.\u2026\u00b7_\- \t]+)?"
    )
    updated, count = pattern.subn("", paragraph.text, count=1)
    if not count:
        return False
    _replace_paragraph_text_preserve_runs(paragraph, updated.strip())
    return True


def apply_scoped_english_period_dates(
    *,
    source_content: bytes,
    output_content: bytes,
    scope: PeriodScopeResult,
    drafting_date: date,
) -> ScopedPeriodDateOverlayResult:
    from io import BytesIO
    from docx import Document

    source_document = Document(BytesIO(source_content))
    output_document = Document(BytesIO(output_content))
    if len(source_document.tables) != len(output_document.tables):
        raise ValueError("DOCUMENT_TABLE_SHAPE_MISMATCH")

    occurrence_map = scope.occurrence_map()
    selected = set(scope.selected_periods)
    planning_owner_period = min(selected) if selected else None
    warnings = list(scope.warnings)
    applied = []

    for source_table, output_table in zip(source_document.tables, output_document.tables):
        period = _table_period(source_table)
        if period is None:
            continue

        # Restore the source header first. The legacy single-value overlay may
        # have changed every Period/date/title in the document.
        _copy_table_paragraph_text(source_table=source_table, output_table=output_table)

        # Keep Date of planning only in the first selected period header.
        if period in selected and period != planning_owner_period:
            visited_planning_cells = set()
            for planning_row in output_table.rows:
                for planning_cell in planning_row.cells:
                    planning_identity = id(planning_cell._tc)
                    if planning_identity in visited_planning_cells:
                        continue
                    visited_planning_cells.add(planning_identity)
                    for planning_paragraph in planning_cell.paragraphs:
                        _remove_english_planning_field(planning_paragraph)

        if period not in selected:
            continue
        occurrences = tuple(occurrence_map.get(period, ()))
        unique_dates = tuple(sorted({item.teaching_date for item in occurrences}))
        if len(unique_dates) != 1:
            warnings.append(PeriodScopeWarning(
                code="AMBIGUOUS_PERIOD_TEACHING_DATE",
                curriculum_period=period,
                message=f"Period {period} does not have exactly one teaching date.",
            ))
            continue

        teaching_replaced = False
        visited = set()
        for row in output_table.rows:
            for cell in row.cells:
                identity = id(cell._tc)
                if identity in visited:
                    continue
                visited.add(identity)
                for paragraph in cell.paragraphs:

                    if period == planning_owner_period:
                        _replace_english_date_field(
                            paragraph,
                            label="Date of planning",
                            value=drafting_date,
                        )
                    teaching_replaced = _replace_english_date_field(
                        paragraph, label="Date of teaching", value=unique_dates[0]
                    ) or teaching_replaced

        if teaching_replaced:
            applied.append(period)
        else:
            warnings.append(PeriodScopeWarning(
                code="ENGLISH_DATE_LABEL_NOT_FOUND",
                curriculum_period=period,
                message=f"Period {period} is missing Date of teaching.",
            ))

    buffer = BytesIO()
    output_document.save(buffer)
    return ScopedPeriodDateOverlayResult(
        content=buffer.getvalue(),
        applied_periods=tuple(applied),
        warnings=tuple(warnings),
    )
