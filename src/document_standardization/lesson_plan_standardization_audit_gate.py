from __future__ import annotations

from dataclasses import dataclass
from enum import Enum
from hashlib import sha256
from io import BytesIO
from typing import Any, Mapping

from docx import Document

from document_standardization.lesson_plan_standardization_quality_gate import (
    LessonPlanQualityGateResult,
    LessonPlanStandardizationQualityGate,
    QualityGateStatus,
)


class AuditStatus(str, Enum):
    PASS = "pass"
    WARNING = "warning"
    FAIL = "fail"
    UNVERIFIED = "unverified"


@dataclass(frozen=True)
class AuditEvidence:
    code: str
    status: AuditStatus
    message: str
    evidence: tuple[str, ...] = ()


@dataclass(frozen=True)
class LessonPlanAuditResult:
    status: AuditStatus
    trust_score: int
    evidence: tuple[AuditEvidence, ...]
    source_sha256: str
    output_sha256: str
    source_paragraph_count: int
    output_paragraph_count: int
    source_table_count: int
    output_table_count: int
    quality_gate: LessonPlanQualityGateResult | None

    @property
    def passed(self) -> bool:
        return self.status is AuditStatus.PASS


def _digest(content: bytes) -> str:
    return sha256(content).hexdigest()


def _load_docx(content: bytes):
    if not isinstance(content, (bytes, bytearray)) or not content:
        raise ValueError("DOCX content is empty.")
    return Document(BytesIO(bytes(content)))


def _quality_to_audit(status: QualityGateStatus) -> AuditStatus:
    if status is QualityGateStatus.PASS:
        return AuditStatus.PASS
    if status is QualityGateStatus.REVIEW:
        return AuditStatus.WARNING
    return AuditStatus.FAIL


def _score(evidence: tuple[AuditEvidence, ...]) -> int:
    if not evidence:
        return 0
    weights = {
        AuditStatus.PASS: 1.0,
        AuditStatus.WARNING: 0.5,
        AuditStatus.FAIL: 0.0,
        AuditStatus.UNVERIFIED: 0.0,
    }
    value = sum(weights[item.status] for item in evidence) / len(evidence)
    return max(0, min(100, round(value * 100)))


class LessonPlanStandardizationAuditGate:
    """
    Independent read-only audit gate.

    It never mutates source or output DOCX bytes. It verifies that the
    produced artifact is readable and independently aggregates the
    existing canonical quality-gate result into a user-facing audit
    status and deterministic evidence-coverage score.
    """

    def __init__(
        self,
        *,
        quality_gate: LessonPlanStandardizationQualityGate | None = None,
    ):
        self._quality_gate = quality_gate or LessonPlanStandardizationQualityGate()

    def evaluate(
        self,
        *,
        original_content: bytes,
        standardized_content: bytes,
        canonical_context: Any,
        validated_analysis: Any,
        context_result: Any,
        standardization_report: Mapping[str, Any] | None,
    ) -> LessonPlanAuditResult:
        evidence: list[AuditEvidence] = []

        try:
            source_doc = _load_docx(original_content)
        except Exception as error:
            return LessonPlanAuditResult(
                status=AuditStatus.UNVERIFIED,
                trust_score=0,
                evidence=(
                    AuditEvidence(
                        code="SOURCE_DOCX_UNREADABLE",
                        status=AuditStatus.UNVERIFIED,
                        message="Original DOCX could not be independently read.",
                        evidence=(type(error).__name__ + ": " + str(error),),
                    ),
                ),
                source_sha256=_digest(bytes(original_content or b"")),
                output_sha256=_digest(bytes(standardized_content or b"")),
                source_paragraph_count=0,
                output_paragraph_count=0,
                source_table_count=0,
                output_table_count=0,
                quality_gate=None,
            )

        try:
            output_doc = _load_docx(standardized_content)
        except Exception as error:
            return LessonPlanAuditResult(
                status=AuditStatus.FAIL,
                trust_score=0,
                evidence=(
                    AuditEvidence(
                        code="OUTPUT_DOCX_UNREADABLE",
                        status=AuditStatus.FAIL,
                        message="Standardized DOCX is empty, corrupt, or unreadable.",
                        evidence=(type(error).__name__ + ": " + str(error),),
                    ),
                ),
                source_sha256=_digest(bytes(original_content)),
                output_sha256=_digest(bytes(standardized_content or b"")),
                source_paragraph_count=len(source_doc.paragraphs),
                output_paragraph_count=0,
                source_table_count=len(source_doc.tables),
                output_table_count=0,
                quality_gate=None,
            )

        evidence.append(
            AuditEvidence(
                code="SOURCE_DOCX_READABLE",
                status=AuditStatus.PASS,
                message="Original DOCX is independently readable.",
            )
        )
        evidence.append(
            AuditEvidence(
                code="OUTPUT_DOCX_READABLE",
                status=AuditStatus.PASS,
                message="Standardized DOCX is independently readable.",
            )
        )

        source_hash = _digest(bytes(original_content))
        output_hash = _digest(bytes(standardized_content))
        evidence.append(
            AuditEvidence(
                code="ARTIFACT_HASH_EVIDENCE",
                status=AuditStatus.PASS,
                message="Source and output artifact hashes were recorded independently.",
                evidence=("source=" + source_hash, "output=" + output_hash),
            )
        )

        quality = self._quality_gate.evaluate(
            canonical_context=canonical_context,
            validated_analysis=validated_analysis,
            context_result=context_result,
            standardization_report=standardization_report,
        )

        for criterion in quality.criteria:
            evidence.append(
                AuditEvidence(
                    code="QUALITY_GATE_" + criterion.code,
                    status=_quality_to_audit(criterion.status),
                    message=criterion.message,
                    evidence=tuple(criterion.evidence),
                )
            )

        if any(item.status is AuditStatus.FAIL for item in evidence):
            overall = AuditStatus.FAIL
        elif any(item.status is AuditStatus.UNVERIFIED for item in evidence):
            overall = AuditStatus.UNVERIFIED
        elif any(item.status is AuditStatus.WARNING for item in evidence):
            overall = AuditStatus.WARNING
        else:
            overall = AuditStatus.PASS

        frozen = tuple(evidence)
        return LessonPlanAuditResult(
            status=overall,
            trust_score=_score(frozen),
            evidence=frozen,
            source_sha256=source_hash,
            output_sha256=output_hash,
            source_paragraph_count=len(source_doc.paragraphs),
            output_paragraph_count=len(output_doc.paragraphs),
            source_table_count=len(source_doc.tables),
            output_table_count=len(output_doc.tables),
            quality_gate=quality,
        )

    def evaluate_artifact_only(
        self,
        *,
        original_content: bytes,
        standardized_content: bytes,
    ) -> LessonPlanAuditResult:
        """
        Runtime-safe artifact verification used before full canonical evidence
        is wired through the standardization call chain.

        A readable artifact is WARNING, never PASS, because metadata and
        canonical business rules have not yet been independently verified.
        """
        try:
            source_doc = _load_docx(original_content)
        except Exception as error:
            return LessonPlanAuditResult(
                status=AuditStatus.UNVERIFIED,
                trust_score=0,
                evidence=(
                    AuditEvidence(
                        code="SOURCE_DOCX_UNREADABLE",
                        status=AuditStatus.UNVERIFIED,
                        message="Original DOCX could not be independently read.",
                        evidence=(type(error).__name__ + ": " + str(error),),
                    ),
                ),
                source_sha256=_digest(bytes(original_content or b"")),
                output_sha256=_digest(bytes(standardized_content or b"")),
                source_paragraph_count=0,
                output_paragraph_count=0,
                source_table_count=0,
                output_table_count=0,
                quality_gate=None,
            )

        try:
            output_doc = _load_docx(standardized_content)
        except Exception as error:
            return LessonPlanAuditResult(
                status=AuditStatus.FAIL,
                trust_score=0,
                evidence=(
                    AuditEvidence(
                        code="OUTPUT_DOCX_UNREADABLE",
                        status=AuditStatus.FAIL,
                        message="Standardized DOCX is empty, corrupt, or unreadable.",
                        evidence=(type(error).__name__ + ": " + str(error),),
                    ),
                ),
                source_sha256=_digest(bytes(original_content)),
                output_sha256=_digest(bytes(standardized_content or b"")),
                source_paragraph_count=len(source_doc.paragraphs),
                output_paragraph_count=0,
                source_table_count=len(source_doc.tables),
                output_table_count=0,
                quality_gate=None,
            )

        source_hash = _digest(bytes(original_content))
        output_hash = _digest(bytes(standardized_content))
        evidence = (
            AuditEvidence(
                code="SOURCE_DOCX_READABLE",
                status=AuditStatus.PASS,
                message="Original DOCX is independently readable.",
            ),
            AuditEvidence(
                code="OUTPUT_DOCX_READABLE",
                status=AuditStatus.PASS,
                message="Standardized DOCX is independently readable.",
            ),
            AuditEvidence(
                code="ARTIFACT_HASH_EVIDENCE",
                status=AuditStatus.PASS,
                message="Source and output artifact hashes were recorded independently.",
                evidence=("source=" + source_hash, "output=" + output_hash),
            ),
            AuditEvidence(
                code="CANONICAL_EVIDENCE_NOT_WIRED",
                status=AuditStatus.WARNING,
                message=(
                    "Artifact verification passed, but canonical field evidence "
                    "is not yet carried through the current runtime contract."
                ),
            ),
        )
        return LessonPlanAuditResult(
            status=AuditStatus.WARNING,
            trust_score=_score(evidence),
            evidence=evidence,
            source_sha256=source_hash,
            output_sha256=output_hash,
            source_paragraph_count=len(source_doc.paragraphs),
            output_paragraph_count=len(output_doc.paragraphs),
            source_table_count=len(source_doc.tables),
            output_table_count=len(output_doc.tables),
            quality_gate=None,
        )

