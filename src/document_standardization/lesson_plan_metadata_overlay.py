from __future__ import annotations

from dataclasses import dataclass
from datetime import date
import re

from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

from document_standardization.lesson_plan_metadata import (
    LessonPlanMetadata,
)
from document_standardization.lesson_plan_metadata_locator import (
    LessonPlanMetadataLocator,
    MetadataField,
    MetadataLocation,
    MetadataLocationKind,
    MetadataMatchStrategy,
    _fold_text,
)


class MetadataOverlayError(
    ValueError
):
    pass


@dataclass(
    frozen=True,
    slots=True,
)
class MetadataOverlayChange:
    field: MetadataField
    old_value: str
    new_value: str
    kind: MetadataLocationKind
    strategy: MetadataMatchStrategy
    confidence: float
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    paragraph_index: int | None = None


@dataclass(
    frozen=True,
    slots=True,
)
class MetadataOverlayResult:
    changes: tuple[
        MetadataOverlayChange,
        ...,
    ] = ()

    unresolved_fields: tuple[
        MetadataField,
        ...,
    ] = ()

    skipped_low_confidence: tuple[
        MetadataField,
        ...,
    ] = ()

    @property
    def changed(self) -> bool:
        return bool(
            self.changes
        )

    @property
    def change_count(self) -> int:
        return len(
            self.changes
        )


_FIELD_ATTRIBUTE = {
    MetadataField.SCHOOL_NAME:
        "school_name",

    MetadataField.TEACHER_NAME:
        "teacher_name",

    MetadataField.SUBJECT_NAME:
        "subject_name",

    MetadataField.CLASS_NAME:
        "class_name",

    MetadataField.LESSON_TITLE:
        "lesson_title",

    MetadataField.CURRICULUM_PERIOD:
        "curriculum_period",

    MetadataField.DRAFTING_DATE:
        "drafting_date",

    MetadataField.TEACHING_DATE:
        "teaching_date",
}


def _display_value(
    value,
) -> str:

    if isinstance(
        value,
        date,
    ):
        return value.strftime(
            "%d/%m/%Y"
        )

    return str(
        value
    )


def _lesson_heading_value(
    *,
    current: str,
    replacement: str,
) -> str:
    """
    Preserve an existing numbered lesson prefix
    when canonical lesson_title only contains the
    title itself.

    Example:
        current:
            Bai 7. Old title

        replacement:
            New title

        output:
            Bai 7. New title

    If replacement already contains a numbered
    Bai prefix, use it exactly.
    """

    folded_replacement = (
        _fold_text(
            replacement
        )
    )

    if re.match(
        r"^bai\s+"
        r"(?:\d+|[ivxlcdm]+)"
        r"(?:\s*[.:\-]|\s+)",
        folded_replacement,
    ):
        return replacement

    pattern = re.compile(
        "^("
        "\\s*"
        "B\\u00e0i"
        "\\s+"
        "(?:\\d+|[IVXLCDMivxlcdm]+)"
        "\\s*"
        "[.:\\-]?"
        "\\s*"
        ")"
        "(.+)$"
    )

    match = pattern.match(
        current
    )

    if match is None:
        # The locator may intentionally provide
        # only the editable title span rather than
        # the complete "B?i N." heading.
        #
        # Preserve presentation casing from the
        # existing visible title.
        letters = [
            character
            for character in current
            if character.isalpha()
        ]

        if (
            letters
            and all(
                character.isupper()
                for character in letters
            )
        ):
            return replacement.upper()

        return replacement

    title = match.group(2)

    letters = [
        character
        for character in title
        if character.isalpha()
    ]

    if (
        letters
        and all(
            character.isupper()
            for character in letters
        )
    ):
        replacement = (
            replacement.upper()
        )

    return (
        match.group(1)
        + replacement
    )


def _text_nodes(
    paragraph: Paragraph,
):
    """
    Return ordinary w:t nodes inside paragraph runs.

    We edit the text nodes directly rather than
    assigning paragraph.text or run.text.

    This preserves:
      - run properties
      - bold/italic/font
      - surrounding XML
      - non-text children
    """

    nodes = []

    for run in paragraph.runs:
        nodes.extend(
            run._r.xpath(
                ".//w:t"
            )
        )

    return nodes


def _replace_text_nodes(
    *,
    paragraph: Paragraph,
    old_value: str,
    new_value: str,
) -> None:
    """
    Replace one exact visible-text span while
    preserving the surrounding runs/XML.

    The new value inherits the formatting of the
    first text node occupied by the old value.
    """

    if not old_value:
        raise MetadataOverlayError(
            "old metadata value must not be empty"
        )

    nodes = _text_nodes(
        paragraph
    )

    if not nodes:
        raise MetadataOverlayError(
            "target paragraph has no editable text nodes"
        )

    values = [
        node.text or ""
        for node in nodes
    ]

    combined = "".join(
        values
    )

    start = combined.find(
        old_value
    )

    if start < 0:
        raise MetadataOverlayError(
            "metadata value not found in target paragraph"
        )

    end = (
        start
        + len(
            old_value
        )
    )

    positions = []

    cursor = 0

    for index, value in enumerate(
        values
    ):
        node_start = cursor
        node_end = (
            cursor
            + len(
                value
            )
        )

        positions.append(
            (
                index,
                node_start,
                node_end,
            )
        )

        cursor = node_end

    affected = [
        item
        for item in positions
        if (
            item[2] > start
            and item[1] < end
        )
    ]

    if not affected:
        raise MetadataOverlayError(
            "metadata span could not be mapped to text nodes"
        )

    first_index = affected[
        0
    ][0]

    last_index = affected[
        -1
    ][0]

    first_node_start = affected[
        0
    ][1]

    last_node_start = affected[
        -1
    ][1]

    first_value = values[
        first_index
    ]

    last_value = values[
        last_index
    ]

    prefix_length = (
        start
        - first_node_start
    )

    suffix_offset = (
        end
        - last_node_start
    )

    prefix = first_value[
        :prefix_length
    ]

    suffix = last_value[
        suffix_offset:
    ]

    if (
        first_index
        == last_index
    ):
        nodes[
            first_index
        ].text = (
            prefix
            + new_value
            + suffix
        )

        return

    nodes[
        first_index
    ].text = (
        prefix
        + new_value
    )

    for index in range(
        first_index + 1,
        last_index,
    ):
        nodes[
            index
        ].text = ""

    nodes[
        last_index
    ].text = suffix


class LessonPlanMetadataOverlay:
    """
    Apply canonical metadata only to locations
    discovered by LessonPlanMetadataLocator.

    Safety contract:
      - high-confidence locations only
      - no paragraph.text assignment
      - no cell.text assignment
      - no document reconstruction
      - existing run/XML formatting preserved
      - missing metadata does not erase source data
    """

    def __init__(
        self,
        locator: (
            LessonPlanMetadataLocator
            | None
        ) = None,
    ) -> None:

        self.locator = (
            locator
            or LessonPlanMetadataLocator()
        )

    def apply(
        self,
        *,
        document: DocxDocument,
        metadata: LessonPlanMetadata,
    ) -> MetadataOverlayResult:

        if not isinstance(
            metadata,
            LessonPlanMetadata,
        ):
            raise TypeError(
                "metadata must be LessonPlanMetadata"
            )

        requested_values = (
            metadata.overlay_values()
        )

        if not requested_values:
            return MetadataOverlayResult()

        locations = (
            self.locator.locate(
                document
            )
        )

        changes = []

        resolved_fields = set()
        skipped_low_confidence = set()

        for field in MetadataField:

            attribute_name = (
                _FIELD_ATTRIBUTE[
                    field
                ]
            )

            if (
                attribute_name
                not in requested_values
            ):
                continue

            canonical_value = (
                _display_value(
                    requested_values[
                        attribute_name
                    ]
                )
            )

            matches = [
                location
                for location
                in locations
                if location.field == field
            ]

            for location in matches:

                if (
                    not location
                    .is_high_confidence
                ):
                    skipped_low_confidence.add(
                        field
                    )
                    continue

                if not location.value_text:
                    continue

                paragraph = (
                    self._resolve_target_paragraph(
                        document=document,
                        location=location,
                    )
                )

                replacement = (
                    canonical_value
                )

                if (
                    location.strategy
                    == MetadataMatchStrategy
                    .LESSON_HEADING
                ):
                    replacement = (
                        _lesson_heading_value(
                            current=(
                                location
                                .value_text
                            ),
                            replacement=(
                                canonical_value
                            ),
                        )
                    )

                # Matching the target is sufficient
                # to resolve the field even if the
                # value is already canonical.
                resolved_fields.add(
                    field
                )

                if (
                    location.value_text
                    == replacement
                ):
                    continue

                _replace_text_nodes(
                    paragraph=paragraph,
                    old_value=(
                        location
                        .value_text
                    ),
                    new_value=(
                        replacement
                    ),
                )

                changes.append(
                    MetadataOverlayChange(
                        field=field,
                        old_value=(
                            location
                            .value_text
                        ),
                        new_value=(
                            replacement
                        ),
                        kind=(
                            location.kind
                        ),
                        strategy=(
                            location.strategy
                        ),
                        confidence=(
                            location.confidence
                        ),
                        table_index=(
                            location
                            .value_table_index
                        ),
                        row_index=(
                            location
                            .value_row_index
                        ),
                        cell_index=(
                            location
                            .value_cell_index
                        ),
                        paragraph_index=(
                            location
                            .value_paragraph_index
                            if (
                                location
                                .kind
                                == MetadataLocationKind
                                .PARAGRAPH
                            )
                            else (
                                location
                                .value_cell_paragraph_index
                            )
                        ),
                    )
                )

        requested_fields = {
            field
            for field in MetadataField
            if (
                _FIELD_ATTRIBUTE[
                    field
                ]
                in requested_values
            )
        }

        unresolved = (
            requested_fields
            - resolved_fields
        )

        return MetadataOverlayResult(
            changes=tuple(
                changes
            ),
            unresolved_fields=tuple(
                sorted(
                    unresolved,
                    key=lambda item:
                        item.value,
                )
            ),
            skipped_low_confidence=tuple(
                sorted(
                    skipped_low_confidence,
                    key=lambda item:
                        item.value,
                )
            ),
        )

    @staticmethod
    def _resolve_target_paragraph(
        *,
        document: DocxDocument,
        location: MetadataLocation,
    ) -> Paragraph:

        if (
            location.kind
            == MetadataLocationKind
            .PARAGRAPH
        ):
            index = (
                location
                .value_paragraph_index
            )

            if index is None:
                raise MetadataOverlayError(
                    "paragraph target is missing"
                )

            return document.paragraphs[
                index
            ]

        table_index = (
            location
            .value_table_index
        )

        row_index = (
            location
            .value_row_index
        )

        cell_index = (
            location
            .value_cell_index
        )

        paragraph_index = (
            location
            .value_cell_paragraph_index
        )

        if any(
            value is None
            for value in (
                table_index,
                row_index,
                cell_index,
                paragraph_index,
            )
        ):
            raise MetadataOverlayError(
                "table target coordinates are incomplete"
            )

        return (
            document.tables[
                table_index
            ]
            .rows[
                row_index
            ]
            .cells[
                cell_index
            ]
            .paragraphs[
                paragraph_index
            ]
        )
