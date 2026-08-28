from __future__ import annotations

from hashlib import sha256
from io import BytesIO
from math import isfinite
from typing import Mapping, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from .dynamic_document_renderer import (
    AssessmentDocumentRenderPlan,
)


class AssessmentDocxRendererError(ValueError):
    """Raised when a render plan cannot be converted to DOCX."""


_ALIGNMENTS = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_PAGE_SIZES_MM = {
    "A4": (210.0, 297.0),
    "LETTER": (215.9, 279.4),
}


def _is_sequence(value: object) -> bool:
    return isinstance(value, Sequence) and not isinstance(
        value,
        (str, bytes, bytearray),
    )


def _object(value: object, field_name: str) -> Mapping[str, object]:
    if not isinstance(value, Mapping):
        raise AssessmentDocxRendererError(
            f"{field_name} must be an object"
        )
    return value


def _array(value: object, field_name: str) -> tuple[object, ...]:
    if not _is_sequence(value):
        raise AssessmentDocxRendererError(
            f"{field_name} must be an array"
        )
    return tuple(value)


def _text(
    value: object,
    field_name: str,
    *,
    allow_empty: bool = False,
) -> str:
    if not isinstance(value, str):
        raise AssessmentDocxRendererError(
            f"{field_name} must be text"
        )
    normalized = value.strip()
    if not normalized and not allow_empty:
        raise AssessmentDocxRendererError(
            f"{field_name} is required"
        )
    return normalized


def _number(
    value: object,
    field_name: str,
    *,
    minimum: float,
    maximum: float,
) -> float:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        raise AssessmentDocxRendererError(
            f"{field_name} must be numeric"
        )
    number = float(value)
    if not isfinite(number) or number < minimum or number > maximum:
        raise AssessmentDocxRendererError(
            f"{field_name} is outside the supported range"
        )
    return number


def _value_at_path(value: object, path: str) -> object:
    current = value
    for segment in _text(path, "value_path").split("."):
        if isinstance(current, Mapping):
            if segment not in current:
                raise AssessmentDocxRendererError(
                    f"value_path does not resolve: {path}"
                )
            current = current[segment]
            continue
        if _is_sequence(current) and segment.isdigit():
            index = int(segment)
            if index >= len(current):
                raise AssessmentDocxRendererError(
                    f"value_path index is out of range: {path}"
                )
            current = current[index]
            continue
        raise AssessmentDocxRendererError(
            f"value_path does not resolve: {path}"
        )
    return current


def _display(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Có" if value else "Không"
    if isinstance(value, float):
        return f"{value:g}"
    if isinstance(value, Mapping):
        return "; ".join(
            f"{key}: {_display(item)}"
            for key, item in value.items()
        )
    if _is_sequence(value):
        return "; ".join(_display(item) for item in value)
    return str(value)


def _set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    properties = run._element.get_or_add_rPr()
    fonts = properties.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)


def _set_cell_margins(cell, margins: Mapping[str, int]) -> None:
    properties = cell._tc.get_or_add_tcPr()
    node = properties.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        properties.append(node)
    for side in ("top", "start", "bottom", "end"):
        child = node.find(qn(f"w:{side}"))
        if child is None:
            child = OxmlElement(f"w:{side}")
            node.append(child)
        child.set(qn("w:w"), str(margins[side]))
        child.set(qn("w:type"), "dxa")


class AssessmentDocxRenderPlanRenderer:
    """Convert a safe assessment render plan to deterministic DOCX bytes."""

    def render(
        self,
        *,
        plan: AssessmentDocumentRenderPlan,
        template_asset: bytes | None = None,
    ) -> bytes:
        if not isinstance(plan, AssessmentDocumentRenderPlan):
            raise AssessmentDocxRendererError(
                "assessment render plan is required"
            )
        if plan.renderer_code != "DOCX_JSON_V1":
            raise AssessmentDocxRendererError(
                "DOCX_JSON_V1 render plan is required"
            )

        document = self._open_document(
            plan=plan,
            template_asset=template_asset,
        )
        self._configure_document(document, plan)

        for index, section in enumerate(plan.sections):
            self._render_section(
                document=document,
                section=section,
                styles=plan.styles,
                section_index=index,
            )

        output = BytesIO()
        document.save(output)
        content = output.getvalue()
        if not content.startswith(b"PK"):
            raise AssessmentDocxRendererError(
                "renderer did not produce a DOCX package"
            )
        return content

    @staticmethod
    def _open_document(
        *,
        plan: AssessmentDocumentRenderPlan,
        template_asset: bytes | None,
    ) -> DocumentObject:
        expects_asset = plan.template_asset_path is not None
        if expects_asset != (template_asset is not None):
            raise AssessmentDocxRendererError(
                "template asset bytes must match the render plan"
            )
        if template_asset is None:
            return Document()
        actual_hash = sha256(template_asset).hexdigest()
        if actual_hash != plan.template_asset_hash:
            raise AssessmentDocxRendererError(
                "template asset hash does not match"
            )
        try:
            return Document(BytesIO(template_asset))
        except Exception as error:
            raise AssessmentDocxRendererError(
                "template asset is not a valid DOCX"
            ) from error

    def _configure_document(
        self,
        document: DocumentObject,
        plan: AssessmentDocumentRenderPlan,
    ) -> None:
        layout = plan.layout
        styles = plan.styles
        page_size = str(layout.get("paper_size", "A4")).upper()
        if page_size not in _PAGE_SIZES_MM:
            raise AssessmentDocxRendererError(
                "unsupported paper size"
            )
        orientation = str(
            layout.get("page_orientation", "PORTRAIT")
        ).upper()
        if orientation not in {"PORTRAIT", "LANDSCAPE"}:
            raise AssessmentDocxRendererError(
                "unsupported page orientation"
            )
        width_mm, height_mm = _PAGE_SIZES_MM[page_size]
        if orientation == "LANDSCAPE":
            width_mm, height_mm = height_mm, width_mm

        margins = _object(
            layout.get("margins_mm", {}),
            "layout.margins_mm",
        )
        top = _number(
            margins.get("top", 20),
            "margins_mm.top",
            minimum=5,
            maximum=60,
        )
        right = _number(
            margins.get("right", 15),
            "margins_mm.right",
            minimum=5,
            maximum=60,
        )
        bottom = _number(
            margins.get("bottom", 20),
            "margins_mm.bottom",
            minimum=5,
            maximum=60,
        )
        left = _number(
            margins.get("left", 20),
            "margins_mm.left",
            minimum=5,
            maximum=60,
        )

        for section in document.sections:
            section.orientation = (
                WD_ORIENT.LANDSCAPE
                if orientation == "LANDSCAPE"
                else WD_ORIENT.PORTRAIT
            )
            section.page_width = Mm(width_mm)
            section.page_height = Mm(height_mm)
            section.top_margin = Mm(top)
            section.right_margin = Mm(right)
            section.bottom_margin = Mm(bottom)
            section.left_margin = Mm(left)

        font_name = _text(
            styles.get("font_family", "Times New Roman"),
            "styles.font_family",
        )
        font_size = _number(
            styles.get("font_size", 12),
            "styles.font_size",
            minimum=8,
            maximum=24,
        )
        normal = document.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(font_size)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(
            _number(
                styles.get("paragraph_space_after_pt", 3),
                "styles.paragraph_space_after_pt",
                minimum=0,
                maximum=36,
            )
        )
        normal.paragraph_format.line_spacing = _number(
            styles.get("line_spacing", 1.15),
            "styles.line_spacing",
            minimum=1,
            maximum=2,
        )

    def _render_section(
        self,
        *,
        document: DocumentObject,
        section: Mapping[str, object],
        styles: Mapping[str, object],
        section_index: int,
    ) -> None:
        section_type = _text(
            section.get("section_type"),
            f"sections[{section_index}].section_type",
        ).upper()
        if section_type == "PAGE_BREAK":
            document.add_page_break()
            return

        title = section.get("title")
        if title is not None:
            self._add_paragraph(
                document,
                _text(title, f"sections[{section_index}].title"),
                styles,
                bold=True,
                alignment=str(
                    section.get("title_alignment", "LEFT")
                ),
                size_key="heading_size",
                keep_with_next=True,
            )

        data = _object(
            section.get("data", {}),
            f"sections[{section_index}].data",
        )
        if section_type == "FIELDS":
            self._render_fields(document, section, data, styles)
        elif section_type == "TEXT":
            self._render_text(document, section, data, styles)
        elif section_type == "TABLE":
            self._render_table(document, section, data, styles)
        elif section_type == "REPEAT":
            self._render_repeat(document, section, data, styles)
        else:
            raise AssessmentDocxRendererError(
                f"unsupported DOCX section type: {section_type}"
            )

    def _render_fields(self, document, section, data, styles) -> None:
        labels = section.get("field_labels", {})
        labels = _object(labels, "field_labels")
        for name, value in data.items():
            label = str(labels.get(name, name))
            self._add_paragraph(
                document,
                f"{label}: {_display(value)}",
                styles,
                bold=False,
                alignment=str(section.get("alignment", "LEFT")),
            )

    def _render_text(self, document, section, data, styles) -> None:
        literal = section.get("text")
        if literal is not None:
            self._add_paragraph(
                document,
                _text(literal, "section.text", allow_empty=True),
                styles,
                bold=bool(section.get("bold", False)),
                alignment=str(section.get("alignment", "LEFT")),
            )
        for value in data.values():
            self._add_paragraph(
                document,
                _display(value),
                styles,
                bold=False,
                alignment=str(section.get("alignment", "LEFT")),
            )

    def _render_repeat(self, document, section, data, styles) -> None:
        if len(data) != 1:
            raise AssessmentDocxRendererError(
                "REPEAT section requires exactly one binding"
            )
        rows = _array(next(iter(data.values())), "repeat binding")
        fields_value = section.get("fields", ())
        fields = _array(fields_value, "repeat fields")
        for item_index, item in enumerate(rows):
            row = _object(item, f"repeat item {item_index}")
            if fields:
                parts = []
                for field_index, field_value in enumerate(fields):
                    field = _object(
                        field_value,
                        f"repeat fields[{field_index}]",
                    )
                    path = _text(field.get("value_path"), "value_path")
                    label = str(field.get("label", "")).strip()
                    value = _display(_value_at_path(row, path))
                    parts.append(f"{label}: {value}" if label else value)
                content = " | ".join(parts)
            else:
                content = _display(row)
            self._add_paragraph(
                document,
                content,
                styles,
                bold=False,
                alignment=str(section.get("alignment", "LEFT")),
            )

    def _render_table(self, document, section, data, styles) -> None:
        if len(data) != 1:
            raise AssessmentDocxRendererError(
                "TABLE section requires exactly one binding"
            )
        rows = _array(next(iter(data.values())), "table binding")
        columns = tuple(
            _object(item, f"columns[{index}]")
            for index, item in enumerate(
                _array(section.get("columns"), "columns")
            )
        )
        if not columns:
            raise AssessmentDocxRendererError(
                "TABLE section requires columns"
            )

        usable_dxa = self._usable_width_dxa(document)
        raw_widths = [column.get("width_dxa") for column in columns]
        if all(width is None for width in raw_widths):
            base, remainder = divmod(usable_dxa, len(columns))
            widths = [base] * len(columns)
            widths[-1] += remainder
        elif any(width is None for width in raw_widths):
            raise AssessmentDocxRendererError(
                "column widths must be all present or all omitted"
            )
        else:
            widths = [
                int(
                    _number(
                        width,
                        "column.width_dxa",
                        minimum=240,
                        maximum=20000,
                    )
                )
                for width in raw_widths
            ]
            if sum(widths) != usable_dxa:
                raise AssessmentDocxRendererError(
                    "column widths must equal usable page width"
                )

        table = document.add_table(rows=1, cols=len(columns))
        table.autofit = False
        table.style = "Table Grid"
        header = table.rows[0]
        self._mark_header_row(header)
        for column_index, column in enumerate(columns):
            heading = str(
                column.get("heading", column.get("column_code", ""))
            )
            self._set_cell(
                header.cells[column_index],
                heading,
                styles,
                width=widths[column_index],
                bold=True,
                alignment=str(column.get("alignment", "CENTER")),
            )

        for row_index, item in enumerate(rows):
            row_data = _object(item, f"table row {row_index}")
            row = table.add_row()
            for column_index, column in enumerate(columns):
                path = _text(column.get("value_path"), "value_path")
                self._set_cell(
                    row.cells[column_index],
                    _display(_value_at_path(row_data, path)),
                    styles,
                    width=widths[column_index],
                    bold=False,
                    alignment=str(column.get("alignment", "LEFT")),
                )
        self._apply_table_geometry(table, widths)

    def _add_paragraph(
        self,
        document,
        text,
        styles,
        *,
        bold,
        alignment,
        size_key="font_size",
        keep_with_next=False,
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = self._alignment(alignment)
        paragraph.paragraph_format.keep_with_next = keep_with_next
        run = paragraph.add_run(text)
        run.bold = bold
        _set_run_font(
            run,
            _text(
                styles.get("font_family", "Times New Roman"),
                "styles.font_family",
            ),
            _number(
                styles.get(size_key, styles.get("font_size", 12)),
                f"styles.{size_key}",
                minimum=8,
                maximum=32,
            ),
        )
        return paragraph

    def _set_cell(
        self,
        cell,
        text,
        styles,
        *,
        width,
        bold,
        alignment,
    ) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = self._alignment(alignment)
        run = paragraph.add_run(text)
        run.bold = bold
        _set_run_font(
            run,
            _text(
                styles.get("font_family", "Times New Roman"),
                "styles.font_family",
            ),
            _number(
                styles.get("table_font_size", styles.get("font_size", 12)),
                "styles.table_font_size",
                minimum=8,
                maximum=20,
            ),
        )
        tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
        tc_width.set(qn("w:w"), str(width))
        tc_width.set(qn("w:type"), "dxa")
        _set_cell_margins(
            cell,
            {"top": 80, "bottom": 80, "start": 120, "end": 120},
        )

    @staticmethod
    def _usable_width_dxa(document: DocumentObject) -> int:
        section = document.sections[-1]
        return round(
            (
                section.page_width
                - section.left_margin
                - section.right_margin
            )
            / 635
        )

    @staticmethod
    def _mark_header_row(row) -> None:
        properties = row._tr.get_or_add_trPr()
        marker = properties.find(qn("w:tblHeader"))
        if marker is None:
            marker = OxmlElement("w:tblHeader")
            properties.append(marker)
        marker.set(qn("w:val"), "true")

    @staticmethod
    def _apply_table_geometry(table, widths: list[int]) -> None:
        properties = table._tbl.tblPr
        table_width = properties.first_child_found_in("w:tblW")
        if table_width is None:
            table_width = OxmlElement("w:tblW")
            properties.insert(0, table_width)
        table_width.set(qn("w:w"), str(sum(widths)))
        table_width.set(qn("w:type"), "dxa")

        indent = properties.first_child_found_in("w:tblInd")
        if indent is None:
            indent = OxmlElement("w:tblInd")
            properties.append(indent)
        indent.set(qn("w:w"), "120")
        indent.set(qn("w:type"), "dxa")

        grid_columns = list(table._tbl.tblGrid.gridCol_lst)
        for column, width in zip(grid_columns, widths):
            column.set(qn("w:w"), str(width))

        for row in table.rows:
            for index, cell in enumerate(row.cells):
                tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_width.set(qn("w:w"), str(widths[index]))
                tc_width.set(qn("w:type"), "dxa")

    @staticmethod
    def _alignment(value: object):
        normalized = str(value).strip().upper()
        if normalized not in _ALIGNMENTS:
            raise AssessmentDocxRendererError(
                "unsupported paragraph alignment"
            )
        return _ALIGNMENTS[normalized]
