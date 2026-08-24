from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document

from document_intelligence.contracts import (
    DocumentAnalysis,
)
from document_intelligence.hybrid_analyzer import (
    HybridAnalysisResult,
)


@dataclass(frozen=True)
class LessonPlanIntelligencePreview:
    source: Path
    document_text: str
    analysis: DocumentAnalysis
    ai_used: bool
    ai_failed: bool

    @property
    def has_proposals(self) -> bool:
        return bool(
            self.analysis.proposals
        )


class LessonPlanIntelligencePreviewService:
    """
    Read-only intelligence preview for lesson-plan DOCX.

    This service:
    - reads the source document;
    - extracts text;
    - runs document intelligence;
    - never modifies the DOCX;
    - never applies proposals automatically.
    """

    def __init__(
        self,
        *,
        analyzer,
    ) -> None:
        self._analyzer = analyzer

    def preview(
        self,
        *,
        source: Path,
    ) -> LessonPlanIntelligencePreview:
        source = Path(source)

        if not source.exists():
            raise FileNotFoundError(
                source
            )

        if source.suffix.casefold() != ".docx":
            raise ValueError(
                "source must be a DOCX document"
            )

        document_text = (
            self._extract_document_text(
                source
            )
        )

        result = self._analyzer.analyze(
            document_text=document_text
        )

        if not isinstance(
            result,
            HybridAnalysisResult,
        ):
            raise TypeError(
                "analyzer must return "
                "HybridAnalysisResult"
            )

        return LessonPlanIntelligencePreview(
            source=source,
            document_text=document_text,
            analysis=result.analysis,
            ai_used=result.ai_used,
            ai_failed=result.ai_error is not None,
        )

    @staticmethod
    def _extract_document_text(
        source: Path,
    ) -> str:
        document = Document(source)

        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()

                    if text:
                        parts.append(text)

        return "\n".join(parts)
