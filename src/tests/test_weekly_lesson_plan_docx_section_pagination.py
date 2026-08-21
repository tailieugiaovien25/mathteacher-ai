from pathlib import Path
import importlib.util

from docx import Document

from lesson_planning_v2.services.weekly_lesson_plan_docx_renderer import (
    WeeklyLessonPlanDocxRenderer,
)
from lesson_planning_v2.weekly_lesson_plan_docx_layout import (
    WeeklyLessonPlanDocxLayoutProfile,
)


def _fixture_module():
    path = Path(
        "src/tests/"
        "test_weekly_lesson_plan_docx_display_e2e.py"
    )

    spec = importlib.util.spec_from_file_location(
        "weekly_pagination_fixture",
        path,
    )

    if spec is None or spec.loader is None:
        raise RuntimeError(
            "fixture module cannot be loaded"
        )

    module = importlib.util.module_from_spec(
        spec
    )

    spec.loader.exec_module(
        module
    )

    return module


def _render(tmp_path):
    fixture = _fixture_module()

    output = (
        tmp_path
        / "section-pagination.docx"
    )

    WeeklyLessonPlanDocxRenderer().render(
        document=fixture._document(),
        output_path=output,
        presentation_profile=(
            fixture._presentation()
        ),
        layout_profile=(
            WeeklyLessonPlanDocxLayoutProfile
            .default()
        ),
    )

    return Document(output)


def _period_titles(docx):
    return [
        paragraph
        for paragraph in docx.paragraphs
        if paragraph.text.startswith(
            "TI\u1ebeT "
        )
    ]


def _content_headings(docx):
    return [
        paragraph
        for paragraph in docx.paragraphs
        if paragraph.text.strip()
        in (
            "I. M\u1ee5c ti\u00eau",
            (
                "II. Thi\u1ebft b\u1ecb "
                "v\u00e0 h\u1ecdc li\u1ec7u"
            ),
            (
                "III. Ti\u1ebfn tr\u00ecnh "
                "d\u1ea1y h\u1ecdc"
            ),
        )
    ]


def _period_metadata_tables(docx):
    result = []

    for table in docx.tables:
        text = "\n".join(
            cell.text
            for row in table.rows
            for cell in row.cells
        )

        if (
            "Ti\u1ebft PPCT:" in text
            and "Ng\u00e0y so\u1ea1n:" in text
            and "Ng\u00e0y d\u1ea1y:" in text
        ):
            result.append(table)

    return result


def test_each_period_title_is_kept_with_metadata(
    tmp_path,
):
    docx = _render(tmp_path)

    titles = _period_titles(docx)

    assert len(titles) == 3

    for paragraph in titles:
        assert (
            paragraph.paragraph_format
            .keep_with_next
            is True
        )


def test_period_metadata_paragraphs_continue_keep_chain(
    tmp_path,
):
    docx = _render(tmp_path)

    tables = _period_metadata_tables(
        docx
    )

    assert len(tables) == 3

    for table in tables:
        non_empty_paragraphs = [
            paragraph
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            if paragraph.text.strip()
        ]

        assert non_empty_paragraphs

        for paragraph in non_empty_paragraphs:
            assert (
                paragraph.paragraph_format
                .keep_with_next
                is True
            )


def test_first_content_heading_is_kept_with_body(
    tmp_path,
):
    docx = _render(tmp_path)

    first_headings = [
        paragraph
        for paragraph in docx.paragraphs
        if paragraph.text.strip()
        == "I. M\u1ee5c ti\u00eau"
    ]

    assert len(first_headings) == 3

    for paragraph in first_headings:
        assert (
            paragraph.paragraph_format
            .keep_with_next
            is True
        )


def test_all_content_headings_are_kept_with_following_content(
    tmp_path,
):
    docx = _render(tmp_path)

    headings = _content_headings(
        docx
    )

    assert len(headings) == 9

    for paragraph in headings:
        assert (
            paragraph.paragraph_format
            .keep_with_next
            is True
        )


def test_body_font_remains_times_new_roman_14(
    tmp_path,
):
    docx = _render(tmp_path)

    normal = docx.styles["Normal"]

    assert (
        normal.font.name
        == "Times New Roman"
    )

    assert (
        normal.font.size.pt
        == 14
    )


def test_no_forced_page_break_between_periods(
    tmp_path,
):
    docx = _render(tmp_path)

    titles = _period_titles(docx)

    assert len(titles) == 3

    for paragraph in titles:
        assert (
            paragraph.paragraph_format
            .page_break_before
            in (
                None,
                False,
            )
        )
