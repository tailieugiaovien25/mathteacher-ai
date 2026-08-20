import os
from datetime import date
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from lxml import etree

from document_standardization import (
    LessonPlanDocumentContextApplier,
)
from educational_planning_v2.models import (
    TeachingSession,
)
from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)


ENV_NAME = "REAL_TEACHER_DOCX"


def teacher_docx() -> Path:
    raw = os.environ.get(
        ENV_NAME
    )

    if not raw:
        pytest.skip(
            f"{ENV_NAME} is not configured"
        )

    path = Path(raw)

    if not path.exists():
        pytest.skip(
            f"Teacher DOCX does not exist: {path}"
        )

    return path


def make_context():
    return ScheduledLessonContext(
        teaching_date=date(
            2026,
            9,
            30,
        ),
        drafting_date=date(
            2026,
            9,
            29,
        ),
        class_id="6A3",
        subject_ref="MATHEMATICS",
        component_ref="ALGEBRA",
        curriculum_period=20,
        lesson_id="REAL-LESSON-001",
        lesson_title=(
            "Thứ tự thực hiện "
            "các phép tính"
        ),
        session=TeachingSession.MORNING,
        timetable_period=1,
        period_in_lesson=1,
    )


def count_omml(path: Path) -> int:
    with ZipFile(path) as archive:
        xml = archive.read(
            "word/document.xml"
        )

    root = etree.fromstring(xml)

    namespace = {
        "m": (
            "http://schemas.openxmlformats.org/"
            "officeDocument/2006/math"
        )
    }

    return len(
        root.xpath(
            ".//m:oMath",
            namespaces=namespace,
        )
    )


def paragraph_texts(path: Path):
    document = Document(path)

    return [
        paragraph.text
        for paragraph
        in document.paragraphs
    ]


def test_context_applier_preserves_real_teacher_document_structure(
    tmp_path,
):
    source = teacher_docx()
    output = (
        tmp_path
        / "KHBD.DAI6.TUAN4.modified.docx"
    )

    before = Document(source)

    before_paragraph_count = len(
        before.paragraphs
    )
    before_table_count = len(
        before.tables
    )
    before_shape_count = len(
        before.inline_shapes
    )
    before_omml_count = count_omml(
        source
    )

    source_bytes = source.read_bytes()

    result = (
        LessonPlanDocumentContextApplier()
        .apply(
            source=source,
            output=output,
            context=make_context(),
        )
    )

    assert output.exists()

    # Source file must remain untouched.
    assert (
        source.read_bytes()
        == source_bytes
    )

    after = Document(output)

    assert (
        len(after.paragraphs)
        == before_paragraph_count
    )

    assert (
        len(after.tables)
        == before_table_count
    )

    assert (
        len(after.inline_shapes)
        == before_shape_count
    )

    assert (
        count_omml(output)
        == before_omml_count
    )

    assert result is not None


def test_context_applier_updates_first_lesson_metadata(
    tmp_path,
):
    source = teacher_docx()
    output = (
        tmp_path
        / "KHBD.DAI6.TUAN4.modified.docx"
    )

    LessonPlanDocumentContextApplier().apply(
        source=source,
        output=output,
        context=make_context(),
    )

    text = "\n".join(
        paragraph_texts(output)
    )

    assert "29/09/2026" in text
    assert "30/09/2026" in text
    assert "6A3" in text

    assert (
        "TIẾT 20"
        in text
        or "Tiết 20"
        in text
    )


def test_context_applier_preserves_second_lesson_marker(
    tmp_path,
):
    source = teacher_docx()
    output = (
        tmp_path
        / "KHBD.DAI6.TUAN4.modified.docx"
    )

    LessonPlanDocumentContextApplier().apply(
        source=source,
        output=output,
        context=make_context(),
    )

    text = "\n".join(
        paragraph_texts(output)
    )

    assert (
        "Ngày dạy: 02/10/2025"
        in text
    )

    assert "TIẾT 12" in text

    assert (
        "LUYỆN TẬP CHUNG"
        in text
    )


def test_context_applier_preserves_math_omml(
    tmp_path,
):
    source = teacher_docx()
    output = (
        tmp_path
        / "KHBD.DAI6.TUAN4.modified.docx"
    )

    before = count_omml(source)

    LessonPlanDocumentContextApplier().apply(
        source=source,
        output=output,
        context=make_context(),
    )

    after = count_omml(output)

    assert before == 9
    assert after == before
