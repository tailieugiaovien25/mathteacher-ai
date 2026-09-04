
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image

import portal_v2.ui.weekly_schedule_streamlit as module


def image_bytes(
    width=1200,
    height=600,
):
    image = Image.new(
        "RGB",
        (
            width,
            height,
        ),
        "white",
    )

    buffer = BytesIO()

    image.save(
        buffer,
        format="PNG",
    )

    buffer.seek(0)

    return buffer


def ratio(shape):
    return (
        float(shape.width)
        / float(shape.height)
    )


def test_body_image_fits_page_and_keeps_ratio():
    document = Document()

    paragraph = (
        document.add_paragraph()
    )

    run = paragraph.add_run()

    run.add_picture(
        image_bytes(
            1200,
            600,
        ),
        width=Inches(1),
    )

    before_shape = (
        document.inline_shapes[0]
    )

    before_ratio = ratio(
        before_shape
    )

    source = BytesIO()

    document.save(
        source
    )

    result = (
        module._mt_autofit_images_bytes(
            source.getvalue()
        )
    )

    processed = Document(
        BytesIO(result)
    )

    shape = (
        processed.inline_shapes[0]
    )

    page_width, _ = (
        module._mt_image_page_content_box(
            processed
        )
    )

    assert (
        int(shape.width)
        <= int(
            page_width
            * 0.95
        )
    )

    assert abs(
        ratio(shape)
        - before_ratio
    ) < 0.02

    # Small images remain at their authored size; autofit does not upscale.
    assert int(shape.width) == int(before_shape.width)
    assert processed.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_oversized_body_image_is_shrunk_to_printable_page():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(
        image_bytes(2400, 1200),
        width=Inches(12),
    )
    source = BytesIO()
    document.save(source)

    processed = Document(
        BytesIO(
            module._mt_autofit_images_bytes(source.getvalue())
        )
    )
    shape = processed.inline_shapes[0]
    page_width, _ = module._mt_image_page_content_box(processed)

    assert int(shape.width) <= int(page_width * 0.95)
    assert int(shape.width) < int(Inches(12))
    assert processed.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_table_image_fits_cell_and_keeps_ratio():
    document = Document()

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.columns[0].width = (
        Inches(2.5)
    )

    table.columns[1].width = (
        Inches(2.5)
    )

    cell = table.cell(
        0,
        0,
    )

    paragraph = (
        cell.paragraphs[0]
    )

    run = paragraph.add_run()

    run.add_picture(
        image_bytes(
            1000,
            500,
        ),
        width=Inches(5),
    )

    source = BytesIO()

    document.save(
        source
    )

    result = (
        module._mt_autofit_images_bytes(
            source.getvalue()
        )
    )

    processed = Document(
        BytesIO(result)
    )

    shape = (
        processed.inline_shapes[0]
    )

    assert ratio(shape) > 1.9
    assert ratio(shape) < 2.1

    # Must be substantially smaller than original 5 inches.
    assert (
        int(shape.width)
        < int(
            Inches(5)
        )
    )


def test_multiple_images_share_available_width():
    document = Document()

    paragraph = (
        document.add_paragraph()
    )

    run1 = paragraph.add_run()

    run1.add_picture(
        image_bytes(
            800,
            600,
        ),
        width=Inches(1),
    )

    run2 = paragraph.add_run()

    run2.add_picture(
        image_bytes(
            800,
            600,
        ),
        width=Inches(1),
    )

    source = BytesIO()

    document.save(
        source
    )

    result = (
        module._mt_autofit_images_bytes(
            source.getvalue()
        )
    )

    processed = Document(
        BytesIO(result)
    )

    assert (
        len(
            processed.inline_shapes
        )
        == 2
    )

    page_width, _ = (
        module._mt_image_page_content_box(
            processed
        )
    )

    total_width = sum(
        int(shape.width)
        for shape
        in processed.inline_shapes
    )

    assert (
        total_width
        <= page_width
    )


def test_text_and_metadata_survive_image_fit():
    document = Document()

    document.add_paragraph(
        (
            "Ng\u00e0y so\u1ea1n: "
            "25/09/2026"
        )
    )

    document.add_paragraph(
        (
            "Ng\u00e0y d\u1ea1y:\n"
            "7A1 - 29/09/2026\n"
            "7A2 - 30/09/2026"
        )
    )

    paragraph = (
        document.add_paragraph(
            "H\u00ecnh minh h\u1ecda: "
        )
    )

    paragraph.add_run().add_picture(
        image_bytes(),
        width=Inches(1),
    )

    document.add_paragraph(
        (
            "N\u1ed9i dung b\u00e0i h\u1ecdc "
            "ph\u1ea3i gi\u1eef nguy\u00ean."
        )
    )

    source = BytesIO()

    document.save(
        source
    )

    result = (
        module._mt_autofit_images_bytes(
            source.getvalue()
        )
    )

    processed = Document(
        BytesIO(result)
    )

    text = "\n".join(
        paragraph.text
        for paragraph
        in processed.paragraphs
    )

    assert (
        "Ng\u00e0y so\u1ea1n: 25/09/2026"
        in text
    )

    assert (
        "7A1 - 29/09/2026"
        in text
    )

    assert (
        "7A2 - 30/09/2026"
        in text
    )

    assert (
        "N\u1ed9i dung b\u00e0i h\u1ecdc "
        "ph\u1ea3i gi\u1eef nguy\u00ean."
        in text
    )


def test_document_without_images_is_identical_bytes():
    document = Document()

    document.add_paragraph(
        "Kh\u00f4ng c\u00f3 h\u00ecnh \u1ea3nh."
    )

    source = BytesIO()

    document.save(
        source
    )

    original = source.getvalue()

    result = (
        module._mt_autofit_images_bytes(
            original
        )
    )

    assert result == original


def test_control_panel_contains_image_option():
    source = open(
        "src/portal_v2/ui/"
        "weekly_schedule_streamlit.py",
        encoding="utf-8",
    ).read()

    assert (
        "standardization_image_autofit_enabled"
        in source
    )

    assert "def _mt_autofit_images_bytes(" in source
    assert "# Image autofit runs LAST." in source


def test_picture_table_row_fixed_height_is_removed():
    document = Document()
    table = document.add_table(rows=1, cols=1)
    row = table.rows[0]
    row.height = Inches(6)
    row.cells[0].paragraphs[0].add_run().add_picture(
        image_bytes(1200, 600),
        width=Inches(2),
    )
    source = BytesIO()
    document.save(source)

    processed = Document(
        BytesIO(module._mt_autofit_images_bytes(source.getvalue()))
    )

    assert not processed.tables[0].rows[0]._tr.xpath(
        "./w:trPr/w:trHeight"
    )
    assert (
        processed.tables[0].cell(0, 0).paragraphs[0].alignment
        == WD_ALIGN_PARAGRAPH.CENTER
    )


def test_all_table_rows_reflow_and_redundant_empty_paragraphs_collapse():
    document = Document()
    table = document.add_table(rows=1, cols=1)
    row = table.rows[0]
    row.height = Inches(8)
    cell = row.cells[0]
    cell.paragraphs[0].text = "Nội dung đầu."
    cell.add_paragraph("")
    cell.add_paragraph("")
    cell.add_paragraph("")
    cell.add_paragraph("Nội dung tiếp theo.")
    source = BytesIO()
    document.save(source)

    result = module._mt_format_lesson_document_layout_bytes(
        source.getvalue(),
        ensure_end_rule=False,
    )
    processed = Document(BytesIO(result))
    processed_cell = processed.tables[0].cell(0, 0)

    assert not processed.tables[0].rows[0]._tr.xpath(
        "./w:trPr/w:trHeight"
    )
    assert [p.text for p in processed_cell.paragraphs] == [
        "Nội dung đầu.",
        "",
        "Nội dung tiếp theo.",
    ]


def test_image_wrapper_updates_legacy_tuple_result(monkeypatch):
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(
        image_bytes(2400, 1200),
        width=Inches(12),
    )
    source = BytesIO()
    document.save(source)

    def base_processor(*args, **kwargs):
        return ("giao-an.docx", source.getvalue(), ())

    monkeypatch.setattr(
        module,
        "_mt_original_process_lesson_plan_upload_3c",
        base_processor,
    )
    monkeypatch.setattr(
        module,
        "st",
        type(
            "FakeStreamlit",
            (),
            {
                "session_state": {
                    module._MT_IMAGE_AUTOFIT_ENABLED: True,
                }
            },
        )(),
    )

    result = module._process_lesson_plan_upload()
    processed = Document(BytesIO(result[1]))

    assert result[0] == "giao-an.docx"
    assert result[2] == ()
    assert int(processed.inline_shapes[0].width) < int(Inches(12))
    assert processed.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
