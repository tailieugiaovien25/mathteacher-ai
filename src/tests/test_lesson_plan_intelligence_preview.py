from pathlib import Path

import pytest
from docx import Document

from document_intelligence import (
    build_document_analyzer,
)
from document_intelligence.lesson_plan_preview import (
    LessonPlanIntelligencePreviewService,
)


def make_docx(
    path: Path,
) -> None:
    document = Document()

    document.add_paragraph(
        "Ngày soạn: 07/09/2026"
    )

    document.add_paragraph(
        "Ngày dạy: 08/09/2026"
    )

    document.add_paragraph(
        "Lớp: 8A2"
    )

    document.add_paragraph(
        "Tiết 4. Bài 3. Đơn thức"
    )

    document.add_paragraph(
        "Nội dung bài học phải được giữ nguyên."
    )

    table = document.add_table(
        rows=1,
        cols=1,
    )

    table.cell(
        0,
        0,
    ).text = "Hoạt động học tập"

    document.save(path)


def test_preview_reads_docx_without_modifying_it(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    source = tmp_path / "lesson.docx"

    make_docx(source)

    original = source.read_bytes()

    service = (
        LessonPlanIntelligencePreviewService(
            analyzer=build_document_analyzer()
        )
    )

    preview = service.preview(
        source=source
    )

    assert source.read_bytes() == original

    assert (
        "Ngày soạn: 07/09/2026"
        in preview.document_text
    )

    assert (
        "Hoạt động học tập"
        in preview.document_text
    )

    assert preview.ai_used is False
    assert preview.ai_failed is False
    assert preview.has_proposals is True


def test_preview_recognizes_lesson_metadata(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    source = tmp_path / "lesson.docx"

    make_docx(source)

    service = (
        LessonPlanIntelligencePreviewService(
            analyzer=build_document_analyzer()
        )
    )

    preview = service.preview(
        source=source
    )

    values = {
        proposal.field.value:
        proposal.value
        for proposal
        in preview.analysis.proposals
    }

    assert (
        values["drafting_date"]
        == "07/09/2026"
    )

    assert (
        values["teaching_date"]
        == "08/09/2026"
    )

    assert (
        values["class_name"]
        == "8A2"
    )

    assert (
        values["curriculum_period"]
        == "4"
    )

    assert (
        values["lesson_title"]
        == "Đơn thức"
    )


def test_preview_rejects_non_docx(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    source = tmp_path / "lesson.txt"

    source.write_text(
        "test",
        encoding="utf-8",
    )

    service = (
        LessonPlanIntelligencePreviewService(
            analyzer=build_document_analyzer()
        )
    )

    with pytest.raises(
        ValueError,
        match="DOCX",
    ):
        service.preview(
            source=source
        )


def test_preview_rejects_missing_file(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    service = (
        LessonPlanIntelligencePreviewService(
            analyzer=build_document_analyzer()
        )
    )

    with pytest.raises(
        FileNotFoundError
    ):
        service.preview(
            source=(
                tmp_path
                / "missing.docx"
            )
        )
