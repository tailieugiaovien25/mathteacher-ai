from datetime import date
from pathlib import Path

from docx import Document

from document_standardization.lesson_plan_document_context_applier import (
    LessonPlanDocumentContextApplier,
)
from educational_planning_v2.models.weekly_teaching_schedule import (
    TeachingSession,
)
from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)


def _context():
    return ScheduledLessonContext(
        subject_ref="TOAN",
        component_ref=None,
        lesson_id="TOAN6-012",
        class_id="6A2",
        curriculum_period=12,
        lesson_title=(
            "Ph\u00e2n s\u1ed1"
        ),
        drafting_date=date(
            2026,
            9,
            10,
        ),
        teaching_date=date(
            2026,
            9,
            15,
        ),
        period_in_lesson=1,
        session=TeachingSession.MORNING,
        timetable_period=2,
    )


def _build_realistic_source(
    path: Path,
) -> None:

    document = Document()

    # -----------------------------------------------------
    # 1. Header/footer must survive.
    # -----------------------------------------------------

    section = document.sections[0]

    header = section.header

    header.paragraphs[0].text = (
        "MATH-TEACHER-AI HEADER"
    )

    footer = section.footer

    footer.paragraphs[0].text = (
        "MATH-TEACHER-AI FOOTER"
    )


    # -----------------------------------------------------
    # 2. Multi-metadata paragraph.
    # -----------------------------------------------------

    paragraph = document.add_paragraph()

    run1 = paragraph.add_run(
        "Ng\u00e0y so\u1ea1n:"
    )
    run1.bold = True

    run2 = paragraph.add_run(
        " 01/01/2020    "
    )
    run2.italic = True

    run3 = paragraph.add_run(
        "Ng\u00e0y d\u1ea1y:"
    )
    run3.bold = True

    run4 = paragraph.add_run(
        " 02/01/2020    "
    )
    run4.italic = True

    run5 = paragraph.add_run(
        "L\u1edbp:"
    )
    run5.bold = True

    run6 = paragraph.add_run(
        " 8A1"
    )
    run6.italic = True


    # -----------------------------------------------------
    # 3. Composite heading.
    # -----------------------------------------------------

    heading = document.add_paragraph()

    heading_run = heading.add_run(
        "Ti\u1ebft 9. "
        "B\u00c0I 7. PH\u00c2N S\u1ed0 C\u0168"
    )

    heading_run.bold = True


    # -----------------------------------------------------
    # 4. Paired-cell metadata.
    # -----------------------------------------------------

    table = document.add_table(
        rows=1,
        cols=2,
    )

    label_cell = table.cell(
        0,
        0,
    )

    label_paragraph = (
        label_cell.paragraphs[0]
    )

    label_run = (
        label_paragraph.add_run(
            "Ng\u00e0y gi\u1ea3ng:"
        )
    )

    label_run.bold = True

    value_cell = table.cell(
        0,
        1,
    )

    value_paragraph = (
        value_cell.paragraphs[0]
    )

    value_run = (
        value_paragraph.add_run(
            "03/01/2020"
        )
    )

    value_run.italic = True


    # -----------------------------------------------------
    # 5. Ordinary body content must survive.
    # -----------------------------------------------------

    document.add_paragraph(
        "I. M\u1ee4C TI\u00caU"
    )

    body = document.add_paragraph()

    body_run_1 = body.add_run(
        "H\u1ecdc sinh "
    )
    body_run_1.bold = True

    body_run_2 = body.add_run(
        "nh\u1eadn bi\u1ebft "
    )
    body_run_2.italic = True

    body_run_3 = body.add_run(
        "\u0111\u01b0\u1ee3c ph\u00e2n s\u1ed1."
    )

    document.add_paragraph(
        "II. THI\u1ebeT B\u1eca V\u00c0 H\u1eccC LI\u1ec6U"
    )

    document.add_paragraph(
        "SGK, b\u1ea3ng ph\u1ee5, m\u00e1y chi\u1ebfu."
    )

    document.add_paragraph(
        "III. TI\u1ebeN TR\u00ccNH D\u1ea0Y H\u1eccC"
    )

    document.add_paragraph(
        "Ho\u1ea1t \u0111\u1ed9ng 1: M\u1edf \u0111\u1ea7u"
    )


    # -----------------------------------------------------
    # 6. Extra table representing lesson content.
    # Must not be interpreted as metadata.
    # -----------------------------------------------------

    content_table = document.add_table(
        rows=2,
        cols=2,
    )

    content_table.cell(
        0,
        0,
    ).text = (
        "Ho\u1ea1t \u0111\u1ed9ng"
    )

    content_table.cell(
        0,
        1,
    ).text = (
        "S\u1ea3n ph\u1ea9m"
    )

    content_table.cell(
        1,
        0,
    ).text = (
        "B\u00e0i t\u1eadp 1"
    )

    content_table.cell(
        1,
        1,
    ).text = (
        "C\u00e2u tr\u1ea3 l\u1eddi"
    )


    document.save(
        path
    )


def test_real_docx_metadata_overlay_end_to_end(
    tmp_path: Path,
):

    source = (
        tmp_path
        / "source.docx"
    )

    output = (
        tmp_path
        / "output.docx"
    )

    _build_realistic_source(
        source
    )


    # -----------------------------------------------------
    # Snapshot source structural facts.
    # -----------------------------------------------------

    before = Document(
        source
    )

    before_paragraph_count = len(
        before.paragraphs
    )

    before_table_count = len(
        before.tables
    )

    before_section_count = len(
        before.sections
    )

    before_header = (
        before.sections[0]
        .header
        .paragraphs[0]
        .text
    )

    before_footer = (
        before.sections[0]
        .footer
        .paragraphs[0]
        .text
    )


    # -----------------------------------------------------
    # Execute the REAL production compatibility entry.
    # -----------------------------------------------------

    result = (
        LessonPlanDocumentContextApplier()
        .apply(
            source=source,
            output=output,
            context=_context(),
        )
    )


    assert output.exists()

    assert output.stat().st_size > 0


    # -----------------------------------------------------
    # Re-open generated DOCX from disk.
    # -----------------------------------------------------

    processed = Document(
        output
    )


    # -----------------------------------------------------
    # A. Multi-metadata paragraph.
    # -----------------------------------------------------

    first = (
        processed.paragraphs[0]
    )

    expected_first = (
        "Ng\u00e0y so\u1ea1n: "
        "10/09/2026    "
        "Ng\u00e0y d\u1ea1y: "
        "15/09/2026    "
        "L\u1edbp: 6A2"
    )

    assert (
        first.text
        == expected_first
    )


    # Labels retain bold.
    assert (
        first.runs[0].bold
        is True
    )

    assert (
        first.runs[2].bold
        is True
    )

    assert (
        first.runs[4].bold
        is True
    )


    # Values retain italic.
    assert (
        first.runs[1].italic
        is True
    )

    assert (
        first.runs[3].italic
        is True
    )

    assert (
        first.runs[5].italic
        is True
    )


    # -----------------------------------------------------
    # B. Composite heading.
    # -----------------------------------------------------

    heading = (
        processed.paragraphs[1]
    )

    expected_heading = (
        "Ti\u1ebft 12. "
        "B\u00c0I 7. PH\u00c2N S\u1ed0"
    )

    assert (
        heading.text
        == expected_heading
    )

    assert (
        heading.runs[0].bold
        is True
    )


    # -----------------------------------------------------
    # C. Paired-cell teaching date.
    # -----------------------------------------------------

    metadata_table = (
        processed.tables[0]
    )

    assert (
        metadata_table.cell(
            0,
            0,
        ).text
        == "Ng\u00e0y gi\u1ea3ng:"
    )

    assert (
        metadata_table.cell(
            0,
            1,
        ).text
        == "15/09/2026"
    )

    assert (
        metadata_table.cell(
            0,
            0,
        ).paragraphs[
            0
        ].runs[
            0
        ].bold
        is True
    )

    assert (
        metadata_table.cell(
            0,
            1,
        ).paragraphs[
            0
        ].runs[
            0
        ].italic
        is True
    )


    # -----------------------------------------------------
    # D. Unrelated body text must survive.
    # -----------------------------------------------------

    body_text = [
        paragraph.text
        for paragraph
        in processed.paragraphs
    ]

    assert (
        "I. M\u1ee4C TI\u00caU"
        in body_text
    )

    assert (
        "II. THI\u1ebeT B\u1eca V\u00c0 H\u1eccC LI\u1ec6U"
        in body_text
    )

    assert (
        "III. TI\u1ebeN TR\u00ccNH D\u1ea0Y H\u1eccC"
        in body_text
    )

    assert (
        "Ho\u1ea1t \u0111\u1ed9ng 1: M\u1edf \u0111\u1ea7u"
        in body_text
    )


    # -----------------------------------------------------
    # E. Unrelated body formatting must survive.
    # -----------------------------------------------------

    objective_body = next(
        paragraph
        for paragraph
        in processed.paragraphs
        if (
            "H\u1ecdc sinh"
            in paragraph.text
        )
    )

    assert (
        objective_body.runs[
            0
        ].bold
        is True
    )

    assert (
        objective_body.runs[
            1
        ].italic
        is True
    )


    # -----------------------------------------------------
    # F. Non-metadata lesson table must survive.
    # -----------------------------------------------------

    content_table = (
        processed.tables[
            1
        ]
    )

    assert (
        content_table.cell(
            0,
            0,
        ).text
        == "Ho\u1ea1t \u0111\u1ed9ng"
    )

    assert (
        content_table.cell(
            0,
            1,
        ).text
        == "S\u1ea3n ph\u1ea9m"
    )

    assert (
        content_table.cell(
            1,
            0,
        ).text
        == "B\u00e0i t\u1eadp 1"
    )

    assert (
        content_table.cell(
            1,
            1,
        ).text
        == "C\u00e2u tr\u1ea3 l\u1eddi"
    )


    # -----------------------------------------------------
    # G. Header/footer survive.
    # -----------------------------------------------------

    assert (
        processed.sections[
            0
        ].header.paragraphs[
            0
        ].text
        == before_header
    )

    assert (
        processed.sections[
            0
        ].footer.paragraphs[
            0
        ].text
        == before_footer
    )


    # -----------------------------------------------------
    # H. Major structural counts survive.
    # -----------------------------------------------------

    assert (
        len(
            processed.paragraphs
        )
        == before_paragraph_count
    )

    assert (
        len(
            processed.tables
        )
        == before_table_count
    )

    assert (
        len(
            processed.sections
        )
        == before_section_count
    )


    # -----------------------------------------------------
    # I. Source file remains unchanged.
    # -----------------------------------------------------

    original = Document(
        source
    )

    assert (
        original.paragraphs[
            0
        ].text
        == (
            "Ng\u00e0y so\u1ea1n: "
            "01/01/2020    "
            "Ng\u00e0y d\u1ea1y: "
            "02/01/2020    "
            "L\u1edbp: 8A1"
        )
    )

    assert (
        original.paragraphs[
            1
        ].text
        == (
            "Ti\u1ebft 9. "
            "B\u00c0I 7. "
            "PH\u00c2N S\u1ed0 C\u0168"
        )
    )


    # -----------------------------------------------------
    # J. Context result reports success.
    # -----------------------------------------------------

    assert (
        result.unresolved_fields
        == ()
    )

    assert (
        set(
            result.applied_fields
        )
        >= {
            "drafting_date",
            "teaching_date",
            "class_id",
            "curriculum_period",
            "lesson_title",
        }
    )
