from io import BytesIO
from types import SimpleNamespace

from docx import Document

from lesson_planning_v2.services.lesson_plan_workspace_v1_service import (
    LessonPlanDocxTextImporter,
    LessonPlanLibrarySourceService,
)


class FakeRepository:
    def list_all(self):
        return (
            SimpleNamespace(
                document_id="doc-1",
                title="Giáo án A",
                file_name="lesson.docx",
                mime_type=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml."
                    "document"
                ),
                storage_file_id="file-1",
                web_view_link=None,
                academic_year="2026-2027",
                subject="math",
                grade_level="6",
                class_name="6A2",
            ),
            SimpleNamespace(
                document_id="doc-2",
                title="PDF",
                file_name="other.pdf",
                mime_type="application/pdf",
                storage_file_id="file-2",
                web_view_link=None,
                academic_year="2026-2027",
                subject="math",
                grade_level="6",
                class_name=None,
            ),
        )


class FakeStorage:
    def __init__(self, content):
        self.content = content
        self.calls = []

    def download(self, file_id):
        self.calls.append(file_id)
        return self.content


def _docx():
    document = Document()
    document.add_paragraph(
        "I. MỤC TIÊU"
    )
    document.add_paragraph(
        "Nhận biết nội dung."
    )

    document.add_paragraph(
        "II. THIẾT BỊ VÀ HỌC LIỆU"
    )
    document.add_paragraph(
        "Máy chiếu."
    )

    document.add_paragraph(
        "III. TIẾN TRÌNH DẠY HỌC"
    )
    document.add_paragraph(
        "Hoạt động học tập."
    )

    output = BytesIO()
    document.save(output)

    return output.getvalue()


def test_library_lists_only_docx():
    service = (
        LessonPlanLibrarySourceService(
            repository=FakeRepository(),
        )
    )

    items = service.list_docx(
        academic_year="2026-2027",
    )

    assert len(items) == 1
    assert items[0].title == "Giáo án A"


def test_library_can_load_existing_plan():
    storage = FakeStorage(
        _docx()
    )

    service = (
        LessonPlanLibrarySourceService(
            repository=FakeRepository(),
            storage=storage,
        )
    )

    item = service.list_docx()[0]

    content = service.load_bytes(
        item
    )

    imported = (
        LessonPlanDocxTextImporter()
        .import_bytes(content)
    )

    assert storage.calls == [
        "file-1"
    ]

    assert (
        imported.objectives_text
        == "Nhận biết nội dung."
    )

    assert (
        imported.materials_text
        == "Máy chiếu."
    )

    assert (
        imported.teaching_process_text
        == "Hoạt động học tập."
    )
