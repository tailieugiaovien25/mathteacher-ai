from datetime import date
from pathlib import Path

from docx import Document

from educational_planning_v2.models import (
    TeachingSession,
)
from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)


def make_context():
    return ScheduledLessonContext(
        teaching_date=date(2026, 9, 28),
        drafting_date=date(2026, 9, 27),
        class_id="8A2",
        subject_ref="MATHEMATICS",
        component_ref="ALGEBRA",
        curriculum_period=9,
        lesson_id="LESSON-009",
        lesson_title="\u0110\u01a1n th\u1ee9c",
        session=TeachingSession.MORNING,
        timetable_period=1,
        period_in_lesson=1,
    )


def make_real_docx(path: Path):
    document = Document()

    document.add_paragraph(
        "K\u1ebe HO\u1ea0CH B\u00c0I D\u1ea0Y"
    )

    table = document.add_table(
        rows=5,
        cols=2,
    )

    values = (
        ("L\u1edbp", "8A1"),
        ("Ti\u1ebft PPCT", "1"),
        ("T\u00ean b\u00e0i", "B\u00e0i c\u0169"),
        ("Ng\u00e0y so\u1ea1n", "07/09/2026"),
        ("Ng\u00e0y d\u1ea1y", "08/09/2026"),
    )

    for row, (label, value) in zip(
        table.rows,
        values,
    ):
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_paragraph(
        "I. M\u1ee4C TI\u00caU"
    )

    document.add_paragraph(
        "N\u1ed9i dung b\u00e0i h\u1ecdc "
        "ph\u1ea3i \u0111\u01b0\u1ee3c b\u1ea3o to\u00e0n."
    )

    document.save(path)


def all_document_text(path: Path):
    document = Document(path)

    parts = []

    for paragraph in document.paragraphs:
        parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def test_real_docx_context_application(tmp_path):
    from document_standardization import (
        LessonPlanDocumentContextApplier,
    )

    source = tmp_path / "lesson.docx"
    output = tmp_path / "lesson.modified.docx"

    make_real_docx(source)

    applier = LessonPlanDocumentContextApplier()

    result = applier.apply(
        source=source,
        output=output,
        context=make_context(),
    )

    assert output.exists()

    text = all_document_text(output)

    assert "8A2" in text
    assert "9" in text
    assert "\u0110\u01a1n th\u1ee9c" in text
    assert "27/09/2026" in text
    assert "28/09/2026" in text

    assert "8A1" not in text
    assert "B\u00e0i c\u0169" not in text

    assert (
        "N\u1ed9i dung b\u00e0i h\u1ecdc "
        "ph\u1ea3i \u0111\u01b0\u1ee3c b\u1ea3o to\u00e0n."
        in text
    )

    assert result is not None
