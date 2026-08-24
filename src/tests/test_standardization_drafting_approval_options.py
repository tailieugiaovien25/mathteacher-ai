
from datetime import date
from io import BytesIO
from types import SimpleNamespace

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import portal_v2.ui.weekly_schedule_streamlit as module


def test_monday_is_week_reference():
    assert (
        module._mt_monday_of_teaching_week(
            date(2026, 9, 30)
        )
        == date(2026, 9, 28)
    )


def test_drafting_date_before_monday_any_days():
    teaching_date = date(
        2026,
        9,
        30,
    )

    assert (
        module._mt_date_before_teaching_week(
            teaching_date,
            3,
        )
        == date(
            2026,
            9,
            25,
        )
    )

    assert (
        module._mt_date_before_teaching_week(
            teaching_date,
            100,
        )
        == date(
            2026,
            6,
            20,
        )
    )


def test_approval_date_is_overwritten_in_word():
    document = Document()

    paragraph = document.add_paragraph()

    label = paragraph.add_run(
        "Ng\u00e0y duy\u1ec7t: "
    )
    label.bold = True

    value = paragraph.add_run(
        "01/01/2020"
    )
    value.italic = True

    document.add_paragraph(
        "N\u1ed9i dung gi\u00e1o \u00e1n "
        "gi\u1eef nguy\u00ean."
    )

    source = BytesIO()
    document.save(source)

    result = (
        module._mt_overlay_approval_date_bytes(
            source.getvalue(),
            date(
                2026,
                9,
                27,
            ),
        )
    )

    processed = Document(
        BytesIO(result)
    )

    assert (
        processed.paragraphs[0].text
        ==
        "Ng\u00e0y duy\u1ec7t: 27/09/2026"
    )

    assert (
        processed.paragraphs[0]
        .runs[0].bold
        is True
    )

    assert (
        processed.paragraphs[0]
        .runs[1].italic
        is True
    )

    assert (
        processed.paragraphs[1].text
        ==
        "N\u1ed9i dung gi\u00e1o \u00e1n "
        "gi\u1eef nguy\u00ean."
    )


def test_approval_date_before_school_approval_label_is_overwritten():
    document = Document()
    document.add_paragraph("Ngày 13 tháng 09 năm 2025")
    document.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    document.save(source)

    result = module._mt_overlay_approval_date_bytes(
        source.getvalue(),
        date(2026, 9, 13),
    )
    processed = Document(BytesIO(result))

    assert processed.paragraphs[0].text == "Ngày 13 tháng 09 năm 2026"
    assert processed.paragraphs[1].text == "Tổ CM duyệt:"
    assert processed.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert processed.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_missing_approval_date_is_added_before_approval_label():
    document = Document()
    document.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    document.save(source)

    result = module._mt_overlay_approval_date_bytes(
        source.getvalue(),
        date(2026, 9, 13),
    )
    processed = Document(BytesIO(result))

    assert processed.paragraphs[0].text == "Ngày 13 tháng 09 năm 2026"
    assert processed.paragraphs[1].text == "Tổ CM duyệt:"
    assert processed.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert processed.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_lesson_end_rule_is_inserted_before_approval_block():
    document = Document()
    document.add_paragraph("Nội dung cuối bài.")
    document.add_paragraph("Ngày 13 tháng 09 năm 2026")
    document.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    document.save(source)

    result = module._mt_ensure_lesson_end_rule_bytes(source.getvalue())
    processed = Document(BytesIO(result))

    approval_index = next(
        index
        for index, paragraph in enumerate(processed.paragraphs)
        if paragraph.text == "Ngày 13 tháng 09 năm 2026"
    )
    rule = processed.paragraphs[approval_index - 1]

    assert rule._p.xpath("./w:pPr/w:pBdr/w:bottom")


def test_existing_lesson_end_rule_is_not_duplicated():
    first = Document()
    first.add_paragraph("Nội dung cuối bài.")
    first.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    first.save(source)

    once = module._mt_ensure_lesson_end_rule_bytes(source.getvalue())
    twice = module._mt_ensure_lesson_end_rule_bytes(once)
    processed = Document(BytesIO(twice))

    rules = [
        paragraph
        for paragraph in processed.paragraphs
        if paragraph._p.xpath("./w:pPr/w:pBdr/w:bottom")
    ]
    assert len(rules) == 1


def test_final_layout_centers_headings_and_italicizes_date_blocks():
    document = Document()
    document.add_paragraph("CHỦ ĐỀ 1: NGÀY KHAI TRƯỜNG (T1)")
    document.add_paragraph("Ngày soạn: 04/09/2026")
    document.add_paragraph("Ngày dạy:")
    document.add_paragraph("7A1 - 08/09/2026")
    document.add_paragraph("TIẾT 1: HỌC BÀI HÁT: KHAI TRƯỜNG")
    document.add_paragraph("Ngày 06 tháng 09 năm 2026")
    document.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    document.save(source)

    result = module._mt_format_lesson_document_layout_bytes(
        source.getvalue()
    )
    processed = Document(BytesIO(result))

    assert processed.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    lesson_heading = next(
        p for p in processed.paragraphs if p.text.startswith("TIẾT 1:")
    )
    assert lesson_heading.alignment == WD_ALIGN_PARAGRAPH.CENTER

    for paragraph in processed.paragraphs:
        if (
            paragraph.text.startswith("Ngày ")
            or paragraph.text.startswith("7A1 -")
        ):
            assert paragraph.runs
            assert all(run.italic is True for run in paragraph.runs)


def test_final_approval_block_has_one_gap_and_four_lines_below():
    document = Document()
    document.add_paragraph("Nội dung cuối bài.")
    document.add_paragraph("")
    document.add_paragraph("Ngày 06 tháng 09 năm 2026")
    document.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    document.save(source)

    result = module._mt_format_lesson_document_layout_bytes(
        source.getvalue()
    )
    processed = Document(BytesIO(result))
    paragraphs = processed.paragraphs
    date_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.startswith("Ngày 06 tháng")
    )
    marker_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text == "Tổ CM duyệt:"
    )

    assert paragraphs[date_index - 1].text == ""
    assert paragraphs[date_index - 2]._p.xpath(
        "./w:pPr/w:pBdr/w:bottom"
    )
    assert all(
        paragraph.text == ""
        for paragraph in paragraphs[marker_index + 1:marker_index + 5]
    )
    assert len(paragraphs[marker_index + 1:]) >= 4


def test_obsolete_incomplete_approval_date_line_is_removed():
    document = Document()
    document.add_paragraph("Ngày ...../09/2025")
    document.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    document.save(source)

    result = module._mt_overlay_approval_date_bytes(
        source.getvalue(),
        date(2026, 9, 6),
    )
    processed = Document(BytesIO(result))
    text = "\n".join(p.text for p in processed.paragraphs)

    assert "Ngày 06 tháng 09 năm 2026" in text
    assert "Ngày ...../09/2025" not in text
    assert text.count("Ngày ") == 1


def test_drafting_date_is_overwritten_from_monday_rule():
    document = Document()
    document.add_paragraph("Ngày soạn: 17/09/2026")
    document.add_paragraph("Ngày dạy: 7A1 - 18/09/2026")
    source = BytesIO()
    document.save(source)

    drafting_date = module._mt_date_before_teaching_week(
        date(2026, 9, 18),
        3,
    )
    result = module._mt_overlay_drafting_date_bytes(
        source.getvalue(),
        drafting_date,
    )
    processed = Document(BytesIO(result))

    assert drafting_date == date(2026, 9, 11)
    assert processed.paragraphs[0].text == "Ngày soạn: 11/09/2026"


def test_week_one_drafting_date_three_days_before_monday():
    assert module._mt_date_before_teaching_week(
        date(2026, 9, 11),
        3,
    ) == date(2026, 9, 4)


def test_missing_drafting_date_is_added_before_teaching_date():
    document = Document()
    document.add_paragraph("Ngày dạy: 7A1 - 18/09/2026")
    source = BytesIO()
    document.save(source)

    result = module._mt_overlay_drafting_date_bytes(
        source.getvalue(),
        date(2026, 9, 11),
    )
    processed = Document(BytesIO(result))

    assert processed.paragraphs[0].text == "Ngày soạn: 11/09/2026"
    assert processed.paragraphs[1].text == "Ngày dạy: 7A1 - 18/09/2026"


def test_date_wrapper_applies_drafting_and_approval_to_tuple_result(
    monkeypatch,
):
    document = Document()
    document.add_paragraph("Ngày soạn: 17/09/2026")
    document.add_paragraph("Ngày 13 tháng 09 năm 2025")
    document.add_paragraph("Tổ CM duyệt:")
    source = BytesIO()
    document.save(source)

    def base_processor(
        *,
        row,
        drafting_date,
        content,
        original_name,
        modification_plan=None,
        options=None,
        original_content=None,
        ai_revised_text="",
    ):
        return (original_name, content, ())

    monkeypatch.setattr(
        module,
        "_mt_original_process_lesson_plan_upload_dates",
        base_processor,
    )
    monkeypatch.setattr(
        module,
        "st",
        SimpleNamespace(
            session_state={
                module._MT_DRAFTING_ENABLED: True,
                module._MT_DRAFTING_DAYS: 3,
                module._MT_APPROVAL_ENABLED: True,
                module._MT_APPROVAL_DAYS: 1,
            }
        ),
    )

    result = module._mt_original_process_lesson_plan_upload_3b(
        row=SimpleNamespace(
            teaching_date=date(2026, 9, 18),
        ),
        drafting_date=date(2026, 9, 17),
        content=source.getvalue(),
        original_name="giao-an.docx",
    )
    processed = Document(BytesIO(result[1]))
    text = "\n".join(p.text for p in processed.paragraphs)

    assert "Ngày soạn: 11/09/2026" in text
    assert "Ngày 13 tháng 09 năm 2026" in text
    assert "Ngày soạn: 17/09/2026" not in text
    assert "Ngày 13 tháng 09 năm 2025" not in text


def test_control_panel_has_date_options():
    source = open(
        "src/portal_v2/ui/"
        "weekly_schedule_streamlit.py",
        encoding="utf-8",
    ).read()

    assert (
        "D\\u00e1n \\u0111\\u00e8 "
        "Ng\\u00e0y so\\u1ea1n"
        in source
    )

    assert (
        "D\\u00e1n \\u0111\\u00e8 "
        "Ng\\u00e0y duy\\u1ec7t"
        in source
    )
