
from io import BytesIO

from docx import Document

from lesson_planning_v2.services.lesson_plan_workspace_v1_service import (
    LessonPlanFullDocumentDocxAdapter,
)


def test_internal_docx_adapter_creates_valid_docx():
    raw = (
        LessonPlanFullDocumentDocxAdapter()
        .build_bytes(
            "TITLE\n\nFIRST\nSECOND"
        )
    )

    assert isinstance(
        raw,
        bytes,
    )

    assert raw

    document = Document(
        BytesIO(raw)
    )

    values = [
        paragraph.text
        for paragraph
        in document.paragraphs
    ]

    assert "TITLE" in values
    assert "FIRST" in values
    assert "SECOND" in values


def test_internal_docx_adapter_preserves_line_order():
    raw = (
        LessonPlanFullDocumentDocxAdapter()
        .build_bytes(
            "A\nB\nC"
        )
    )

    document = Document(
        BytesIO(raw)
    )

    values = [
        paragraph.text
        for paragraph
        in document.paragraphs
        if paragraph.text
    ]

    assert values == [
        "A",
        "B",
        "C",
    ]


def test_internal_docx_adapter_rejects_empty_text():
    adapter = (
        LessonPlanFullDocumentDocxAdapter()
    )

    try:
        adapter.build_bytes(
            "   "
        )
    except ValueError:
        return

    raise AssertionError(
        "Expected ValueError"
    )
