from pathlib import Path

from docx import Document

from document_standardization import LessonPlanAiRevisionOverlay


def _source(path: Path) -> None:
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Đầu trang gốc"
    document.add_paragraph("Nội dung giữ nguyên")
    document.add_paragraph("Nội dung AI sẽ sửa")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Bảng gốc"
    table.cell(0, 1).text = "Không đổi"
    document.save(path)


def test_overlay_changes_only_safely_mapped_paragraphs(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _source(source)

    result = LessonPlanAiRevisionOverlay().apply(
        source=source,
        output=output,
        revised_text=(
            "Nội dung giữ nguyên\n"
            "Nội dung AI đã sửa"
        ),
    )

    document = Document(output)
    assert document.paragraphs[0].text == "Nội dung giữ nguyên"
    assert document.paragraphs[1].text == "Nội dung AI đã sửa"
    assert document.tables[0].cell(0, 0).text == "Bảng gốc"
    assert document.sections[0].header.paragraphs[0].text == "Đầu trang gốc"
    assert result.changed_paragraphs == 1


def test_overlay_preserves_ambiguous_deletions(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    _source(source)

    result = LessonPlanAiRevisionOverlay().apply(
        source=source,
        output=output,
        revised_text="Nội dung giữ nguyên",
    )

    document = Document(output)
    assert document.paragraphs[1].text == "Nội dung AI sẽ sửa"
    assert result.preserved_ambiguous_blocks >= 1
    assert result.warnings
