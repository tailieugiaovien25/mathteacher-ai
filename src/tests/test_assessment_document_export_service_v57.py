from __future__ import annotations

from hashlib import sha256
from io import BytesIO
from json import loads
from zipfile import ZipFile

import pytest
from docx import Document

from assessment_generation_v2.documents import (
    CanonicalAssessmentDocument,
)
from assessment_generation_v2.renderers import (
    AssessmentTemplateDefinition,
)
from assessment_generation_v2.services import (
    ApprovedAssessmentTemplate,
    AssessmentDocumentExportError,
    AssessmentDocumentExportRequest,
    AssessmentDocumentExportService,
    AssessmentDocumentExportValidationError,
    PublishedAssessmentRenderSource,
)


EXAM_VERSION_ID = "11111111-1111-4111-8111-111111111111"
VARIANT_ID = "22222222-2222-4222-8222-222222222222"
OWNER_ID = "33333333-3333-4333-8333-333333333333"
TEMPLATE_IDS = {
    value: f"44444444-4444-4444-8{index:03d}-444444444444"
    for index, value in enumerate(
        (
            "MATRIX",
            "SPECIFICATION",
            "STUDENT_EXAM",
            "ANSWER_KEY",
            "SCORING_GUIDE",
        ),
        start=1,
    )
}


def _canonical() -> CanonicalAssessmentDocument:
    return CanonicalAssessmentDocument(
        schema_version=1,
        metadata={
            "exam": {
                "exam_title": "Đề kiểm tra giữa học kỳ I",
                "variant_code": "101",
            }
        },
        matrix=(
            {
                "topic_name": "Số và Đại số",
                "cognitive_level_name": "Nhận biết",
                "question_count": 1,
            },
        ),
        specification=(
            {
                "requirement_code": "REQ-01",
                "requirement_text": "Nhận biết số tự nhiên.",
            },
        ),
        questions=(
            {
                "display_number": 1,
                "prompt_text": "Số tự nhiên nhỏ nhất là số nào?",
            },
        ),
        answer_key=(
            {
                "display_number": 1,
                "answer_text": "0",
            },
        ),
        scoring_guide=(
            {
                "display_number": 1,
                "scoring_text": "Trả lời đúng được 1 điểm.",
            },
        ),
    )


class FakeBuilder:
    def __init__(self) -> None:
        self.calls = 0

    def build(self, **payloads) -> CanonicalAssessmentDocument:
        self.calls += 1
        assert set(payloads) == {
            "snapshot_document",
            "student_exam_payload",
            "answer_key_payload",
            "scoring_guide_payload",
        }
        return _canonical()


def _template(document_type: str) -> AssessmentTemplateDefinition:
    roots = {
        "MATRIX": ("matrix_rows", "matrix"),
        "SPECIFICATION": (
            "specification_rows",
            "specification",
        ),
        "STUDENT_EXAM": ("question_rows", "questions"),
        "ANSWER_KEY": ("answer_rows", "answer_key"),
        "SCORING_GUIDE": ("scoring_rows", "scoring_guide"),
    }
    binding_name, binding_path = roots[document_type]
    field_paths = {
        "MATRIX": ("topic_name", "cognitive_level_name"),
        "SPECIFICATION": (
            "requirement_code",
            "requirement_text",
        ),
        "STUDENT_EXAM": ("display_number", "prompt_text"),
        "ANSWER_KEY": ("display_number", "answer_text"),
        "SCORING_GUIDE": (
            "display_number",
            "scoring_text",
        ),
    }
    fields = field_paths[document_type]
    return AssessmentTemplateDefinition(
        document_type_code=document_type,
        renderer_code="DOCX_JSON_V1",
        layout_schema={
            "paper_size": "A4",
            "page_orientation": "PORTRAIT",
            "margins_mm": {
                "top": 20,
                "right": 20,
                "bottom": 20,
                "left": 20,
            },
        },
        style_schema={
            "font_family": "Times New Roman",
            "font_size": 12,
            "table_font_size": 11,
            "heading_size": 14,
        },
        binding_schema={binding_name: binding_path},
        section_schema=(
            {
                "section_code": "TITLE",
                "section_type": "TEXT",
                "text": document_type.replace("_", " "),
                "title": "HỒ SƠ ĐÁNH GIÁ",
                "title_alignment": "CENTER",
            },
            {
                "section_code": "CONTENT",
                "section_type": "TABLE",
                "bindings": (binding_name,),
                "columns": tuple(
                    {
                        "column_code": f"C{index}",
                        "heading": field.replace("_", " "),
                        "value_path": field,
                    }
                    for index, field in enumerate(fields, start=1)
                ),
            },
        ),
    )


class FakeGateway:
    def __init__(self) -> None:
        self.source = PublishedAssessmentRenderSource(
            exam_version_id=EXAM_VERSION_ID,
            variant_id=VARIANT_ID,
            owner_user_id=OWNER_ID,
            snapshot_hash="a" * 64,
            hash_verified=True,
            snapshot_document={"source": "snapshot"},
            student_exam_payload={"source": "student"},
            answer_key_payload={"source": "answers"},
            scoring_guide_payload={"source": "scoring"},
        )
        self.templates = {
            value: ApprovedAssessmentTemplate(
                template_version_id=TEMPLATE_IDS[value],
                template_set_code="PHONG-MAU-2026",
                review_status="APPROVED",
                lifecycle_status="ACTIVE",
                definition=_template(value),
            )
            for value in TEMPLATE_IDS
        }
        self.assets: dict[str, bytes] = {}
        self.calls: list[tuple[str, str]] = []

    def load_published_render_source(
        self, *, exam_version_id: str, variant_id: str
    ):
        self.calls.append(("source", variant_id))
        return self.source

    def find_active_template(
        self, *, template_set_code: str, document_type: str
    ):
        self.calls.append(("template", document_type))
        return self.templates.get(document_type)

    def load_template_asset(self, *, asset_path: str) -> bytes:
        self.calls.append(("asset", asset_path))
        return self.assets[asset_path]


def _request(**changes) -> AssessmentDocumentExportRequest:
    values = {
        "exam_version_id": EXAM_VERSION_ID,
        "variant_id": VARIANT_ID,
        "owner_user_id": OWNER_ID,
        "template_set_code": "PHONG-MAU-2026",
        "bundle_name": "Toán 6 - Giữa kỳ I - Mã 101",
    }
    values.update(changes)
    return AssessmentDocumentExportRequest(**values)


def _service(gateway=None, builder=None):
    return AssessmentDocumentExportService(
        gateway=gateway or FakeGateway(),
        builder=builder or FakeBuilder(),
    )


def test_exports_complete_five_document_bundle() -> None:
    builder = FakeBuilder()
    result = _service(builder=builder).export(request=_request())

    assert builder.calls == 1
    assert result.bundle_filename == (
        "to-n-6-gi-a-k-i-m-101.zip"
    )
    assert len(result.documents) == 5
    assert result.bundle_hash == sha256(
        result.bundle_content
    ).hexdigest()

    with ZipFile(BytesIO(result.bundle_content)) as archive:
        assert set(archive.namelist()) == {
            "ma-tran.docx",
            "ban-dac-ta.docx",
            "de-kiem-tra.docx",
            "dap-an.docx",
            "huong-dan-cham.docx",
            "manifest.json",
        }
        manifest = loads(archive.read("manifest.json"))

    assert manifest["snapshot_hash"] == "a" * 64
    assert len(manifest["documents"]) == 5


def test_each_document_is_a_valid_docx() -> None:
    result = _service().export(request=_request())

    for item in result.documents:
        assert item.content.startswith(b"PK")
        assert item.content_hash == sha256(item.content).hexdigest()
        assert Document(BytesIO(item.content)).tables


def test_student_exam_never_contains_answer_root() -> None:
    result = _service().export(
        request=_request(document_types=("STUDENT_EXAM",))
    )
    document = Document(BytesIO(result.documents[0].content))
    content = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )

    assert "Số tự nhiên nhỏ nhất" in content
    assert "answer_text" not in content
    assert "Đáp án" not in content


def test_bundle_is_deterministic() -> None:
    service = _service()
    first = service.export(request=_request())
    second = service.export(request=_request())

    assert first.bundle_content == second.bundle_content
    assert first.bundle_hash == second.bundle_hash
    assert tuple(
        item.content for item in first.documents
    ) == tuple(
        item.content for item in second.documents
    )


def test_docx_zip_entries_have_fixed_timestamps() -> None:
    result = _service().export(
        request=_request(document_types=("MATRIX",))
    )

    with ZipFile(BytesIO(result.documents[0].content)) as archive:
        assert archive.infolist()
        assert all(
            entry.date_time == (1980, 1, 1, 0, 0, 0)
            for entry in archive.infolist()
        )


def test_rejects_unverified_snapshot_before_template_lookup() -> None:
    gateway = FakeGateway()
    gateway.source = PublishedAssessmentRenderSource(
        exam_version_id=EXAM_VERSION_ID,
        variant_id=VARIANT_ID,
        owner_user_id=OWNER_ID,
        snapshot_hash="a" * 64,
        hash_verified=False,
        snapshot_document={},
        student_exam_payload={},
        answer_key_payload={},
        scoring_guide_payload={},
    )

    with pytest.raises(AssessmentDocumentExportError, match="integrity"):
        _service(gateway=gateway).export(request=_request())

    assert not any(call[0] == "template" for call in gateway.calls)


def test_rejects_owner_mismatch() -> None:
    with pytest.raises(PermissionError, match="owner"):
        _service().export(
            request=_request(
                owner_user_id=(
                    "99999999-9999-4999-8999-999999999999"
                )
            )
        )


@pytest.mark.parametrize(
    ("review_status", "lifecycle_status", "message"),
    (
        ("PENDING_REVIEW", "ACTIVE", "approved"),
        ("APPROVED", "INACTIVE", "active"),
    ),
)
def test_requires_approved_active_template(
    review_status: str,
    lifecycle_status: str,
    message: str,
) -> None:
    gateway = FakeGateway()
    original = gateway.templates["MATRIX"]
    gateway.templates["MATRIX"] = ApprovedAssessmentTemplate(
        template_version_id=original.template_version_id,
        template_set_code=original.template_set_code,
        review_status=review_status,
        lifecycle_status=lifecycle_status,
        definition=original.definition,
    )

    with pytest.raises(AssessmentDocumentExportError, match=message):
        _service(gateway=gateway).export(
            request=_request(document_types=("MATRIX",))
        )


def test_rejects_missing_template() -> None:
    gateway = FakeGateway()
    gateway.templates.pop("ANSWER_KEY")

    with pytest.raises(AssessmentDocumentExportError, match="unavailable"):
        _service(gateway=gateway).export(
            request=_request(document_types=("ANSWER_KEY",))
        )


def test_verifies_template_asset_hash() -> None:
    gateway = FakeGateway()
    definition = _template("MATRIX")
    definition = AssessmentTemplateDefinition(
        document_type_code=definition.document_type_code,
        renderer_code=definition.renderer_code,
        layout_schema=definition.layout_schema,
        style_schema=definition.style_schema,
        binding_schema=definition.binding_schema,
        section_schema=definition.section_schema,
        template_asset_path="templates/matrix.docx",
        template_asset_hash="b" * 64,
    )
    gateway.templates["MATRIX"] = ApprovedAssessmentTemplate(
        template_version_id=TEMPLATE_IDS["MATRIX"],
        template_set_code="PHONG-MAU-2026",
        review_status="APPROVED",
        lifecycle_status="ACTIVE",
        definition=definition,
    )
    gateway.assets["templates/matrix.docx"] = b"wrong"

    with pytest.raises(AssessmentDocumentExportError, match="integrity"):
        _service(gateway=gateway).export(
            request=_request(document_types=("MATRIX",))
        )


def test_rejects_duplicate_document_types() -> None:
    with pytest.raises(
        AssessmentDocumentExportValidationError,
        match="unique",
    ):
        _request(document_types=("MATRIX", "MATRIX"))


def test_subset_export_contains_only_requested_document() -> None:
    result = _service().export(
        request=_request(document_types=("MATRIX",))
    )

    with ZipFile(BytesIO(result.bundle_content)) as archive:
        assert set(archive.namelist()) == {
            "ma-tran.docx",
            "manifest.json",
        }


def test_export_service_has_no_ui_or_database_imports() -> None:
    from pathlib import Path

    source = Path(
        "src/assessment_generation_v2/services/"
        "assessment_document_export_service.py"
    ).read_text(encoding="utf-8")

    for forbidden in (
        "import streamlit",
        "from streamlit",
        "import supabase",
        "from supabase",
        "portal_v2",
    ):
        assert forbidden not in source
