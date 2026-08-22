from datetime import date
from pathlib import Path

from docx import Document

from document_standardization.lesson_plan_document_context_applier import (
    LessonPlanDocumentContextApplier,
)
from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)
from educational_planning_v2.models.weekly_teaching_schedule import (
    TeachingSession,
)


def _build_context():
    return ScheduledLessonContext(
        subject_ref="TOAN",
        component_ref=None,
        lesson_id="TOAN6-012",
        class_id="6A2",
        curriculum_period=12,
        lesson_title="Ph\u00e2n s\u1ed1",
        drafting_date=date(
            2026,
            9,
            10,
        ),
        teaching_date=date(
            2026,
            9,
            15,
        ),
        period_in_lesson=1,
        session=TeachingSession.MORNING,
        timetable_period=2,
    )


def test_context_applier_uses_preservation_overlay(
    tmp_path: Path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()

    paragraph = document.add_paragraph()

    label = paragraph.add_run(
        "Lớp:"
    )
    label.bold = True

    value = paragraph.add_run(
        " 6A1"
    )
    value.italic = True

    document.add_paragraph(
        "Ngày soạn: 01/09/2026"
    )

    document.add_paragraph(
        "Ngày dạy: 02/09/2026"
    )

    document.add_paragraph(
        "Tiết PPCT: 1"
    )

    document.add_paragraph(
        "Tên bài: Bài cũ"
    )

    document.save(source)

    result = (
        LessonPlanDocumentContextApplier()
        .apply(
            source=source,
            output=output,
            context=_build_context(),
        )
    )

    processed = Document(output)

    assert (
        processed.paragraphs[0].text
        == "Lớp: 6A2"
    )

    assert (
        processed.paragraphs[0]
        .runs[0]
        .bold
        is True
    )

    assert (
        processed.paragraphs[0]
        .runs[1]
        .italic
        is True
    )

    texts = [
        paragraph.text
        for paragraph
        in processed.paragraphs
    ]

    assert (
        "Ngày soạn: 10/09/2026"
        in texts
    )

    assert (
        "Ngày dạy: 15/09/2026"
        in texts
    )

    assert (
        "Tiết PPCT: 12"
        in texts
    )

    assert (
        "Tên bài: Phân số"
        in texts
    )

    assert (
        result.unresolved_fields
        == ()
    )
