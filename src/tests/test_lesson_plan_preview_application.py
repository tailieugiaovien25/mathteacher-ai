from pathlib import Path

import pytest
from docx import Document

from document_intelligence import (
    build_document_analyzer,
)
from document_intelligence.lesson_plan_preview import (
    LessonPlanIntelligencePreviewService,
)
from document_intelligence.lesson_plan_preview_application import (
    LessonPlanPreviewApplicationService,
)
from document_intelligence.lesson_plan_preview_presenter import (
    PreviewReviewState,
)
from document_intelligence.validation import (
    CanonicalDocumentContext,
)


def make_docx(
    path: Path,
    *,
    class_name="8A2",
    curriculum_period=4,
    lesson_title="Đơn thức",
    drafting_date="07/09/2026",
    teaching_date="08/09/2026",
) -> None:
    document = Document()

    document.add_paragraph(
        f"Ngày soạn: {drafting_date}"
    )

    document.add_paragraph(
        f"Ngày dạy: {teaching_date}"
    )

    document.add_paragraph(
        f"Lớp: {class_name}"
    )

    document.add_paragraph(
        f"Tiết {curriculum_period}. "
        f"Bài 3. {lesson_title}"
    )

    document.save(path)


def make_service():
    return LessonPlanPreviewApplicationService(
        preview_service=(
            LessonPlanIntelligencePreviewService(
                analyzer=build_document_analyzer()
            )
        )
    )


def test_application_prepares_accepted_view(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    source = tmp_path / "lesson.docx"

    make_docx(source)

    original = source.read_bytes()

    view = make_service().prepare(
        source=source,
        canonical=CanonicalDocumentContext(
            class_name="8A2",
            curriculum_period=4,
            lesson_title="Đơn thức",
            drafting_date="07/09/2026",
            teaching_date="08/09/2026",
        ),
    )

    assert source.read_bytes() == original

    assert len(view.items) == 5
    assert view.requires_review is False
    assert view.conflict_count == 0
    assert view.ai_used is False
    assert view.ai_failed is False

    assert all(
        item.review_state
        is PreviewReviewState.ACCEPTED
        for item in view.items
    )


def test_application_exposes_conflict(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    source = tmp_path / "lesson.docx"

    make_docx(
        source,
        class_name="8A2",
    )

    view = make_service().prepare(
        source=source,
        canonical=CanonicalDocumentContext(
            class_name="8A1",
            curriculum_period=4,
            lesson_title="Đơn thức",
            drafting_date="07/09/2026",
            teaching_date="08/09/2026",
        ),
    )

    assert view.requires_review is True
    assert view.conflict_count == 1

    conflicts = tuple(
        item
        for item in view.items
        if item.review_state
        is PreviewReviewState.CONFLICT
    )

    assert len(conflicts) == 1
    assert conflicts[0].value == "8A2"


def test_application_rejects_invalid_canonical(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    source = tmp_path / "lesson.docx"

    make_docx(source)

    with pytest.raises(
        TypeError,
        match="canonical must be CanonicalDocumentContext",
    ):
        make_service().prepare(
            source=source,
            canonical=object(),
        )


def test_application_rejects_non_docx(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    source = tmp_path / "lesson.txt"

    source.write_text(
        "test",
        encoding="utf-8",
    )

    with pytest.raises(
        ValueError,
        match="DOCX",
    ):
        make_service().prepare(
            source=source,
            canonical=CanonicalDocumentContext(),
        )
