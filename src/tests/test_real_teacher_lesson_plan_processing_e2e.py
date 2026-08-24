import os
from datetime import date
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from lxml import etree

from document_intelligence.contracts import (
    DocumentField,
)
from document_intelligence.lesson_plan_modification_plan import (
    LessonPlanFieldModification,
    LessonPlanModificationPlan,
)
from educational_planning_v2.models.weekly_teaching_schedule import (
    TeachingSession,
    WeeklyTeachingScheduleEntry,
)
from lesson_planning_v2.services.lesson_plan_document_processing_service import (
    LessonPlanDocumentProcessingService,
)


ENV_NAME = "REAL_TEACHER_DOCX"

PROFILE = Path(
    "scripts/word_standardizer/"
    "lesson_plan_profile.json"
)


def teacher_docx() -> Path:
    raw = os.environ.get(
        ENV_NAME
    )

    if not raw:
        pytest.skip(
            f"{ENV_NAME} is not configured"
        )

    path = Path(raw)

    if not path.exists():
        pytest.skip(
            f"Teacher DOCX does not exist: {path}"
        )

    return path


def make_row():
    return WeeklyTeachingScheduleEntry(
        teaching_date=date(
            2026,
            9,
            30,
        ),
        weekday=3,
        timetable_period=1,
        session=TeachingSession.MORNING,
        teacher_id="GV001",
        class_id="6A2",
        subject_ref="MATHEMATICS",
        curriculum_period=10,
        lesson_id="REAL-LESSON-001",
        lesson_title=(
            "Thứ tự thực hiện "
            "các phép tính"
        ),
        component_ref="ALGEBRA",
        period_in_lesson=1,
        total_lesson_periods=2,
        teaching_equipment=(),
    )


def make_plan():
    return LessonPlanModificationPlan(
        modifications=(
            LessonPlanFieldModification(
                field=DocumentField.CLASS_NAME,
                value="6A3",
            ),
            LessonPlanFieldModification(
                field=DocumentField.CURRICULUM_PERIOD,
                value="20",
            ),
            LessonPlanFieldModification(
                field=DocumentField.LESSON_TITLE,
                value=(
                    "Thứ tự thực hiện "
                    "các phép tính"
                ),
            ),
            LessonPlanFieldModification(
                field=DocumentField.DRAFTING_DATE,
                value="29/09/2026",
            ),
            LessonPlanFieldModification(
                field=DocumentField.TEACHING_DATE,
                value="30/09/2026",
            ),
        )
    )


def count_omml(content: bytes) -> int:
    with ZipFile(
        BytesIO(content)
    ) as archive:
        xml = archive.read(
            "word/document.xml"
        )

    root = etree.fromstring(xml)

    namespace = {
        "m": (
            "http://schemas.openxmlformats.org/"
            "officeDocument/2006/math"
        )
    }

    return len(
        root.xpath(
            ".//m:oMath",
            namespaces=namespace,
        )
    )


def read_document(content: bytes):
    return Document(
        BytesIO(content)
    )


def all_text(content: bytes) -> str:
    document = read_document(
        content
    )

    parts = []

    for paragraph in document.paragraphs:
        parts.append(
            paragraph.text
        )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(
                    cell.text
                )

    return "\n".join(parts)


def test_real_teacher_docx_reaches_processing_service():
    source = teacher_docx()

    assert PROFILE.exists()

    original_bytes = source.read_bytes()

    original_document = Document(
        source
    )

    original_paragraphs = len(
        original_document.paragraphs
    )

    original_tables = len(
        original_document.tables
    )

    original_shapes = len(
        original_document.inline_shapes
    )

    original_omml = count_omml(
        original_bytes
    )

    service = (
        LessonPlanDocumentProcessingService(
            profile_path=PROFILE
        )
    )

    result = service.process(
        row=make_row(),
        drafting_date=date(
            2026,
            9,
            29,
        ),
        content=original_bytes,
        original_name=source.name,
        modification_plan=make_plan(),
    )

    assert result.output_bytes.startswith(
        b"PK"
    )

    output_document = read_document(
        result.output_bytes
    )

    assert (
        len(output_document.paragraphs)
        == original_paragraphs
    )

    assert (
        len(output_document.tables)
        == original_tables
    )

    assert (
        len(output_document.inline_shapes)
        == original_shapes
    )

    assert (
        count_omml(
            result.output_bytes
        )
        == original_omml
    )

    text = all_text(
        result.output_bytes
    )

    assert "29/09/2026" in text
    assert "30/09/2026" in text
    assert "6A3" in text

    assert (
        "TIẾT 20 + 21"
        in text
        or "Tiết 20 + 21"
        in text
    )

    assert (
        "Ngày dạy: 02/10/2025"
        in text
    )

    assert "TIẾT 12" in text

    assert (
        "LUYỆN TẬP CHUNG"
        in text
    )

    assert result.unresolved_fields == ()

    # The real source must remain untouched.
    assert (
        source.read_bytes()
        == original_bytes
    )
