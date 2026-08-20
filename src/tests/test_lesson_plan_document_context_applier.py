from datetime import date

from docx import Document

from educational_planning_v2.models import (
    TeachingSession,
)
from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)
from document_standardization import (
    LessonPlanDocumentContextApplier,
)


def make_context():
    return ScheduledLessonContext(
        teaching_date=date(2026, 9, 28),
        drafting_date=date(2026, 9, 27),
        class_id="6A1",
        subject_ref="MATHEMATICS",
        component_ref="ALGEBRA",
        curriculum_period=9,
        lesson_id="LESSON-009",
        lesson_title="Phân số",
        session=TeachingSession.MORNING,
        timetable_period=1,
        period_in_lesson=1,
    )


def test_applier_updates_paragraph_fields(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()
    document.add_paragraph(
        "Ngày soạn: 01/01/2020"
    )
    document.add_paragraph(
        "Ngày dạy: 02/01/2020"
    )
    document.add_paragraph(
        "Lớp: 6A2"
    )
    document.add_paragraph(
        "Tiết PPCT: 1"
    )
    document.add_paragraph(
        "Bài: Bài cũ"
    )
    document.add_paragraph(
        "Nội dung bài học phải được giữ nguyên."
    )
    document.save(source)

    result = (
        LessonPlanDocumentContextApplier()
        .apply(
            source,
            output,
            make_context(),
        )
    )

    updated = Document(output)

    texts = [
        paragraph.text
        for paragraph in updated.paragraphs
    ]

    assert "Ngày soạn: 27/09/2026" in texts
    assert "Ngày dạy: 28/09/2026" in texts
    assert "Lớp: 6A1" in texts
    assert "Tiết PPCT: 9" in texts
    assert "Bài: Phân số" in texts

    assert (
        "Nội dung bài học phải được giữ nguyên."
        in texts
    )

    assert result.unresolved_fields == ()


def test_applier_updates_fields_inside_table(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(0, 0).text = (
        "Ngày soạn: 01/01/2020"
    )
    table.cell(0, 1).text = (
        "Ngày giảng: 02/01/2020"
    )
    table.cell(1, 0).text = "Lớp: 6A2"
    table.cell(1, 1).text = "PPCT: 1"

    document.add_paragraph(
        "Bài: Bài cũ"
    )

    document.save(source)

    result = (
        LessonPlanDocumentContextApplier()
        .apply(
            source,
            output,
            make_context(),
        )
    )

    updated = Document(output)

    assert (
        updated.tables[0]
        .cell(0, 0)
        .text
        == "Ngày soạn: 27/09/2026"
    )

    assert (
        updated.tables[0]
        .cell(0, 1)
        .text
        == "Ngày giảng: 28/09/2026"
    )

    assert (
        updated.tables[0]
        .cell(1, 0)
        .text
        == "Lớp: 6A1"
    )

    assert (
        updated.tables[0]
        .cell(1, 1)
        .text
        == "PPCT: 9"
    )

    assert result.unresolved_fields == ()


def test_applier_reports_unresolved_fields(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()
    document.add_paragraph(
        "Nội dung không có metadata."
    )
    document.save(source)

    result = (
        LessonPlanDocumentContextApplier()
        .apply(
            source,
            output,
            make_context(),
        )
    )

    assert set(
        result.unresolved_fields
    ) == {
        "drafting_date",
        "teaching_date",
        "class_id",
        "curriculum_period",
        "lesson_title",
    }


def test_applier_does_not_overwrite_source(
    tmp_path,
):
    source = tmp_path / "source.docx"

    document = Document()
    document.add_paragraph(
        "Ngày soạn: 01/01/2020"
    )
    document.save(source)

    try:
        LessonPlanDocumentContextApplier().apply(
            source,
            source,
            make_context(),
        )
    except ValueError as exc:
        assert "ghi đè" in str(exc)
    else:
        raise AssertionError(
            "Expected overwrite protection"
        )
