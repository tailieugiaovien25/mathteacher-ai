from pathlib import Path

from docx import Document

from document_standardization import (
    LessonPlanStandardizationOptions,
    LessonPlanWordStandardizer,
)


PROFILE = {
    "profile_name": "selective-test",
    "page": {
        "margin_left_cm": 3,
        "margin_right_cm": 2,
        "margin_top_cm": 2,
        "margin_bottom_cm": 2,
    },
    "body": {
        "font": "Times New Roman",
        "size_pt": 14,
        "line_spacing": 1.15,
    },
    "title": {"size_pt": 14},
    "table": {"size_pt": 14},
    "header_footer": {"remove_existing": True, "page_number": True},
    "equations": {"mode": "safe", "text_font": "Times New Roman"},
}


def test_unselected_formatting_operations_leave_format_untouched(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    report = tmp_path / "report.json"
    document = Document()
    run = document.add_paragraph().add_run("Nội dung gốc")
    run.font.name = "Arial"
    document.sections[0].header.paragraphs[0].text = "Đầu trang gốc"
    document.save(source)

    options = LessonPlanStandardizationOptions(
        preserve_original_maximum=True,
        sync_context=False,
        normalize_font=False,
        normalize_equations=False,
        normalize_tables=False,
        normalize_page_layout=False,
        normalize_spacing=False,
        normalize_header_footer=False,
    )
    result = LessonPlanWordStandardizer(PROFILE).standardize(
        source,
        output,
        report,
        options=options,
    )

    standardized = Document(output)
    assert standardized.paragraphs[0].runs[0].font.name == "Arial"
    assert standardized.sections[0].header.paragraphs[0].text == "Đầu trang gốc"
    assert result["selected_options"]["normalize_font"] is False


def test_empty_selection_is_detectable():
    options = LessonPlanStandardizationOptions(
        preserve_original_maximum=False,
        sync_context=False,
        normalize_font=False,
        normalize_equations=False,
        normalize_tables=False,
        normalize_page_layout=False,
        normalize_spacing=False,
        normalize_header_footer=False,
    )
    assert options.has_selected_operation is False
