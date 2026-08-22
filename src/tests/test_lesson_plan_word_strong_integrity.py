from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document

from document_standardization.lesson_plan_standardizer import (
    LessonPlanWordStandardizer,
    inventory,
)


def _profile():
    return {
        "header_footer": {
            "remove_existing": False,
        }
    }


def test_inventory_contains_semantic_fingerprints(
    tmp_path: Path,
):
    source = tmp_path / "source.docx"

    document = Document()
    document.add_paragraph(
        "Lesson content"
    )

    table = document.add_table(
        rows=1,
        cols=1,
    )
    table.cell(
        0,
        0,
    ).text = "Table content"

    document.save(
        source
    )

    result = inventory(
        source
    )

    assert result.body_text_sha256
    assert result.header_footer_text_sha256

    assert isinstance(
        result.media_fingerprints,
        tuple,
    )

    assert isinstance(
        result.embedded_fingerprints,
        tuple,
    )

    assert isinstance(
        result.external_relationships,
        tuple,
    )


def test_integrity_gate_rejects_body_text_change(
    tmp_path: Path,
):
    source = tmp_path / "source.docx"

    document = Document()
    document.add_paragraph(
        "Original content"
    )
    document.save(
        source
    )

    before = inventory(
        source
    )

    after = replace(
        before,
        body_text_sha256=(
            "different-body-fingerprint"
        ),
    )

    standardizer = (
        LessonPlanWordStandardizer(
            _profile()
        )
    )

    with pytest.raises(
        ValueError,
        match=(
            "body or table text changed"
        ),
    ):
        standardizer._validate_integrity(
            before,
            after,
        )


def test_integrity_gate_rejects_media_change(
    tmp_path: Path,
):
    source = tmp_path / "source.docx"

    document = Document()
    document.add_paragraph(
        "Document"
    )
    document.save(
        source
    )

    before = inventory(
        source
    )

    after = replace(
        before,
        media_fingerprints=(
            "word/media/image1.png:changed",
        ),
    )

    standardizer = (
        LessonPlanWordStandardizer(
            _profile()
        )
    )

    with pytest.raises(
        ValueError,
        match="media_fingerprints",
    ):
        standardizer._validate_integrity(
            before,
            after,
        )


def test_header_footer_change_rejected_in_preserve_mode(
    tmp_path: Path,
):
    source = tmp_path / "source.docx"

    document = Document()
    document.sections[
        0
    ].header.paragraphs[
        0
    ].text = "School header"

    document.save(
        source
    )

    before = inventory(
        source
    )

    after = replace(
        before,
        header_footer_text_sha256=(
            "different-header"
        ),
    )

    standardizer = (
        LessonPlanWordStandardizer(
            _profile()
        )
    )

    with pytest.raises(
        ValueError,
        match="header/footer",
    ):
        standardizer._validate_integrity(
            before,
            after,
        )


def test_header_footer_change_allowed_when_explicit(
    tmp_path: Path,
):
    source = tmp_path / "source.docx"

    document = Document()
    document.save(
        source
    )

    before = inventory(
        source
    )

    after = replace(
        before,
        header_footer_text_sha256=(
            "explicitly-changed-header"
        ),
    )

    standardizer = (
        LessonPlanWordStandardizer(
            {
                "header_footer": {
                    "remove_existing": True,
                }
            }
        )
    )

    standardizer._validate_integrity(
        before,
        after,
    )
