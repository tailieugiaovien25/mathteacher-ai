from __future__ import annotations

from hashlib import sha256
from io import BytesIO

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from assessment_generation_v2.renderers import (
    AssessmentDocumentRenderPlan,
    AssessmentDocxRendererError,
    AssessmentDocxRenderPlanRenderer,
)


def _plan(
    *,
    document_type: str = "MATRIX",
    sections: tuple[dict[str, object], ...] | None = None,
    layout: dict[str, object] | None = None,
    styles: dict[str, object] | None = None,
    asset_path: str | None = None,
    asset_hash: str | None = None,
) -> AssessmentDocumentRenderPlan:
    if sections is None:
        sections = (
            {
                "section_code": "HEADER",
                "section_type": "FIELDS",
                "field_labels": {
                    "exam_title": "Tên đề",
                    "variant_code": "Mã đề",
                },
                "data": {
                    "exam_title": "Đề kiểm tra giữa học kỳ I",
                    "variant_code": "101",
                },
            },
            {
                "section_code": "MATRIX",
                "section_type": "TABLE",
                "title": "MA TRẬN ĐỀ KIỂM TRA",
                "title_alignment": "CENTER",
                "columns": (
                    {
                        "column_code": "TOPIC",
                        "heading": "Chủ đề",
                        "value_path": "topic_name",
                    },
                    {
                        "column_code": "LEVEL",
                        "heading": "Mức độ",
                        "value_path": "cognitive_level_name",
                        "alignment": "CENTER",
                    },
                    {
                        "column_code": "COUNT",
                        "heading": "Số câu",
                        "value_path": "question_count",
                        "alignment": "CENTER",
                    },
                ),
                "data": {
                    "matrix_rows": (
                        {
                            "topic_name": "Số và Đại số",
                            "cognitive_level_name": "Nhận biết",
                            "question_count": 1,
                        },
                        {
                            "topic_name": "Hình học",
                            "cognitive_level_name": "Vận dụng",
                            "question_count": 2,
                        },
                    ),
                },
            },
        )
    return AssessmentDocumentRenderPlan(
        schema_version=1,
        document_type_code=document_type,
        renderer_code="DOCX_JSON_V1",
        layout=(
            layout
            if layout is not None
            else {
                "paper_size": "A4",
                "page_orientation": "LANDSCAPE",
                "margins_mm": {
                    "top": 20,
                    "right": 15,
                    "bottom": 20,
                    "left": 20,
                },
            }
        ),
        styles=(
            styles
            if styles is not None
            else {
                "font_family": "Times New Roman",
                "font_size": 12,
                "table_font_size": 11,
                "heading_size": 14,
                "line_spacing": 1.15,
                "paragraph_space_after_pt": 3,
            }
        ),
        bindings={},
        sections=sections,
        template_asset_path=asset_path,
        template_asset_hash=asset_hash,
    )


def _render(plan: AssessmentDocumentRenderPlan | None = None) -> bytes:
    return AssessmentDocxRenderPlanRenderer().render(
        plan=plan or _plan()
    )


def test_returns_valid_docx_bytes() -> None:
    content = _render()

    assert content.startswith(b"PK")
    document = Document(BytesIO(content))
    text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    assert "Đề kiểm tra giữa học kỳ I" in text
    assert "MA TRẬN ĐỀ KIỂM TRA" in text
    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 0).text == (
        "Số và Đại số"
    )


def test_applies_a4_landscape_and_margins() -> None:
    document = Document(BytesIO(_render()))
    section = document.sections[0]

    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert round(section.page_width.mm, 1) == 297.0
    assert round(section.page_height.mm, 1) == 210.0
    assert round(section.left_margin.mm, 1) == 20.0
    assert round(section.right_margin.mm, 1) == 15.0


def test_applies_vietnamese_font_to_runs() -> None:
    document = Document(BytesIO(_render()))
    run = document.paragraphs[0].runs[0]
    fonts = run._element.get_or_add_rPr().rFonts

    assert run.font.name == "Times New Roman"
    assert run.font.size.pt == 12
    assert fonts.get(qn("w:eastAsia")) == "Times New Roman"


def test_table_uses_fixed_dxa_geometry() -> None:
    document = Document(BytesIO(_render()))
    table = document.tables[0]
    widths = [
        int(column.get(qn("w:w")))
        for column in table._tbl.tblGrid.gridCol_lst
    ]

    assert table.autofit is False
    section = document.sections[0]
    usable_width = int(
        section.page_width.twips
        - section.left_margin.twips
        - section.right_margin.twips
    )

    assert widths == [4951, 4951, 4952]
    assert sum(widths) == usable_width
    table_width = table._tbl.tblPr.first_child_found_in("w:tblW")
    assert table_width.get(qn("w:type")) == "dxa"
    assert int(table_width.get(qn("w:w"))) == usable_width
    indentation = table._tbl.tblPr.first_child_found_in(
        "w:tblInd"
    )
    assert indentation.get(qn("w:w")) == "120"


def test_table_marks_repeating_header_without_fixed_heights() -> None:
    table = Document(BytesIO(_render())).tables[0]
    header_properties = table.rows[0]._tr.get_or_add_trPr()

    assert header_properties.find(qn("w:tblHeader")) is not None
    assert all(
        row._tr.get_or_add_trPr().find(qn("w:trHeight")) is None
        for row in table.rows
    )


def test_renders_text_repeat_and_page_break_sections() -> None:
    plan = _plan(
        document_type="STUDENT_EXAM",
        sections=(
            {
                "section_code": "INTRO",
                "section_type": "TEXT",
                "text": "Làm bài trực tiếp vào đề.",
                "data": {},
            },
            {
                "section_code": "QUESTIONS",
                "section_type": "REPEAT",
                "fields": (
                    {
                        "label": "Câu",
                        "value_path": "display_number",
                    },
                    {
                        "value_path": "prompt_text",
                    },
                ),
                "data": {
                    "questions": (
                        {
                            "display_number": 1,
                            "prompt_text": (
                                "Số tự nhiên nhỏ nhất là số nào?"
                            ),
                        },
                    ),
                },
            },
            {
                "section_code": "BREAK",
                "section_type": "PAGE_BREAK",
                "data": {},
            },
            {
                "section_code": "END",
                "section_type": "TEXT",
                "text": "Hết",
                "data": {},
            },
        ),
    )
    document = Document(BytesIO(_render(plan)))
    text = "\n".join(p.text for p in document.paragraphs)

    assert "Làm bài trực tiếp vào đề." in text
    assert "Câu: 1 | Số tự nhiên nhỏ nhất là số nào?" in text
    assert "Hết" in text
    assert "w:type=\"page\"" in document._element.xml


@pytest.mark.parametrize(
    "document_type",
    (
        "MATRIX",
        "SPECIFICATION",
        "STUDENT_EXAM",
        "ANSWER_KEY",
        "SCORING_GUIDE",
    ),
)
def test_renders_every_document_type(document_type: str) -> None:
    content = _render(_plan(document_type=document_type))

    assert Document(BytesIO(content)).tables


def test_opens_verified_template_asset() -> None:
    template = Document()
    template.add_paragraph("CƠ QUAN QUẢN LÝ")
    stream = BytesIO()
    template.save(stream)
    asset = stream.getvalue()
    plan = _plan(
        asset_path="templates/matrix.docx",
        asset_hash=sha256(asset).hexdigest(),
    )

    content = AssessmentDocxRenderPlanRenderer().render(
        plan=plan,
        template_asset=asset,
    )
    text = "\n".join(
        paragraph.text
        for paragraph in Document(BytesIO(content)).paragraphs
    )

    assert "CƠ QUAN QUẢN LÝ" in text
    assert "Đề kiểm tra giữa học kỳ I" in text


def test_rejects_template_hash_mismatch() -> None:
    plan = _plan(
        asset_path="templates/matrix.docx",
        asset_hash="a" * 64,
    )

    with pytest.raises(
        AssessmentDocxRendererError,
        match="hash does not match",
    ):
        AssessmentDocxRenderPlanRenderer().render(
            plan=plan,
            template_asset=b"not-the-approved-template",
        )


def test_rejects_missing_or_unexpected_template_bytes() -> None:
    with pytest.raises(
        AssessmentDocxRendererError,
        match="must match the render plan",
    ):
        AssessmentDocxRenderPlanRenderer().render(
            plan=_plan(
                asset_path="templates/matrix.docx",
                asset_hash="a" * 64,
            )
        )

    with pytest.raises(
        AssessmentDocxRendererError,
        match="must match the render plan",
    ):
        AssessmentDocxRenderPlanRenderer().render(
            plan=_plan(),
            template_asset=b"unexpected",
        )


def test_rejects_table_width_mismatch() -> None:
    sections = (
        {
            "section_code": "BAD_TABLE",
            "section_type": "TABLE",
            "columns": (
                {
                    "column_code": "TOPIC",
                    "value_path": "topic_name",
                    "width_dxa": 1000,
                },
            ),
            "data": {
                "rows": (
                    {"topic_name": "Số và Đại số"},
                ),
            },
        },
    )

    with pytest.raises(
        AssessmentDocxRendererError,
        match="usable page width",
    ):
        _render(_plan(sections=sections))


def test_student_exam_plan_contains_no_answer_content() -> None:
    plan = _plan(
        document_type="STUDENT_EXAM",
        sections=(
            {
                "section_code": "QUESTIONS",
                "section_type": "REPEAT",
                "fields": (
                    {"value_path": "prompt_text"},
                ),
                "data": {
                    "questions": (
                        {
                            "prompt_text": "Tính 2 + 2.",
                        },
                    ),
                },
            },
        ),
    )

    content = _render(plan)
    text = "\n".join(
        paragraph.text
        for paragraph in Document(BytesIO(content)).paragraphs
    )

    assert "Tính 2 + 2." in text
    assert "Đáp án" not in text
    assert "correct_options" not in text


def test_renderer_has_no_authority_specific_layout() -> None:
    from pathlib import Path

    source = Path(
        "src/assessment_generation_v2/renderers/"
        "docx_render_plan_renderer.py"
    ).read_text(encoding="utf-8")

    forbidden = (
        "DIEN_BIEN",
        "PHONG_GIAO_DUC",
        "SO_GIAO_DUC",
        "Mau_Phong",
        "Mau_So",
    )

    for value in forbidden:
        assert value not in source
