from pathlib import Path

import pytest
from docx import Document

from document_standardization import LessonPlanWordStandardizer


PROFILE = {
    "profile_name": "test", "page": {"margin_left_cm": 3, "margin_right_cm": 2, "margin_top_cm": 2, "margin_bottom_cm": 2},
    "body": {"font": "Times New Roman", "size_pt": 13, "line_spacing": 1.15},
    "title": {"size_pt": 14}, "table": {"size_pt": 12},
    "header_footer": {"remove_existing": True, "page_number": True},
    "equations": {"mode": "safe", "text_font": "Times New Roman"},
}


def make_docx(path: Path):
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Đầu trang cũ"
    document.sections[0].footer.paragraphs[0].text = "Chân trang cũ"
    document.add_paragraph("BÀI 1: BÀI HỌC MẪU")
    document.add_paragraph("I. MỤC TIÊU")
    document.add_paragraph("Nội dung không được thay đổi.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text, table.cell(0, 1).text = "HOẠT ĐỘNG", "SẢN PHẨM"
    table.cell(1, 0).text, table.cell(1, 1).text = "GV giao nhiệm vụ", "Câu trả lời"
    document.save(path)


def test_standardizer_preserves_source_and_content(tmp_path):
    source, output, report = tmp_path / "source.docx", tmp_path / "out.docx", tmp_path / "report.json"
    make_docx(source)
    source_bytes = source.read_bytes()
    result = LessonPlanWordStandardizer(PROFILE).standardize(source, output, report)
    assert source.read_bytes() == source_bytes
    assert output.exists() and report.exists() and result["source_preserved"]
    document = Document(output)
    assert "Nội dung không được thay đổi." in "\n".join(p.text for p in document.paragraphs)
    assert len(document.tables) == 1
    assert document.paragraphs[0].runs[0].font.name == "Times New Roman"


def test_standardizer_refuses_to_overwrite_source(tmp_path):
    source = tmp_path / "source.docx"; make_docx(source)
    with pytest.raises(ValueError, match="ghi đè"):
        LessonPlanWordStandardizer(PROFILE).standardize(source, source, tmp_path / "report.json")


def test_standardizer_scales_wide_tables_to_printable_width(tmp_path):
    source, output = tmp_path / "source.docx", tmp_path / "out.docx"
    make_docx(source)
    result = LessonPlanWordStandardizer(PROFILE).standardize(source, output, tmp_path / "report.json")
    document = Document(output)
    section = document.sections[0]
    usable_dxa = round((section.page_width - section.left_margin - section.right_margin) / 635)
    widths = [int(column.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) for column in document.tables[0]._tbl.tblGrid.gridCol_lst]
    assert sum(widths) <= usable_dxa
    assert result["before"]["table_cells"] == result["after"]["table_cells"]


def test_standardizer_allows_long_data_rows_to_continue_on_next_page(tmp_path):
    source, output = tmp_path / "source.docx", tmp_path / "out.docx"
    make_docx(source)
    LessonPlanWordStandardizer(PROFILE).standardize(source, output, tmp_path / "report.json")
    rows = Document(output).tables[0].rows
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    assert rows[0]._tr.trPr.find(namespace + "cantSplit") is not None
    assert rows[1]._tr.trPr.find(namespace + "cantSplit") is None


def test_standardizer_replaces_headers_and_footers_with_automatic_page_number(tmp_path):
    source, output = tmp_path / "source.docx", tmp_path / "out.docx"
    make_docx(source)
    result = LessonPlanWordStandardizer(PROFILE).standardize(
        source, output, tmp_path / "report.json"
    )
    document = Document(output)
    section = document.sections[0]
    assert "Đầu trang cũ" not in section.header.paragraphs[0].text
    assert "Chân trang cũ" not in section.footer.paragraphs[0].text
    footer_xml = section.footer._element.xml
    assert " PAGE " in footer_xml
    assert 'w:val="center"' in footer_xml
    assert result["changes"]["automatic_page_numbers_added"] >= 1
