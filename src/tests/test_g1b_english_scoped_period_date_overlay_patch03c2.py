from datetime import date
from io import BytesIO

from docx import Document

from document_standardization.lesson_plan_multi_period_scope import (
    apply_scoped_english_period_dates,
    resolve_multi_period_scope,
)


def build(periods, *, corrupted=False):
    document = Document()
    for period in periods:
        table = document.add_table(rows=1, cols=2)
        left, right = table.rows[0].cells
        left.text = "Date of planning:..........\nDate of teaching: ..........\nWEEK: 4"
        shown = 10 if corrupted else period
        right.text = f"Period {shown} : UNIT 2: MY HOUSE\nLesson for source period {period}"
    value = BytesIO()
    document.save(value)
    return value.getvalue()


def context():
    return {
        "curriculum_periods": [10, 11, 12],
        "occurrences": [
            {"curriculum_period": 10, "teaching_date": "2026-09-14", "class_id": "6A1"},
            {"curriculum_period": 10, "teaching_date": "2026-09-14", "class_id": "6A2"},
            {"curriculum_period": 11, "teaching_date": "2026-09-15", "class_id": "6A1"},
            {"curriculum_period": 11, "teaching_date": "2026-09-15", "class_id": "6A2"},
            {"curriculum_period": 12, "teaching_date": "2026-09-16", "class_id": "6A1"},
            {"curriculum_period": 12, "teaching_date": "2026-09-16", "class_id": "6A2"},
        ],
    }


def table_texts(content):
    document = Document(BytesIO(content))
    return ["\n".join(cell.text for row in table.rows for cell in row.cells) for table in document.tables]


def test_scoped_overlay_restores_periods_and_updates_only_10_to_12():
    periods = tuple(range(9, 16))
    scope = resolve_multi_period_scope(group_context=context(), document_periods=periods)
    result = apply_scoped_english_period_dates(
        source_content=build(periods),
        output_content=build(periods, corrupted=True),
        scope=scope,
        drafting_date=date(2026, 9, 11),
    )
    texts = table_texts(result.content)
    assert result.applied_periods == (10, 11, 12)
    for index, period in enumerate(periods):
        assert f"Period {period} " in texts[index]
        assert f"Lesson for source period {period}" in texts[index]
    assert "Date of planning:11/09/2026" in texts[1]
    assert "Date of teaching: 14/09/2026" in texts[1]
    assert "Date of teaching: 15/09/2026" in texts[2]
    assert "Date of teaching: 16/09/2026" in texts[3]
    for index in (0, 4, 5, 6):
        assert "Date of planning:.........." in texts[index]
        assert "Date of teaching: .........." in texts[index]


def test_conflicting_dates_fail_closed_for_that_period():
    value = context()
    value["occurrences"].append(
        {"curriculum_period": 11, "teaching_date": "2026-09-16", "class_id": "6A3"}
    )
    periods = tuple(range(9, 16))
    scope = resolve_multi_period_scope(group_context=value, document_periods=periods)
    result = apply_scoped_english_period_dates(
        source_content=build(periods),
        output_content=build(periods, corrupted=True),
        scope=scope,
        drafting_date=date(2026, 9, 11),
    )
    assert result.applied_periods == (10, 12)
    assert any(
        warning.code == "AMBIGUOUS_PERIOD_TEACHING_DATE"
        and warning.curriculum_period == 11
        for warning in result.warnings
    )
