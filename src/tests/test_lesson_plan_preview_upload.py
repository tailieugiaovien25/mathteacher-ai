from pathlib import Path

import pytest
from docx import Document

from document_intelligence.lesson_plan_preview_presenter import (
    PreviewReviewState,
)
from document_intelligence.lesson_plan_preview_upload import (
    LessonPlanPreviewUploadService,
)
from document_intelligence.validation import (
    CanonicalDocumentContext,
)


def make_docx_bytes(
    tmp_path: Path,
) -> bytes:
    path = tmp_path / "lesson.docx"

    document = Document()

    document.add_paragraph(
        "Ngày soạn: 07/09/2026"
    )

    document.add_paragraph(
        "Ngày dạy: 08/09/2026"
    )

    document.add_paragraph(
        "Lớp: 8A2"
    )

    document.add_paragraph(
        "Tiết 4. Bài 3. Đơn thức"
    )

    document.save(path)

    return path.read_bytes()


def canonical():
    return CanonicalDocumentContext(
        class_name="8A2",
        curriculum_period=4,
        lesson_title="Đơn thức",
        drafting_date="07/09/2026",
        teaching_date="08/09/2026",
    )


def test_upload_service_prepares_preview(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    view = (
        LessonPlanPreviewUploadService()
        .prepare(
            content=make_docx_bytes(
                tmp_path
            ),
            canonical=canonical(),
        )
    )

    assert len(view.items) == 5
    assert view.requires_review is False
    assert view.conflict_count == 0

    assert all(
        item.review_state
        is PreviewReviewState.ACCEPTED
        for item in view.items
    )


def test_upload_service_exposes_conflict(
    tmp_path,
    monkeypatch,
):
    monkeypatch.setenv(
        "DOCUMENT_AI_ENABLED",
        "false",
    )

    view = (
        LessonPlanPreviewUploadService()
        .prepare(
            content=make_docx_bytes(
                tmp_path
            ),
            canonical=(
                CanonicalDocumentContext(
                    class_name="8A1",
                    curriculum_period=4,
                    lesson_title="Đơn thức",
                    drafting_date="07/09/2026",
                    teaching_date="08/09/2026",
                )
            ),
        )
    )

    assert view.requires_review is True
    assert view.conflict_count == 1


def test_upload_service_rejects_empty_bytes():
    with pytest.raises(
        ValueError,
        match="content must not be empty",
    ):
        (
            LessonPlanPreviewUploadService()
            .prepare(
                content=b"",
                canonical=canonical(),
            )
        )


def test_upload_service_rejects_invalid_content():
    with pytest.raises(
        TypeError,
        match="content must be bytes",
    ):
        (
            LessonPlanPreviewUploadService()
            .prepare(
                content="not-bytes",
                canonical=canonical(),
            )
        )


def test_upload_service_rejects_invalid_canonical(
    tmp_path,
):
    with pytest.raises(
        TypeError,
        match="canonical must be CanonicalDocumentContext",
    ):
        (
            LessonPlanPreviewUploadService()
            .prepare(
                content=make_docx_bytes(
                    tmp_path
                ),
                canonical=object(),
            )
        )
