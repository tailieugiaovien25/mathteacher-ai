
from datetime import date
from io import BytesIO
from types import SimpleNamespace

from docx import Document

import portal_v2.ui.weekly_schedule_streamlit as module


def assignment(
    class_id,
    *,
    subject="MATHEMATICS",
    component="ALGEBRA",
):
    return SimpleNamespace(
        class_id=class_id,
        subject_ref=subject,
        component_ref=component,
    )


def row(
    class_id,
    teaching_date,
    *,
    lesson_id="LESSON-009",
    curriculum_period=9,
):
    return SimpleNamespace(
        class_id=class_id,
        subject_ref="MATHEMATICS",
        component_ref="ALGEBRA",
        lesson_id=lesson_id,
        curriculum_period=curriculum_period,
        lesson_title="\u0110\u01a1n th\u1ee9c",
        teaching_date=teaching_date,
    )


def test_multiple_assigned_classes_use_week_dates():
    current = row(
        "7A2",
        date(2026, 9, 30),
    )

    pairs = (
        module._mt_resolve_teaching_date_pairs(
            current_row=current,
            assignments=(
                assignment("7A1"),
                assignment("7A2"),
                assignment("7A3"),
            ),
            weekly_rows=(
                row(
                    "7A1",
                    date(2026, 9, 29),
                ),
                current,
                row(
                    "7A3",
                    date(2026, 10, 2),
                ),
            ),
        )
    )

    assert pairs == (
        (
            "7A1",
            date(2026, 9, 29),
        ),
        (
            "7A2",
            date(2026, 9, 30),
        ),
        (
            "7A3",
            date(2026, 10, 2),
        ),
    )


def test_same_ppct_lesson_can_have_different_record_ids_by_class():
    current = row(
        "7A2",
        date(2026, 9, 30),
        lesson_id="LESSON-7A2-009",
        curriculum_period=9,
    )

    pairs = module._mt_resolve_teaching_date_pairs(
        current_row=current,
        assignments=(assignment("7A1"), assignment("7A2")),
        weekly_rows=(
            row(
                "7A1",
                date(2026, 9, 29),
                lesson_id="LESSON-7A1-009",
                curriculum_period=9,
            ),
            current,
        ),
    )

    assert pairs == (
        ("7A1", date(2026, 9, 29)),
        ("7A2", date(2026, 9, 30)),
    )


def test_processor_tuple_output_bytes_can_be_replaced():
    result = (
        "giao-an.docx",
        b"old-docx",
        ("warning",),
    )

    assert module._mt_result_output_bytes(result) == b"old-docx"
    assert module._mt_result_with_output_bytes(
        result,
        b"new-docx",
    ) == (
        "giao-an.docx",
        b"new-docx",
        ("warning",),
    )


def test_teaching_date_wrapper_updates_legacy_tuple_result(monkeypatch):
    current = row(
        "7A2",
        date(2026, 9, 30),
    )

    def base_processor(
        *,
        row,
        drafting_date,
        content,
        original_name,
        modification_plan=None,
        options=None,
        original_content=None,
        ai_revised_text="",
    ):
        return (original_name, b"old-docx", ())

    monkeypatch.setattr(
        module,
        "_mt_original_process_lesson_plan_upload_3b",
        base_processor,
    )
    monkeypatch.setattr(
        module,
        "_mt_load_active_assignments",
        lambda **kwargs: (
            assignment("7A1"),
            assignment("7A2"),
        ),
    )
    monkeypatch.setattr(
        module,
        "_mt_sync_teaching_date_bytes",
        lambda content, *, pairs: b"new-docx",
    )
    monkeypatch.setattr(
        module,
        "st",
        SimpleNamespace(
            session_state={
                module._MT_TEACHING_SYNC_ENABLED: True,
                "lesson_standardization_supabase_client": object(),
                "lesson_standardization_teacher_user_id": "teacher-1",
                "_standardization_current_academic_year": "2026-2027",
                "_standardization_current_week_view": SimpleNamespace(
                    rows=(
                        row("7A1", date(2026, 9, 29)),
                        current,
                    )
                ),
            }
        ),
    )

    result = module._mt_original_process_lesson_plan_upload_3c(
        row=current,
        drafting_date=date(2026, 9, 28),
        content=b"source-docx",
        original_name="giao-an.docx",
    )

    assert result == (
        "giao-an.docx",
        b"new-docx",
        (),
    )

def test_unassigned_class_is_removed():
    current = row(
        "7A2",
        date(2026, 9, 30),
    )

    pairs = (
        module._mt_resolve_teaching_date_pairs(
            current_row=current,
            assignments=(
                assignment("7A2"),
            ),
            weekly_rows=(
                row(
                    "7A1",
                    date(2026, 9, 29),
                ),
                current,
                row(
                    "7A3",
                    date(2026, 10, 2),
                ),
            ),
        )
    )

    assert pairs == (
        (
            "7A2",
            date(2026, 9, 30),
        ),
    )


def test_assigned_current_class_without_week_timetable_keeps_only_current():
    current = row(
        "7A2",
        date(2026, 9, 30),
    )

    pairs = (
        module._mt_resolve_teaching_date_pairs(
            current_row=current,
            assignments=(
                assignment("7A1"),
                assignment("7A2"),
                assignment("7A3"),
            ),
            weekly_rows=(
                row(
                    "7A1",
                    date(2026, 9, 29),
                ),
                row(
                    "7A3",
                    date(2026, 10, 2),
                ),
            ),
        )
    )

    assert pairs == (
        (
            "7A2",
            date(2026, 9, 30),
        ),
    )


def test_different_lesson_is_not_added():
    current = row(
        "7A2",
        date(2026, 9, 30),
    )

    pairs = (
        module._mt_resolve_teaching_date_pairs(
            current_row=current,
            assignments=(
                assignment("7A1"),
                assignment("7A2"),
            ),
            weekly_rows=(
                row(
                    "7A1",
                    date(2026, 9, 29),
                    lesson_id="LESSON-008",
                    curriculum_period=8,
                ),
                current,
            ),
        )
    )

    assert pairs == (
        (
            "7A2",
            date(2026, 9, 30),
        ),
    )


def test_word_block_removes_extra_and_adds_missing_classes():
    document = Document()

    paragraph = document.add_paragraph()

    label = paragraph.add_run(
        "Ng\u00e0y d\u1ea1y:"
    )
    label.bold = True

    value = paragraph.add_run(
        "\n7A2 - 30/09/2025"
    )
    value.italic = True

    document.add_paragraph(
        "N\u1ed9i dung gi\u00e1o \u00e1n "
        "gi\u1eef nguy\u00ean."
    )

    source = BytesIO()

    document.save(
        source
    )

    result = (
        module._mt_sync_teaching_date_bytes(
            source.getvalue(),
            pairs=(
                (
                    "7A1",
                    date(2026, 9, 29),
                ),
                (
                    "7A2",
                    date(2026, 9, 30),
                ),
                (
                    "7A3",
                    date(2026, 10, 2),
                ),
            ),
        )
    )

    processed = Document(
        BytesIO(result)
    )

    text = "\n".join(
        p.text
        for p in processed.paragraphs
    )

    assert (
        "7A1 - 29/09/2026"
        in text
    )

    assert (
        "7A2 - 30/09/2026"
        in text
    )

    assert (
        "7A3 - 02/10/2026"
        in text
    )

    assert (
        "7A2 - 30/09/2025"
        not in text
    )

    assert (
        "N\u1ed9i dung gi\u00e1o \u00e1n "
        "gi\u1eef nguy\u00ean."
        in text
    )


def test_word_block_single_class_deletes_other_classes():
    document = Document()

    document.add_paragraph(
        (
            "Ng\u00e0y d\u1ea1y:\n"
            "7A1 - 29/09/2025\n"
            "7A2 - 30/09/2025\n"
            "7A3 - 02/10/2025"
        )
    )

    source = BytesIO()

    document.save(
        source
    )

    result = (
        module._mt_sync_teaching_date_bytes(
            source.getvalue(),
            pairs=(
                (
                    "7A2",
                    date(2026, 9, 30),
                ),
            ),
        )
    )

    processed = Document(
        BytesIO(result)
    )

    text = processed.paragraphs[
        0
    ].text

    assert (
        "7A2 - 30/09/2026"
        in text
    )

    assert "7A1 -" not in text
    assert "7A3 -" not in text


def test_word_block_removes_old_dates_stored_in_following_paragraphs():
    document = Document()
    document.add_paragraph("Ngày dạy: 8A1 – 10/09/2025")
    document.add_paragraph("8A2 – 09/09/2025")
    document.add_paragraph("Nội dung bài dạy giữ nguyên.")

    source = BytesIO()
    document.save(source)

    result = module._mt_sync_teaching_date_bytes(
        source.getvalue(),
        pairs=(
            ("8A1", date(2026, 9, 10)),
            ("8A2", date(2026, 9, 9)),
        ),
    )

    processed = Document(BytesIO(result))
    text = "\n".join(p.text for p in processed.paragraphs)

    assert "8A1 - 10/09/2026" in text
    assert "8A2 - 09/09/2026" in text
    assert "10/09/2025" not in text
    assert "09/09/2025" not in text
    assert "Nội dung bài dạy giữ nguyên." in text


def test_processor_prefers_selected_display_class_date_pairs(monkeypatch):
    current = row("internal-class-id", date(2026, 9, 10))

    def base_processor(**kwargs):
        return (kwargs["original_name"], b"old-docx", ())

    captured = {}

    monkeypatch.setattr(
        module,
        "_mt_original_process_lesson_plan_upload_3b",
        base_processor,
    )
    monkeypatch.setattr(
        module,
        "_mt_sync_teaching_date_bytes",
        lambda content, *, pairs: captured.setdefault("pairs", pairs)
        and b"new-docx",
    )
    monkeypatch.setattr(
        module,
        "st",
        SimpleNamespace(
            session_state={
                module._MT_TEACHING_SYNC_ENABLED: True,
                "_standardization_selected_teaching_date_pairs": (
                    ("8A1", date(2026, 9, 10)),
                    ("8A2", date(2026, 9, 9)),
                ),
            }
        ),
    )

    result = module._mt_original_process_lesson_plan_upload_3c(
        row=current,
        drafting_date=date(2026, 9, 4),
        content=b"source-docx",
        original_name="giao-an.docx",
    )

    assert result[1] == b"new-docx"
    assert captured["pairs"] == (
        ("8A1", date(2026, 9, 10)),
        ("8A2", date(2026, 9, 9)),
    )


def test_processor_reads_row_through_real_var_keyword_wrapper(monkeypatch):
    current = row("internal-class-id", date(2026, 9, 10))

    def var_keyword_base(*args, **kwargs):
        return (kwargs["original_name"], b"old-docx", ())

    monkeypatch.setattr(
        module,
        "_mt_original_process_lesson_plan_upload_3b",
        var_keyword_base,
    )
    monkeypatch.setattr(
        module,
        "_mt_sync_teaching_date_bytes",
        lambda content, *, pairs: b"new-docx",
    )
    monkeypatch.setattr(
        module,
        "st",
        SimpleNamespace(
            session_state={
                module._MT_TEACHING_SYNC_ENABLED: True,
                "_standardization_selected_teaching_date_pairs": (
                    ("8A1", date(2026, 9, 10)),
                    ("8A2", date(2026, 9, 9)),
                ),
            }
        ),
    )

    result = module._mt_original_process_lesson_plan_upload_3c(
        row=current,
        drafting_date=date(2026, 9, 4),
        content=b"source-docx",
        original_name="giao-an.docx",
    )

    assert result[1] == b"new-docx"


def test_control_panel_contains_sync_option():
    source = open(
        "src/portal_v2/ui/"
        "weekly_schedule_streamlit.py",
        encoding="utf-8",
    ).read()

    assert (
        "standardization_assignment_timetable_sync_enabled"
        in source
    )

    assert (
        "\\u0110\\u1ed3ng b\\u1ed9 "
        "Ng\\u00e0y d\\u1ea1y"
        in source
    )
