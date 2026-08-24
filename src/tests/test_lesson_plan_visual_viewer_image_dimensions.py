from io import BytesIO

from docx import Document
from docx.shared import Inches
from PIL import Image

from scripts.teacher_portal.lesson_plan_visual_viewer import (
    build_document_html,
)


def _image_bytes():
    image = Image.new("RGB", (1200, 600), "white")
    output = BytesIO()
    image.save(output, format="PNG")
    output.seek(0)
    return output


def _document_bytes(width):
    document = Document()
    document.add_paragraph().add_run().add_picture(
        _image_bytes(),
        width=width,
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_viewer_uses_word_inline_shape_width_instead_of_full_width():
    html = build_document_html(_document_bytes(Inches(2)))

    assert "Hình ảnh trong tài liệu" not in html
    assert 'class="doc-inline-image"' in html
    assert "width:144.0pt" in html
    assert "max-width:100%;height:auto" in html


def test_smaller_word_image_renders_smaller_in_preview():
    small = build_document_html(_document_bytes(Inches(1)))
    large = build_document_html(_document_bytes(Inches(4)))

    assert small != large
    assert "width:72.0pt" in small
    assert "width:288.0pt" in large


def test_image_stays_between_surrounding_content_in_preview():
    document = Document()
    document.add_paragraph("Trước hình")
    document.add_paragraph().add_run().add_picture(
        _image_bytes(),
        width=Inches(2),
    )
    document.add_paragraph("Sau hình")
    output = BytesIO()
    document.save(output)

    rendered = build_document_html(output.getvalue())

    image_index = rendered.index('<img class="doc-inline-image"')
    assert rendered.index("Trước hình") < image_index
    assert image_index < rendered.index("Sau hình")


def test_viewer_renders_word_paragraph_bottom_border():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document()
    paragraph = document.add_paragraph("")
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    p_bdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(p_bdr)
    output = BytesIO()
    document.save(output)

    rendered = build_document_html(output.getvalue())

    assert "border-bottom:1px solid #222" in rendered
