from datetime import date
from pathlib import Path

from docx import Document

from lesson_planning_v2.lesson_plan_teaching_scope import (
    LessonPlanTeachingScope,
)
from lesson_planning_v2.services.weekly_lesson_plan_docx_renderer import (
    WeeklyLessonPlanDocxRenderer,
)
from lesson_planning_v2.weekly_lesson_plan_identity import (
    WeeklyLessonPlanIdentity,
)
from lesson_planning_v2.weekly_lesson_plan_word_document import (
    WeeklyLessonPlanWordApproval,
    WeeklyLessonPlanWordDocument,
    WeeklyLessonPlanWordHeader,
    WeeklyLessonPlanWordSection,
)


def word_document():
    identity = WeeklyLessonPlanIdentity(
        teacher_id="GV002",
        academic_year="2026-2027",
        week_number=8,
        subject_ref="FOREIGN-LANGUAGE-1",
        teaching_scope=(
            LessonPlanTeachingScope.for_class(
                class_id="CLASS-6A1",
            )
        ),
    )

    return WeeklyLessonPlanWordDocument(
        identity=identity,
        header=WeeklyLessonPlanWordHeader(
            teacher_id="GV002",
            academic_year="2026-2027",
            week_number=8,
            subject_ref="FOREIGN-LANGUAGE-1",
            scope_label="Lớp 6A1",
        ),
        sections=(
            WeeklyLessonPlanWordSection(
                period_number=1,
                curriculum_period=22,
                preparation_date=date(
                    2026,
                    10,
                    12,
                ),
                teaching_date=date(
                    2026,
                    10,
                    13,
                ),
                title="Lesson 1",
                class_id="CLASS-6A1",
                component_ref="COMP-A",
                content={
                    "objectives": "Objectives 1",
                    "materials": "Materials 1",
                    "teaching_process": "Process 1",
                },
            ),
            WeeklyLessonPlanWordSection(
                period_number=2,
                curriculum_period=23,
                preparation_date=date(
                    2026,
                    10,
                    14,
                ),
                teaching_date=date(
                    2026,
                    10,
                    15,
                ),
                title="Lesson 2",
                class_id="CLASS-6A1",
                component_ref="COMP-B",
                content={
                    "objectives": "Objectives 2",
                    "materials": "Materials 2",
                    "teaching_process": "Process 2",
                },
            ),
            WeeklyLessonPlanWordSection(
                period_number=3,
                curriculum_period=24,
                preparation_date=date(
                    2026,
                    10,
                    16,
                ),
                teaching_date=date(
                    2026,
                    10,
                    17,
                ),
                title="Lesson 3",
                class_id="CLASS-6A1",
                component_ref="COMP-A",
                content={
                    "objectives": "Objectives 3",
                    "materials": "Materials 3",
                    "teaching_process": "Process 3",
                },
            ),
        ),
        approval=WeeklyLessonPlanWordApproval(
            approver_role="Tổ chuyên môn",
        ),
    )


def read_text(path):
    document = Document(path)

    parts = [
        paragraph.text
        for paragraph in document.paragraphs
    ]

    parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )

    return "\n".join(parts)


def test_renderer_creates_docx_file(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    result = renderer.render(
        document=word_document(),
        output_path=output,
    )

    assert result == output
    assert output.exists()
    assert output.stat().st_size > 0


def test_rendered_docx_can_be_opened(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    loaded = Document(output)

    assert loaded is not None


def test_renderer_writes_weekly_header(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    text = read_text(output)

    assert "GV002" in text
    assert "2026-2027" in text
    assert "8" in text

    assert (
        "FOREIGN-LANGUAGE-1"
        in text
    )

    assert "Lớp 6A1" in text


def test_renderer_writes_all_three_periods(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    text = read_text(output)

    assert "Lesson 1" in text
    assert "Lesson 2" in text
    assert "Lesson 3" in text

    assert "22" in text
    assert "23" in text
    assert "24" in text


def test_renderer_writes_each_preparation_and_teaching_date(
    tmp_path,
):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    text = read_text(output)

    assert "12/10/2026" in text
    assert "13/10/2026" in text

    assert "14/10/2026" in text
    assert "15/10/2026" in text

    assert "16/10/2026" in text
    assert "17/10/2026" in text


def test_renderer_writes_section_content(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    text = read_text(output)

    assert "Objectives 1" in text
    assert "Materials 2" in text
    assert "Process 3" in text


def test_renderer_preserves_component_ref(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    text = read_text(output)

    assert "COMP-A" in text
    assert "COMP-B" in text


def test_renderer_writes_approval_once(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    text = read_text(output)

    assert (
        text.count("Tổ chuyên môn")
        == 1
    )


def test_approval_is_after_last_lesson(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    text = read_text(output)

    assert (
        text.index("Tổ chuyên môn")
        > text.index("Lesson 3")
    )


def test_renderer_creates_parent_directory(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = (
        tmp_path
        / "nested"
        / "weekly-plan.docx"
    )

    renderer.render(
        document=word_document(),
        output_path=output,
    )

    assert output.exists()


def test_renderer_returns_path_object(tmp_path):
    renderer = WeeklyLessonPlanDocxRenderer()

    output = tmp_path / "weekly-plan.docx"

    result = renderer.render(
        document=word_document(),
        output_path=output,
    )

    assert isinstance(
        result,
        Path,
    )
