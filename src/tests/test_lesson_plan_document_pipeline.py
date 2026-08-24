from datetime import date

from docx import Document

from educational_planning_v2.models import (
    TeachingSession,
)
from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)
from document_standardization import (
    LessonPlanDocumentPipeline,
    LessonPlanWordStandardizer,
)


PROFILE = {
    "profile_name": "pipeline-test",
    "page": {
        "margin_left_cm": 3,
        "margin_right_cm": 2,
        "margin_top_cm": 2,
        "margin_bottom_cm": 2,
    },
    "body": {
        "font": "Times New Roman",
        "size_pt": 13,
        "line_spacing": 1.15,
    },
    "title": {
        "size_pt": 14,
    },
    "table": {
        "size_pt": 12,
    },
    "header_footer": {
        "remove_existing": False,
        "page_number": False,
    },
    "equations": {
        "mode": "safe",
        "text_font": "Times New Roman",
    },
}


def make_context():
    return ScheduledLessonContext(
        teaching_date=date(
            2026,
            9,
            28,
        ),
        drafting_date=date(
            2026,
            9,
            27,
        ),
        class_id="6A1",
        subject_ref="MATHEMATICS",
        component_ref="ALGEBRA",
        curriculum_period=9,
        lesson_id="LESSON-009",
        lesson_title="Phân số",
        session=TeachingSession.MORNING,
        timetable_period=1,
        period_in_lesson=1,
    )


def test_pipeline_applies_context_then_standardizes(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    report = tmp_path / "report.json"

    document = Document()

    document.add_paragraph(
        "Ngày soạn: 01/01/2020"
    )
    document.add_paragraph(
        "Ngày dạy: 02/01/2020"
    )
    document.add_paragraph(
        "Lớp: 6A2"
    )
    document.add_paragraph(
        "Tiết PPCT: 1"
    )
    document.add_paragraph(
        "Bài: Bài cũ"
    )
    document.add_paragraph(
        "Nội dung bài học giữ nguyên."
    )

    document.save(source)

    source_bytes = source.read_bytes()

    pipeline = LessonPlanDocumentPipeline(
        standardizer=(
            LessonPlanWordStandardizer(
                PROFILE
            )
        )
    )

    result = pipeline.process(
        source=source,
        output=output,
        report_path=report,
        context=make_context(),
    )

    assert source.read_bytes() == source_bytes
    assert output.exists()
    assert report.exists()

    updated = Document(output)

    texts = [
        paragraph.text
        for paragraph in updated.paragraphs
    ]

    assert "Ngày soạn: 27/09/2026" in texts
    assert "Ngày dạy: 28/09/2026" in texts
    assert "Lớp: 6A1" in texts
    assert "Tiết PPCT: 9" in texts
    assert "Bài: Phân số" in texts

    assert (
        "Nội dung bài học giữ nguyên."
        in texts
    )

    assert (
        result.context_result
        .unresolved_fields
        == ()
    )

    assert (
        result.standardization_report[
            "source_preserved"
        ]
        is True
    )


def test_pipeline_keeps_unresolved_fields_visible(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    report = tmp_path / "report.json"

    document = Document()
    document.add_paragraph(
        "Nội dung không có metadata."
    )
    document.save(source)

    pipeline = LessonPlanDocumentPipeline(
        standardizer=(
            LessonPlanWordStandardizer(
                PROFILE
            )
        )
    )

    result = pipeline.process(
        source=source,
        output=output,
        report_path=report,
        context=make_context(),
    )

    assert set(
        result.context_result
        .unresolved_fields
    ) == {
        "drafting_date",
        "teaching_date",
        "class_id",
        "curriculum_period",
        "lesson_title",
    }


def test_pipeline_rejects_source_overwrite(
    tmp_path,
):
    source = tmp_path / "source.docx"
    report = tmp_path / "report.json"

    document = Document()
    document.add_paragraph("Bài: cũ")
    document.save(source)

    pipeline = LessonPlanDocumentPipeline(
        standardizer=(
            LessonPlanWordStandardizer(
                PROFILE
            )
        )
    )

    try:
        pipeline.process(
            source=source,
            output=source,
            report_path=report,
            context=make_context(),
        )
    except ValueError as exc:
        assert "ghi đè" in str(exc)
    else:
        raise AssertionError(
            "Expected overwrite protection"
        )
