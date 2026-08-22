from docx import Document

from document_standardization.lesson_plan_metadata_locator import (
    LessonPlanMetadataLocator,
    MetadataField,
    MetadataLocationKind,
    MetadataMatchStrategy,
)


def _u(value):
    return value.encode(
        "ascii"
    ).decode(
        "unicode_escape"
    )


def _by_field(
    locations,
    field,
):
    return [
        item
        for item in locations
        if item.field == field
    ]


def test_inline_paragraph_metadata():
    document = Document()

    document.add_paragraph(
        _u(
            "Ng\\u00e0y so\\u1ea1n: "
            "01/09/2026"
        )
    )

    document.add_paragraph(
        _u(
            "Ng\\u00e0y d\\u1ea1y: "
            "03/09/2026"
        )
    )

    document.add_paragraph(
        _u(
            "L\\u1edbp: 6A1"
        )
    )

    document.add_paragraph(
        _u(
            "Ti\\u1ebft PPCT: 12"
        )
    )

    document.add_paragraph(
        _u(
            "M\\u00f4n: To\\u00e1n"
        )
    )

    locations = (
        LessonPlanMetadataLocator()
        .locate(document)
    )

    assert (
        _by_field(
            locations,
            MetadataField.DRAFTING_DATE,
        )[0].value_text
        == "01/09/2026"
    )

    assert (
        _by_field(
            locations,
            MetadataField.TEACHING_DATE,
        )[0].value_text
        == "03/09/2026"
    )

    assert (
        _by_field(
            locations,
            MetadataField.CLASS_NAME,
        )[0].value_text
        == "6A1"
    )

    assert (
        _by_field(
            locations,
            MetadataField.CURRICULUM_PERIOD,
        )[0].value_text
        == "12"
    )

    assert (
        _by_field(
            locations,
            MetadataField.SUBJECT_NAME,
        )[0].value_text
        == _u("To\\u00e1n")
    )


def test_no_colon_variants_are_supported():
    document = Document()

    document.add_paragraph(
        _u(
            "NG\\u00c0Y SO\\u1ea0N "
            "10/09/2026"
        )
    )

    document.add_paragraph(
        _u(
            "Ng\\u00e0y d\\u1ea1y "
            "12/09/2026"
        )
    )

    document.add_paragraph(
        _u(
            "L\\u1edaP 8A2"
        )
    )

    document.add_paragraph(
        _u(
            "Ti\\u1ebft 15"
        )
    )

    locations = (
        LessonPlanMetadataLocator()
        .locate(document)
    )

    assert (
        _by_field(
            locations,
            MetadataField.CLASS_NAME,
        )[0].value_text
        == "8A2"
    )

    assert (
        _by_field(
            locations,
            MetadataField.CURRICULUM_PERIOD,
        )[0].value_text
        == "15"
    )


def test_paired_table_cells_have_real_value_target():
    document = Document()

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.cell(
        0,
        0,
    ).text = _u(
        "Ng\\u00e0y d\\u1ea1y:"
    )

    table.cell(
        0,
        1,
    ).text = "15/09/2026"

    locations = (
        LessonPlanMetadataLocator()
        .locate(document)
    )

    item = _by_field(
        locations,
        MetadataField.TEACHING_DATE,
    )[0]

    assert (
        item.kind
        == MetadataLocationKind
        .TABLE_PAIRED_CELL
    )

    assert (
        item.strategy
        == MetadataMatchStrategy
        .PAIRED_CELL
    )

    assert item.confidence == 0.99
    assert item.is_high_confidence

    assert item.cell_index == 0
    assert item.value_cell_index == 1

    assert (
        item.value_text
        == "15/09/2026"
    )


def test_multiple_pairs_in_same_row():
    document = Document()

    table = document.add_table(
        rows=1,
        cols=4,
    )

    table.cell(
        0,
        0,
    ).text = _u(
        "Ng\\u00e0y d\\u1ea1y:"
    )

    table.cell(
        0,
        1,
    ).text = "15/09/2026"

    table.cell(
        0,
        2,
    ).text = _u(
        "L\\u1edbp:"
    )

    table.cell(
        0,
        3,
    ).text = "6A2"

    locations = (
        LessonPlanMetadataLocator()
        .locate(document)
    )

    teaching = _by_field(
        locations,
        MetadataField.TEACHING_DATE,
    )[0]

    class_name = _by_field(
        locations,
        MetadataField.CLASS_NAME,
    )[0]

    assert (
        teaching.value_cell_index
        == 1
    )

    assert (
        class_name.value_cell_index
        == 3
    )

    assert (
        class_name.value_text
        == "6A2"
    )


def test_inline_table_metadata_targets_same_cell():
    document = Document()

    table = document.add_table(
        rows=1,
        cols=1,
    )

    table.cell(
        0,
        0,
    ).text = _u(
        "L\\u1edbp: 6A3"
    )

    locations = (
        LessonPlanMetadataLocator()
        .locate(document)
    )

    item = _by_field(
        locations,
        MetadataField.CLASS_NAME,
    )[0]

    assert (
        item.kind
        == MetadataLocationKind.TABLE_CELL
    )

    assert item.cell_index == 0
    assert item.value_cell_index == 0

    assert (
        item.value_text
        == "6A3"
    )


def test_numbered_lesson_heading_detected():
    document = Document()

    document.add_paragraph(
        _u(
            "B\\u00e0i 7. "
            "Th\\u1ee9 t\\u1ef1 "
            "th\\u1ef1c hi\\u1ec7n "
            "c\\u00e1c ph\\u00e9p t\\u00ednh"
        )
    )

    locations = (
        LessonPlanMetadataLocator()
        .locate(document)
    )

    item = _by_field(
        locations,
        MetadataField.LESSON_TITLE,
    )[0]

    assert (
        item.strategy
        == MetadataMatchStrategy
        .LESSON_HEADING
    )

    assert item.confidence == 0.95


def test_lesson_content_false_positives_are_rejected():
    document = Document()

    document.add_paragraph(
        _u(
            "B\\u00e0i t\\u1eadp 1"
        )
    )

    document.add_paragraph(
        _u(
            "Ti\\u1ebft h\\u1ecdc "
            "h\\u00f4m nay"
        )
    )

    document.add_paragraph(
        _u(
            "L\\u1edbp h\\u1ecdc "
            "tham gia ho\\u1ea1t "
            "\\u0111\\u1ed9ng"
        )
    )

    document.add_paragraph(
        _u(
            "M\\u00f4n h\\u1ecdc "
            "n\\u00e0y c\\u00f3 "
            "nhi\\u1ec1u n\\u1ed9i dung"
        )
    )

    locations = (
        LessonPlanMetadataLocator()
        .locate(document)
    )

    assert locations == ()


def test_locator_is_strictly_read_only():
    document = Document()

    paragraph = document.add_paragraph()

    first = paragraph.add_run(
        _u(
            "Ng\\u00e0y d\\u1ea1y:"
        )
    )
    first.bold = True

    second = paragraph.add_run(
        " 01/09/2026"
    )
    second.italic = True

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.cell(
        0,
        0,
    ).text = _u(
        "L\\u1edbp:"
    )

    table.cell(
        0,
        1,
    ).text = "6A1"

    before_xml = (
        document._element.xml
    )

    LessonPlanMetadataLocator().locate(
        document
    )

    after_xml = (
        document._element.xml
    )

    assert before_xml == after_xml

    assert (
        paragraph.runs[
            0
        ].bold
        is True
    )

    assert (
        paragraph.runs[
            1
        ].italic
        is True
    )
