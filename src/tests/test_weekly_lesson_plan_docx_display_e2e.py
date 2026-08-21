from datetime import date

from docx import Document

from lesson_planning_v2.services.weekly_lesson_plan_display_resolver import (
    WeeklyLessonPlanDisplayResolver,
)
from lesson_planning_v2.services.weekly_lesson_plan_docx_renderer import (
    WeeklyLessonPlanDocxRenderer,
)
from lesson_planning_v2.weekly_lesson_plan_docx_presentation import (
    WeeklyLessonPlanDocxPresentationProfile,
)
from lesson_planning_v2.weekly_lesson_plan_identity import (
    WeeklyLessonPlanIdentity,
)
from lesson_planning_v2.lesson_plan_teaching_scope import (
    LessonPlanTeachingScope,
)
from lesson_planning_v2.weekly_lesson_plan_word_document import (
    WeeklyLessonPlanWordApproval,
    WeeklyLessonPlanWordDocument,
    WeeklyLessonPlanWordHeader,
    WeeklyLessonPlanWordSection,
)


def _presentation():
    return WeeklyLessonPlanDocxPresentationProfile(
        document_title=(
            "GI\u00c1O \u00c1N TU\u1ea6N"
        ),
        teacher_label=(
            "Gi\u00e1o vi\u00ean"
        ),
        subject_label=(
            "M\u00f4n h\u1ecdc"
        ),
        academic_year_label=(
            "N\u0103m h\u1ecdc"
        ),
        week_label=(
            "Tu\u1ea7n"
        ),
        scope_label=(
            "L\u1edbp / Kh\u1ed1i"
        ),
        curriculum_period_label=(
            "Ti\u1ebft PPCT"
        ),
        preparation_date_label=(
            "Ng\u00e0y so\u1ea1n"
        ),
        teaching_date_label=(
            "Ng\u00e0y d\u1ea1y"
        ),
        class_label="L\u1edbp",
        component_label=(
            "Ph\u00e2n m\u00f4n"
        ),
        objectives_label=(
            "I. M\u1ee5c ti\u00eau"
        ),
        materials_label=(
            "II. Thi\u1ebft b\u1ecb "
            "v\u00e0 h\u1ecdc li\u1ec7u"
        ),
        teaching_process_label=(
            "III. Ti\u1ebfn tr\u00ecnh "
            "d\u1ea1y h\u1ecdc"
        ),
        show_document_title=True,
        show_component=True,
        page_break_between_sections=False,
        approval_blank_lines=3,
    )


def _display():
    return (
        WeeklyLessonPlanDisplayResolver()
        .resolve(
            teacher_id="GV002",
            subject_ref=(
                "FOREIGN-LANGUAGE-1"
            ),
            class_id="CLASS-6A1",
            component_ref="COMP-A",
            teacher_name_resolver=(
                lambda _value: (
                    "Nguy\u1ec5n V\u0103n A"
                )
            ),
            subject_name_resolver=(
                lambda _value: (
                    "Ngo\u1ea1i ng\u1eef 1"
                )
            ),
            class_name_resolver=(
                lambda _value: (
                    "L\u1edbp 6A1"
                )
            ),
            component_name_resolver=(
                lambda _value: (
                    "Ti\u1ebfng Anh"
                )
            ),
        )
    )


def _document():
    display = _display()

    identity = WeeklyLessonPlanIdentity(
        teacher_id="GV002",
        academic_year="2026-2027",
        week_number=8,
        subject_ref=(
            "FOREIGN-LANGUAGE-1"
        ),
        teaching_scope=(
            LessonPlanTeachingScope.for_class(
                class_id="CLASS-6A1",
            )
        ),
    )

    sections = tuple(
        WeeklyLessonPlanWordSection(
            period_number=number,
            curriculum_period=(
                21 + number
            ),
            preparation_date=date(
                2026,
                10,
                10 + number,
            ),
            teaching_date=date(
                2026,
                10,
                11 + number,
            ),
            title=(
                f"Lesson {number}"
            ),
            class_id=(
                display.class_name
            ),
            component_ref=(
                display.component_name
            ),
            content={
                "objectives": (
                    f"Objectives {number}"
                ),
                "materials": (
                    f"Materials {number}"
                ),
                "teaching_process": (
                    f"Process {number}"
                ),
            },
        )
        for number in (1, 2, 3)
    )

    return WeeklyLessonPlanWordDocument(
        identity=identity,
        header=WeeklyLessonPlanWordHeader(
            teacher_id=(
                display.teacher_name
            ),
            academic_year="2026-2027",
            week_number=8,
            subject_ref=(
                display.subject_name
            ),
            scope_label=(
                display.class_name
            ),
        ),
        sections=sections,
        approval=(
            WeeklyLessonPlanWordApproval(
                approver_role=(
                    "T\u1ed5 chuy\u00ean m\u00f4n"
                ),
            )
        ),
    )


def _document_text(path):
    docx = Document(path)

    parts = []

    for paragraph in docx.paragraphs:
        parts.append(
            paragraph.text
        )

    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(
                    cell.text
                )

    return "\n".join(parts)


def test_e2e_docx_uses_display_names(
    tmp_path,
):
    output = (
        tmp_path
        / "weekly-display.docx"
    )

    WeeklyLessonPlanDocxRenderer().render(
        document=_document(),
        output_path=output,
        presentation_profile=(
            _presentation()
        ),
    )

    assert output.exists()
    assert output.stat().st_size > 0

    text = _document_text(output)

    required = (
        "Nguy\u1ec5n V\u0103n A",
        "Ngo\u1ea1i ng\u1eef 1",
        "L\u1edbp 6A1",
        "Ti\u1ebfng Anh",
    )

    for value in required:
        assert value in text


def test_e2e_docx_hides_canonical_ids(
    tmp_path,
):
    output = (
        tmp_path
        / "weekly-display.docx"
    )

    WeeklyLessonPlanDocxRenderer().render(
        document=_document(),
        output_path=output,
        presentation_profile=(
            _presentation()
        ),
    )

    text = _document_text(output)

    forbidden = (
        "GV002",
        "FOREIGN-LANGUAGE-1",
        "CLASS-6A1",
        "COMP-A",
    )

    for value in forbidden:
        assert value not in text


def test_approval_appears_once(
    tmp_path,
):
    output = (
        tmp_path
        / "weekly-display.docx"
    )

    WeeklyLessonPlanDocxRenderer().render(
        document=_document(),
        output_path=output,
        presentation_profile=(
            _presentation()
        ),
    )

    text = _document_text(output)

    assert (
        text.count(
            "T\u1ed5 chuy\u00ean m\u00f4n"
        )
        == 1
    )


def test_three_weekly_periods_are_present(
    tmp_path,
):
    output = (
        tmp_path
        / "weekly-display.docx"
    )

    WeeklyLessonPlanDocxRenderer().render(
        document=_document(),
        output_path=output,
        presentation_profile=(
            _presentation()
        ),
    )

    text = _document_text(output)

    assert "TI\u1ebeT 1: Lesson 1" in text
    assert "TI\u1ebeT 2: Lesson 2" in text
    assert "TI\u1ebeT 3: Lesson 3" in text
