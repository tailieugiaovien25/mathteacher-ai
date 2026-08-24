from datetime import date

from docx import Document

from document_standardization.lesson_plan_metadata import (
    LessonPlanMetadata,
)
from document_standardization.lesson_plan_metadata_locator import (
    MetadataField,
)
from document_standardization.lesson_plan_metadata_overlay import (
    LessonPlanMetadataOverlay,
)


def _u(value):
    return value.encode(
        "ascii"
    ).decode(
        "unicode_escape"
    )


def test_inline_overlay_preserves_run_formatting():
    document = Document()

    paragraph = (
        document.add_paragraph()
    )

    label = paragraph.add_run(
        _u(
            "Ng\\u00e0y d\\u1ea1y:"
        )
    )
    label.bold = True

    value = paragraph.add_run(
        " 01/09/2026"
    )
    value.italic = True

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    teaching_date=date(
                        2026,
                        9,
                        15,
                    )
                )
            ),
        )
    )

    assert (
        paragraph.text
        == _u(
            "Ng\\u00e0y d\\u1ea1y: "
            "15/09/2026"
        )
    )

    assert (
        paragraph.runs[
            0
        ].bold
        is True
    )

    assert (
        paragraph.runs[
            1
        ].italic
        is True
    )

    assert result.change_count == 1

    assert (
        result.changes[
            0
        ].field
        == MetadataField.TEACHING_DATE
    )


def test_paired_cell_overlay_changes_value_cell_only():
    document = Document()

    table = document.add_table(
        rows=1,
        cols=2,
    )

    label_paragraph = (
        table.cell(
            0,
            0,
        ).paragraphs[
            0
        ]
    )

    label_run = (
        label_paragraph.add_run(
            _u(
                "Ng\\u00e0y d\\u1ea1y:"
            )
        )
    )
    label_run.bold = True

    value_paragraph = (
        table.cell(
            0,
            1,
        ).paragraphs[
            0
        ]
    )

    value_run = (
        value_paragraph.add_run(
            "01/09/2026"
        )
    )
    value_run.italic = True

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    teaching_date=date(
                        2026,
                        9,
                        21,
                    )
                )
            ),
        )
    )

    assert (
        table.cell(
            0,
            0,
        ).text
        == _u(
            "Ng\\u00e0y d\\u1ea1y:"
        )
    )

    assert (
        table.cell(
            0,
            1,
        ).text
        == "21/09/2026"
    )

    assert (
        label_run.bold
        is True
    )

    assert (
        value_run.italic
        is True
    )

    assert result.change_count == 1


def test_multiple_occurrences_are_all_updated():
    document = Document()

    document.add_paragraph(
        _u(
            "L\\u1edbp: 6A1"
        )
    )

    document.add_paragraph(
        _u(
            "L\\u1edbp: 6A1"
        )
    )

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    class_name="6A2"
                )
            ),
        )
    )

    assert (
        document.paragraphs[
            0
        ].text
        == _u(
            "L\\u1edbp: 6A2"
        )
    )

    assert (
        document.paragraphs[
            1
        ].text
        == _u(
            "L\\u1edbp: 6A2"
        )
    )

    assert result.change_count == 2


def test_missing_metadata_is_preserved():
    document = Document()

    paragraph = (
        document.add_paragraph(
            _u(
                "L\\u1edbp: 6A1"
            )
        )
    )

    before_xml = (
        document._element.xml
    )

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata()
            ),
        )
    )

    assert (
        document._element.xml
        == before_xml
    )

    assert (
        paragraph.text
        == _u(
            "L\\u1edbp: 6A1"
        )
    )

    assert result.change_count == 0


def test_requested_field_without_location_is_unresolved():
    document = Document()

    document.add_paragraph(
        _u(
            "I. M\\u1ee4C TI\\u00caU"
        )
    )

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    teacher_name="Teacher A"
                )
            ),
        )
    )

    assert result.change_count == 0

    assert (
        MetadataField.TEACHER_NAME
        in result.unresolved_fields
    )


def test_low_confidence_label_only_is_not_modified():
    document = Document()

    table = document.add_table(
        rows=1,
        cols=1,
    )

    table.cell(
        0,
        0,
    ).text = _u(
        "L\\u1edbp:"
    )

    before_xml = (
        document._element.xml
    )

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    class_name="6A9"
                )
            ),
        )
    )

    assert (
        document._element.xml
        == before_xml
    )

    assert (
        MetadataField.CLASS_NAME
        in result.unresolved_fields
    )


def test_numbered_lesson_prefix_is_preserved():
    document = Document()

    paragraph = (
        document.add_paragraph(
            _u(
                "B\\u00e0i 7. "
                "T\\u00ean c\\u0169"
            )
        )
    )

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    lesson_title=_u(
                        "Th\\u1ee9 t\\u1ef1 "
                        "th\\u1ef1c hi\\u1ec7n "
                        "c\\u00e1c ph\\u00e9p "
                        "t\\u00ednh"
                    )
                )
            ),
        )
    )

    assert (
        paragraph.text
        == _u(
            "B\\u00e0i 7. "
            "Th\\u1ee9 t\\u1ef1 "
            "th\\u1ef1c hi\\u1ec7n "
            "c\\u00e1c ph\\u00e9p "
            "t\\u00ednh"
        )
    )

    assert result.change_count == 1


def test_already_canonical_value_is_resolved_without_change():
    document = Document()

    document.add_paragraph(
        _u(
            "L\\u1edbp: 6A1"
        )
    )

    before_xml = (
        document._element.xml
    )

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    class_name="6A1"
                )
            ),
        )
    )

    assert (
        document._element.xml
        == before_xml
    )

    assert result.change_count == 0

    assert (
        MetadataField.CLASS_NAME
        not in result.unresolved_fields
    )


def test_overlay_does_not_destroy_unrelated_runs():
    document = Document()

    paragraph = document.add_paragraph()

    first = paragraph.add_run(
        _u(
            "L\\u1edbp:"
        )
    )
    first.bold = True

    second = paragraph.add_run(
        " 6A1"
    )
    second.italic = True

    third = paragraph.add_run(
        " "
    )
    third.underline = True

    result = (
        LessonPlanMetadataOverlay()
        .apply(
            document=document,
            metadata=(
                LessonPlanMetadata(
                    class_name="6A2"
                )
            ),
        )
    )

    assert (
        paragraph.text
        == _u(
            "L\\u1edbp: 6A2 "
        )
    )

    assert first.bold is True
    assert second.italic is True
    assert third.underline is True

    assert result.change_count == 1
