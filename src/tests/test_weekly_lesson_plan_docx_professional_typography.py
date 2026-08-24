from pathlib import Path
import importlib.util

from docx import Document

from lesson_planning_v2.services.weekly_lesson_plan_docx_renderer import (
    WeeklyLessonPlanDocxRenderer,
)
from lesson_planning_v2.weekly_lesson_plan_docx_layout import (
    WeeklyLessonPlanDocxLayoutProfile,
)


def _fixture_module():
    path = Path(
        "src/tests/"
        "test_weekly_lesson_plan_docx_display_e2e.py"
    )

    spec = importlib.util.spec_from_file_location(
        "weekly_display_fixture",
        path,
    )

    if spec is None or spec.loader is None:
        raise RuntimeError(
            "fixture module cannot be loaded"
        )

    module = importlib.util.module_from_spec(
        spec
    )
    spec.loader.exec_module(module)

    return module


def _render(tmp_path):
    fixture = _fixture_module()

    output = (
        tmp_path
        / "professional-typography.docx"
    )

    WeeklyLessonPlanDocxRenderer().render(
        document=fixture._document(),
        output_path=output,
        presentation_profile=(
            fixture._presentation()
        ),
        layout_profile=(
            WeeklyLessonPlanDocxLayoutProfile
            .default()
        ),
    )

    return Document(output)


def _all_text(docx):
    parts = []

    for paragraph in docx.paragraphs:
        parts.append(paragraph.text)

    for table in docx.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def test_document_title_contains_week_number(
    tmp_path,
):
    docx = _render(tmp_path)
    text = _all_text(docx)

    assert "GIÁO ÁN TUẦN 8" in text

    # Week is integrated into the title, so the old
    # standalone metadata line is no longer needed.
    assert "Tuần: 8" not in text


def test_document_metadata_uses_two_column_table(
    tmp_path,
):
    docx = _render(tmp_path)

    matching_tables = []

    for table in docx.tables:
        text = "\n".join(
            cell.text
            for row in table.rows
            for cell in row.cells
        )

        if (
            "Giáo viên:" in text
            and "Môn học:" in text
            and "Lớp / Khối:" in text
            and "Năm học:" in text
        ):
            matching_tables.append(table)

    assert len(matching_tables) == 1

    table = matching_tables[0]

    assert len(table.columns) == 2


def test_period_metadata_uses_two_column_layout(
    tmp_path,
):
    docx = _render(tmp_path)

    period_tables = []

    for table in docx.tables:
        text = "\n".join(
            cell.text
            for row in table.rows
            for cell in row.cells
        )

        if (
            "Tiết PPCT:" in text
            and "Ngày soạn:" in text
            and "Ngày dạy:" in text
        ):
            period_tables.append(table)

    # One metadata table for each of 3 weekly periods.
    assert len(period_tables) == 3

    for table in period_tables:
        assert len(table.columns) == 2


def test_title_uses_title_typography_role(
    tmp_path,
):
    docx = _render(tmp_path)

    title = next(
        paragraph
        for paragraph in docx.paragraphs
        if paragraph.text.strip()
        == "GIÁO ÁN TUẦN 8"
    )

    assert title.alignment is not None

    runs = [
        run
        for run in title.runs
        if run.text.strip()
    ]

    assert runs

    run = runs[0]

    assert run.bold is True
    assert run.font.size is not None
    assert run.font.size.pt == 16


def test_period_titles_are_bold_and_kept_with_next(
    tmp_path,
):
    docx = _render(tmp_path)

    period_titles = [
        paragraph
        for paragraph in docx.paragraphs
        if paragraph.text.startswith(
            "TIẾT "
        )
    ]

    assert len(period_titles) == 3

    for paragraph in period_titles:
        assert (
            paragraph.paragraph_format
            .keep_with_next
            is True
        )

        runs = [
            run
            for run in paragraph.runs
            if run.text.strip()
        ]

        assert runs
        assert runs[0].bold is True


def test_content_headings_use_heading_role(
    tmp_path,
):
    docx = _render(tmp_path)

    prefixes = (
        "I. Mục tiêu",
        "II. Thiết bị và học liệu",
        "III. Tiến trình dạy học",
    )

    headings = [
        paragraph
        for paragraph in docx.paragraphs
        if paragraph.text.strip()
        in prefixes
    ]

    assert len(headings) == 9

    for paragraph in headings:
        runs = [
            run
            for run in paragraph.runs
            if run.text.strip()
        ]

        assert runs
        assert runs[0].bold is True
        assert runs[0].font.size is not None
        assert runs[0].font.size.pt == 14


def test_body_still_uses_times_new_roman_14(
    tmp_path,
):
    docx = _render(tmp_path)

    normal = docx.styles["Normal"]

    assert (
        normal.font.name
        == "Times New Roman"
    )

    assert normal.font.size.pt == 14


def test_approval_still_appears_once(
    tmp_path,
):
    docx = _render(tmp_path)
    text = _all_text(docx)

    assert (
        text.count("Tổ chuyên môn")
        == 1
    )
