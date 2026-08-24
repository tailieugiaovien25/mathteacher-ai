import os
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from lxml import etree


ENV_NAME = "REAL_TEACHER_DOCX"


def teacher_docx() -> Path:
    raw = os.environ.get(
        ENV_NAME
    )

    if not raw:
        pytest.skip(
            f"{ENV_NAME} is not configured"
        )

    path = Path(raw)

    if not path.exists():
        pytest.skip(
            f"Teacher DOCX does not exist: {path}"
        )

    return path


def all_paragraph_text(document):
    return "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )


def count_omml(path: Path) -> int:
    with ZipFile(path) as archive:
        xml = archive.read(
            "word/document.xml"
        )

    root = etree.fromstring(xml)

    namespace = {
        "m": (
            "http://schemas.openxmlformats.org/"
            "officeDocument/2006/math"
        )
    }

    return len(
        root.xpath(
            ".//m:oMath",
            namespaces=namespace,
        )
    )


def test_real_teacher_document_opens():
    path = teacher_docx()

    document = Document(path)

    assert len(document.paragraphs) == 214
    assert len(document.tables) == 1
    assert len(document.inline_shapes) == 3


def test_real_teacher_document_has_expected_first_lesson_metadata():
    document = Document(
        teacher_docx()
    )

    text = all_paragraph_text(
        document
    )

    assert (
        "Ngày soạn: 27/09/2025"
        in text
    )

    assert (
        "Ngày dạy: 30/09/2025"
        in text
    )

    assert "Lớp 6A2" in text

    assert (
        "TIẾT 10 + 11"
        in text
    )

    assert (
        "THỨ TỰ THỰC HIỆN "
        "CÁC PHÉP TÍNH"
        in text
    )


def test_real_teacher_document_contains_second_lesson():
    document = Document(
        teacher_docx()
    )

    text = all_paragraph_text(
        document
    )

    assert (
        "Ngày dạy: 02/10/2025"
        in text
    )

    assert "TIẾT 12" in text

    assert (
        "LUYỆN TẬP CHUNG"
        in text
    )


def test_real_teacher_document_contains_math_omml():
    assert (
        count_omml(
            teacher_docx()
        )
        == 9
    )


def test_real_teacher_document_has_nontrivial_visual_content():
    document = Document(
        teacher_docx()
    )

    assert len(
        document.inline_shapes
    ) == 3

    assert len(
        document.tables
    ) == 1


def test_characterization_does_not_modify_source():
    path = teacher_docx()

    before = path.read_bytes()

    Document(path)

    after = path.read_bytes()

    assert after == before
