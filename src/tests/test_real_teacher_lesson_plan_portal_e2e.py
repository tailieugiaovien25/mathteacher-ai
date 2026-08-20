import os
from datetime import date
from pathlib import Path
from types import SimpleNamespace

import pytest

from document_intelligence.contracts import (
    DocumentField,
)
from document_intelligence.lesson_plan_workflow_state import (
    LessonPlanWorkflowState,
)
from document_intelligence.lesson_plan_teacher_review_resolver import (
    LessonPlanTeacherReviewResolution,
    ResolvedLessonPlanMetadata,
)
from educational_planning_v2.models.weekly_teaching_schedule import (
    TeachingSession,
)

import portal_v2.ui.weekly_schedule_streamlit as module


ENV_NAME = "REAL_TEACHER_DOCX"


def teacher_docx() -> Path:
    raw = os.environ.get(
        ENV_NAME
    )

    if not raw:
        pytest.skip(
            f"{ENV_NAME} is not configured"
        )

    path = Path(raw)

    if not path.exists():
        pytest.skip(
            f"Teacher DOCX does not exist: {path}"
        )

    return path


class FakeUpload:
    def __init__(
        self,
        *,
        name,
        content,
    ):
        self.name = name
        self._content = content

    def getvalue(self):
        return self._content


class FakeStreamlit:
    def __init__(
        self,
        *,
        uploaded,
    ):
        self.uploaded = uploaded

        self.session_state = {}

        self.downloads = []
        self.errors = []
        self.warnings = []
        self.successes = []

        self.process_clicked = True

    def divider(self):
        pass

    def subheader(self, *args, **kwargs):
        pass

    def caption(self, *args, **kwargs):
        pass

    def info(self, *args, **kwargs):
        pass

    def success(self, message, *args, **kwargs):
        self.successes.append(message)

    def warning(self, message, *args, **kwargs):
        self.warnings.append(message)

    def error(self, message, *args, **kwargs):
        self.errors.append(message)

    def selectbox(
        self,
        label,
        *,
        options,
        format_func=None,
        key=None,
        **kwargs,
    ):
        return options[0]

    def date_input(
        self,
        label,
        *,
        value,
        max_value=None,
        key=None,
        **kwargs,
    ):
        return date(
            2026,
            9,
            29,
        )

    def file_uploader(
        self,
        *args,
        **kwargs,
    ):
        return self.uploaded

    def button(
        self,
        *args,
        **kwargs,
    ):
        return self.process_clicked

    def spinner(
        self,
        *args,
        **kwargs,
    ):
        class Spinner:
            def __enter__(self):
                return None

            def __exit__(
                self,
                exc_type,
                exc_value,
                traceback,
            ):
                return False

        return Spinner()

    def download_button(
        self,
        label,
        *,
        data,
        file_name,
        mime,
        **kwargs,
    ):
        self.downloads.append(
            {
                "label": label,
                "data": data,
                "file_name": file_name,
                "mime": mime,
            }
        )

        return False


def make_view():
    row = SimpleNamespace(
        teaching_date=date(
            2026,
            9,
            30,
        ),
        weekday=3,
        timetable_period=1,
        session=TeachingSession.MORNING,
        teacher_id="GV001",
        class_id="6A2",
        subject_ref="MATHEMATICS",
        curriculum_period=10,
        lesson_id="REAL-LESSON-001",
        lesson_title=(
            "Th\u1ee9 t\u1ef1 th\u1ef1c hi\u1ec7n "
            "c\u00e1c ph\u00e9p t\u00ednh"
        ),
        component_ref="ALGEBRA",
        period_in_lesson=1,
        total_lesson_periods=2,
        teaching_equipment=(),
    )

    return SimpleNamespace(
        week_number=5,
        rows=(row,),
    )


def accepted_resolution():
    return LessonPlanTeacherReviewResolution(
        accepted=True,
        metadata=ResolvedLessonPlanMetadata(
            values=(
                (
                    DocumentField.CLASS_NAME,
                    "6A3",
                ),
                (
                    DocumentField.CURRICULUM_PERIOD,
                    "20",
                ),
                (
                    DocumentField.LESSON_TITLE,
                    (
                        "Th\u1ee9 t\u1ef1 th\u1ef1c hi\u1ec7n "
                        "c\u00e1c ph\u00e9p t\u00ednh"
                    ),
                ),
                (
                    DocumentField.DRAFTING_DATE,
                    "29/09/2026",
                ),
                (
                    DocumentField.TEACHING_DATE,
                    "30/09/2026",
                ),
            )
        ),
        rejected_fields=(),
    )

def test_real_teacher_docx_runs_through_portal(
    monkeypatch,
):
    source = teacher_docx()

    original_bytes = (
        source.read_bytes()
    )

    uploaded = FakeUpload(
        name=source.name,
        content=original_bytes,
    )

    st = FakeStreamlit(
        uploaded=uploaded
    )

    resolution = (
        accepted_resolution()
    )

    review = object()

    class FakePreviewService:
        def prepare(
            self,
            *,
            content,
            canonical,
        ):
            assert content == original_bytes

            return SimpleNamespace(
                document=object()
            )

    class FakePresenter:
        def present(
            self,
            *,
            preview,
            canonical_values,
        ):
            return object()

    class FakeResolver:
        def resolve(
            self,
            *,
            preview,
            review,
        ):
            return resolution

    monkeypatch.setattr(
        module,
        "st",
        st,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanPreviewUploadService",
        FakePreviewService,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_preview",
        lambda **kwargs: None,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewPresenter",
        FakePresenter,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_teacher_review",
        lambda **kwargs: review,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewResolver",
        FakeResolver,
    )

    module._render_lesson_plan_standardization_workspace(
        make_view()
    )

    assert st.errors == []

    assert len(st.downloads) == 1

    download = st.downloads[0]

    assert (
        download["file_name"]
        .endswith(
            ".lbg-standardized.docx"
        )
    )

    assert (
        download["data"]
        .startswith(b"PK")
    )

    assert (
        "wordprocessingml.document"
        in download["mime"]
    )

    workflow_states = [
        value
        for value
        in st.session_state.values()
        if isinstance(
            value,
            LessonPlanWorkflowState,
        )
    ]

    assert len(workflow_states) == 1

    state = workflow_states[0]

    assert state.result is not None

    (
        output_name,
        output_bytes,
        unresolved_fields,
    ) = state.result

    assert (
        output_name
        == download["file_name"]
    )

    assert (
        output_bytes
        == download["data"]
    )

    assert unresolved_fields == ()

    # Source teacher file must remain unchanged.
    assert (
        source.read_bytes()
        == original_bytes
    )


def test_real_teacher_portal_rerun_reuses_result(
    monkeypatch,
):
    source = teacher_docx()

    uploaded = FakeUpload(
        name=source.name,
        content=source.read_bytes(),
    )

    st = FakeStreamlit(
        uploaded=uploaded
    )

    resolution = (
        accepted_resolution()
    )

    review = object()

    preview_calls = []

    class FakePreviewService:
        def prepare(
            self,
            *,
            content,
            canonical,
        ):
            preview_calls.append(
                content
            )

            return SimpleNamespace(
                document=object()
            )

    class FakePresenter:
        def present(
            self,
            *,
            preview,
            canonical_values,
        ):
            return object()

    class FakeResolver:
        def resolve(
            self,
            *,
            preview,
            review,
        ):
            return resolution

    monkeypatch.setattr(
        module,
        "st",
        st,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanPreviewUploadService",
        FakePreviewService,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_preview",
        lambda **kwargs: None,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewPresenter",
        FakePresenter,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_teacher_review",
        lambda **kwargs: review,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewResolver",
        FakeResolver,
    )

    module._render_lesson_plan_standardization_workspace(
        make_view()
    )

    assert len(
        st.downloads
    ) == 1

    first_download = (
        st.downloads[0]
    )

    st.downloads.clear()
    st.process_clicked = False

    module._render_lesson_plan_standardization_workspace(
        make_view()
    )

    assert len(
        preview_calls
    ) == 1

    assert len(
        st.downloads
    ) == 1

    second_download = (
        st.downloads[0]
    )

    assert (
        second_download["file_name"]
        == first_download["file_name"]
    )

    assert (
        second_download["data"]
        == first_download["data"]
    )

from io import BytesIO
from zipfile import ZipFile

from docx import Document
from lxml import etree


def portal_docx_omml_count(
    content: bytes,
) -> int:
    with ZipFile(
        BytesIO(content)
    ) as archive:
        xml = archive.read(
            "word/document.xml"
        )

    root = etree.fromstring(
        xml
    )

    namespace = {
        "m": (
            "http://schemas.openxmlformats.org/"
            "officeDocument/2006/math"
        )
    }

    return len(
        root.xpath(
            ".//m:oMath",
            namespaces=namespace,
        )
    )


def portal_docx_text(
    content: bytes,
) -> str:
    document = Document(
        BytesIO(content)
    )

    parts = []

    for paragraph in document.paragraphs:
        parts.append(
            paragraph.text
        )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(
                    cell.text
                )

    return "\n".join(parts)


def test_portal_download_preserves_real_docx_structure_and_content(
    monkeypatch,
):
    source = teacher_docx()

    original_bytes = (
        source.read_bytes()
    )

    original_document = Document(
        source
    )

    uploaded = FakeUpload(
        name=source.name,
        content=original_bytes,
    )

    st = FakeStreamlit(
        uploaded=uploaded
    )

    resolution = (
        accepted_resolution()
    )

    review = object()

    class FakePreviewService:
        def prepare(
            self,
            *,
            content,
            canonical,
        ):
            assert (
                content
                == original_bytes
            )

            return SimpleNamespace(
                document=object()
            )

    class FakePresenter:
        def present(
            self,
            *,
            preview,
            canonical_values,
        ):
            return object()

    class FakeResolver:
        def resolve(
            self,
            *,
            preview,
            review,
        ):
            return resolution

    monkeypatch.setattr(
        module,
        "st",
        st,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanPreviewUploadService",
        FakePreviewService,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_preview",
        lambda **kwargs: None,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewPresenter",
        FakePresenter,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_teacher_review",
        lambda **kwargs: review,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewResolver",
        FakeResolver,
    )

    module._render_lesson_plan_standardization_workspace(
        make_view()
    )

    assert st.errors == []

    assert len(
        st.downloads
    ) == 1

    content = (
        st.downloads[0]["data"]
    )

    assert content.startswith(
        b"PK"
    )

    output_document = Document(
        BytesIO(content)
    )

    assert (
        len(output_document.paragraphs)
        == len(
            original_document.paragraphs
        )
        == 214
    )

    assert (
        len(output_document.tables)
        == len(
            original_document.tables
        )
        == 1
    )

    assert (
        len(output_document.inline_shapes)
        == len(
            original_document.inline_shapes
        )
        == 3
    )

    assert (
        portal_docx_omml_count(
            original_bytes
        )
        == 9
    )

    assert (
        portal_docx_omml_count(
            content
        )
        == 9
    )

    text = portal_docx_text(
        content
    )

    assert "6A3" in text

    assert (
        "29/09/2026"
        in text
    )

    assert (
        "30/09/2026"
        in text
    )

    assert (
        "TI\u1ebeT 20 + 21"
        in text
    )

    assert (
        "TI\u1ebeT 12"
        in text
    )

    assert (
        "LUY\u1ec6N T\u1eacP CHUNG"
        in text
    )

    assert (
        source.read_bytes()
        == original_bytes
    )
