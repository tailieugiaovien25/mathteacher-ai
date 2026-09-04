from collections import Counter

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor

from document_standardization.lesson_plan_standardizer import (
    LessonPlanWordStandardizer,
)


def _profile():
    return {
        "profile_name": "a8b-r2-test",
        "page": {
            "margin_left_cm": 3.0,
            "margin_right_cm": 2.0,
            "margin_top_cm": 1.5,
            "margin_bottom_cm": 1.5,
        },
        "body": {
            "font": "Times New Roman",
            "size_pt": 14,
            "line_spacing": 1.0,
        },
        "title": {"size_pt": 14},
        "table": {"size_pt": 14},
        "header_footer": {
            "remove_existing": False,
            "page_number": False,
        },
        "equations": {
            "mode": "safe",
            "text_font": "Times New Roman",
        },
    }


def _set_grid(table, widths):
    for column, width in zip(
        table._tbl.tblGrid.gridCol_lst,
        widths,
    ):
        column.set(qn("w:w"), str(width))


def test_global_text_color_and_highlight_are_normalized():
    doc = Document()
    paragraph = doc.add_paragraph()
    outside = paragraph.add_run("Ngoài bảng")
    outside.font.color.rgb = RGBColor(255, 0, 0)

    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    outside._r.get_or_add_rPr().append(highlight)

    table = doc.add_table(rows=1, cols=1)
    inside = table.cell(0, 0).paragraphs[0].add_run("Trong bảng")
    inside.font.color.rgb = RGBColor(0, 0, 255)

    changes = Counter()
    LessonPlanWordStandardizer(
        _profile()
    )._normalize_global_text_color(
        doc,
        changes,
    )

    for run in (outside, inside):
        properties = run._r.rPr
        color = properties.find(qn("w:color"))
        assert color.get(qn("w:val")) == "000000"
        assert properties.find(qn("w:highlight")) is None

    assert changes["text_colors_normalized"] == 2
    assert changes["text_highlights_removed"] == 1


def test_wide_complex_table_uses_content_autofit():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    table = doc.add_table(rows=2, cols=3)
    _set_grid(table, [4000, 4000, 4000])

    changes = Counter()
    LessonPlanWordStandardizer(
        _profile()
    )._normalize_tables(
        doc,
        changes,
    )

    assert table.autofit is True
    layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    assert layout.get(qn("w:type")) == "autofit"

    width = table._tbl.tblPr.first_child_found_in("w:tblW")
    assert width.get(qn("w:type")) == "pct"
    assert width.get(qn("w:w")) == "5000"
    assert changes["tables_content_autofit"] == 1


def test_two_column_activity_table_balances_toward_content_demand():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    table = doc.add_table(rows=5, cols=2)
    _set_grid(table, [5400, 5400])

    table.cell(0, 0).text = "Hoạt động của giáo viên"
    table.cell(0, 1).text = "Hoạt động của học sinh"

    long_left = (
        "Giáo viên tổ chức hoạt động, đặt câu hỏi, "
        "hướng dẫn học sinh thảo luận và tổng hợp kết quả. "
    )

    for row in range(1, 5):
        table.cell(row, 0).text = long_left * 4
        table.cell(row, 1).text = "Học sinh trả lời."

    standardizer = LessonPlanWordStandardizer(_profile())
    ratio = standardizer._two_column_balance_ratio(table)

    assert ratio is not None
    assert 0.50 < ratio <= 0.64

    changes = Counter()
    standardizer._normalize_tables(doc, changes)

    widths = [
        int(column.get(qn("w:w")))
        for column in table._tbl.tblGrid.gridCol_lst
    ]

    assert widths[0] > widths[1]
    assert table.autofit is False
    assert changes["two_column_tables_balanced"] == 1


def test_short_two_column_table_is_not_forced_into_balance():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    _set_grid(table, [5000, 5000])

    ratio = LessonPlanWordStandardizer(
        _profile()
    )._two_column_balance_ratio(table)

    assert ratio is None