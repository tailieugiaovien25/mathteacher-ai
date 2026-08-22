from io import BytesIO

from docx import Document

from lesson_planning_v2.adapters.in_memory_lesson_plan_workspace_draft_repository import (
    InMemoryLessonPlanWorkspaceDraftRepository,
)
from lesson_planning_v2.services.lesson_plan_draft_workspace_service import (
    LessonPlanDraftWorkspaceService,
)
from lesson_planning_v2.services.lesson_plan_workspace_v1_service import (
    LessonPlanDocxTextImporter,
    LessonPlanSimpleDocxExporter,
    LessonPlanWorkspaceContent,
    LessonPlanWorkspaceContext,
    LessonPlanWorkspaceV1Service,
)


def _context(
    *,
    unit="lesson-007",
):
    return LessonPlanWorkspaceContext(
        teacher_user_id="teacher-001",
        academic_year="2026-2027",
        week_number=1,
        subject_ref="math",
        selection_mode="LESSON",
        selection_unit_id=unit,
        class_or_grade_ref="6A2",
        lesson_id=unit,
        title="Bài kiểm thử",
    )


def _service():
    repository = (
        InMemoryLessonPlanWorkspaceDraftRepository()
    )

    return LessonPlanWorkspaceV1Service(
        draft_service=(
            LessonPlanDraftWorkspaceService(
                repository
            )
        )
    )


def test_draft_identity_is_stable():
    first = _context()
    second = _context()

    assert first.draft_id == second.draft_id
    assert first.widget_prefix == second.widget_prefix


def test_different_lessons_have_different_identity():
    first = _context(
        unit="lesson-001",
    )

    second = _context(
        unit="lesson-007",
    )

    assert first.draft_id != second.draft_id
    assert (
        first.widget_prefix
        != second.widget_prefix
    )


def test_save_then_load_round_trip():
    service = _service()

    context = _context()

    content = LessonPlanWorkspaceContent(
        objectives_text="Mục tiêu",
        materials_text="Thiết bị",
        teaching_process_text="Tiến trình",
    )

    saved = service.save(
        context=context,
        content=content,
    )

    loaded = service.load(
        context=context,
    )

    assert loaded == saved

    assert (
        loaded.selection_unit_id
        == "lesson-007"
    )


def test_docx_import_recognizes_sections():
    document = Document()

    document.add_paragraph(
        "I. MỤC TIÊU"
    )
    document.add_paragraph(
        "Mục tiêu A"
    )

    document.add_paragraph(
        "II. THIẾT BỊ VÀ HỌC LIỆU"
    )
    document.add_paragraph(
        "Máy chiếu"
    )

    document.add_paragraph(
        "III. TIẾN TRÌNH DẠY HỌC"
    )
    document.add_paragraph(
        "Hoạt động 1"
    )

    stream = BytesIO()
    document.save(stream)

    imported = (
        LessonPlanDocxTextImporter()
        .import_bytes(
            stream.getvalue()
        )
    )

    assert (
        imported.objectives_text
        == "Mục tiêu A"
    )

    assert (
        imported.materials_text
        == "Máy chiếu"
    )

    assert (
        imported.teaching_process_text
        == "Hoạt động 1"
    )


def test_docx_export_uses_tnr14():
    data = (
        LessonPlanSimpleDocxExporter()
        .export(
            context=_context(),
            content=(
                LessonPlanWorkspaceContent(
                    objectives_text="A",
                    materials_text="B",
                    teaching_process_text="C",
                )
            ),
        )
    )

    loaded = Document(
        BytesIO(data)
    )

    normal = loaded.styles[
        "Normal"
    ]

    assert (
        normal.font.name
        == "Times New Roman"
    )

    assert normal.font.size.pt == 14
