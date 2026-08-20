from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document

from document_intelligence.contracts import (
    DocumentField,
)
from document_intelligence.lesson_plan_modification_plan import (
    LessonPlanFieldModification,
    LessonPlanModificationPlan,
)
from educational_planning_v2.models.weekly_teaching_schedule import (
    TeachingSession,
    WeeklyTeachingScheduleEntry,
)
from lesson_planning_v2.services.lesson_plan_document_processing_service import (
    LessonPlanDocumentProcessingService,
)


PROFILE = Path(
    "scripts/word_standardizer/"
    "lesson_plan_profile.json"
)


def make_source_docx() -> bytes:
    document = Document()

    document.add_paragraph(
        "K\u1ebe HO\u1ea0CH B\u00c0I D\u1ea0Y"
    )

    table = document.add_table(
        rows=5,
        cols=2,
    )

    values = (
        (
            "L\u1edbp",
            "8A1",
        ),
        (
            "Ti\u1ebft PPCT",
            "1",
        ),
        (
            "T\u00ean b\u00e0i",
            "B\u00e0i c\u0169",
        ),
        (
            "Ng\u00e0y so\u1ea1n",
            "07/09/2026",
        ),
        (
            "Ng\u00e0y d\u1ea1y",
            "08/09/2026",
        ),
    )

    for row, (label, value) in zip(
        table.rows,
        values,
    ):
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_paragraph(
        "I. M\u1ee4C TI\u00caU"
    )

    document.add_paragraph(
        "N\u1ed9i dung b\u00e0i h\u1ecdc "
        "ph\u1ea3i \u0111\u01b0\u1ee3c b\u1ea3o to\u00e0n."
    )

    buffer = BytesIO()

    document.save(buffer)

    return buffer.getvalue()


def make_schedule_row():
    return WeeklyTeachingScheduleEntry(
        teaching_date=date(
            2026,
            9,
            8,
        ),
        weekday=2,
        timetable_period=1,
        session=TeachingSession.MORNING,
        teacher_id="GV001",
        class_id="8A1",
        subject_ref="MATHEMATICS",
        curriculum_period=1,
        lesson_id="LESSON-001",
        lesson_title="B\u00e0i c\u0169",
        component_ref="ALGEBRA",
        period_in_lesson=1,
        total_lesson_periods=1,
        teaching_equipment=(),
    )


def make_modification_plan():
    return LessonPlanModificationPlan(
        modifications=(
            LessonPlanFieldModification(
                field=DocumentField.CLASS_NAME,
                value="8A2",
            ),
            LessonPlanFieldModification(
                field=DocumentField.CURRICULUM_PERIOD,
                value="9",
            ),
            LessonPlanFieldModification(
                field=DocumentField.LESSON_TITLE,
                value="\u0110\u01a1n th\u1ee9c",
            ),
            LessonPlanFieldModification(
                field=DocumentField.DRAFTING_DATE,
                value="27/09/2026",
            ),
            LessonPlanFieldModification(
                field=DocumentField.TEACHING_DATE,
                value="28/09/2026",
            ),
        )
    )


def document_text(content: bytes) -> str:
    document = Document(
        BytesIO(content)
    )

    parts = []

    for paragraph in document.paragraphs:
        parts.append(
            paragraph.text
        )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(
                    cell.text
                )

    return "\n".join(parts)


def test_teacher_modification_plan_reaches_real_docx():
    assert PROFILE.exists()

    service = (
        LessonPlanDocumentProcessingService(
            profile_path=PROFILE
        )
    )

    result = service.process(
        row=make_schedule_row(),
        drafting_date=date(
            2026,
            9,
            7,
        ),
        content=make_source_docx(),
        original_name="lesson.docx",
        modification_plan=(
            make_modification_plan()
        ),
    )

    assert result.output_name.endswith(
        ".docx"
    )

    assert result.output_bytes.startswith(
        b"PK"
    )

    text = document_text(
        result.output_bytes
    )

    assert "8A2" in text
    assert "9" in text
    assert "\u0110\u01a1n th\u1ee9c" in text
    assert "27/09/2026" in text
    assert "28/09/2026" in text

    assert "8A1" not in text
    assert "B\u00e0i c\u0169" not in text
    assert "07/09/2026" not in text
    assert "08/09/2026" not in text

    assert (
        "N\u1ed9i dung b\u00e0i h\u1ecdc "
        "ph\u1ea3i \u0111\u01b0\u1ee3c b\u1ea3o to\u00e0n."
        in text
    )

    assert result.unresolved_fields == ()
