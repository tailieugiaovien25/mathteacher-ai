
from io import BytesIO

from docx import Document

from lesson_planning_v2.services.lesson_plan_workspace_v1_service import (
    LessonPlanDocxWholeDocumentImporter,
)


def _make_document() -> bytes:
    stream = BytesIO()

    document = Document()

    document.add_paragraph(
        "LESSON TITLE"
    )

    document.add_paragraph(
        "FIRST PARAGRAPH"
    )

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(
        0,
        0,
    ).text = "A1"

    table.cell(
        0,
        1,
    ).text = "B1"

    table.cell(
        1,
        0,
    ).text = "A2"

    table.cell(
        1,
        1,
    ).text = "B2"

    document.add_paragraph(
        "AFTER TABLE"
    )

    document.save(
        stream
    )

    return stream.getvalue()


def test_whole_document_importer_preserves_block_order():
    content = (
        LessonPlanDocxWholeDocumentImporter()
        .import_bytes(
            _make_document()
        )
    )

    assert (
        content.index(
            "LESSON TITLE"
        )
        <
        content.index(
            "FIRST PARAGRAPH"
        )
        <
        content.index(
            "A1 | B1"
        )
        <
        content.index(
            "A2 | B2"
        )
        <
        content.index(
            "AFTER TABLE"
        )
    )


def test_whole_document_importer_does_not_split_sections():
    stream = BytesIO()

    document = Document()

    document.add_paragraph(
        "I. MUC TIEU"
    )

    document.add_paragraph(
        "Objective content"
    )

    document.add_paragraph(
        "II. THIET BI VA HOC LIEU"
    )

    document.add_paragraph(
        "Material content"
    )

    document.add_paragraph(
        "III. TIEN TRINH DAY HOC"
    )

    document.add_paragraph(
        "Teaching process content"
    )

    document.save(
        stream
    )

    content = (
        LessonPlanDocxWholeDocumentImporter()
        .import_bytes(
            stream.getvalue()
        )
    )

    assert "I. MUC TIEU" in content
    assert "Objective content" in content

    assert (
        "II. THIET BI VA HOC LIEU"
        in content
    )

    assert (
        "Material content"
        in content
    )

    assert (
        "III. TIEN TRINH DAY HOC"
        in content
    )

    assert (
        "Teaching process content"
        in content
    )


def test_whole_document_importer_rejects_empty_content():
    importer = (
        LessonPlanDocxWholeDocumentImporter()
    )

    try:
        importer.import_bytes(
            b""
        )
    except ValueError:
        return

    raise AssertionError(
        "Expected ValueError"
    )
