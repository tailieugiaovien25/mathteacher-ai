from docx import Document
from io import BytesIO
from contextlib import nullcontext
from datetime import date
from types import SimpleNamespace

import portal_v2.ui.weekly_schedule_streamlit as module

from document_intelligence.lesson_plan_workflow_state import (
    LessonPlanWorkflowState,
)
from educational_planning_v2.models.weekly_teaching_schedule import (
    TeachingSession,
)
from portal_v2.ui.weekly_schedule_portal import (
    WeeklySchedulePortalPreviewRow,
)


def make_docx_bytes(
    marker="lesson-plan-content",
):
    document = Document()

    document.add_paragraph(
        "I. M?C TI?U"
    )

    document.add_paragraph(
        marker
    )

    document.add_paragraph(
        "II. THI?T B? V? H?C LI?U"
    )

    document.add_paragraph(
        "M?y chi?u"
    )

    document.add_paragraph(
        "III. TI?N TR?NH D?Y H?C"
    )

    document.add_paragraph(
        "Ho?t ??ng 1"
    )

    stream = BytesIO()

    document.save(
        stream
    )

    return stream.getvalue()




class FakeUpload:
    def __init__(
        self,
        *,
        name: str,
        content: bytes,
    ):
        self.content = content
        self.name = name
        self._content = content

    def getvalue(self):
        return self._content




class _FakeColumn:
    def __enter__(self):
        return self

    def __exit__(
        self,
        exc_type,
        exc,
        traceback,
    ):
        return False

    def markdown(
        self,
        *args,
        **kwargs,
    ):
        return None

    def caption(
        self,
        *args,
        **kwargs,
    ):
        return None

    def write(
        self,
        *args,
        **kwargs,
    ):
        return None

    def metric(
        self,
        *args,
        **kwargs,
    ):
        return None





class _FakeComponentsV1:
    def html(
        self,
        *args,
        **kwargs,
    ):
        return None


class _FakeComponents:
    def __init__(self):
        self.v1 = _FakeComponentsV1()


class FakeStreamlit:
    def __init__(
        self,
        *,
        uploaded,
        process_clicked=True,
    ):
        self.components = _FakeComponents()
        self.uploaded = uploaded
        self.process_clicked = process_clicked
        self.session_state = {}

        # C?c ki?m th? g?i tr?c ti?p workspace n?n ph?i m? ph?ng
        # tr?ng th?i ???c t?o sau khi gi?o vi?n b?m n?t x?c nh?n.
        if process_clicked:
            self.session_state[
                "lesson_plan_standardization_confirmed_options"
            ] = module.LessonPlanStandardizationOptions()
            self.session_state[
                "lesson_plan_standardization_execute_requested"
            ] = True

        self.successes = []
        self.warnings = []
        self.errors = []
        self.downloads = []

    def divider(self):
        pass

    def subheader(self, value):
        pass

    def caption(self, value):
        pass

    def selectbox(
        self,
        label,
        *,
        options,
        format_func,
        key,
    ):
        return options[0]

    def info(self, value):
        pass

    def date_input(
        self,
        label,
        *,
        value,
        max_value=None,
        key,
    ):
        return value

    def file_uploader(
        self,
        label,
        *,
        type,
        accept_multiple_files,
        key,
    ):
        return self.uploaded

    def success(self, value):
        self.successes.append(value)

    def warning(self, value):
        self.warnings.append(value)

    def error(self, value):
        self.errors.append(value)

    def button(
        self,
        label,
        *,
        type,
        width=None,
        key,
        **kwargs,
    ):
        # The real workspace now contains several independent
        # buttons.  process_clicked represents only the
        # standardization/process action used by these E2E
        # workflow tests.
        # The production workspace now exposes a dedicated
        # five-action standardization toolbar.  Keep support
        # for the legacy process key while treating the real
        # "T?o gi?o ?n chu?n" action as the process trigger.
        button_key = str(
            key
            or ""
        )

        is_process_button = (
            button_key.startswith(
                "lbg_lesson_plan_process_"
            )
            or button_key
            == "standardization_action_create"
            or button_key
            == "standardization_control_panel_confirm"
        )

        if (
            is_process_button
            and self.process_clicked
        ):
            callback = kwargs.get(
                "on_click"
            )

            if callable(callback):
                callback(
                    *kwargs.get(
                        "args",
                        (),
                    ),
                    **kwargs.get(
                        "kwargs",
                        {},
                    ),
                )

            return True

        return False



    def spinner(self, value):
        return nullcontext()

    def download_button(
        self,
        label,
        *,
        data,
        file_name,
        mime,
        width,
        key,
    ):
        self.downloads.append(
            {
                "label": label,
                "data": data,
                "file_name": file_name,
                "mime": mime,
                "key": key,
            }
        )


    def columns(self, count):
        return tuple(
            _FakeColumn()
            for _ in range(count)
        )

    def markdown(
        self,
        *args,
        **kwargs,
    ):
        return None

    def write(
        self,
        *args,
        **kwargs,
    ):
        return None

    def radio(
        self,
        label,
        options,
        horizontal=False,
        key=None,
        **kwargs,
    ):
        """
        Minimal Streamlit radio test double.

        Existing workflow tests exercise the upload path,
        so default to the first available option.
        """
        values = tuple(
            options
        )

        if not values:
            return None

        return values[0]

    def text_input(
        self,
        label,
        *,
        value="",
        key=None,
        **kwargs,
    ):
        return value
    def number_input(
        self,
        label,
        *,
        value=0,
        key=None,
        **kwargs,
    ):
        return value

    def expander(
        self,
        label,
        *,
        expanded=False,
        **kwargs,
    ):
        return self

    def __enter__(self):
        return self


    def __exit__(
        self,
        exc_type,
        exc_value,
        traceback,
    ):
        return False







def make_row():
    return WeeklySchedulePortalPreviewRow(
        teaching_date=date(
            2026,
            9,
            28,
        ),
        weekday=1,
        timetable_period=2,
        session=TeachingSession.MORNING,
        class_id="8A1",
        subject_ref="TOAN",
        component_ref=None,
        curriculum_period=9,
        lesson_id="TOAN8-009",
        lesson_title="Đơn thức",
        period_in_lesson=1,
        teaching_equipment=(),
    )


def make_view():
    return SimpleNamespace(
        week_number=5,
        rows=(
            make_row(),
        ),
    )


def test_teacher_workflow_happy_path_reaches_download(
    monkeypatch,
):
    uploaded = FakeUpload(
        name="lesson.docx",
        content=make_docx_bytes("lesson-plan-content"),
    )

    st = FakeStreamlit(
        uploaded=uploaded,
        process_clicked=True,
    )

    preview = object()
    review = object()

    resolution = SimpleNamespace(
        accepted=True,
        metadata=SimpleNamespace(
            values=(
                (
                    module.DocumentField.CLASS_NAME,
                    "8A1",
                ),
                (
                    module.DocumentField.CURRICULUM_PERIOD,
                    "9",
                ),
                (
                    module.DocumentField.LESSON_TITLE,
                    "Đơn thức",
                ),
                (
                    module.DocumentField.DRAFTING_DATE,
                    "28/09/2026",
                ),
                (
                    module.DocumentField.TEACHING_DATE,
                    "28/09/2026",
                ),
            )
        ),
    )

    class FakePreviewService:
        def prepare(
            self,
            *,
            content,
            canonical,
        ):
            assert content == uploaded.content
            assert canonical.class_name == "8A1"
            assert canonical.curriculum_period == 9
            assert canonical.lesson_title == "Đơn thức"
            return preview

    class FakePresenter:
        def present(
            self,
            *,
            preview,
            canonical_values,
        ):
            return object()

    class FakeResolver:
        def resolve(
            self,
            *,
            preview,
            review,
        ):
            return resolution

    processed = []

    def fake_process(
        *,
        row,
        drafting_date,
        content,
        original_name,
        modification_plan=None,
        options=None,
        original_content=None,
        ai_revised_text="",
    ):
        processed.append(
            {
                "row": row,
                "drafting_date": drafting_date,
                "content": content,
                "original_name": original_name,
            }
        )

        return (
            "lesson-standardized.docx",
            b"standardized-docx",
            (),
        )

    monkeypatch.setattr(
        module,
        "st",
        st,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanPreviewUploadService",
        FakePreviewService,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_preview",
        lambda **kwargs: None,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewPresenter",
        FakePresenter,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_teacher_review",
        lambda **kwargs: review,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewResolver",
        FakeResolver,
    )

    class FakeModificationPlanner:
        def build(
            self,
            *,
            resolution,
        ):
            return object()

        def build_from_values(
            self,
            *,
            values,
        ):
            return object()

    monkeypatch.setattr(
        module,
        "LessonPlanModificationPlanner",
        FakeModificationPlanner,
    )

    monkeypatch.setattr(
        module,
        "_process_lesson_plan_upload",
        fake_process,
    )

    module._render_lesson_plan_standardization_workspace(
        make_view(),
        workspace_focus="STANDARDIZE",
    )

    assert st.errors == []

    assert len(processed) == 1

    assert (
        processed[0]["content"]
        == uploaded.content
    )

    assert (
        processed[0]["original_name"]
        == "lesson.docx"
    )

    assert (
        processed[0]["row"].class_id
        == "8A1"
    )

    assert (
        processed[0]["row"].lesson_title
        == "Đơn thức"
    )

    assert len(st.downloads) == 1

    assert (
        st.downloads[0]["file_name"]
        == "lesson-standardized.docx"
    )

    assert (
        st.downloads[0]["data"]
        == b"standardized-docx"
    )

    workflow_states = [
        value
        for value
        in st.session_state.values()
        if isinstance(
            value,
            LessonPlanWorkflowState,
        )
    ]

    assert len(workflow_states) == 1

    state = workflow_states[0]

    assert state.result == (
        "lesson-standardized.docx",
        b"standardized-docx",
        (),
    )


def test_teacher_override_reaches_processing_row(
    monkeypatch,
):
    uploaded = FakeUpload(
        name="lesson.docx",
        content=make_docx_bytes("lesson-plan-override"),
    )

    st = FakeStreamlit(
        uploaded=uploaded,
        process_clicked=True,
    )

    preview = object()
    review = object()

    override_title = (
        "\u0110\u01a1n th\u1ee9c m\u1edbi"
    )

    resolution = SimpleNamespace(
        accepted=True,
        metadata=SimpleNamespace(
            values=(
                (
                    module.DocumentField.CLASS_NAME,
                    "8A1",
                ),
                (
                    module.DocumentField.CURRICULUM_PERIOD,
                    "9",
                ),
                (
                    module.DocumentField.LESSON_TITLE,
                    override_title,
                ),
                (
                    module.DocumentField.DRAFTING_DATE,
                    "28/09/2026",
                ),
                (
                    module.DocumentField.TEACHING_DATE,
                    "28/09/2026",
                ),
            )
        ),
    )

    class FakePreviewService:
        def prepare(
            self,
            *,
            content,
            canonical,
        ):
            assert (
                canonical.lesson_title
                == "\u0110\u01a1n th\u1ee9c"
            )

            return preview

    class FakePresenter:
        def present(
            self,
            *,
            preview,
            canonical_values,
        ):
            assert (
                canonical_values[
                    module.DocumentField.LESSON_TITLE
                ]
                == "\u0110\u01a1n th\u1ee9c"
            )

            return object()

    class FakeResolver:
        def resolve(
            self,
            *,
            preview,
            review,
        ):
            return resolution

    processed = []

    def fake_process(
        *,
        row,
        drafting_date,
        content,
        original_name,
        modification_plan=None,
        options=None,
        original_content=None,
        ai_revised_text="",
    ):
        processed.append(row)

        return (
            "lesson-override-standardized.docx",
            b"override-standardized-docx",
            (),
        )

    monkeypatch.setattr(
        module,
        "st",
        st,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanPreviewUploadService",
        FakePreviewService,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_preview",
        lambda **kwargs: None,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewPresenter",
        FakePresenter,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_teacher_review",
        lambda **kwargs: review,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewResolver",
        FakeResolver,
    )

    class FakeModificationPlanner:
        def build(
            self,
            *,
            resolution,
        ):
            return object()

        def build_from_values(
            self,
            *,
            values,
        ):
            return object()

    monkeypatch.setattr(
        module,
        "LessonPlanModificationPlanner",
        FakeModificationPlanner,
    )

    monkeypatch.setattr(
        module,
        "_process_lesson_plan_upload",
        fake_process,
    )

    module._render_lesson_plan_standardization_workspace(
        make_view(),
        workspace_focus="STANDARDIZE",
    )

    assert st.errors == []

    assert len(processed) == 1

    reviewed_row = processed[0]

    assert (
        reviewed_row.lesson_title
        == "\u0110\u01a1n th\u1ee9c"
    )


    assert reviewed_row.class_id == "8A1"

    assert (
        reviewed_row.curriculum_period
        == 9
    )

    assert len(st.downloads) == 1

    assert (
        st.downloads[0]["file_name"]
        == "lesson-override-standardized.docx"
    )

    assert (
        st.downloads[0]["data"]
        == b"override-standardized-docx"
    )


def test_legacy_rejected_review_does_not_block_direct_canonical_processing(
    monkeypatch,
):
    uploaded = FakeUpload(
        name="lesson.docx",
        content=make_docx_bytes("lesson-plan-rejected"),
    )

    st = FakeStreamlit(
        uploaded=uploaded,
        process_clicked=True,
    )

    preview = object()
    review = object()

    resolution = SimpleNamespace(
        accepted=False,
        metadata=SimpleNamespace(
            values=()
        ),
    )

    class FakePreviewService:
        def prepare(
            self,
            *,
            content,
            canonical,
        ):
            return preview

    class FakePresenter:
        def present(
            self,
            *,
            preview,
            canonical_values,
        ):
            return object()

    class FakeResolver:
        def resolve(
            self,
            *,
            preview,
            review,
        ):
            return resolution

    process_calls = []

    def fake_process(
        *,
        row,
        drafting_date,
        content,
        original_name,
        modification_plan=None,
        options=None,
        original_content=None,
        ai_revised_text="",
    ):
        process_calls.append(
            {
                "row": row,
                "drafting_date": drafting_date,
                "content": content,
                "original_name": original_name,
            }
        )

        return (
            "lesson-rejected-standardized.docx",
            b"rejected-standardized-docx",
            (),
        )

    monkeypatch.setattr(
        module,
        "st",
        st,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanPreviewUploadService",
        FakePreviewService,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_preview",
        lambda **kwargs: None,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewPresenter",
        FakePresenter,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_teacher_review",
        lambda **kwargs: review,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewResolver",
        FakeResolver,
    )

    class FakeModificationPlanner:
        def build(
            self,
            *,
            resolution,
        ):
            return object()

        def build_from_values(
            self,
            *,
            values,
        ):
            return object()

    monkeypatch.setattr(
        module,
        "LessonPlanModificationPlanner",
        FakeModificationPlanner,
    )

    monkeypatch.setattr(
        module,
        "_process_lesson_plan_upload",
        fake_process,
    )

    module._render_lesson_plan_standardization_workspace(
        make_view(),
        workspace_focus="STANDARDIZE",
    )

    assert len(process_calls) == 1

    assert len(st.downloads) == 1

    assert st.errors == []

    workflow_states = [
        value
        for value
        in st.session_state.values()
        if isinstance(
            value,
            LessonPlanWorkflowState,
        )
    ]

    assert len(workflow_states) == 1

    state = workflow_states[0]

    assert state.result == (
        "lesson-rejected-standardized.docx",
        b"rejected-standardized-docx",
        (),
    )


def test_rerun_preserves_result_without_reprocessing(
    monkeypatch,
):
    uploaded = FakeUpload(
        name="lesson.docx",
        content=make_docx_bytes("lesson-plan-rerun"),
    )

    st = FakeStreamlit(
        uploaded=uploaded,
        process_clicked=True,
    )

    preview = object()
    review = object()

    resolution = SimpleNamespace(
        accepted=True,
        metadata=SimpleNamespace(
            values=(
                (
                    module.DocumentField.CLASS_NAME,
                    "8A1",
                ),
                (
                    module.DocumentField.CURRICULUM_PERIOD,
                    "9",
                ),
                (
                    module.DocumentField.LESSON_TITLE,
                    "\u0110\u01a1n th\u1ee9c",
                ),
                (
                    module.DocumentField.DRAFTING_DATE,
                    "28/09/2026",
                ),
                (
                    module.DocumentField.TEACHING_DATE,
                    "28/09/2026",
                ),
            )
        ),
    )

    preview_calls = []
    process_calls = []

    class FakePreviewService:
        def prepare(
            self,
            *,
            content,
            canonical,
        ):
            preview_calls.append(content)
            return preview

    class FakePresenter:
        def present(
            self,
            *,
            preview,
            canonical_values,
        ):
            return object()

    class FakeResolver:
        def resolve(
            self,
            *,
            preview,
            review,
        ):
            return resolution

    def fake_process(
        *,
        row,
        drafting_date,
        content,
        original_name,
        modification_plan=None,
        options=None,
        original_content=None,
        ai_revised_text="",
    ):
        process_calls.append(
            {
                "row": row,
                "content": content,
                "original_name": original_name,
            }
        )

        return (
            "lesson-rerun-standardized.docx",
            b"rerun-standardized-docx",
            (),
        )

    monkeypatch.setattr(
        module,
        "st",
        st,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanPreviewUploadService",
        FakePreviewService,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_preview",
        lambda **kwargs: None,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewPresenter",
        FakePresenter,
    )

    monkeypatch.setattr(
        module,
        "render_lesson_plan_teacher_review",
        lambda **kwargs: review,
    )

    monkeypatch.setattr(
        module,
        "LessonPlanTeacherReviewResolver",
        FakeResolver,
    )

    class FakeModificationPlanner:
        def build(
            self,
            *,
            resolution,
        ):
            return object()

        def build_from_values(
            self,
            *,
            values,
        ):
            return object()

    monkeypatch.setattr(
        module,
        "LessonPlanModificationPlanner",
        FakeModificationPlanner,
    )

    monkeypatch.setattr(
        module,
        "_process_lesson_plan_upload",
        fake_process,
    )

    # First render: process and store result.
    module._render_lesson_plan_standardization_workspace(
        make_view(),
        workspace_focus="STANDARDIZE",
    )

    assert len(process_calls) == 1
    assert len(st.downloads) == 1

    first_state_values = [
        value
        for value
        in st.session_state.values()
        if isinstance(
            value,
            LessonPlanWorkflowState,
        )
    ]

    assert len(first_state_values) == 1

    first_state = first_state_values[0]

    assert first_state.result == (
        "lesson-rerun-standardized.docx",
        b"rerun-standardized-docx",
        (),
    )

    # Simulate Streamlit rerun without clicking process again.
    st.process_clicked = False
    st.downloads.clear()

    module._render_lesson_plan_standardization_workspace(
        make_view(),
        workspace_focus="STANDARDIZE",
    )

    # Canonical preparation may run again during a Streamlit
    # rerender, but expensive document processing must not run
    # again when a completed workflow result already exists.
    assert len(process_calls) == 1

    # Existing result remains downloadable.
    assert len(st.downloads) == 1

    assert (
        st.downloads[0]["file_name"]
        == "lesson-rerun-standardized.docx"
    )

    assert (
        st.downloads[0]["data"]
        == b"rerun-standardized-docx"
    )

    second_state_values = [
        value
        for value
        in st.session_state.values()
        if isinstance(
            value,
            LessonPlanWorkflowState,
        )
    ]

    assert len(second_state_values) == 1

    second_state = second_state_values[0]

    assert (
        second_state.result
        == first_state.result
    )
