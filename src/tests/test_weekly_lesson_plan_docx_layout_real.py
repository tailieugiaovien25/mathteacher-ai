from docx import Document

from lesson_planning_v2.services.weekly_lesson_plan_docx_renderer import (
    WeeklyLessonPlanDocxRenderer,
)
from lesson_planning_v2.weekly_lesson_plan_docx_layout import (
    WeeklyLessonPlanDocxLayoutProfile,
)

from test_weekly_lesson_plan_docx_display_e2e import (
    _document,
    _presentation,
)


def _cm(value):
    return round(value.cm, 2)


def test_real_docx_uses_default_tnr_14(
    tmp_path,
):
    output = tmp_path / "default-layout.docx"

    WeeklyLessonPlanDocxRenderer().render(
        document=_document(),
        output_path=output,
        presentation_profile=_presentation(),
        layout_profile=(
            WeeklyLessonPlanDocxLayoutProfile.default()
        ),
    )

    loaded = Document(output)
    normal = loaded.styles["Normal"]

    assert normal.font.name == "Times New Roman"
    assert normal.font.size.pt == 14


def test_real_docx_uses_default_margins(
    tmp_path,
):
    output = tmp_path / "default-margins.docx"

    WeeklyLessonPlanDocxRenderer().render(
        document=_document(),
        output_path=output,
        presentation_profile=_presentation(),
        layout_profile=(
            WeeklyLessonPlanDocxLayoutProfile.default()
        ),
    )

    loaded = Document(output)
    section = loaded.sections[0]

    assert _cm(section.top_margin) == 2.0
    assert _cm(section.bottom_margin) == 2.0
    assert _cm(section.left_margin) == 3.0
    assert _cm(section.right_margin) == 2.0


def test_custom_layout_changes_real_docx(
    tmp_path,
):
    output = tmp_path / "custom-layout.docx"

    profile = WeeklyLessonPlanDocxLayoutProfile(
        body_font="Arial",
        body_size=12,
        title_size=17,
        heading_size=13,
        top_margin_cm=1.5,
        bottom_margin_cm=1.5,
        left_margin_cm=2.5,
        right_margin_cm=1.5,
        line_spacing=1.0,
        space_before_pt=0,
        space_after_pt=3,
        header_alignment="center",
        approval_alignment="right",
    )

    WeeklyLessonPlanDocxRenderer().render(
        document=_document(),
        output_path=output,
        presentation_profile=_presentation(),
        layout_profile=profile,
    )

    loaded = Document(output)
    normal = loaded.styles["Normal"]
    section = loaded.sections[0]

    assert normal.font.name == "Arial"
    assert normal.font.size.pt == 12

    assert _cm(section.top_margin) == 1.5
    assert _cm(section.bottom_margin) == 1.5
    assert _cm(section.left_margin) == 2.5
    assert _cm(section.right_margin) == 1.5
