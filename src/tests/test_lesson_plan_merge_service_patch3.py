from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path
import zipfile

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from lesson_planning_v2.services.lesson_plan_merge_service import (
    LessonPlanMergeError,
    LessonPlanMergeService,
    LessonPlanMergeSource,
)


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAusB9Y9Z1Z8AAAAASUVORK5CYII="
)


def _add_formula(paragraph) -> None:
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = "x+1=2"
    run.append(text)
    math.append(run)
    paragraph._p.append(math)


def _make_docx(
    *,
    title: str,
    approval: str | None = "TỔ CHUYÊN MÔN DUYỆT",
    with_image: bool = False,
    with_table: bool = False,
    with_formula: bool = False,
) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph("Nội dung " + title)

    if with_table:
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = title
        table.cell(1, 1).text = "Giữ bảng"

    if with_image:
        document.add_picture(
            BytesIO(_PNG_1X1),
            width=Inches(0.2),
        )

    if with_formula:
        paragraph = document.add_paragraph()
        _add_formula(paragraph)

    if approval is not None:
        document.add_paragraph(
            "Ngày 20 tháng 8 năm 2026"
        )
        document.add_paragraph(approval)
        document.add_paragraph("Tổ trưởng")
        document.add_paragraph("")
        document.add_paragraph("Nguyễn Văn A")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _source(source_id: str, title: str, **kwargs):
    return LessonPlanMergeSource(
        source_id=source_id,
        file_name=source_id + ".docx",
        content=_make_docx(
            title=title,
            **kwargs,
        ),
    )


def _document_xml(content: bytes) -> str:
    with zipfile.ZipFile(BytesIO(content)) as archive:
        return archive.read(
            "word/document.xml"
        ).decode("utf-8")


def test_merge_preserves_source_order_and_provenance():
    service = LessonPlanMergeService()

    result = service.merge(
        [
            _source("a", "Bài A"),
            _source("b", "Bài B"),
            _source("c", "Bài C"),
        ]
    )

    document = Document(BytesIO(result.content))
    text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    assert text.index("Bài A") < text.index("Bài B")
    assert text.index("Bài B") < text.index("Bài C")

    assert result.source_ids == ("a", "b", "c")
    assert result.source_file_names == (
        "a.docx",
        "b.docx",
        "c.docx",
    )


def test_merge_keeps_exactly_one_final_approval_block():
    service = LessonPlanMergeService()

    result = service.merge(
        [
            _source("a", "Bài A"),
            _source("b", "Bài B"),
            _source("c", "Bài C"),
        ]
    )

    document = Document(BytesIO(result.content))
    text = "\n".join(
        paragraph.text.upper()
        for paragraph in document.paragraphs
    )

    assert text.count("TỔ CHUYÊN MÔN DUYỆT") == 1
    assert text.count("NGÀY 20 THÁNG 8 NĂM 2026") == 1
    assert text.rstrip().endswith("NGUYỄN VĂN A")


def test_merge_preserves_tables_images_and_formula_xml():
    service = LessonPlanMergeService()

    result = service.merge(
        [
            _source(
                "a",
                "Bài có bảng",
                with_table=True,
            ),
            _source(
                "b",
                "Bài có ảnh",
                with_image=True,
            ),
            _source(
                "c",
                "Bài có công thức",
                with_formula=True,
            ),
        ]
    )

    document = Document(BytesIO(result.content))

    assert len(document.tables) == 1
    assert document.tables[0].cell(1, 1).text == "Giữ bảng"
    assert len(document.inline_shapes) == 1

    xml = _document_xml(result.content)
    assert "<m:oMath" in xml
    assert "x+1=2" in xml


def test_merge_accepts_short_approval_marker_variant():
    service = LessonPlanMergeService()

    result = service.merge(
        [
            _source(
                "a",
                "Bài A",
                approval="Tổ CM duyệt",
            ),
            _source(
                "b",
                "Bài B",
                approval="Tổ CM duyệt",
            ),
        ]
    )

    document = Document(BytesIO(result.content))
    text = "\n".join(
        paragraph.text.casefold()
        for paragraph in document.paragraphs
    )

    assert text.count("tổ cm duyệt".casefold()) == 1


def test_merge_requires_final_approval_block():
    service = LessonPlanMergeService()

    with pytest.raises(
        LessonPlanMergeError,
        match="final lesson plan",
    ):
        service.merge(
            [
                _source("a", "Bài A"),
                _source(
                    "b",
                    "Bài B",
                    approval=None,
                ),
            ]
        )


def test_merge_rejects_duplicate_source_ids():
    service = LessonPlanMergeService()

    with pytest.raises(
        LessonPlanMergeError,
        match="Duplicate source_id",
    ):
        service.merge(
            [
                _source("same", "Bài A"),
                _source("same", "Bài B"),
            ]
        )


def test_merge_does_not_generate_consecutive_trailing_blank_paragraphs():
    service = LessonPlanMergeService()

    result = service.merge(
        [
            _source("a", "Bài A"),
            _source("b", "Bài B"),
        ]
    )

    document = Document(BytesIO(result.content))
    texts = [
        paragraph.text
        for paragraph in document.paragraphs
    ]

    trailing_blank_count = 0
    for text in reversed(texts):
        if text.strip():
            break
        trailing_blank_count += 1

    assert trailing_blank_count == 0
