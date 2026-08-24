from datetime import date

from docx import Document

from educational_planning_v2.models import TeachingSession
from lesson_planning_v2.contexts import ScheduledLessonContext
from document_standardization import (
    LessonPlanDocumentContextApplier,
)


def make_context():
    return ScheduledLessonContext(
        teaching_date=date(2026, 9, 8),
        drafting_date=date(2026, 9, 7),
        class_id="8A2",
        subject_ref="MATHEMATICS",
        component_ref="ALGEBRA",
        curriculum_period=1,
        lesson_id="LESSON-001",
        lesson_title="Đơn thức",
        session=TeachingSession.MORNING,
        timetable_period=1,
        period_in_lesson=1,
    )


def test_applier_supports_real_lesson_plan_layout(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()

    document.add_paragraph(
        "Ngày soạn: 01/01/2020    "
        "Ngày dạy: 02/01/2020    "
        "Lớp: 8A1"
    )

    document.add_paragraph(
        "Tiết 99. BÀI 99. BÀI CŨ"
    )

    document.add_paragraph(
        "I. MỤC TIÊU"
    )

    document.add_paragraph(
        "Nội dung bài học phải được giữ nguyên."
    )

    document.save(source)

    result = (
        LessonPlanDocumentContextApplier()
        .apply(
            source,
            output,
            make_context(),
        )
    )

    updated = Document(output)

    texts = [
        paragraph.text
        for paragraph in updated.paragraphs
    ]

    combined = "\n".join(texts)

    assert "07/09/2026" in combined
    assert "08/09/2026" in combined
    assert "8A2" in combined

    assert (
        "Tiết 1" in combined
        or "TIẾT 1" in combined
    )

    assert "ĐƠN THỨC" in combined.upper()

    assert (
        "Nội dung bài học phải được giữ nguyên."
        in combined
    )

    assert result.unresolved_fields == ()


def test_real_layout_keeps_unrelated_content(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()

    document.add_paragraph(
        "Ngày soạn: 01/01/2020    "
        "Ngày dạy: 02/01/2020    "
        "Lớp: 8A1"
    )

    document.add_paragraph(
        "Tiết 99. BÀI 99. BÀI CŨ"
    )

    original_content = (
        "Hoạt động 1: Giáo viên tổ chức "
        "cho học sinh thực hiện nhiệm vụ."
    )

    document.add_paragraph(original_content)

    document.save(source)

    LessonPlanDocumentContextApplier().apply(
        source,
        output,
        make_context(),
    )

    updated = Document(output)

    texts = [
        paragraph.text
        for paragraph in updated.paragraphs
    ]

    assert original_content in texts
