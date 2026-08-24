
from datetime import date

from docx import Document

from document_standardization import (
    LessonPlanDocumentContextApplier,
)
from educational_planning_v2.models import (
    TeachingSession,
)
from lesson_planning_v2.contexts import (
    ScheduledLessonContext,
)


def _context():
    return ScheduledLessonContext(
        teaching_date=date(
            2026,
            9,
            29,
        ),
        drafting_date=date(
            2026,
            9,
            28,
        ),
        class_id="7A2",
        subject_ref="MATHEMATICS",
        component_ref="ALGEBRA",
        curriculum_period=9,
        lesson_id="LESSON-009",
        lesson_title="??n th?c",
        session=TeachingSession.MORNING,
        timetable_period=1,
        period_in_lesson=1,
    )


def test_multiclass_teaching_date_overwrites_only_selected_class(
    tmp_path,
):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()

    document.add_paragraph(
        "Ng?y so?n: 01/01/2020"
    )

    p = document.add_paragraph()

    label = p.add_run(
        "Ng?y d?y:\n"
    )
    label.bold = True

    class_1 = p.add_run(
        "7A1 - 02/10/2025\n"
    )
    class_1.italic = True

    class_2 = p.add_run(
        "7A2 - 30/09/2025"
    )
    class_2.italic = True

    document.add_paragraph(
        "L?p: 7A2"
    )

    document.add_paragraph(
        "Ti?t PPCT: 1"
    )

    document.add_paragraph(
        "B?i: B?i c?"
    )

    document.add_paragraph(
        "N?i dung b?i h?c gi? nguy?n."
    )

    document.save(source)

    LessonPlanDocumentContextApplier().apply(
        source,
        output,
        _context(),
    )

    processed = Document(output)

    combined = "\n".join(
        paragraph.text
        for paragraph
        in processed.paragraphs
    )

    assert (
        "7A1 - 02/10/2025"
        in combined
    )

    assert (
        "7A2 - 29/09/2026"
        in combined
    )

    assert (
        "7A2 - 30/09/2025"
        not in combined
    )

    assert (
        "N?i dung b?i h?c gi? nguy?n."
        in combined
    )

    # Source file must remain untouched.
    original = Document(source)

    original_text = "\n".join(
        paragraph.text
        for paragraph
        in original.paragraphs
    )

    assert (
        "7A2 - 30/09/2025"
        in original_text
    )
